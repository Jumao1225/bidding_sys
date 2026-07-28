import glob
import io
import re
import sys
from docx import Document
from app.db.session import SessionLocal
from app.services.bid_format_filler_service import bid_format_filler_service
from app.agents.bid_filler_agent import bid_filler_agent
from app.schemas.bid_filler_schema import CompanyProfile

sys.stdout.reconfigure(encoding='utf-8')

doc_files = glob.glob('d:/Myproject/bidding_sys/backend/uploads/*.docx')
if not doc_files:
    print("No docx files found")
    exit(0)

target_file = doc_files[0]
print(f"Testing on file: {target_file}")

with open(target_file, "rb") as f:
    docx_bytes = f.read()

detected = bid_format_filler_service.scan_detected_placeholders(docx_bytes)
print(f"Detected placeholders count: {len(detected)}")
for item in detected:
    print("  Detected:", item)

db = SessionLocal()
try:
    profile = CompanyProfile()
    replacement_map, report = bid_filler_agent.process_filling_tasks(
        db=db,
        document_id="test-doc",
        profile=profile,
        detected_placeholders=detected
    )
    print("\nReplacement Map:")
    for k, v in replacement_map.items():
        print(f"  {k} -> {v}")

    filled_bytes = bid_format_filler_service.fill_docx_with_audit_trail(
        docx_bytes=docx_bytes,
        replacement_map=replacement_map,
        audit_items=report.audit_items
    )
    doc = Document(io.BytesIO(filled_bytes))

    print("\n=== Remaining unfilled placeholders/underscores in Paragraphs ===")
    unfilled_count = 0
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if re.search(r'_{2,}|\[[^\]]+\]|［[^］]+］', txt):
            unfilled_count += 1
            print(f"L{idx:04d}: {txt}")

    print(f"\nTotal remaining unfilled paragraphs: {unfilled_count}")

    print("\n=== Remaining unfilled placeholders/underscores in Table Cells ===")
    table_unfilled_count = 0
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                txt = cell.text.strip()
                if re.search(r'_{2,}|\[[^\]]+\]|［[^］]+］', txt):
                    table_unfilled_count += 1
                    print(f"Table {t_idx} Row {r_idx} Col {c_idx}: {txt}")

    print(f"\nTotal remaining unfilled table cells: {table_unfilled_count}")

finally:
    db.close()
