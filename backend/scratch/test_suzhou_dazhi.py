import glob
import io
import re
import sys
from docx import Document
from app.db.session import SessionLocal
from app.services.bid_format_filler_service import bid_format_filler_service
from app.agents.bid_filler_agent import bid_filler_agent
from app.schemas.bid_filler_schema import CompanyProfile

sys.stdout.reconfigure(encoding='utf-8')

def custom_replace_in_place(p, replacement_map, has_underline, font_ctx):
    text = p.text
    if not text:
        return False

    # 1. 判断是否为日期模板段落 (如 "日    期：     年      月     日")
    if re.search(r'日\s*期\s*[:：]?\s*[\s_]*年[\s_]*月[\s_]*日', text):
        date_val = replacement_map.get("投标日期") or "2026年06月30日"
        m = re.search(r'(日\s*期\s*[:：]?\s*)', text)
        prefix = m.group(1) if m else "日    期："
        p._element.clear_content()
        r_pre = p.add_run(prefix)
        bid_format_filler_service._apply_run_style(r_pre, font_context=font_ctx, underline=False)
        r_val = p.add_run(date_val)
        bid_format_filler_service._apply_run_style(r_val, font_context=font_ctx, underline=True)
        return True

    # 2. 判断是否为短字段标签行 (如 "地    址："、"电    话："、"投标单位名称：")
    txt_stripped = text.strip()
    is_label_line = (
        len(txt_stripped) < 25 and 
        txt_stripped.endswith((':', '：')) and 
        not txt_stripped.startswith(('1', '2', '3', '4', '5', '6', '7', '8', '9', '一', '二', '三', '四', '五', '致', '据此', '根据'))
    )

    # 查找匹配的占位符模式
    if is_label_line:
        placeholder_regex = re.compile(r'(_{2,}|\[[^\]]+\]|［[^］]+］|\s{2,}|(?<=[:：])\s*$)')
    else:
        placeholder_regex = re.compile(r'(_{2,}|\[[^\]]+\]|［[^］]+］|\s{3,})')

    matches = list(placeholder_regex.finditer(text))
    if not matches:
        return False

    p_name = replacement_map.get("项目名称", "")
    p_code = replacement_map.get("项目编号", "")
    c_name = replacement_map.get("投标人名称") or replacement_map.get("投标人") or ""
    legal_rep = replacement_map.get("法定代表人", "")
    delegate = replacement_map.get("授权代理人", "")
    credit_code = replacement_map.get("统一社会信用代码", "")
    address = replacement_map.get("地址", "")
    phone = replacement_map.get("电话", "")
    fax = replacement_map.get("传真", "")
    postcode = replacement_map.get("邮编", "")
    email = replacement_map.get("电子邮箱", "")
    bank_name = replacement_map.get("开户银行", "")
    bank_account = replacement_map.get("银行账号", "")
    id_card = replacement_map.get("身份证号码", "")
    period = replacement_map.get("工期", "")
    quality = replacement_map.get("质量标准", "")
    deposit = replacement_map.get("投标保证金", "")
    total_price_num = replacement_map.get("投标总价", "")
    total_price_words = replacement_map.get("投标总价大写", "")
    date_str = replacement_map.get("投标日期", "")

    last_idx = 0
    new_runs_data = []

    for m in matches:
        start, end = m.span()
        ph_str = m.group(0)
        before_text = text[last_idx:start]

        if before_text:
            new_runs_data.append((before_text, False))

        fill_val = ""
        ph_clean = ph_str.strip("[]［］_ ")

        # 结合前文与后文分析
        ctx_text = text[:start]
        ctx_clean = re.sub(r'\s+', '', ctx_text)
        after_text = text[end:]

        if re.search(r'招标文件', after_text):
            fill_val = p_code
        elif re.search(r'姓名和职务', after_text):
            fill_val = delegate or legal_rep or "张三"
        elif re.search(r'投标人的名称', after_text):
            fill_val = c_name
        elif re.search(r'(大写|人民币)', ctx_clean):
            fill_val = total_price_words
        elif re.search(r'(¥|￥|小写)', ctx_clean):
            fill_val = str(total_price_num) if total_price_num else ""
        elif re.search(r'(项目名称)', ctx_clean):
            fill_val = p_name
        elif re.search(r'(项目编号)', ctx_clean):
            fill_val = p_code
        elif re.search(r'(投标人|单位名称)', ctx_clean) and not re.search(r'^\s*致\s*[:：]', ctx_clean):
            fill_val = c_name
        elif re.search(r'(法定代表人|法人代表)', ctx_clean):
            fill_val = legal_rep
        elif re.search(r'(授权代理人|委托代理人|代表姓名)', ctx_clean):
            fill_val = delegate or legal_rep
        elif re.search(r'(开户银行|开户行)', ctx_clean):
            fill_val = bank_name
        elif re.search(r'(银行账号|账号)', ctx_clean):
            fill_val = bank_account
        elif re.search(r'(信用代码|注册号)', ctx_clean):
            fill_val = credit_code
        elif re.search(r'(地址|通信地址|注册地址)', ctx_clean):
            fill_val = address
        elif re.search(r'(电话|联系电话)', ctx_clean):
            fill_val = phone
        elif re.search(r'(传真)', ctx_clean):
            fill_val = fax
        elif re.search(r'(邮编|邮政编码)', ctx_clean):
            fill_val = postcode
        elif re.search(r'(邮箱|Email)', ctx_clean):
            fill_val = email
        elif re.search(r'(身份证)', ctx_clean):
            fill_val = id_card
        elif re.search(r'(工期|交货期)', ctx_clean):
            fill_val = period
        elif re.search(r'(质量)', ctx_clean):
            fill_val = quality
        elif re.search(r'(保证金)', ctx_clean):
            fill_val = deposit
        elif re.search(r'(日期|时间)', ctx_clean):
            fill_val = date_str

        if fill_val:
            new_runs_data.append((fill_val, True))
        else:
            new_runs_data.append((ph_str, False))

        last_idx = end

    if last_idx < len(text):
        new_runs_data.append((text[last_idx:], False))

    p._element.clear_content()
    for chunk_text, is_filled in new_runs_data:
        if not chunk_text:
            continue
        r = p.add_run(chunk_text)
        bid_format_filler_service._apply_run_style(
            r,
            font_context=font_ctx,
            underline=(has_underline or is_filled)
        )

    return True


fpath = 'd:/Myproject/bidding_sys/backend/uploads/08d69824-df66-444d-b132-248b73c3c806_（终）招标文件--和烁热能（400kW)光伏发电项目(1).docx'

with open(fpath, "rb") as f:
    docx_bytes = f.read()

db = SessionLocal()
try:
    detected = bid_format_filler_service.scan_detected_placeholders(docx_bytes)
    profile = CompanyProfile()
    replacement_map, report = bid_filler_agent.process_filling_tasks(
        db=db,
        document_id="test-doc",
        profile=profile,
        detected_placeholders=detected
    )

    doc = Document(io.BytesIO(docx_bytes))
    print("\n=== TESTING REFINED REPLACE IN '二、投标函格式' SECTION ===")
    in_section = False
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if "二、投标函格式" in txt:
            in_section = True
        if in_section and ("三、" in txt or ("3、" in txt and len(txt) < 15 and "授权" in txt)):
            in_section = False
            
        if in_section and txt:
            has_u = bid_format_filler_service._has_original_underline(p, txt)
            font_ctx = bid_format_filler_service._extract_font_context(p)
            
            p_copy = doc.add_paragraph()
            p_copy.text = txt
            replaced = custom_replace_in_place(p_copy, replacement_map, has_u, font_ctx)
            print(f"L{idx:04d} (Replaced={replaced}): {p_copy.text}")

finally:
    db.close()
