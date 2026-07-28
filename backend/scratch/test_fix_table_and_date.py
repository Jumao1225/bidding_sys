import io
import re
import sys
from docx import Document
from app.db.session import SessionLocal
from app.services.bid_format_filler_service import bid_format_filler_service
from app.agents.bid_filler_agent import bid_filler_agent
from app.schemas.bid_filler_schema import CompanyProfile

sys.stdout.reconfigure(encoding='utf-8')

fpath = 'd:/Myproject/bidding_sys/backend/uploads/08d69824-df66-444d-b132-248b73c3c806_（终）招标文件--和烁热能（400kW)光伏发电项目(1).docx'

with open(fpath, "rb") as f:
    docx_bytes = f.read()

db = SessionLocal()
try:
    detected = bid_format_filler_service.scan_detected_placeholders(docx_bytes)
    profile = CompanyProfile()
    replacement_map, report = bid_filler_agent.process_filling_tasks(
        db=db,
        document_id="test-doc",
        profile=profile,
        detected_placeholders=detected
    )

    doc = Document(io.BytesIO(docx_bytes))
    
    # 模拟填报测试
    filled_bytes = bid_format_filler_service.fill_docx_with_audit_trail(
        docx_bytes=docx_bytes,
        replacement_map=replacement_map,
        audit_items=[]
    )

    res_doc = Document(io.BytesIO(filled_bytes))

    print("=== INSPECTING FILLED '三、开标一览表' TABLE ===")
    for idx, t in enumerate(res_doc.tables):
        for r_idx, row in enumerate(t.rows):
            row_txt = [c.text.strip() for c in row.cells]
            print(f"Table {idx} Row {r_idx}: {row_txt}")

finally:
    db.close()
