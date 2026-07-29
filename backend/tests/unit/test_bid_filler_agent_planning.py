"""自动填报 Agent 原文理解与填报计划测试。"""

import io

from docx import Document

from app.agents.bid_filler_agent import bid_filler_agent
from app.schemas.bid_filler_schema import AgentFillPlanItem, BidFillPlan
from app.services.bid_format_filler_service import bid_format_filler_service


def _build_template_bytes() -> bytes:
    """构造包含正文和表格语境的最小 Word 模板。"""
    document = Document()
    document.add_paragraph("根据贵方的______号招标文件，我方项目名称为：[项目名称]。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "投标总价（大写）"
    table.cell(0, 1).text = "________________"
    table.cell(1, 0).text = "投标人名称"
    table.cell(1, 1).text = "________________"

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def test_extract_original_document_context_should_include_paragraph_and_table():
    """Agent 上下文必须同时包含正文整句和表格行列信息。"""
    context = bid_format_filler_service.extract_original_document_context(_build_template_bytes())

    assert "根据贵方的______号招标文件" in context
    assert "投标总价（大写）" in context
    assert "第 1 行" in context



def test_empty_company_profile_should_not_contain_demo_values():
    """缺少企业档案时不允许自动填入演示公司数据。"""
    from app.schemas.bid_filler_schema import CompanyProfile

    profile = CompanyProfile()

    assert profile.company_name == ""
    assert profile.legal_representative == ""
    assert profile.bank_account == ""
