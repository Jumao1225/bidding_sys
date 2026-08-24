"""
单元测试：Supervisor 表格深度质检与自愈管线 (test_supervisor_table_audit.py)
"""

import pytest
import os
import tempfile
from unittest.mock import patch, MagicMock
from docx import Document
from types import SimpleNamespace

from app.agents.bid_filler_agent import supervisor_audit_node, write_docx_node


def test_supervisor_table_audit_summary_table_auto_heal():
    """测试 Supervisor 深度表格质检对开标一览表大写金额留空的自动自愈能力"""
    # 创建一个模拟的 Word 文档，包含开标一览表
    doc = Document()
    doc.add_heading("二、开标一览表", level=1)
    
    # 模拟 3 行的开标一览表：表头 + 数据行 + 大写总报价行
    table = doc.add_table(rows=3, cols=4)
    # Row 1: 表头
    headers = ["项目名称", "技术要求", "总价（元）", "备注"]
    for c_i, h in enumerate(headers):
        table.rows[0].cells[c_i].text = h
    # Row 2: 数据行
    table.rows[1].cells[0].text = "某光伏采购项目"
    table.rows[1].cells[1].text = "详见采购参数"
    table.rows[1].cells[2].text = "1020000.00"
    table.rows[1].cells[3].text = "无"
    # Row 3: 原模板大写总报价行（右侧单元格留空）
    table.rows[2].cells[0].text = "投标总报价（大写）"
    table.rows[2].cells[1].text = ""  # 留空待自愈
    table.rows[2].cells[2].text = ""
    table.rows[2].cells[3].text = ""

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
        doc.save(tmp_path)

    try:
        # Mock 数据库查询 CostEstimate
        mock_items = [
            SimpleNamespace(calculated_total=1020000.0, manufacturer="隆基", brand="隆基")
        ]
        mock_db = MagicMock()
        mock_query = MagicMock()
        mock_query.filter.return_value.all.return_value = mock_items
        mock_db.query.return_value = mock_query

        with patch("app.agents.bid_filler_agent.SessionLocal", return_value=mock_db), \
             patch("app.services.office_cli_service.office_cli_service.query_structure", return_value="tbl[1]"):

            state = {
                "document_id": "test_doc_123",
                "docx_temp_path": tmp_path,
                "original_context": "",
                "repair_count": 0,
                "max_repair_rounds": 2,
                "repair_instructions_map": {},
                "audit_items": [],
            }

            res = supervisor_audit_node(state)

            # 验证 Word 文件已被自动自愈修复
            doc_healed = Document(tmp_path)
            healed_table = doc_healed.tables[0]
            # 第 3 行第 2 个单元格应已被原位填入大写金额
            caps_text = healed_table.rows[2].cells[1].text
            assert "壹佰零贰万元整" in caps_text or "壹" in caps_text

    finally:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass


def test_supervisor_table_audit_manufacturer_column_auto_heal():
    """测试 Supervisor 深度质检对分项报价表中生产厂家整列留白的自动自愈能力"""
    doc = Document()
    doc.add_heading("三、投标配置及分项报价表", level=1)

    # 模拟 3 行的分项报价表：表头 + 2 个数据行（生产厂家列留空）
    table = doc.add_table(rows=3, cols=6)
    headers = ["序号", "货物名称", "规格型号", "生产厂家", "数量", "合价"]
    for c_i, h in enumerate(headers):
        table.rows[0].cells[c_i].text = h

    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "光伏组件"
    table.rows[1].cells[2].text = "550W"
    table.rows[1].cells[3].text = ""  # 留空
    table.rows[1].cells[4].text = "800"
    table.rows[1].cells[5].text = "520000.00"

    table.rows[2].cells[0].text = "2"
    table.rows[2].cells[1].text = "逆变器"
    table.rows[2].cells[2].text = "110kW"
    table.rows[2].cells[3].text = ""  # 留空
    table.rows[2].cells[4].text = "4"
    table.rows[2].cells[5].text = "100000.00"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
        doc.save(tmp_path)

    try:
        mock_items = [
            SimpleNamespace(item_name="光伏组件", manufacturer="隆基乐叶", brand="隆基", calculated_total=520000.0),
            SimpleNamespace(item_name="逆变器", manufacturer="", brand="华为技术", calculated_total=100000.0),
        ]
        mock_db = MagicMock()
        mock_query = MagicMock()
        mock_query.filter.return_value.all.return_value = mock_items
        mock_db.query.return_value = mock_query

        with patch("app.agents.bid_filler_agent.SessionLocal", return_value=mock_db), \
             patch("app.services.office_cli_service.office_cli_service.query_structure", return_value="tbl[1]"):

            state = {
                "document_id": "test_doc_456",
                "docx_temp_path": tmp_path,
                "original_context": "",
                "repair_count": 0,
                "max_repair_rounds": 2,
                "repair_instructions_map": {},
                "audit_items": [],
            }

            res = supervisor_audit_node(state)

            doc_healed = Document(tmp_path)
            healed_table = doc_healed.tables[0]
            # 第 1 个数据行第 4 列应填入 "隆基乐叶"
            assert healed_table.rows[1].cells[3].text == "隆基乐叶"
            # 第 2 个数据行第 4 列应回溯填入 brand "华为技术"
            assert healed_table.rows[2].cells[3].text == "华为技术"

    finally:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass


def test_supervisor_paragraph_blank_slots_auto_heal():
    """测试 Supervisor 深度段落质检对地址/电话/全称/法人等留白段落的原位自愈能力，同时验证对公文导语/正文条款 100% 保护不误填"""
    doc = Document()
    doc.add_heading("二、投标函格式", level=1)
    doc.add_paragraph("致某招标代理机构：")
    doc.add_paragraph("据此函，签字人兹宣布同意如下：")
    doc.add_paragraph("5、与本投标有关的正式通讯地址为：")
    p_addr = doc.add_paragraph("地    址：                                ")
    p_phone = doc.add_paragraph("电    话：                                ")
    p_name = doc.add_paragraph("投标单位名称：                            ")
    p_rep = doc.add_paragraph("法定代表人或授权代表签字（或盖章）：                       年     月    日")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
        doc.save(tmp_path)

    try:
        mock_prof = SimpleNamespace(
            company_name="测试工程有限公司",
            registered_address="某省某市高新技术开发区100号",
            contact_phone="028-12345678",
            postal_code="610000",
            email="test@example.com",
            legal_representative="李四",
        )
        mock_tl = SimpleNamespace(
            planned_delivery_date="2026年06月30日"
        )
        mock_db = MagicMock()
        mock_query_prof = MagicMock()
        mock_query_prof.first.return_value = mock_prof

        mock_query_tl = MagicMock()
        mock_query_tl.filter.return_value.first.return_value = mock_tl

        def mock_query_router(model):
            from app.db.models.business import CompanyProfileModel
            from app.db.models.metadata import TimelineMetadata
            if model == CompanyProfileModel:
                return mock_query_prof
            elif model == TimelineMetadata:
                return mock_query_tl
            return MagicMock()

        mock_db.query.side_effect = mock_query_router

        with patch("app.agents.bid_filler_agent.SessionLocal", return_value=mock_db), \
             patch("app.services.office_cli_service.office_cli_service.query_structure", return_value=""):

            state = {
                "document_id": "test_doc_para_heal",
                "docx_temp_path": tmp_path,
                "original_context": "",
                "repair_count": 0,
                "max_repair_rounds": 2,
                "repair_instructions_map": {},
                "audit_items": [],
            }

            res = supervisor_audit_node(state)

            doc_healed = Document(tmp_path)
            all_text = "\n".join(p.text for p in doc_healed.paragraphs)

            # 1. 验证致函单位、声明引导句、条款导语等固定正文 100% 原样保留，绝无被误填人名或地址
            assert doc_healed.paragraphs[1].text.strip() == "致某招标代理机构："
            assert doc_healed.paragraphs[2].text.strip() == "据此函，签字人兹宣布同意如下："
            assert doc_healed.paragraphs[3].text.strip() == "5、与本投标有关的正式通讯地址为："
            assert "李四" not in doc_healed.paragraphs[2].text
            assert "某省某市高新技术开发区100号" not in doc_healed.paragraphs[3].text

            # 2. 验证真正留白的地址、电话、企业名称、法定代表人已自愈填入
            assert "某省某市高新技术开发区100号" in all_text
            assert "028-12345678" in all_text
            assert "测试工程有限公司" in all_text
            assert "李四" in all_text
            assert "2026年06月30日" in all_text

    finally:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass


def test_supervisor_summary_table_duplicate_row_cleanup():
    """测试 Supervisor 对开标一览表中多余重复汇总行进行自愈清理的能力"""
    doc = Document()
    doc.add_heading("三、开标一览表", level=1)
    table = doc.add_table(rows=4, cols=4)
    # Row 0: 表头
    for c_i, h in enumerate(["项目名称", "技术要求", "总价（元）", "备注"]):
        table.rows[0].cells[c_i].text = h
    # Row 1: 数据行
    table.rows[1].cells[0].text = "某光伏采购项目"
    table.rows[1].cells[1].text = "详见采购参数"
    table.rows[1].cells[2].text = "1020000.00"
    table.rows[1].cells[3].text = "无"
    # Row 2: 误插入的重复汇总行
    table.rows[2].cells[0].text = "投标总报价（大写）：壹佰零贰万元整"
    table.rows[2].cells[1].text = "1020000.00"
    table.rows[2].cells[2].text = ""
    table.rows[2].cells[3].text = ""
    # Row 3: 模板原生汇总行
    table.rows[3].cells[0].text = "投标总报价（大写）"
    table.rows[3].cells[1].text = ""
    table.rows[3].cells[2].text = ""
    table.rows[3].cells[3].text = ""

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
        doc.save(tmp_path)

    try:
        mock_items = [
            SimpleNamespace(calculated_total=1020000.0, manufacturer="隆基", brand="隆基")
        ]
        mock_db = MagicMock()
        mock_query = MagicMock()
        mock_query.filter.return_value.all.return_value = mock_items
        mock_db.query.return_value = mock_query

        with patch("app.agents.bid_filler_agent.SessionLocal", return_value=mock_db), \
             patch("app.services.office_cli_service.office_cli_service.query_structure", return_value="tbl[1]"):

            state = {
                "document_id": "test_doc_cleanup",
                "docx_temp_path": tmp_path,
                "original_context": "",
                "repair_count": 0,
                "max_repair_rounds": 2,
                "repair_instructions_map": {},
                "audit_items": [],
            }

            res = supervisor_audit_node(state)

            doc_healed = Document(tmp_path)
            healed_table = doc_healed.tables[0]
            # 原 4 行表格清理后应为 3 行
            assert len(healed_table.rows) == 3
            # 第 3 行应为原生汇总行且已填入大写金额
            assert "投标总报价（大写）" in healed_table.rows[2].cells[0].text
            assert "壹佰零贰万元整" in healed_table.rows[2].cells[1].text

    finally:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass
