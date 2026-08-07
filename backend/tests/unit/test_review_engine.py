"""
Review Engine 深度质检引擎单元测试

覆盖 5 条规则审查管线的核心逻辑：
  R1: 必填项遗漏检测
  R2: 数据交叉一致性校验
  R3: 财务数据准确性校验
  R4: 格式合规性校验
  R5: 评分项覆盖率检查
  R6: 自动修正管线
"""

import pytest
import os
import io
from docx import Document

from app.agents.review_engine import (
    check_unfilled_slots,
    check_data_consistency,
    check_financial_accuracy,
    check_format_compliance,
    check_scoring_coverage,
    auto_fix_proposals,
    make_finding,
    _extract_score_tree_keywords,
)


# ============================================================
# 测试辅助工具
# ============================================================

def _create_test_docx_with_blanks(tmp_path: str) -> str:
    """创建包含占位符和空白单元格的测试 Word 文档"""
    doc = Document()
    # 正文段落：包含占位符
    doc.add_paragraph("投标人名称：________")
    doc.add_paragraph("法定代表人：张三")
    doc.add_paragraph("项目编号：[待补充]")
    doc.add_paragraph("投标总价：967840.36 元")

    # 表格：包含空白单元格
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "序号"
    table.cell(0, 1).text = "设备名称"
    table.cell(0, 2).text = "单价"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "核心主设备"
    table.cell(1, 2).text = "500000"
    table.cell(2, 0).text = "2"
    table.cell(2, 1).text = ""       # 空白单元格
    table.cell(2, 2).text = ""       # 空白单元格

    filepath = os.path.join(tmp_path, "test_review.docx")
    doc.save(filepath)
    return filepath


def _create_test_docx_all_filled(tmp_path: str) -> str:
    """创建所有字段均已填满的测试 Word 文档"""
    doc = Document()
    doc.add_paragraph("投标人名称：四川石楠建设工程有限公司")
    doc.add_paragraph("法定代表人：张三")
    doc.add_paragraph("项目编号：HNYT-2026-001")
    doc.add_paragraph("投标总价：967840.36 元")

    filepath = os.path.join(tmp_path, "test_review_filled.docx")
    doc.save(filepath)
    return filepath


# ============================================================
# R1: 必填项遗漏检测
# ============================================================

class TestCheckUnfilledSlots:
    """R1: 必填项遗漏检测测试"""

    def test_check_unfilled_slots_should_detect_underline_blanks(self, tmp_path):
        """正常场景：检测连续下划线占位符"""
        docx_path = _create_test_docx_with_blanks(str(tmp_path))
        findings = check_unfilled_slots(docx_path)

        # 应检测到 "________" 和 "[待补充]"
        unfilled = [f for f in findings if f["rule_id"] == "R1-UNFILLED"]
        assert len(unfilled) >= 2, f"应至少检测到 2 个占位符，实际: {len(unfilled)}"

    def test_check_unfilled_slots_should_detect_empty_cells(self, tmp_path):
        """正常场景：检测空白表格单元格"""
        docx_path = _create_test_docx_with_blanks(str(tmp_path))
        findings = check_unfilled_slots(docx_path)

        empty_cells = [f for f in findings if f["rule_id"] == "R1-EMPTY-CELL"]
        assert len(empty_cells) >= 2, f"应至少检测到 2 个空白单元格，实际: {len(empty_cells)}"

    def test_check_unfilled_slots_should_pass_when_all_filled(self, tmp_path):
        """边界场景：所有字段均已填满，不应报错"""
        docx_path = _create_test_docx_all_filled(str(tmp_path))
        findings = check_unfilled_slots(docx_path)

        errors_and_warnings = [f for f in findings if f["severity"] in ("error", "warning")]
        assert len(errors_and_warnings) == 0, f"全部已填满不应有 error/warning，实际: {errors_and_warnings}"

    def test_check_unfilled_slots_should_handle_missing_file(self):
        """异常场景：文件不存在时应优雅降级"""
        findings = check_unfilled_slots("/non/existent/path.docx")
        assert len(findings) == 1
        assert findings[0]["severity"] == "error"


# ============================================================
# R2: 数据交叉一致性校验
# ============================================================

class TestCheckDataConsistency:
    """R2: 数据交叉一致性校验测试"""

    def test_check_data_consistency_should_detect_mismatch(self):
        """异常场景：同一字段在不同章节填写值不一致"""
        proposals = [
            {
                "path": "/body/p[5]",
                "proposed_text": "投标人名称（盖章）：四川石楠建设工程有限公司",
                "original_context": "投标人名称（盖章）：____",
                "source_tool": "query_company_profile_tool",
            },
            {
                "path": "/body/tbl[2]/tr[3]/tc[1]/p[1]",
                "proposed_text": "投标人名称：四川石楠建设有限公司",  # 少了 "工程" 两个字
                "original_context": "投标人名称",
                "source_tool": "query_company_profile_tool",
            },
        ]
        findings = check_data_consistency(proposals)

        inconsistent = [f for f in findings if f["rule_id"] == "R2-INCONSISTENT"]
        assert len(inconsistent) >= 1, "应检测到公司名不一致"

    def test_check_data_consistency_should_pass_when_consistent(self):
        """正常场景：同一字段在不同章节填写值一致"""
        proposals = [
            {
                "path": "/body/p[5]",
                "proposed_text": "投标人名称（盖章）：四川石楠建设工程有限公司",
                "original_context": "投标人名称（盖章）：____",
                "source_tool": "query_company_profile_tool",
            },
            {
                "path": "/body/tbl[2]/tr[3]/tc[1]/p[1]",
                "proposed_text": "四川石楠建设工程有限公司",
                "original_context": "投标人名称",
                "source_tool": "query_company_profile_tool",
            },
        ]
        findings = check_data_consistency(proposals)

        inconsistent = [f for f in findings if f["rule_id"] == "R2-INCONSISTENT"]
        assert len(inconsistent) == 0, f"一致性校验应通过，实际: {inconsistent}"

    def test_check_data_consistency_should_skip_empty_proposals(self):
        """边界场景：空提案列表不应报错"""
        findings = check_data_consistency([])
        assert len(findings) == 0


# ============================================================
# R3: 财务数据准确性校验 (不依赖 DB 的纯逻辑测试)
# ============================================================

class TestCheckFinancialAccuracy:
    """R3: 财务数据准确性校验测试"""

    def test_check_financial_accuracy_should_validate_chinese_rmb(self):
        """正常场景：验证大写金额转换工具的正确性"""
        from app.utils.rmb_formatter import number_to_chinese_rmb

        # 已知正确转换
        assert number_to_chinese_rmb(967840.36) == "玖拾陆万柒仟捌佰肆拾元叁角陆分"
        assert number_to_chinese_rmb(0.0) == "零元整"
        assert number_to_chinese_rmb(100.50) == "壹佰元伍角"


# ============================================================
# R4: 格式合规性校验
# ============================================================

class TestCheckFormatCompliance:
    """R4: 格式合规性校验测试"""

    def test_check_format_compliance_should_detect_bad_date_format(self):
        """异常场景：检测不规范的日期格式 (YYYY-MM-DD)"""
        proposals = [
            {
                "path": "/body/p[10]",
                "proposed_text": "投标日期：2026-08-04",
                "original_context": "日期：____",
                "source_tool": "query_project_metadata_tool",
            },
        ]
        findings = check_format_compliance(proposals)

        date_issues = [f for f in findings if f["rule_id"] == "R4-DATE-FORMAT"]
        assert len(date_issues) >= 1, "应检测到日期格式不规范"
        # 验证自动修正方案
        assert date_issues[0]["auto_fixable"] is True
        fix = date_issues[0].get("fix_proposal", {})
        assert "2026年8月4日" in fix.get("proposed_text", ""), f"修正后应为中文日期，实际: {fix}"

    def test_check_format_compliance_should_pass_good_date_format(self):
        """正常场景：中文日期格式不应报错"""
        proposals = [
            {
                "path": "/body/p[10]",
                "proposed_text": "投标日期：2026年8月4日",
                "original_context": "日期：____",
                "source_tool": "query_project_metadata_tool",
            },
        ]
        findings = check_format_compliance(proposals)

        date_issues = [f for f in findings if f["rule_id"] == "R4-DATE-FORMAT"]
        assert len(date_issues) == 0, f"规范日期不应报错，实际: {date_issues}"

    def test_check_format_compliance_should_detect_incomplete_company_name(self):
        """异常场景：公司名缺少法律实体后缀"""
        proposals = [
            {
                "path": "/body/p[5]",
                "proposed_text": "四川石楠建设工程",  # 缺少 "有限公司"
                "original_context": "投标人名称：____",
                "source_tool": "query_company_profile_tool",
            },
        ]
        findings = check_format_compliance(proposals)

        name_issues = [f for f in findings if f["rule_id"] == "R4-COMPANY-NAME"]
        assert len(name_issues) >= 1, "应检测到公司名不完整"


# ============================================================
# R5: 评分项覆盖率检查 — score_tree 关键词提取测试
# ============================================================

class TestScoreTreeExtraction:
    """R5: 评分项覆盖率检查 — score_tree 结构解析"""

    def test_extract_score_tree_keywords_should_parse_nested_structure(self):
        """正常场景：递归解析嵌套的 score_tree 结构"""
        score_tree = {
            "name": "综合评分",
            "score": 100,
            "children": [
                {
                    "name": "技术方案评分",
                    "score": 40,
                    "criteria": "投标人提供的技术方案完整性、可行性、创新性",
                    "children": [
                        {"name": "设计方案", "score": 20, "criteria": "整体设计方案的合理性与先进性"},
                        {"name": "施工组织", "score": 20, "criteria": "施工组织设计的科学性与完整性"},
                    ],
                },
                {
                    "name": "商务评分",
                    "score": 30,
                    "criteria": "投标人的企业资质、业绩、人员配置",
                },
                {
                    "name": "价格评分",
                    "score": 30,
                    "criteria": "投标报价的合理性",
                },
            ],
        }

        items = _extract_score_tree_keywords(score_tree)

        # 应展开为至少 5 个评分项（含嵌套）
        assert len(items) >= 5, f"应至少提取 5 个评分项，实际: {len(items)}"

        # 验证关键词提取
        all_keywords = set()
        for item in items:
            all_keywords.update(item["keywords"])

        assert any("技术方案" in kw for kw in all_keywords), f"应包含含 '技术方案' 的关键词，实际: {all_keywords}"
        assert any("设计方案" in kw for kw in all_keywords), f"应包含含 '设计方案' 的关键词，实际: {all_keywords}"
        assert any("投标报价" in kw or "价格" in kw for kw in all_keywords), f"应包含含 '投标报价' 或 '价格' 的关键词，实际: {all_keywords}"

    def test_extract_score_tree_keywords_should_handle_empty_tree(self):
        """边界场景：空 score_tree 不应报错"""
        items = _extract_score_tree_keywords({})
        assert len(items) == 0

        items = _extract_score_tree_keywords([])
        assert len(items) == 0

    def test_extract_score_tree_keywords_should_handle_flat_list(self):
        """边界场景：扁平列表格式的 score_tree"""
        score_tree = [
            {"name": "技术评分", "score": 50, "criteria": "技术方案"},
            {"name": "报价评分", "score": 50, "criteria": "价格合理性"},
        ]
        items = _extract_score_tree_keywords(score_tree)
        assert len(items) == 2


# ============================================================
# R6: 自动修正管线
# ============================================================

class TestAutoFixProposals:
    """R6: 自动修正管线测试"""

    def test_auto_fix_should_apply_rule_fix_proposal(self):
        """正常场景：应用规则引擎提供的确定性修正方案"""
        proposals = [
            {
                "path": "/body/p[10]",
                "proposed_text": "投标日期：2026-08-04",
                "original_context": "日期：____",
                "source_tool": "query_project_metadata_tool",
            },
        ]
        findings = [
            make_finding(
                rule_id="R4-DATE-FORMAT", severity="warning",
                path="/body/p[10]",
                description="日期格式不规范",
                current_value="投标日期：2026-08-04",
                expected_value="投标日期：2026年8月4日",
                auto_fixable=True,
                fix_proposal={
                    "path": "/body/p[10]",
                    "proposed_text": "投标日期：2026年8月4日",
                    "original_context": "日期：____",
                    "source_tool": "review_engine_fix",
                },
            ),
        ]

        modified, skipped = auto_fix_proposals(proposals, findings, "test-doc-id")

        assert modified[0]["proposed_text"] == "投标日期：2026年8月4日", \
            f"应修正为中文日期格式，实际: {modified[0]['proposed_text']}"
        assert len(skipped) == 0

    def test_auto_fix_should_skip_unfixable_proposals(self):
        """正常场景：不可修正的字段应被跳过并标记待人工补充"""
        proposals = [
            {
                "path": "/body/p[20]",
                "proposed_text": "四川石楠建设工程",
                "original_context": "投标人名称：____",
                "source_tool": "query_company_profile_tool",
            },
        ]
        findings = [
            make_finding(
                rule_id="R4-COMPANY-NAME", severity="error",
                path="/body/p[20]",
                description="公司名不完整",
                auto_fixable=False,
            ),
        ]

        modified, skipped = auto_fix_proposals(proposals, findings, "test-doc-id")

        assert modified[0]["proposed_text"] == "[待人工补充]", \
            f"不可修正字段应标记为 [待人工补充]，实际: {modified[0]['proposed_text']}"
        assert len(skipped) == 1


# ============================================================
# make_finding 工具函数测试
# ============================================================

class TestMakeFinding:
    """make_finding 构造函数测试"""

    def test_make_finding_should_create_valid_dict(self):
        """正常场景：构造完整的 ReviewFinding 字典"""
        finding = make_finding(
            rule_id="R1-UNFILLED", severity="warning",
            path="/body/p[3]",
            description="残留占位符",
            current_value="____",
        )
        assert finding["rule_id"] == "R1-UNFILLED"
        assert finding["severity"] == "warning"
        assert finding["auto_fixable"] is False
        assert finding["fix_proposal"] is None

    def test_make_finding_should_accept_fix_proposal(self):
        """正常场景：携带修正方案的 finding"""
        finding = make_finding(
            rule_id="R3-SUM-MISMATCH", severity="error",
            path="/body/p[5]",
            description="总价不一致",
            auto_fixable=True,
            fix_proposal={"path": "/body/p[5]", "proposed_text": "967840.36"},
        )
        assert finding["auto_fixable"] is True
        assert finding["fix_proposal"]["proposed_text"] == "967840.36"


# ============================================================
# R9: 模板原文破坏熔断检测测试
# ============================================================

class TestCompareTextBeforeAndAfterFilling:
    """R10: 填写前后 DOM 文本 Diff 比对与自动回滚还原测试"""

    def test_compare_text_before_and_after_filling_should_rollback_when_loss_detected(self):
        """正常场景：填报后若检测到模板长句保留率严重下降，触发自动回滚还原原始 Word 段落文本"""
        import tempfile, os
        from docx import Document
        from app.agents.review_engine import compare_text_before_and_after_filling

        doc = Document()
        # 模拟填报后的 Word 段落 (原文模版被擦除，仅留下了公司名)
        doc.add_paragraph("四川石楠建设工程有限公司")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
            temp_path = tf.name

        try:
            doc.save(temp_path)
            # 写前快照包含完整的法律声明长句
            before_snapshots = {
                "/body/p[1]": "根据贵方的 SZDZ-2026-NG008 号招标文件，正式授权下述签字人李四代表我方四川石楠建设工程有限公司全权处理。"
            }

            findings = compare_text_before_and_after_filling(temp_path, before_snapshots)
            assert len(findings) == 1
            assert findings[0]["rule_id"] == "R10-TEXT-DIFF-LOSS"

            # 验证已被自动回滚还原为原始完整长句
            res_doc = Document(temp_path)
            assert res_doc.paragraphs[0].text == before_snapshots["/body/p[1]"]
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)


# ============================================================
# R7: 表格单元格污染检测测试
# ============================================================

class TestCheckTableCellCleanliness:
    """R7: 表格单元格污染检测与清扫纯化算法测试"""

    def test_clean_cell_text_value_should_strip_pollution_prefixes(self):
        """正常场景：剔除 '序号 1 行，答：'、'序号2：'、'答：' 等污染前缀"""
        from app.agents.review_engine import clean_cell_text_value

        assert clean_cell_text_value("序号 1 行，答：无", "有无偏离") == "无"
        assert clean_cell_text_value("序号 2 行：完全响应招标文件", "承诺") == "完全响应招标文件"
        assert clean_cell_text_value("第 3 行，答：有", "有无偏离") == "有"
        assert clean_cell_text_value("答：满足要求", "偏离说明") == "满足要求"

    def test_check_table_cell_cleanliness_should_detect_and_create_fix_proposals(self):
        """正常场景：检测并生成 R7 污染清扫修正提案"""
        from app.agents.review_engine import check_table_cell_cleanliness

        proposals = [
            {
                "path": "/body/tbl[1]/tr[2]/tc[3]/p[1]",
                "proposed_text": "序号 1 行，答：无",
                "original_context": "表格第 1 个表，第 2 行，第 3 列: 有无偏离",
                "source_tool": "query_company_profile_tool",
            },
            {
                "path": "/body/tbl[1]/tr[3]/tc[3]/p[1]",
                "proposed_text": "无",
                "original_context": "表格第 1 个表，第 3 行，第 3 列: 有无偏离",
                "source_tool": "query_company_profile_tool",
            },
        ]

        findings = check_table_cell_cleanliness(proposals)

        cell_issues = [f for f in findings if f["rule_id"] == "R7-CELL-POLLUTION"]
        assert len(cell_issues) == 1, f"应精确检测到 1 条带有污染前缀的单元格提案，实际: {len(findings)}"
        assert cell_issues[0]["auto_fixable"] is True
        assert cell_issues[0]["fix_proposal"]["proposed_text"] == "无"


# ============================================================
# R8: 短字段长文本错填检测测试
# ============================================================

class TestCheckShortFieldTextOverflow:
    """R8: 短字段长文本错填检测与自动纠偏提纯测试"""

    def test_check_short_field_text_overflow_should_detect_and_fix_overflow_paragraph(self):
        """正常场景：检测并纠正把大段授权话术错填入'投标单位名称:'字段的问题"""
        from app.agents.review_engine import check_short_field_text_overflow

        proposals = [
            {
                "path": "/body/p[25]",
                "proposed_text": "投标单位名称：根据贵方的 SZDZ-2026-NG008 号招标文件，正式授权下述签字人李四代表我方四川石楠建设工程有限公司全权处理",
                "original_context": "投标单位名称：____",
                "source_tool": "query_company_profile_tool",
            },
        ]

        findings = check_short_field_text_overflow(proposals)

        overflow_issues = [f for f in findings if f["rule_id"] == "R8-TEXT-OVERFLOW"]
        assert len(overflow_issues) == 1, f"应精确检测到 1 条短字段长文本溢出错填，实际: {len(findings)}"
        assert overflow_issues[0]["auto_fixable"] is True
        # 断言自动纯化出的期望文本包含纯公司名称，剥离了前面的整句授权长文
        expected = overflow_issues[0]["fix_proposal"]["proposed_text"]
        assert "四川石楠建设工程有限公司" in expected
        assert "根据贵方" not in expected


