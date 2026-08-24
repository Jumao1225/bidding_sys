"""
BidFiller Agent & BidFormatFillerService 单元测试

测试大写金额转化、Agent Tool Calling 查库、下划线继承写 Word 功能。
"""

import pytest
import io
from docx import Document
from app.utils.rmb_formatter import number_to_chinese_rmb
from app.agents.bid_filler_agent import bid_filler_agent
from app.schemas.bid_filler_schema import CompanyProfile
from app.services.bid_format_filler_service import bid_format_filler_service


def test_number_to_chinese_rmb_should_convert_correctly():
    """测试数字转换为人民币大写汉字算法"""
    assert number_to_chinese_rmb(967840.36) == "玖拾陆万柒仟捌佰肆拾元叁角陆分"
    assert number_to_chinese_rmb(20000.00) == "贰万元整"
    assert number_to_chinese_rmb(100.50) == "壹佰元伍角"
    assert number_to_chinese_rmb(0.0) == "零元整"


def test_bid_format_filler_service_underline_inheritance():
    """测试下划线继承逻辑：原处有下划线则保留 underline=True，原处无下划线则不添加"""
    doc = Document()
    
    # 构造原处带下划线的段落
    p1 = doc.add_paragraph()
    r1 = p1.add_run("投标人名称：________")
    r1.underline = True

    # 构造原处不带下划线的段落
    p2 = doc.add_paragraph()
    p2.add_run("项目名称：[项目名称]")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    template_bytes = output.getvalue()

    # 1. 扫描待填位置
    placeholders = bid_format_filler_service.scan_detected_placeholders(template_bytes)
    assert len(placeholders) >= 1

    # 2. 模拟替换
    replacement_map = {
        "投标人名称：________": "投标人名称：四川石楠建设工程有限公司",
        "[项目名称]": "湖南省烟草公司衡阳市公司关于现代终端建设项目"
    }

    filled_bytes = bid_format_filler_service.fill_docx_with_audit_trail(
        docx_bytes=template_bytes,
        replacement_map=replacement_map,
        audit_items=[]
    )

    assert filled_bytes is not None
    res_doc = Document(io.BytesIO(filled_bytes))
    
    # 验证含有“四川石楠建设工程有限公司”的 Run 带有下划线
    target_p = [p for p in res_doc.paragraphs if "四川石楠建设工程有限公司" in p.text][0]
    assert target_p is not None
    val_run = [r for r in target_p.runs if "四川石楠建设工程有限公司" in r.text][0]
    assert val_run.underline is True or 'w:u' in val_run._element.xml


def test_scan_node_temp_path_should_be_in_backend_dir():
    """测试 scan_node 的工作副本路径是否精准定位在 backend 目录下的 uploads/drafts 中"""
    import os
    from app.agents.bid_filler_agent import scan_node

    doc_bytes = io.BytesIO()
    Document().save(doc_bytes)
    state = {
        "document_id": "test_doc_12345678",
        "original_docx": doc_bytes.getvalue(),
        "original_context": "",
        "slot_analysis": None,
        "worker_proposals": None,
        "db_session": None,
        "company_profile": None,
        "docx_temp_path": None,
        "audit_items": [],
        "audit_report": None,
        "filled_docx_bytes": None,
    }
    res = scan_node(state)
    temp_path = res.get("docx_temp_path")
    assert temp_path is not None
    # 判断生成的绝对路径中必定包含 backend
    assert "backend" in os.path.normcase(temp_path)
    assert os.path.exists(temp_path)
    # 测试结束后清理测试文件
    try:
        os.remove(temp_path)
    except Exception:
        pass


def test_proposals_to_commands_should_handle_none_source_and_wrap_tc():
    """测试 proposals_to_commands 是否能够放行 source_tool 为 none/空 的合规提案，并且为表格单元格 /tc[N] 自动补正段落路径 /p[1]"""
    from app.agents.bid_filler_agent import proposals_to_commands

    mock_proposals = [
        # 场景1：source_tool 为 none 或空，但内容合法（如大写金额或自主核算总价），应当放行
        {"source_tool": "none", "path": "/body/p[2]/r[1]", "proposed_text": "玖拾陆万柒仟捌佰肆拾元叁角陆分"},
        {"source_tool": "", "path": "/body/p[3]/r[1]", "proposed_text": "四川某某实业有限公司"},
        # 场景2：表格单元格 XPath /body/tbl[1]/tr[2]/tc[2] 止步于单元格，应自动吸附 /p[1]
        {"source_tool": "query_project_metadata_tool", "path": "/body/tbl[1]/tr[2]/tc[2]", "proposed_text": "15.00万元"},
        # 场景3：携带 OfficeCLI 自带属性的定位路径 /body/p[@paraId=17F154A1]/r[1]，严禁一刀切误拦
        {"source_tool": "query_company_qualification_tool", "path": "/body/p[@paraId=17F154A1]/r[1]", "proposed_text": "电力工程施工总承包二级"},
        # 场景4：占位异常描述，应当过滤
        {"source_tool": "db", "path": "/body/p[4]/r[1]", "proposed_text": "[待补充资质要求]"},
    ]

    cmds, approved, rejected = proposals_to_commands(mock_proposals)
    assert approved == 4
    assert rejected == 1
    assert len(cmds) == 4
    # 断言场景2成功转化为段落路径
    table_cmd = [c for c in cmds if "tbl[1]" in c["path"]][0]
    assert table_cmd["path"] == "/body/tbl[1]/tr[2]/tc[2]/p[1]"
    # 断言场景3原生保持并获悉打进命令集
    para_cmd = [c for c in cmds if "@paraId" in c["path"]][0]
    assert para_cmd["props"]["text"] == "电力工程施工总承包二级"


def test_proposals_to_commands_should_preserve_label_prefix():
    """测试 proposals_to_commands 在提案未携带标签前缀时，能依据 original_context 自动保留并补全字段前缀"""
    from app.agents.bid_filler_agent import proposals_to_commands, _extract_label_prefix

    # 1. 验证 _extract_label_prefix 提炼能力
    assert _extract_label_prefix("投标人名称（单位盖章）：__________________________") == "投标人名称（单位盖章）："
    assert _extract_label_prefix("地    址：__________________________") == "地    址："
    assert _extract_label_prefix("项目编号：[项目编号]") == "项目编号："

    # 2. 验证 proposals_to_commands 前缀自动防护
    mock_proposals = [
        {
            "source_tool": "query_company_profile_tool",
            "path": "/body/p[12]",
            "original_context": "投标人名称（单位盖章）：__________________________",
            "proposed_text": "北京某某科技有限公司",
        },
        {
            "source_tool": "query_company_profile_tool",
            "path": "/body/p[13]",
            "original_context": "地    址：__________________________",
            "proposed_text": "地    址：北京市海淀区中关村南大街1号",
        },
        {
            "source_tool": "query_financial_quotation_tool",
            "path": "/body/tbl[1]/tr[2]/tc[2]",
            "original_context": "表格第 1 个表，第 2 行，第 2 列：",
            "proposed_text": "15.00万元",
        }
    ]

    cmds, approved, rejected = proposals_to_commands(mock_proposals)
    assert approved == 3
    assert cmds[0]["props"]["text"] == "投标人名称（单位盖章）：北京某某科技有限公司"
    # 已自带前缀的不重复拼接
    assert cmds[1]["props"]["text"] == "地    址：北京市海淀区中关村南大街1号"
    # 无标签的数值单元格保持纯数值
    assert cmds[2]["props"]["text"] == "15.00万元"


def test_apply_underline_to_filled_doc_should_set_underline_on_value_run():
    """测试 _apply_underline_to_filled_doc 能给填报值拆成带下划线的 Run，且保持前缀标签无下划线"""
    import tempfile, os
    from docx import Document
    from app.agents.bid_filler_agent import _apply_underline_to_filled_doc

    doc = Document()
    doc.add_paragraph("投标人名称（单位盖章）：北京某某科技有限公司")
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name
    
    try:
        doc.save(temp_path)
        proposals = [{
            "path": "/body/p[1]",
            "original_context": "投标人名称（单位盖章）：__________________________",
            "proposed_text": "北京某某科技有限公司"
        }]
        
        _apply_underline_to_filled_doc(temp_path, proposals)
        
        res_doc = Document(temp_path)
        p = res_doc.paragraphs[0]
        assert len(p.runs) == 2
        assert p.runs[0].text == "投标人名称（单位盖章）："
        assert p.runs[0].underline is False or p.runs[0].underline is None
        assert p.runs[1].text == "北京某某科技有限公司"
        assert p.runs[1].underline is True
        assert "w:u" in p.runs[1]._element.xml
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_fill_docx_proposals_in_dom_should_preserve_real_labels_and_add_underline():
    """测试 DOM 引擎从真实 Word 节点读取原文，保留前缀标签并成功注入 OpenXML 下划线"""
    import tempfile, os
    from docx import Document
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom, _extract_prefix_from_text

    # 1. 验证 _extract_prefix_from_text 各种提炼场景
    assert _extract_prefix_from_text("投标人全称（盖章）__________________________") == "投标人全称（盖章）："
    assert _extract_prefix_from_text("地    址：__________________________") == "地    址："
    assert _extract_prefix_from_text("项目名称") == "项目名称："

    doc = Document()
    doc.add_paragraph("投标人全称（盖章）：__________________________")
    doc.add_paragraph("法定代表人（签字）：__________________________")
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name
    
    try:
        doc.save(temp_path)
        proposals = [
            {
                "path": "/body/p[1]",
                "original_context": "正文段落 1",  # 哪怕 Worker 上下文丢失，DOM 引擎仍能从真实 Word 读取
                "proposed_text": "四川某某实业有限公司"
            },
            {
                "path": "/body/p[2]",
                "original_context": "法定代表人（签字）：__________________________",
                "proposed_text": "李四"
            }
        ]
        
        count = fill_docx_proposals_in_dom(temp_path, proposals)
        assert count == 2
        
        res_doc = Document(temp_path)
        p1 = res_doc.paragraphs[0]
        assert len(p1.runs) == 2
        assert p1.runs[0].text == "投标人全称（盖章）："
        assert p1.runs[0].underline is False or p1.runs[0].underline is None
        assert p1.runs[1].text == "四川某某实业有限公司"
        assert p1.runs[1].underline is True
        assert "w:u" in p1.runs[1]._element.xml

        p2 = res_doc.paragraphs[1]
        assert len(p2.runs) == 2
        assert p2.runs[0].text == "法定代表人（签字）："
        assert p2.runs[1].text == "李四"
        assert "w:u" in p2.runs[1]._element.xml
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_fill_docx_proposals_underline_policy_outside_vs_inside_table():
    """测试 DOM 引擎策略：表格外正文填报施加下划线，表格内单元格填报取消下划线"""
    import tempfile, os
    from docx import Document
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    doc = Document()
    doc.add_paragraph("投标人名称（单位盖章）：__________________________")
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].paragraphs[0].text = "投标总价："
    tbl.rows[0].cells[1].paragraphs[0].text = "__________________________"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)
        proposals = [
            {
                "path": "/body/p[1]",
                "original_context": "投标人名称（单位盖章）：__________________________",
                "proposed_text": "北京某某科技有限公司"
            },
            {
                "path": "/body/tbl[1]/tr[1]/tc[2]",
                "original_context": "表格第 1 个表，第 1 行，第 2 列",
                "proposed_text": "15.00万元"
            }
        ]

        count = fill_docx_proposals_in_dom(temp_path, proposals)
        assert count == 2

        res_doc = Document(temp_path)
        from docx.shared import Pt
        # 1. 验证表格外正文段落：数据值带有下划线，且为【宋体 小四 (12pt)】
        p_body = res_doc.paragraphs[0]
        assert len(p_body.runs) == 2
        assert p_body.runs[1].text == "北京某某科技有限公司"
        assert p_body.runs[1].underline is True
        assert "w:u" in p_body.runs[1]._element.xml
        assert p_body.runs[1].font.name == "宋体"
        assert p_body.runs[1].font.size == Pt(12)

        # 2. 验证表格内单元格：数据值没有下划线，且为【宋体 小五 (9pt)】
        p_cell = res_doc.tables[0].rows[0].cells[1].paragraphs[0]
        assert p_cell.runs[0].text == "15.00万元"
        assert p_cell.runs[0].underline is False or p_cell.runs[0].underline is None
        assert "w:u" not in p_cell.runs[0]._element.xml
        assert p_cell.runs[0].font.name == "宋体"
        assert p_cell.runs[0].font.size == Pt(9)
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_fill_docx_proposals_inplace_sub_replacement_preserves_surrounding_text():
    """测试核心算法：原位切片替换 100% 物理保留占位符前后的所有模板文本，彻底防止丢失原文"""
    import tempfile, os
    from docx import Document
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    doc = Document()
    doc.add_paragraph("根据贵方的 SZDZ-2026-NG008 号招标文件，正式授权下述签字人____代表我方____公司（投标人的名称），全权处理本次项目投标。")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)
        proposals = [
            {
                "path": "/body/p[1]",
                "original_context": "根据贵方的 SZDZ-2026-NG008 号招标文件，正式授权下述签字人____代表我方....",
                "proposed_text": "张三"
            }
        ]

        count = fill_docx_proposals_in_dom(temp_path, proposals)
        assert count == 1

        res_doc = Document(temp_path)
        p = res_doc.paragraphs[0]
        # 断言应该切片为 3 段 Run：[前文模板, 填入值, 后文模板]
        assert len(p.runs) == 3
        assert p.runs[0].text == "根据贵方的 SZDZ-2026-NG008 号招标文件，正式授权下述签字人"
        assert p.runs[0].underline is False or p.runs[0].underline is None
        assert p.runs[1].text == "张三"
        assert p.runs[1].underline is True
        assert p.runs[2].text == "代表我方____公司（投标人的名称），全权处理本次项目投标。"
        assert p.runs[2].underline is False or p.runs[2].underline is None
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)



def test_db_tools_quotation_and_market_price_should_return_full_bom():
    """测试财务报价及指导单价检索 DB 工具能正确支持全量 BOM 成本测算合价与模糊词条组合联查"""
    from app.agents.tools.bid_db_tools import query_market_price_reference_tool, query_financial_quotation_tool

    # 1. 验证 query_market_price_reference_tool 能自动切词并联查 CostEstimate 的单价和合价
    res_inverter = query_market_price_reference_tool.invoke({"item_name": "组串式逆变器 华为 110kW"})
    assert "单价" in res_inverter and "11555.0" in res_inverter

    # 2. 验证并网柜指导单价与实际BOM合价的透查
    res_cabinet = query_market_price_reference_tool.invoke({"item_name": "并网柜"})
    assert "9180.35" in res_cabinet and "18360.7" in res_cabinet

    # 3. 验证财务接口通过 cost_estimates 关键字段一次性直接拉取完整的分项清单表
    res_bom = query_financial_quotation_tool.invoke({"document_id": "06a4abb2-4817-43fc-aa6f-f5f4901025ae", "field_key": "cost_estimates"})
    assert "673492.47" in res_bom and "69330.0" in res_bom and "18360.7" in res_bom


def test_parse_proposals_unescaped_quotes_and_object_recovery():
    """测试 Worker 提案解析器能否容错自动修复内层未转义双引号，并在语法崩溃时实施逐个对象拯救机制"""
    from app.agents.bid_filler_workers import _parse_proposals

    # 包含内层半角双引号 "招标编号：__号" 的日志真实范例片段
    bad_raw_json = (
        '[\n'
        '  {"path": "/body/p[1]", "original_context": "招标", "source_data": "008号", "proposed_text": "008号", '
        '"reasoning": "原文"招标编号：__号"为空白待填"},\n'
        '  {"path": "/body/tbl[1]/tr[2]/tc[1]", "proposed_text": "1", "reasoning": "正常项"}\n'
        ']'
    )

    proposals = _parse_proposals(bad_raw_json)
    assert len(proposals) == 2
    assert proposals[0]["path"] == "/body/p[1]"
    assert "招标编号" in proposals[0]["reasoning"]
    assert proposals[1]["proposed_text"] == "1"


def test_auto_repair_officecli_commands_should_correct_parent_and_paths():
    """测试 Office CLI 指令前置自愈器能自动修复 add row 的 parent 与 set cell 的嵌套路径"""
    from app.agents.bid_filler_agent import auto_repair_officecli_commands

    bad_commands = [
        {"command": "add", "parent": "/body/tbl[3]/tr[2]", "type": "row"},
        {"command": "set", "path": "/body/tbl[3]/tr[2]/tr[last()]/tc[1]", "props": {"text": "测试数据"}},
        {"command": "set", "path": "/body/tbl[1]/tr[2]/tc[3]", "props": {"text": "数值"}}
    ]

    repaired = auto_repair_officecli_commands(bad_commands)
    assert len(repaired) == 3
    # 断言 1: add row 的 parent 被剥离还原为 /body/tbl[3]
    assert repaired[0]["parent"] == "/body/tbl[3]"
    # 断言 2: tr[last()] 嵌套错路径被自愈为 /body/tbl[3]/row[last()]/cell[1]
    assert repaired[1]["path"] == "/body/tbl[3]/row[last()]/cell[1]"
    # 断言 3: 合法规范路径无感放行
    assert repaired[2]["path"] == "/body/tbl[1]/tr[2]/tc[3]"


def test_proposals_to_commands_table_cell_no_prefix():
    """测试 proposals_to_commands 对表格单元格禁止强加标签前缀，并自动纯化单元格文本"""
    from app.agents.bid_filler_agent import proposals_to_commands

    proposals = [
        {
            "path": "/body/p[1]",
            "original_context": "投标人名称：______",
            "proposed_text": "四川某某建设工程有限公司"
        },
        {
            "path": "/body/tbl[1]/tr[2]/tc[1]",
            "original_context": "投标人名称：______",
            "proposed_text": "投标人名称：四川某某建设工程有限公司"
        }
    ]

    cmds, approved, rejected = proposals_to_commands(proposals)
    assert len(cmds) == 2
    # 断言 1: 正文段落保留补全前缀
    assert cmds[0]["props"]["text"] == "投标人名称：四川某某建设工程有限公司"
    # 断言 2: 表格单元格强行剥离纯化，绝无重复前缀
    assert cmds[1]["props"]["text"] == "四川某某建设工程有限公司"


def test_build_dynamic_matrix_for_header_custom_columns():
    """测试 build_dynamic_matrix_for_header 可配置语义映射 + ORM Schema 反射自适应"""
    from app.agents.tools.bid_db_tools import build_dynamic_matrix_for_header
    from unittest.mock import MagicMock

    mock_item = MagicMock()
    mock_item.item_name = "XXX设备"
    mock_item.brand = "XXX品牌"
    mock_item.spec = "XXX规格"
    mock_item.manufacturer = "XXX厂商"
    mock_item.unit = "套"
    mock_item.quantity = 10
    mock_item.unit_price = 100.0
    mock_item.calculated_total = 1000.0
    mock_item.remark = "XXX备注"

    cost_items = [mock_item]

    # 测试场景1: 直接按 ORM 物理字段名提取列矩阵
    hdr_orm = ["item_name", "quantity", "unit_price", "calculated_total"]
    mat_orm = build_dynamic_matrix_for_header(cost_items, hdr_orm)
    assert len(mat_orm[0]) == 4
    assert mat_orm[0] == ["XXX设备", "10", "100.00", "1,000.00"]

    # 测试场景2: 5 列包含 __INDEX__ 的动态 ORM 字段提取
    hdr_5col = ["__INDEX__", "item_name", "spec", "quantity", "calculated_total"]
    mat_5col = build_dynamic_matrix_for_header(cost_items, hdr_5col)
    assert len(mat_5col[0]) == 5
    assert mat_5col[0][0] == "1"          # 序号自动递增
    assert mat_5col[0][1] == "XXX设备"    # item_name
    assert mat_5col[0][2] == "XXX规格"    # spec
    assert mat_5col[0][3] == "10"         # quantity
    assert mat_5col[0][4] == "1,000.00"   # calculated_total

    # 测试场景3: 9 列传统表头 (含 __BRAND_SPEC__ 品牌+规格合并列)
    hdr_9col = ["__INDEX__", "item_name", "__BRAND_SPEC__", "manufacturer", "unit", "quantity", "unit_price", "calculated_total", "remark"]
    mat_9col = build_dynamic_matrix_for_header(cost_items, hdr_9col)
    assert len(mat_9col[0]) == 9
    assert mat_9col[0][0] == "1"
    assert mat_9col[0][1] == "XXX设备"
    assert mat_9col[0][2] == "XXX品牌 XXX规格"  # 品牌+规格合并
    assert mat_9col[0][3] == "XXX厂商"
    assert mat_9col[0][6] == "100.00"    # 单价
    assert mat_9col[0][7] == "1,000.00"  # 总价


def test_auto_repair_officecli_nested_path():
    """测试 auto_repair_officecli_commands 修复嵌套路径 /tbl[N]/tr[M]/tr[last()]/tc[C]"""
    from app.agents.bid_filler_agent import auto_repair_officecli_commands

    # 模式A: /body/tbl[3]/tr[6]/tr[last()]/tc[1] → /body/tbl[3]/row[last()]/cell[1]
    cmds = [
        {"command": "set", "path": "/body/tbl[3]/tr[6]/tr[last()]/tc[1]", "props": {"text": "XXX值"}},
        {"command": "add", "parent": "/body/tbl[3]/tr[6]", "type": "row"},
        {"command": "set", "path": "/body/tbl[3]/row[last()]/cell[1]", "props": {"text": "XXX正常值"}},
    ]
    repaired = auto_repair_officecli_commands(cmds)

    # 嵌套路径被修复
    assert repaired[0]["path"] == "/body/tbl[3]/row[last()]/cell[1]"
    # add row 的 parent 被剥离为表格根路径
    assert repaired[1]["parent"] == "/body/tbl[3]"
    # 已正确的路径保持不变
    assert repaired[2]["path"] == "/body/tbl[3]/row[last()]/cell[1]"


def test_table_pipe_concatenation_auto_split_across_cells():
    """测试自愈引擎：当大模型错误将整行数据用 '｜' 拼接填入首单元格时，自动横向跨列分发"""
    import tempfile, os
    from docx import Document
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    doc = Document()
    tbl = doc.add_table(rows=2, cols=8)
    # 表头
    headers = ["序号", "标的物名称", "规格型号", "生产厂家", "单位", "数量", "单价", "总价"]
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)
        proposals = [{
            "path": "/body/tbl[1]/tr[2]/tc[1]",
            "original_context": "表格第 1 个表，第 2 行，第 1 列",
            "proposed_text": "1 ｜ 光伏组件 ｜ 550W ｜ 某某光伏有限公司 ｜ 块 ｜ 1155 ｜ 550.00 ｜ 635250.00"
        }]

        count = fill_docx_proposals_in_dom(temp_path, proposals)
        assert count == 1

        res_doc = Document(temp_path)
        row2 = res_doc.tables[0].rows[1]
        assert row2.cells[0].text.strip() == "1"
        assert row2.cells[1].text.strip() == "光伏组件"
        assert row2.cells[2].text.strip() == "550W"
        assert row2.cells[3].text.strip() == "某某光伏有限公司"
        assert row2.cells[4].text.strip() == "块"
        assert row2.cells[5].text.strip() == "1155"
        assert row2.cells[6].text.strip() == "550.00"
        assert row2.cells[7].text.strip() == "635250.00"
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_technical_specs_with_slashes_should_not_split_or_overwrite_index_column():
    """测试偏离表防误切分：技术条款中包含多个 '/'（如 IEC61215/IEC61730/CQC/TUV/耐寒/防腐/防火）绝不误触发跨列切分，且绝不覆盖序号列"""
    import tempfile, os
    from docx import Document
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    doc = Document()
    tbl = doc.add_table(rows=2, cols=5)
    headers = ["序号", "招标文件技术要求", "投标文件服务承诺", "有无偏离", "偏离说明"]
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h

    # 模版预设序号 1
    tbl.rows[1].cells[0].text = "1"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)
        long_spec_text = "光伏组件≥630Wp、IEC61215/IEC61730、CQC/TUV认证、耐高温/耐寒/耐紫外线/阻燃、NB/T 32004-2018"
        proposals = [
            {
                "path": "/body/tbl[1]/tr[2]/tc[2]",
                "value": long_spec_text,
            },
            {
                "path": "/body/tbl[1]/tr[2]/tc[3]",
                "value": "拟投天合光能635Wp单晶组件，完全满足全部指标要求。",
            },
            {
                "path": "/body/tbl[1]/tr[2]/tc[4]",
                "value": "无",
            },
            {
                "path": "/body/tbl[1]/tr[2]/tc[5]",
                "value": "完全响应招标文件技术要求，无偏离。",
            }
        ]

        count = fill_docx_proposals_in_dom(temp_path, proposals)
        assert count >= 4

        res_doc = Document(temp_path)
        row = res_doc.tables[0].rows[1]
        # 断言 1: 序号列保持为 1，绝对没有被技术条款覆盖！
        assert row.cells[0].text.strip() == "1"
        # 断言 2: 第 2 列完整保留包含所有斜杠的技术要求全量文字，没有被切碎！
        assert row.cells[1].text.strip() == long_spec_text
        # 断言 3: 第 3 列为服务承诺
        assert "天合光能" in row.cells[2].text.strip()
        # 断言 4: 第 4 列为无偏离
        assert row.cells[3].text.strip() == "无"
        # 断言 5: 第 5 列为无偏离说明
        assert "无偏离" in row.cells[4].text.strip()
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_value_key_normalization_and_open_bid_table_filling():
    """测试 Worker 提案解析与写盘引擎双向兼容 'value' 键名，确保《开标一览表》等表格完整填入"""
    import tempfile, os
    from docx import Document
    from app.agents.bid_filler_workers import _parse_proposals
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    # 1. 验证 _parse_proposals 对纯 value 键名的 JSON 代码块自动归一化
    raw_worker_json = """```json
[
  {"path": "/body/tbl[1]/tr[2]/tc[3]", "value": "1017934.21"},
  {"path": "/body/tbl[1]/tr[2]/tc[4]", "value": "接到采购人进场通知后60日内完工"},
  {"path": "/body/tbl[1]/tr[3]/tc[2]", "value": "壹佰零壹万柒仟玖佰叁拾肆元贰角壹分"}
]
```"""
    parsed = _parse_proposals(raw_worker_json)
    assert len(parsed) == 3
    assert parsed[0]["proposed_text"] == "1017934.21"
    assert parsed[0]["value"] == "1017934.21"
    assert parsed[2]["proposed_text"] == "壹佰零壹万柒仟玖佰叁拾肆元贰角壹分"

    # 2. 验证 fill_docx_proposals_in_dom 将纯 value 键名提案成功写入 Word 表格
    doc = Document()
    tbl = doc.add_table(rows=3, cols=4)
    tbl.rows[0].cells[0].text = "项目名称"
    tbl.rows[0].cells[1].text = "技术要求"
    tbl.rows[0].cells[2].text = "总价(元)"
    tbl.rows[0].cells[3].text = "备注"

    tbl.rows[1].cells[0].text = "某光伏项目"
    tbl.rows[1].cells[1].text = "详见第四章"
    # cells[2] 和 cells[3] 留空待填

    tbl.rows[2].cells[0].text = "投标总报价(大写)"
    # cells[1] 留空待填

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)
        count = fill_docx_proposals_in_dom(temp_path, parsed)
        assert count == 3

        res_doc = Document(temp_path)
        table = res_doc.tables[0]
        # 验证总价数字与备注承诺已成功写入
        assert table.rows[1].cells[2].text.strip() == "1017934.21"
        assert "60日内完工" in table.rows[1].cells[3].text.strip()
        # 验证大写总价已成功写入
        assert table.rows[2].cells[1].text.strip() == "壹佰零壹万柒仟玖佰叁拾肆元贰角壹分"
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_ellipsis_cleanup_in_review_engine():
    """测试质量审核引擎：自动清理省略号截断与残缺标点"""
    from app.agents.review_engine import clean_cell_text_value

    # 场景1: 句子以 ... 结尾，自动补齐为句号
    res1 = clean_cell_text_value("我公司完全响应并承诺，严格按要求提供相关服务及技术支持...")
    assert res1 == "我公司完全响应并承诺，严格按要求提供相关服务及技术支持。"

    # 场景2: 句子以 …… 结尾
    res2 = clean_cell_text_value("由我公司负责相关技术支持……")
    assert res2 == "由我公司负责相关技术支持。"

    # 场景3: 纯省略号残缺行
    res3 = clean_cell_text_value("..")
    assert res3 == ""


def test_clean_zero_change_annotations_and_no_op_detection():
    """测试剥离各类‘（原文无槽位，零改动保留）’说明性元数据注释与无操作提案检测"""
    from app.agents.review_engine import clean_zero_change_annotations, is_zero_change_or_no_op_proposal

    # 1. 测试剥离括号注释
    t1 = "致太湖咨询：（原文无槽位，零改动保留）"
    assert clean_zero_change_annotations(t1) == "致太湖咨询："

    t2 = "据此函，签字人兹宣布同意如下：（固定原文，零改动保留）"
    assert clean_zero_change_annotations(t2) == "据此函，签字人兹宣布同意如下："

    t3 = "1、同意向贵方提供贵方可能另外要求的与投标有关的任何证据或资料。（固定原文，零改动保留）"
    assert clean_zero_change_annotations(t3) == "1、同意向贵方提供贵方可能另外要求的与投标有关的任何证据或资料。"

    t4 = "授权声明段（模板固定原文，100%盲守）"
    assert clean_zero_change_annotations(t4) == "授权声明段"

    t5 = "（正文已写，零改动）"
    assert clean_zero_change_annotations(t5) == ""

    # 2. 保护合法业务括号不被误删
    t_legal = "根据贵方的 SZDZ-2026-HC008 号招标文件，正式授权下述签字人张三（法定代表人）代表我方（单位盖章），全权处理本次项目投标的有关事宜。"
    assert clean_zero_change_annotations(t_legal) == t_legal

    # 3. 测试无操作提案 (No-Op Proposal) 识别
    assert is_zero_change_or_no_op_proposal("致太湖咨询：（原文无槽位，零改动保留）", "致太湖咨询：") is True
    assert is_zero_change_or_no_op_proposal("（固定原文，零改动保留）", "") is True
    assert is_zero_change_or_no_op_proposal("无需写盘", "原内容") is True
    assert is_zero_change_or_no_op_proposal("原样保留", "原内容") is True

    # 4. 实际修改的提案不被误判为 No-Op
    t_filled = "根据贵方的 SZDZ-2026-HC008 号招标文件，正式授权下述签字人张三（法定代表人）代表我方四川某公司，全权处理本次项目投标的有关事宜。（固定原文，零改动保留）"
    orig_template = "根据贵方的 ________ 号招标文件，正式授权下述签字人________（法定代表人）代表我方________，全权处理本次项目投标的有关事宜。"
    assert is_zero_change_or_no_op_proposal(t_filled, orig_template) is False
    assert clean_zero_change_annotations(t_filled) == "根据贵方的 SZDZ-2026-HC008 号招标文件，正式授权下述签字人张三（法定代表人）代表我方四川某公司，全权处理本次项目投标的有关事宜。"


def test_fill_docx_proposals_in_dom_should_filter_zero_change_annotations():
    """测试 DOM 填充引擎对带有零改动注释的提案进行防护，确保原文档不被污染"""
    import tempfile
    import os
    from docx import Document
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    doc = Document()
    p1 = doc.add_paragraph("致苏州太湖项目管理咨询有限公司：")
    p2 = doc.add_paragraph("据此函，签字人兹宣布同意如下：")
    p3 = doc.add_paragraph("根据贵方的 ________ 号招标文件，正式授权下述签字人________代表我方________，全权处理本次项目投标的有关事宜。")

    proposals = [
        {
            "path": "/body/p[1]",
            "original_context": "致苏州太湖项目管理咨询有限公司：",
            "proposed_text": "致苏州太湖项目管理咨询有限公司：（原文无槽位，零改动保留）",
            "type": "sentence_batch"
        },
        {
            "path": "/body/p[2]",
            "original_context": "据此函，签字人兹宣布同意如下：",
            "proposed_text": "据此函，签字人兹宣布同意如下：（固定原文，零改动保留）",
            "type": "sentence_batch"
        },
        {
            "path": "/body/p[3]",
            "original_context": "根据贵方的 ________ 号招标文件，正式授权下述签字人________代表我方________，全权处理本次项目投标的有关事宜。",
            "proposed_text": "根据贵方的 SZDZ-2026-HC008 号招标文件，正式授权下述签字人张三代表我方四川在截建设工程有限公司，全权处理本次项目投标的有关事宜。（固定原文，零改动保留）",
            "type": "sentence_batch"
        }
    ]

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)
        count = fill_docx_proposals_in_dom(temp_path, proposals)

        res_doc = Document(temp_path)
        # p1 与 p2 为零改动段落，不应包含任何“（原文无槽位，零改动保留）”或“（固定原文，零改动保留）”
        assert res_doc.paragraphs[0].text.strip() == "致苏州太湖项目管理咨询有限公司："
        assert "原文无槽位" not in res_doc.paragraphs[0].text
        assert "零改动" not in res_doc.paragraphs[0].text

        assert res_doc.paragraphs[1].text.strip() == "据此函，签字人兹宣布同意如下："
        assert "固定原文" not in res_doc.paragraphs[1].text
        assert "零改动" not in res_doc.paragraphs[1].text

        # p3 有实质性填充，其末尾的“（固定原文，零改动保留）”应被彻底清洗剥离
        assert "固定原文" not in res_doc.paragraphs[2].text
        assert "零改动" not in res_doc.paragraphs[2].text
        assert "SZDZ-2026-HC008" in res_doc.paragraphs[2].text
        assert "张三" in res_doc.paragraphs[2].text
        assert "四川在截建设工程有限公司" in res_doc.paragraphs[2].text
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_is_fixed_slot_form_table_structural_detection():
    """测试基于纯结构与版面形态判定固定单元格填报表（零硬编码）"""
    from docx import Document
    from app.utils.table_utils import is_fixed_slot_form_table

    doc = Document()

    # 1. 结构 1：固定单项汇总表（1 行表头 + 1 行预置标的物数据行 + 1 行跨列合并表尾行）
    tbl_fixed = doc.add_table(rows=3, cols=4)
    tbl_fixed.rows[0].cells[0].text = "标的物"
    tbl_fixed.rows[0].cells[1].text = "技术指标"
    tbl_fixed.rows[0].cells[2].text = "总金额"
    tbl_fixed.rows[0].cells[3].text = "备注"
    tbl_fixed.rows[1].cells[0].text = "某分布式光伏发电工程项目"
    tbl_fixed.rows[1].cells[1].text = "满足招标文件技术规范"
    tbl_fixed.rows[1].cells[2].text = ""  # 待填
    tbl_fixed.rows[1].cells[3].text = ""
    tbl_fixed.rows[2].cells[0].merge(tbl_fixed.rows[2].cells[3])
    tbl_fixed.rows[2].cells[0].text = "总计大写金额："

    assert is_fixed_slot_form_table(tbl_fixed) is True

    # 2. 结构 2：动态多行展开清单表（1 行表头 + 1 行全空占位行 + 1 行合计行，总行数 > 3 或多数据行）
    tbl_dynamic = doc.add_table(rows=5, cols=6)
    for i, h in enumerate(["序号", "设备名称", "规格型号", "数量", "单价", "合价"]):
        tbl_dynamic.rows[0].cells[i].text = h
    # 中间均为待展开行
    assert is_fixed_slot_form_table(tbl_dynamic) is False


def test_is_narrative_clause_or_lead_in_generic_rules():
    """测试基于通用语法标点与篇章结构判定正文叙述句/条款导语（零硬编码）"""
    from app.utils.table_utils import is_narrative_clause_or_lead_in

    # 1. 叙述句/公文导语/条款标题（必须返回 True，严禁被误当成表单属性标签）
    narrative_samples = [
        "据此函，签字人兹宣布同意如下：",
        "5、与本投标有关的正式通讯地址为：",
        "致苏州大智项目管理咨询有限公司：",
        "致某某招标代理机构：",
        "1、同意向贵方提供贵方可能另外要求的与投标有关的任何证据或资料；",
        "2、我们完全理解贵方不一定将合同授予最低报价的投标人；",
        "根据贵方的 SZDZ-2026-NG008 号招标文件，正式授权下述签字人李四代表我方，全权处理本次项目投标的有关事宜。",
        "我方在此声明如下：",
        "本投标人郑重承诺如下：",
        "现授权如下：",
        "特此声明：",
        "（一）关于资格证明文件的书面声明如下：",
    ]

    for s in narrative_samples:
        assert is_narrative_clause_or_lead_in(s) is True, f"'{s}' 应被识别为正文叙述/导语句"

    # 2. 真正的表单属性标签/留白槽位（必须返回 False，允许自愈或填报）
    valid_property_labels = [
        "地    址：                                ",
        "地 址：",
        "电    话：                                ",
        "电 话：",
        "邮    编：                                ",
        "传    真：                                ",
        "投标单位代表姓名（签字）：                ",
        "投标单位名称：                            ",
        "公    章：                               ",
        "日    期：     年      月     日",
        "法定代表人（签字）：",
        "授权代表（签字）：",
        "法定代表人或授权代表签字（或盖章）：                       年     月    日",
        "投标人全称（加盖公章）：",
        "开户银行及账号：",
        "统一社会信用代码：",
    ]

    for lbl in valid_property_labels:
        assert is_narrative_clause_or_lead_in(lbl) is False, f"'{lbl}' 应被识别为合法表单属性标签"


def test_reset_chapter_to_template_should_cleanly_reset_target_chapter():
    """测试单章节重置器能够将工作副本中被污染的段落100%精准重置回纯净模板状态"""
    import tempfile, os
    from docx import Document
    from app.utils.table_utils import reset_chapter_to_template

    # 1. 构建纯净模板
    doc_tpl = Document()
    doc_tpl.add_heading("一、封面", level=1)
    doc_tpl.add_paragraph("项目名称：______")
    doc_tpl.add_heading("二、投标函格式", level=1)
    doc_tpl.add_paragraph("致某某咨询公司：")
    doc_tpl.add_paragraph("根据贵方的______号招标文件，正式授权下述签字人______代表我方______...")
    doc_tpl.add_paragraph("据此函，签字人兹宣布同意如下：")
    doc_tpl.add_paragraph("5、与本投标有关的正式通讯地址为：")
    doc_tpl.add_paragraph("地    址：                                ")
    doc_tpl.add_heading("三、开标一览表", level=1)
    doc_tpl.add_paragraph("开标一览表说明")

    # 2. 构建被历史错误填报污染过的工作副本
    doc_work = Document()
    doc_work.add_heading("一、封面", level=1)
    doc_work.add_paragraph("项目名称：已填充的某光伏项目")
    doc_work.add_heading("二、投标函格式", level=1)
    doc_work.add_paragraph("致某某咨询公司：")
    doc_work.add_paragraph("根据贵方的SZDZ-001号招标文件，正式授权下述签字人李四代表我方四川某公司...")
    doc_work.add_paragraph("据此函，签字人兹宣布同意如下： 张三")
    doc_work.add_paragraph("5、与本投标有关的正式通讯地址为： 四川省成都市高新区128号")
    doc_work.add_paragraph("地    址：四川省成都市高新区128号")
    doc_work.add_heading("三、开标一览表", level=1)
    doc_work.add_paragraph("开标一览表说明")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf_tpl, \
         tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf_work:
        tpl_path = tf_tpl.name
        work_path = tf_work.name

    try:
        doc_tpl.save(tpl_path)
        doc_work.save(work_path)

        # 执行针对“二、投标函格式”的单章节精准重置
        success = reset_chapter_to_template(work_path, tpl_path, "二、投标函格式")
        assert success is True

        res_doc = Document(work_path)
        res_texts = [p.text for p in res_doc.paragraphs]

        # 验证“二、投标函格式”中的污染内容被 100% 清空重置为模板原文
        assert "据此函，签字人兹宣布同意如下：" in res_texts
        assert not any("据此函，签字人兹宣布同意如下： 张三" in t for t in res_texts)
        assert "5、与本投标有关的正式通讯地址为：" in res_texts
        assert not any("5、与本投标有关的正式通讯地址为： 四川省" in t for t in res_texts)
        assert "地    址：                                " in res_texts

        # 验证其余章节（如“一、封面”的已填报内容）未被影响
        assert "项目名称：已填充的某光伏项目" in res_texts
    finally:
        for p in [tpl_path, work_path]:
            if os.path.exists(p):
                os.remove(p)







