"""标书写盘路径与产物回读校验回归测试。"""

import os
import tempfile
from decimal import Decimal

from docx import Document

from app.agents.bid_filler_agent import (
    _normalize_table_path,
    blocked_docx_node,
    fill_docx_proposals_in_dom,
    proposals_to_commands,
    validate_filled_docx_integrity,
)


def _create_opening_summary_doc() -> str:
    """创建最小化开标一览表测试文档。"""
    doc = Document()
    table = doc.add_table(rows=3, cols=4)
    for index, header in enumerate(["项目名称", "技术要求", "总价（元）", "备注"]):
        table.rows[0].cells[index].text = header
    table.rows[1].cells[0].text = "测试光伏项目"
    table.rows[1].cells[1].text = "详见第四章项目需求"
    table.rows[2].cells[0].text = "投标总报价（大写）"

    file_handle = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    file_handle.close()
    doc.save(file_handle.name)
    return file_handle.name


def test_normalize_table_path_should_convert_td_and_cell_to_tc():
    """测试外部 td/cell 路径统一转换为 Word tc 路径。"""
    assert _normalize_table_path("/body/tbl[1]/tr[2]/td[3]/p[1]") == "/body/tbl[1]/tr[2]/tc[3]/p[1]"
    assert _normalize_table_path("/body/tbl[1]/row[2]/cell[3]") == "/body/tbl[1]/row[2]/tc[3]"


def test_proposals_to_commands_should_normalize_td_path_without_row_fallback():
    """测试 td 路径会被规范化，非法表格容器路径不会被放行。"""
    proposals = [
        {
            "path": "/body/tbl[1]/tr[2]/td[3]",
            "proposed_text": "1023645.11",
        },
        {
            "path": "/body/tbl[1]/tr[2]",
            "proposed_text": "错误的整行写入",
        },
    ]

    commands, approved, rejected = proposals_to_commands(proposals)

    assert approved == 1
    assert rejected == 1
    assert commands[0]["path"] == "/body/tbl[1]/tr[2]/tc[3]/p[1]"


def test_fill_docx_proposals_in_dom_should_write_td_path_to_target_cell():
    """测试 td 路径不会退化写入数据行首单元格。"""
    temp_path = _create_opening_summary_doc()
    try:
        proposals = [
            {
                "path": "/body/tbl[1]/tr[2]/td[3]/p[1]",
                "proposed_text": "1023645.11",
                "type": "text",
            },
            {
                "path": "/body/tbl[1]/tr[3]/td[2]/p[1]",
                "proposed_text": "壹佰零贰万叁仟陆佰肆拾伍元壹角壹分",
                "type": "text",
            },
        ]

        written_count = fill_docx_proposals_in_dom(temp_path, proposals)
        result = Document(temp_path)

        assert written_count == 2
        assert result.tables[0].rows[1].cells[0].text == "测试光伏项目"
        assert result.tables[0].rows[1].cells[2].text == "1023645.11"
        assert "壹佰零贰万叁仟陆佰肆拾伍元壹角壹分" in result.tables[0].rows[2].cells[1].text
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_validate_filled_docx_integrity_should_detect_wrong_opening_summary_value():
    """测试产物回读能识别金额错列和大写金额行被公司名称污染。"""
    temp_path = _create_opening_summary_doc()
    try:
        doc = Document(temp_path)
        table = doc.tables[0]
        table.rows[1].cells[2].text = "无偏离"
        table.rows[2].cells[1].text = "四川石楠建设工程有限公司"
        doc.save(temp_path)

        findings = validate_filled_docx_integrity(
            temp_path,
            expected_total=Decimal("1023645.11"),
            expected_total_words="壹佰零贰万叁仟陆佰肆拾伍元壹角壹分",
            company_name="四川石楠建设工程有限公司",
            proposals=[
                {
                    "path": "/body/tbl[1]/tr[2]/td[3]/p[1]",
                    "proposed_text": "1023645.11",
                    "chapter_title": "开标一览表",
                }
            ],
        )

        finding_types = {finding["type"] for finding in findings}
        assert "opening_total_mismatch" in finding_types
        assert "opening_total_semantic_mismatch" in finding_types
        assert "opening_total_words_mismatch" in finding_types
        assert "opening_footer_wrong_value" in finding_types
        assert "proposal_writeback_mismatch" in finding_types
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_validate_filled_docx_integrity_should_ignore_placeholder_and_amount_formatting():
    """回读校验不应把待补充占位符或金额千分位格式差异判为失败。"""
    temp_path = _create_opening_summary_doc()
    try:
        doc = Document(temp_path)
        doc.tables[0].rows[1].cells[2].text = "1,023,645.11"
        doc.tables[0].rows[1].cells[1].text = "实际技术要求"
        doc.tables[0].rows[1].cells[0].text = "张三"
        doc.save(temp_path)

        findings = validate_filled_docx_integrity(
            temp_path,
            proposals=[
                {
                    "path": "/body/tbl[1]/tr[2]/tc[1]",
                    "proposed_text": "[待补充: 姓名]",
                    "chapter_title": "项目负责人及其他人员介绍",
                },
                {
                    "path": "/body/tbl[1]/tr[2]/tc[3]",
                    "proposed_text": "1023645.11",
                    "chapter_title": "开标一览表",
                },
            ],
            expected_total=Decimal("1023645.11"),
        )

        assert not any(item["type"] == "proposal_writeback_mismatch" for item in findings)
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_blocked_docx_node_should_preserve_non_empty_working_copy():
    """阻断发布时仍应返回已填工作副本，避免 API 回退导出空模板。"""
    temp_path = _create_opening_summary_doc()
    try:
        with open(temp_path, "rb") as work_file:
            working_bytes = work_file.read()

        result = blocked_docx_node(
            {
                "document_id": "test-blocked-doc",
                "docx_temp_path": temp_path,
                "original_docx": b"original-template",
                "audit_items": [],
                "review_findings": [],
            }
        )

        assert result["audit_blocked"] is True
        assert result["filled_docx_bytes"] == working_bytes
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)
