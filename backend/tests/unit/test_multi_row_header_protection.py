"""
多行复合表头识别与数据写盘保护单元测试 (test_multi_row_header_protection.py)
"""

import os
import tempfile
import pytest
from docx import Document

from app.utils.table_utils import detect_table_header_rows, get_merged_header_texts
from app.agents.bid_filler_agent import fill_docx_proposals_in_dom
from app.agents.bid_filler_workers import extract_docx_tables_summary


def test_detect_table_header_rows_single_header_should_return_1():
    """测试单行普通表头的表格，识别表头行数为 1"""
    doc = Document()
    table = doc.add_table(rows=3, cols=4)
    # 行 0：表头
    headers = ["序号", "项目名称", "金额", "备注"]
    for c_i, h in enumerate(headers):
        table.rows[0].cells[c_i].text = h
    # 行 1、2：数据行
    table.rows[1].cells[0].text = "1"
    table.rows[2].cells[0].text = "2"

    hdr_count = detect_table_header_rows(table)
    assert hdr_count == 1
    merged = get_merged_header_texts(table, hdr_count)
    assert merged == ["序号", "项目名称", "金额", "备注"]


def test_detect_table_header_rows_multi_row_header_should_return_2():
    """测试 2 行复合表头（如资格证书包含类别与编号子列），识别表头行数为 2"""
    doc = Document()
    table = doc.add_table(rows=5, cols=6)
    # 行 0：父表头（前4列合并与未合并，第4、5列为资格证书跨列合并）
    table.rows[0].cells[0].text = "序号"
    table.rows[0].cells[1].text = "姓名"
    table.rows[0].cells[2].text = "性别"
    table.rows[0].cells[3].text = "技术职称"
    # 第 4、5 列在行 0 横向合并为"资格证书"
    cell_cert = table.rows[0].cells[4]
    cell_cert.merge(table.rows[0].cells[5])
    cell_cert.text = "资格证书"

    # 行 1：子表头（第4列为类别，第5列为编号）
    table.rows[1].cells[0].text = ""
    table.rows[1].cells[1].text = ""
    table.rows[1].cells[2].text = ""
    table.rows[1].cells[3].text = ""
    table.rows[1].cells[4].text = "类别"
    table.rows[1].cells[5].text = "编号"

    # 行 2~4：空数据行
    for r in range(2, 5):
        for c in range(6):
            table.rows[r].cells[c].text = ""

    hdr_count = detect_table_header_rows(table)
    assert hdr_count == 2

    merged = get_merged_header_texts(table, hdr_count)
    assert len(merged) == 6
    assert merged[0] == "序号"
    assert merged[1] == "姓名"
    assert "资格证书(类别)" in merged[4] or "类别" in merged[4]
    assert "资格证书(编号)" in merged[5] or "编号" in merged[5]


def test_fill_docx_proposals_in_dom_should_protect_multi_row_headers_from_overwriting():
    """测试向具有 2 行复合表头的表格写入数据时，第 1 行与第 2 行表头均完好保留，数据从第 3 行正确填充"""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_docx_path = tf.name

    try:
        doc = Document()
        doc.add_heading("十、项目负责人及其他人员介绍", level=1)
        table = doc.add_table(rows=4, cols=5)

        # 行 0 (表头第1行)
        table.rows[0].cells[0].text = "序号"
        table.rows[0].cells[1].text = "姓名"
        table.rows[0].cells[2].text = "职务"
        cell_c = table.rows[0].cells[3]
        cell_c.merge(table.rows[0].cells[4])
        cell_c.text = "资格证书"

        # 行 1 (表头第2行: 子表头)
        table.rows[1].cells[0].text = ""
        table.rows[1].cells[1].text = ""
        table.rows[1].cells[2].text = ""
        table.rows[1].cells[3].text = "类别"
        table.rows[1].cells[4].text = "编号"

        # 行 2、3 为空白数据行
        doc.save(temp_docx_path)

        # 构造 Worker 提交的表格数据行提案 (2 行数据)
        proposals = [
            {
                "path": "/body/tbl[1]",
                "type": "table_rows",
                "proposed_text": '[["1", "人员甲", "项目负责人", "建造师", "BH001"], ["2", "人员乙", "技术负责人", "高级职称", "BH002"]]',
                "status": "success",
            }
        ]

        filled_count = fill_docx_proposals_in_dom(temp_docx_path, proposals)
        assert filled_count > 0

        # 重新读取文档并验证
        res_doc = Document(temp_docx_path)
        res_table = res_doc.tables[0]

        # 验证表头第 0 行未被覆盖
        assert "资格证书" in res_table.rows[0].cells[3].text

        # 验证表头第 1 行（子表头行）未被第一行数据覆盖
        assert res_table.rows[1].cells[3].text.strip() == "类别"
        assert res_table.rows[1].cells[4].text.strip() == "编号"

        # 验证数据从第 2 行（行 index 2）开始写入
        assert res_table.rows[2].cells[0].text.strip() == "1"
        assert res_table.rows[2].cells[1].text.strip() == "人员甲"
        assert res_table.rows[2].cells[2].text.strip() == "项目负责人"
        assert res_table.rows[2].cells[3].text.strip() == "建造师"
        assert res_table.rows[2].cells[4].text.strip() == "BH001"

        # 验证第 3 行数据（行 index 3）
        assert res_table.rows[3].cells[0].text.strip() == "2"
        assert res_table.rows[3].cells[1].text.strip() == "人员乙"
        assert res_table.rows[3].cells[2].text.strip() == "技术负责人"
    finally:
        if os.path.exists(temp_docx_path):
            os.remove(temp_docx_path)
