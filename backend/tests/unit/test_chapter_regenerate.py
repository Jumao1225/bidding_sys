"""
单章节重新生成与 Prompt 微调功能单元测试 (test_chapter_regenerate.py)
"""
import os
import tempfile
from docx import Document

from app.schemas.bid_filler_schema import RegenerateChapterRequest, RegenerateChapterResponse
from app.agents.bid_filler_workers import build_worker_prompt
from app.agents.bid_filler_agent import fill_docx_proposals_in_dom


def test_regenerate_chapter_request_schema_validation():
    """验证单章节微调请求与响应模型结构与默认值。"""
    req = RegenerateChapterRequest(
        chapter_title="商务条款响应及偏差表",
        custom_prompt="所有偏离项全部填无偏离，响应时间不超过1小时",
        category="needs_data",
        mapping_hint="deviation"
    )
    assert req.chapter_title == "商务条款响应及偏差表"
    assert "无偏离" in req.custom_prompt
    assert req.category == "needs_data"
    assert req.mapping_hint == "deviation"

    resp = RegenerateChapterResponse(
        document_id="doc_test_123",
        chapter_title="商务条款响应及偏差表",
        status="success",
        summary="已完成微调",
        proposals_count=3,
        execution_time_ms=1200,
        total_tokens=4500
    )
    assert resp.status == "success"
    assert resp.proposals_count == 3


def test_build_worker_prompt_with_custom_prompt_priority():
    """验证传入微调提示词时，Worker 提示词中能正确注入最高优先级指令。"""
    chapter_title = "技术要求响应及偏离表"
    custom_prompt = "特别声明提供 7x24 小时现场技术支持与应急备件库"

    system_prompt, user_prompt = build_worker_prompt(
        chapter_title=chapter_title,
        category="needs_data",
        template_text="模板原文要求",
        content_hint="填写说明",
        document_id="doc_test_123",
        docx_temp_path="",
        mapping_hint="deviation",
        extra_instructions=custom_prompt,
    )

    assert "【用户单章节专属重新生成与微调指令 — 最高优先级】" in system_prompt
    assert "7x24 小时现场技术支持" in system_prompt
    assert chapter_title in user_prompt


def test_regenerate_chapter_proposals_writeback_in_dom():
    """验证单章节微调生成的提案能精准刷入 Word DOM。"""
    doc = Document()
    doc.add_paragraph("商务条款响应：______")
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
        doc.save(tmp_path)

    try:
        proposals = [
            {
                "path": "/body/p[1]",
                "proposed_text": "完全响应商务条款无偏离",
                "value": "完全响应商务条款无偏离",
                "type": "text"
            }
        ]

        count = fill_docx_proposals_in_dom(tmp_path, proposals)
        assert count > 0

        res_doc = Document(tmp_path)
        assert "完全响应商务条款无偏离" in res_doc.paragraphs[0].text
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def test_pricing_opening_duplicate_path_proposals_should_keep_only_project_info():
    """验证报价章节同一路径重复提案不会拼接标题和旧占位符。"""
    doc = Document()
    doc.add_paragraph("五、投标配置及分项报价表")
    doc.add_paragraph("投标报价分析表")
    doc.add_paragraph("招标编号：号                                 项目名称：")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
        doc.save(tmp_path)

    project_info = "招标编号：SZDZ-2026-NG008号 项目名称：和烁热能公司屋顶（400kW）分布式光伏发电项目"
    proposals = [
        {
            "path": "/body/p[3]",
            "proposed_text": f"投标报价分析表：{project_info}",
            "value": f"投标报价分析表：{project_info}",
            "type": "text",
        },
        {
            "path": "/body/p[3]",
            "proposed_text": project_info,
            "value": project_info,
            "type": "text",
        },
    ]

    try:
        count = fill_docx_proposals_in_dom(tmp_path, proposals)
        assert count > 0

        result_doc = Document(tmp_path)
        result_texts = [paragraph.text.strip() for paragraph in result_doc.paragraphs if paragraph.text.strip()]
        assert result_texts == ["五、投标配置及分项报价表", project_info]
        assert "投标报价分析表" not in "\n".join(result_texts)
        assert "号                                 项目名称" not in "\n".join(result_texts)
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def test_build_worker_prompt_pricing_workflow_direct_matrix():
    """验证造价专家工作流使用抽象通用的 2D 矩阵直通指引，无具体数据硬编码。"""
    chapter_title = "投标配置及分项报价表"
    system_prompt, user_prompt = build_worker_prompt(
        chapter_title=chapter_title,
        category="needs_data",
        template_text="模板原文要求",
        content_hint="填写说明",
        document_id="doc_pricing_test_456",
        docx_temp_path="",
        mapping_hint="pricing",
    )

    # 1. 验证包含 2D 矩阵直通指引与字段说明
    assert "cost_estimates_json_matrix" in system_prompt
    assert "造价工程师与分项报价专家" in system_prompt
    assert "严禁冗余查询" in system_prompt or "绝对禁止同时或重复调用" in system_prompt

    # 2. 验证防重复扫描规则
    assert "严禁重复盲目查询" in system_prompt or "严禁重复查询" in system_prompt

    # 3. 验证严禁硬编码具体设备与数据（杜绝具体锚定）
    for forbidden_word in ["光伏组件", "逆变器", "并网柜", "彩钢瓦", "2235211", "1634971"]:
        assert forbidden_word not in system_prompt
