"""
BidFormatExtractorService 单元测试

测试投标文件格式抽取服务的正向模式、回退托底结构构建与关键词正则切片。
"""

import io
import pytest
from app.services.bid_format_extractor_service import bid_format_extractor_service


def test_build_fallback_structure_should_contain_essential_bid_sections():
    """测试托底架构构建，确保包含投标函、授权书与报价汇总表等核心格式"""
    structure = bid_format_extractor_service._build_fallback_structure("智慧园区建设工程招标文件.pdf")
    
    assert "智慧园区建设工程招标文件" in structure.document_title
    assert len(structure.sections) >= 3
    section_titles = [s.section_title for s in structure.sections]
    assert any("投标函" in t for t in section_titles)
    assert any("授权" in t for t in section_titles)
    assert any("开标一览表" in t for t in section_titles)


def test_slice_text_by_keywords_should_locate_bid_format_chapter():
    """测试在多章节大文本中，正则匹配定位'投标文件格式'章节起止范围"""
    sample_text = """
第一章 招标公告
本工程项目......

第二章 投标人须知
投标人必须具备......

第六章 投标文件格式
附件一：投标函
致：采购人......

附件二：法定代表人授权书
    """
    sliced_text = bid_format_extractor_service._slice_text_by_keywords(sample_text)
    assert "第六章 投标文件格式" in sliced_text
    assert "附件一：投标函" in sliced_text
    assert "第一章 招标公告" not in sliced_text


def test_slice_text_by_keywords_should_prefer_real_yingda_format_chapter_over_toc():
    """测试目录与正文均出现标题时，只截取正文“第四章 应答文件格式”范围。"""
    sample_text = """
目 录
第一章 招标公告 ................................ 1
第四章 应答文件格式 ............................ 88
第五章 评审办法 ................................ 90

第一章 招标公告
公告正文
第四章 应答文件格式
第一册 技术标
有效认证证书
第二册 商务标
法定代表人身份证明
第五章 评审办法
评审正文
"""

    sliced_text = bid_format_extractor_service._slice_text_by_keywords(sample_text)

    assert "第一册 技术标" in sliced_text
    assert "第二册 商务标" in sliced_text
    assert "评审正文" not in sliced_text
    assert "第四章 应答文件格式 ............................ 88" not in sliced_text


def test_slice_text_by_keywords_should_follow_dynamic_toc_chapter_identity():
    """测试章节编号变化时，根据目录条目动态匹配正文而非依赖固定章号。"""
    sample_text = """
目 录
第九章 响应文件格式 ................................ 42
第十章 合同条款 .................................... 56

第一章 项目概况
项目正文
第九章 响应文件格式
响应函
报价明细表
第十章 合同条款
合同正文
"""

    sliced_text = bid_format_extractor_service._slice_text_by_keywords(sample_text)

    assert "第九章 响应文件格式" in sliced_text
    assert "响应函" in sliced_text
    assert "合同正文" not in sliced_text


def test_slice_text_by_keywords_should_return_empty_when_target_chapter_missing():
    """测试未定位到目标章节时返回空文本，禁止把整份文件交给 LLM 猜测。"""
    sample_text = "第一章 招标公告\n附件一：专用资质业绩要求\n评审办法前附表"

    assert bid_format_extractor_service._slice_text_by_keywords(sample_text) == ""


def test_is_toc_line_should_correctly_identify_table_of_contents_lines():
    """测试 _is_toc_line 能否精准判断带-3- / -55- 及第一卷等目录连点页码行，防止误杀切片"""
    toc_line_1 = "第六章 投标文件格式 .................... 40"
    toc_line_2 = "第七章 合同条款................................... 55"
    toc_line_3 = "第六章 投标文件格式\t40"
    toc_line_4 = "第一章 招标公告........................-3-"
    toc_line_5 = "第六章 投标文件格式....................-55-"
    toc_line_vol = "第一卷"

    body_line_1 = "第六章 投标文件格式"
    body_line_2 = "附件一 投标函"

    assert bid_format_extractor_service._is_toc_line(toc_line_1) is True
    assert bid_format_extractor_service._is_toc_line(toc_line_2) is True
    assert bid_format_extractor_service._is_toc_line(toc_line_3) is True
    assert bid_format_extractor_service._is_toc_line(toc_line_4) is True
    assert bid_format_extractor_service._is_toc_line(toc_line_5) is True
    assert bid_format_extractor_service._is_toc_line(toc_line_vol) is True

    assert bid_format_extractor_service._is_toc_line(body_line_1) is False
    assert bid_format_extractor_service._is_toc_line(body_line_2) is False


def test_is_real_next_main_chapter_should_distinguish_main_chapters_from_attachments():
    """测试 _is_real_next_main_chapter 能否区分真实独立大章与格式附件内部子标题"""
    real_chapter_1 = "第七章 评标办法"
    real_chapter_2 = "第七章 合同条款"
    
    internal_sub_1 = "附件七 承诺函"
    internal_sub_2 = "格式七 授权委托书"
    internal_sub_3 = "第七部分 商务响应表"

    assert bid_format_extractor_service._is_real_next_main_chapter(real_chapter_1) is True
    assert bid_format_extractor_service._is_real_next_main_chapter(real_chapter_2) is True

    assert bid_format_extractor_service._is_real_next_main_chapter(internal_sub_1) is False
    assert bid_format_extractor_service._is_real_next_main_chapter(internal_sub_2) is False
    assert bid_format_extractor_service._is_real_next_main_chapter(internal_sub_3) is False


def test_bid_format_structure_should_accept_table_template_and_synonyms():
    """测试 BidFormatStructure 能够兼容 'table_template' 等 LLM 返回的同义词枚举并正常解析"""
    from app.schemas.bid_generator import BidFormatStructure, BidFormatSection, ContentTypeEnum

    raw_data = {
        "document_title": "张家港市渔光互补项目 - 投标文件格式模板",
        "source_chapter_name": "投标文件格式",
        "sections": [
            {
                "section_title": "附件一 投标函",
                "content_type": "text_template",
                "body_markdown": "致：招标人\n我方...",
                "placeholders": ["招标人"]
            },
            {
                "section_title": "附件二 开标一览表",
                "content_type": "table_template",  # 此处测试之前报错的 input_value='table_template'
                "body_markdown": "| 项目 | 金额 |\n| --- | --- |\n| 总计 | ___ |",
                "placeholders": ["金额"]
            },
            {
                "section_title": "附件三 资格审查表",
                "content_type": "checklist",
                "body_markdown": "- [ ] 营业执照",
                "placeholders": None  # 测试 None 占位符容错
            },
            {
                "section_title": "附件四 其他材料",
                "content_type": "custom_unknown_type",  # 测试未知枚举类型回退容错
                "body_markdown": None,  # 测试 None Markdown 容错
                "placeholders": "单一字符串占位符"  # 测试字符串转列表容错
            }
        ]
    }

    structure = BidFormatStructure(**raw_data)
    assert len(structure.sections) == 4
    assert structure.sections[0].content_type == ContentTypeEnum.TEXT_TEMPLATE
    assert structure.sections[1].content_type == ContentTypeEnum.FORM_TABLE
    assert structure.sections[2].content_type == ContentTypeEnum.CHECKLIST
    assert structure.sections[2].placeholders == []
    assert structure.sections[3].content_type == ContentTypeEnum.OTHER
    assert structure.sections[3].body_markdown == ""
    assert structure.sections[3].placeholders == ["单一字符串占位符"]


def test_extract_with_llm_and_rebuild_with_mocked_llm_json(monkeypatch):
    """测试 _extract_with_llm_and_rebuild 在 LLM 返回 table_template 时能成功构建 Word 文档"""
    from unittest.mock import MagicMock

    mock_doc = MagicMock()
    mock_doc.id = "test-doc-id"
    mock_doc.filename = "光伏电站项目.pdf"
    mock_doc.tenant_id = "tenant-bid-format"
    mock_doc.parsed_metadata = {"md_file_path": ""}
    
    mock_db = MagicMock()
    mock_chunk = MagicMock()
    mock_chunk.content = "第六章 投标文件格式\n附件一 投标函\n致招标人：我方愿投标..."
    monkeypatch.setattr("app.services.bid_format_extractor_service.document_crud.get_document_chunks", lambda db, doc_id: [mock_chunk])

    # 模拟 LLM 返回包含 table_template 的 JSON
    mock_llm_service = MagicMock()
    mock_llm_service.is_configured_for_tenant.return_value = True
    mock_llm_service.generate_structured_json.return_value = {
        "document_title": "光伏电站项目 - 投标文件格式模板",
        "source_chapter_name": "第六章 投标文件格式",
        "sections": [
            {
                "section_title": "附件一 投标函",
                "content_type": "text_template",
                "body_markdown": "致招标人：我方愿投标...",
                "placeholders": ["投标报价"]
            },
            {
                "section_title": "附件二 分项报价表",
                "content_type": "table_template",
                "body_markdown": "| 序号 | 名称 | 数量 |\n| :--- | :--- | :--- |\n| 1 | 逆变器 | 10 |",
                "placeholders": []
            }
        ]
    }

    service = bid_format_extractor_service
    monkeypatch.setattr(service, "llm_service", mock_llm_service)
    
    docx_bytes, mode = service._extract_with_llm_and_rebuild(mock_db, mock_doc)
    assert docx_bytes is not None
    assert len(docx_bytes) > 0
    assert mode == "llm_rebuilt"
    mock_llm_service.is_configured_for_tenant.assert_called_once_with("tenant-bid-format")
    assert mock_llm_service.generate_structured_json.call_args.kwargs["tenant_id"] == "tenant-bid-format"


def test_extract_with_llm_should_skip_call_when_target_chapter_missing(monkeypatch):
    """测试目标章节缺失时直接降级，避免将其他章节误交给 LLM。"""
    from unittest.mock import MagicMock

    mock_doc = MagicMock()
    mock_doc.id = "test-missing-format-doc-id"
    mock_doc.filename = "缺少格式章节项目.pdf"
    mock_doc.tenant_id = "tenant-bid-format"
    mock_doc.parsed_metadata = {"md_file_path": ""}

    mock_db = MagicMock()
    mock_chunk = MagicMock()
    mock_chunk.content = "第一章 招标公告\n附件一：专用资质业绩要求\n评审办法前附表"
    monkeypatch.setattr(
        "app.services.bid_format_extractor_service.document_crud.get_document_chunks",
        lambda db, doc_id: [mock_chunk],
    )

    mock_llm_service = MagicMock()
    mock_llm_service.is_configured_for_tenant.return_value = True
    monkeypatch.setattr(bid_format_extractor_service, "llm_service", mock_llm_service)

    docx_bytes, mode = bid_format_extractor_service._extract_with_llm_and_rebuild(mock_db, mock_doc)

    assert docx_bytes
    assert mode == "fallback_template"
    mock_llm_service.generate_structured_json.assert_not_called()


def test_extract_with_llm_exception_should_log_and_return_fallback_mode(monkeypatch):
    """测试当 LLM 出现网络/解析异常时，正确记录异常堆栈并降级返回 fallback_template 模式"""
    from unittest.mock import MagicMock

    mock_doc = MagicMock()
    mock_doc.id = "test-error-doc-id"
    mock_doc.filename = "异常测试项目.pdf"
    mock_doc.tenant_id = "tenant-bid-format"
    mock_doc.parsed_metadata = {"md_file_path": ""}
    
    mock_db = MagicMock()
    mock_chunk = MagicMock()
    mock_chunk.content = "第六章 投标文件格式\n附件一 投标函\n致招标人..."
    monkeypatch.setattr("app.services.bid_format_extractor_service.document_crud.get_document_chunks", lambda db, doc_id: [mock_chunk])

    # 模拟 LLM 抛出异常
    mock_llm_service = MagicMock()
    mock_llm_service.is_configured_for_tenant.return_value = True
    mock_llm_service.generate_structured_json.side_effect = RuntimeError("网络超时或 API 连接异常")

    service = bid_format_extractor_service
    monkeypatch.setattr(service, "llm_service", mock_llm_service)
    
    docx_bytes, mode = service._extract_with_llm_and_rebuild(mock_db, mock_doc)
    assert docx_bytes is not None
    assert len(docx_bytes) > 0
    assert mode == "fallback_template"
    mock_llm_service.is_configured_for_tenant.assert_called_once_with("tenant-bid-format")


def test_slice_docx_natively_should_preserve_exact_word_elements(tmp_path):
    """测试原生 Word 切片模式能够准确切取第六章正文并排除前置章节与目录"""
    import os
    from docx import Document

    test_docx_path = str(tmp_path / "测试招标文件.docx")
    doc = Document()
    
    # 1. 模拟目录
    doc.add_paragraph("目  录")
    doc.add_paragraph("第一章 招标公告 .................... 1")
    doc.add_paragraph("第六章 投标文件格式 ................ 30")
    
    # 2. 模拟第一章正文
    doc.add_paragraph("第一章 招标公告")
    doc.add_paragraph("某某项目进行公开招标...")
    
    # 3. 模拟第六章正文（目标切片范围）
    doc.add_paragraph("第六章 投标文件格式")
    doc.add_paragraph("附件一 投标函")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "品名"
    table.rows[0].cells[1].text = "单价"
    table.rows[1].cells[0].text = "光伏组件"
    table.rows[1].cells[1].text = "1000"
    doc.add_paragraph("附件二 授权书")
    
    # 4. 模拟第七章正文（终止大章）
    doc.add_paragraph("第七章 评标办法")
    doc.add_paragraph("综合评分法细则...")

    doc.save(test_docx_path)

    # 执行原生 Word 切片
    docx_bytes = bid_format_extractor_service._slice_docx_natively(test_docx_path)
    assert docx_bytes is not None
    assert len(docx_bytes) > 0

    # 重新加载切片后的 docx 验证内容
    sliced_doc = Document(io.BytesIO(docx_bytes))
    all_text = "\n".join([p.text for p in sliced_doc.paragraphs])
    
    assert "第六章 投标文件格式" in all_text
    assert "附件一 投标函" in all_text
    assert "附件二 授权书" in all_text
    assert "第一章 招标公告" not in all_text
    assert "第七章 评标办法" not in all_text
    assert len(sliced_doc.tables) == 1


