"""
报价表格单价与分项总价分离及防错列自愈单元测试 (test_pricing_table_column_alignment.py)
"""

import os
import json
import pytest

from app.agents.bid_filler_agent import align_table_row_cells
from app.agents.review_engine import check_pricing_table_column_alignment


def test_align_table_row_cells_should_separate_unit_price_and_total_price_for_package_items():
    """测试 5 列表格中，包干/工程类项目若单价与总价重复，自动纠偏单价为破折号"""
    # 输入为 5 列，但单价与总价相同（如加固工程、防水工程等）
    row_pkg = ["2.14", "加固工程", "500000.00", "500000.00", ""]
    aligned = align_table_row_cells(row_pkg, total_cols=5, row_i=1)

    assert len(aligned) == 5
    assert aligned[0] == "2.14"
    assert aligned[1] == "加固工程"
    assert aligned[2] == "——"  # 单价列纠偏为破折号
    assert aligned[3] == "500000.00"  # 分项总价列保留总金额
    assert aligned[4] == ""


def test_align_table_row_cells_should_keep_unit_and_total_for_equipment_items():
    """测试 5 列表格中，设备采购项（具有独立单价与数量合价）正常保留单价与分项总价"""
    row_equip = ["2.1", "设备标的物", "800.00", "80000.00", ""]
    aligned = align_table_row_cells(row_equip, total_cols=5, row_i=0)

    assert len(aligned) == 5
    assert aligned[0] == "2.1"
    assert aligned[1] == "设备标的物"
    assert aligned[2] == "800.00"  # 单价
    assert aligned[3] == "80000.00"  # 合价
    assert aligned[4] == ""


def test_align_table_row_cells_4_cols_input_should_align_to_total_price():
    """测试 4 列包干工程数据输入到 5 列表格时，自动将金额对齐至分项总价列，单价列补破折号"""
    row_4_cols = ["1", "设计费", "0.00", ""]
    aligned = align_table_row_cells(row_4_cols, total_cols=5, row_i=0)

    assert len(aligned) == 5
    assert aligned[0] == "1"
    assert aligned[1] == "设计费"
    assert aligned[2] == "——"  # 单价列补破折号
    assert aligned[3] == "0.00"  # 分项总价列填入金额
    assert aligned[4] == ""


def test_check_pricing_table_column_alignment_should_detect_duplicate_pricing_and_propose_fix():
    """测试 R11 质检管线对 2D 矩阵中单价与总价重复填报的自动检测与修复"""
    raw_matrix = [
        ["1", "设计费", "0.00", "0.00", ""],
        ["2", "建设费", "—", "2000000.00", ""],
        ["2.1", "主设备", "1000.00", "50000.00", ""],
        ["2.2", "加固工程", "300000.00", "300000.00", ""],
    ]

    proposals = [
        {
            "path": "/body/tbl[2]",
            "type": "table_rows",
            "proposed_text": json.dumps(raw_matrix, ensure_ascii=False),
            "status": "success",
        }
    ]

    findings = check_pricing_table_column_alignment(proposals)
    assert len(findings) == 1
    f = findings[0]
    assert f["rule_id"] == "R11-PRICING-COL-DUPLICATE"
    assert f["auto_fixable"] is True

    # 验证修复后的提案中加固工程的单价已变为破折号
    fix_prop = f["fix_proposal"]
    fixed_matrix = json.loads(fix_prop["proposed_text"])
    assert fixed_matrix[3][1] == "加固工程"
    assert fixed_matrix[3][2] == "——"
    assert fixed_matrix[3][3] == "300000.00"


def test_fill_docx_proposals_in_dom_6cols_grid_5logical_pricing_table():
    """
    测试当原 Word 模板表头存在跨列合并 (底层网格 6 列，但逻辑上 5 列：Col 1 和 Col 2 合并为名称列) 时：
    1. 写入的数据行必须自动将 Col 1 和 Col 2 合并为单列；
    2. 单价必须准确写入 Col 3（对齐表头的单价列）；
    3. 分项总价必须准确写入 Col 4（对齐表头的分项总价列）；
    4. 绝不错位！
    """
    import tempfile
    from docx import Document
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    doc = Document()
    tbl = doc.add_table(rows=5, cols=6)

    # Row 0 (表头): Col 1 和 Col 2 合并
    tbl.rows[0].cells[0].text = "序号"
    tbl.rows[0].cells[1].merge(tbl.rows[0].cells[2])
    tbl.rows[0].cells[1].text = "项目/费用名称"
    tbl.rows[0].cells[3].text = "单价（元）"
    tbl.rows[0].cells[4].text = "分项总价（元）"
    tbl.rows[0].cells[5].text = "备注"

    # Row 1~Row 3: 预置旧数据行，首列有序号，Col 1 和 Col 2 合并
    tbl.rows[1].cells[0].text = "1"
    tbl.rows[1].cells[1].merge(tbl.rows[1].cells[2])
    tbl.rows[1].cells[1].text = "设计费"
    tbl.rows[2].cells[0].text = "2"
    tbl.rows[2].cells[1].merge(tbl.rows[2].cells[2])
    tbl.rows[2].cells[1].text = "建设费"
    tbl.rows[3].cells[1].merge(tbl.rows[3].cells[2])

    # Row 4 (合计行): 前四列合并
    tbl.rows[4].cells[0].merge(tbl.rows[4].cells[3])
    tbl.rows[4].cells[0].text = "合计总价（合计总价=设计费+建设费合计）"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)

        # 5 列报价数据矩阵
        pricing_matrix = [
            ["1", "设计费（含电力设计、建筑设计费用）", "0.00", "0.00", ""],
            ["2", "建设费（含设备购置费）", "——", "2235211.56", ""],
            ["2.1", "光伏组件", "882.69", "1167798.87", ""],
            ["2.2", "逆变器 1", "11555.00", "46220.00", ""],
            ["合计总价（合计总价=设计费+建设费合计）", "", "", "2235211.56", ""]
        ]

        proposals = [
            {
                "path": "/body/tbl[1]",
                "proposed_text": json.dumps(pricing_matrix, ensure_ascii=False),
                "type": "table_rows",
                "chapter_title": "投标配置及分项报价表"
            }
        ]

        count = fill_docx_proposals_in_dom(temp_path, proposals)
        assert count >= 4

        res_doc = Document(temp_path)
        res_tbl = res_doc.tables[0]

        # 校验 2.1 光伏组件行 (Row 3)
        row_comp = res_tbl.rows[3]
        assert row_comp.cells[0].text.strip() == "2.1"
        # 核心断言 1: 名称列必须是单列（Col 1 和 Col 2 已合并）
        assert row_comp.cells[1]._tc == row_comp.cells[2]._tc
        assert row_comp.cells[1].text.strip() == "光伏组件"
        # 核心断言 2: 单价必须在 Col 3（对齐表头的单价列）
        assert row_comp.cells[3].text.strip() == "882.69"
        # 核心断言 3: 分项总价必须在 Col 4（对齐表头的分项总价列）
        assert row_comp.cells[4].text.strip() == "1167798.87"
        # 核心断言 4: 备注在 Col 5
        assert row_comp.cells[5].text.strip() == ""

        # 校验合计行 (Row 5)
        row_sum = res_tbl.rows[5]
        assert "合计总价" in row_sum.cells[0].text
        assert row_sum.cells[4].text.strip() == "2235211.56"

    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_add_rows_with_vmerge_footer_should_clean_vertical_merge_pollution():
    """
    测试当原模板表尾包含纵向合并 (vMerge) 标记时：
    1. 动态新增的多行数据必须彻底清洗 vMerge 标记；
    2. 新增数据行单元格绝不与落款行发生纵向粘连拉伸；
    3. 每一行数据独立、序号连续。
    """
    import tempfile
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    doc = Document()
    tbl = doc.add_table(rows=3, cols=4)

    # Row 0: 表头
    tbl.rows[0].cells[0].text = "序号"
    tbl.rows[0].cells[1].text = "设备名称"
    tbl.rows[0].cells[2].text = "单价"
    tbl.rows[0].cells[3].text = "总价"

    # Row 1: 空数据行
    # Row 2: 表尾落款行，跨列合并且首列带有 vMerge 属性
    tbl.rows[2].cells[0].merge(tbl.rows[2].cells[3])
    tbl.rows[2].cells[0].text = "交货期限及项目所在地说明：由承包人负责完成"
    vmerge_elem = parse_xml(f'<w:vMerge {nsdecls("w")} w:val="restart"/>')
    tcPr = tbl.rows[2].cells[0]._tc.get_or_add_tcPr()
    tcPr.append(vmerge_elem)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)

        # 插入 4 行数据（超出了原模版预置的 1 行数据区）
        data_matrix = [
            ["1", "设备 A", "100.00", "100.00"],
            ["2", "设备 B", "200.00", "200.00"],
            ["3", "设备 C", "300.00", "300.00"],
            ["4", "设备 D", "400.00", "400.00"],
        ]

        proposals = [
            {
                "path": "/body/tbl[1]",
                "proposed_text": json.dumps(data_matrix, ensure_ascii=False),
                "type": "table_rows",
                "chapter_title": "设备清单表"
            }
        ]

        count = fill_docx_proposals_in_dom(temp_path, proposals)
        assert count >= 4

        res_doc = Document(temp_path)
        res_tbl = res_doc.tables[0]

        # 验证所有数据行 (Row 1 ~ Row 4) 均无 vMerge 残留
        for r_i in range(1, 5):
            r = res_tbl.rows[r_i]
            for c in r.cells:
                tcPr = c._tc.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
                if tcPr is not None:
                    assert tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge') is None
            assert r.cells[0].text.strip() == str(r_i)

        # 验证总行数为 6（1 表头 + 4 数据 + 1 表尾）
        assert len(res_tbl.rows) == 6

    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)
