"""
BidFiller Worker & Office CLI MCP 工具直写与闭环修复单元测试 (test_bid_filler_worker.py)

遵循 AGENTS.md 测试规范：
1. 位于 tests/unit/ 目录下；
2. 函数命名遵循 test_<功能>_<场景>_<期望结果> 格式；
3. 包含正常情况、异常情况与边界情况测试；
4. 标记 @pytest.mark.asyncio。
"""

import os
import json
import tempfile
import pytest
from unittest.mock import MagicMock

from app.agents.tools.office_cli_agent_tools import (
    officecli_batch_fill_sentence_tool,
    officecli_fill_table_rows_tool,
    get_all_office_cli_agent_tools,
)
from app.agents.bid_filler_agent import (
    supervisor_audit_node,
    should_repair,
    BidFillerState,
)
from app.agents.bid_filler_workers import _build_worker_tools
from app.agents.tools.writer_tools import WRITER_TOOLS
from app.mcp.office_cli_server import officecli_create_docx
from app.services.office_cli_service import office_cli_service


@pytest.mark.asyncio
async def test_officecli_batch_fill_sentence_tool_should_succeed():
    """测试长句多槽位原子批处理写盘工具"""
    with tempfile.TemporaryDirectory() as tmp_dir:
        test_file = os.path.join(tmp_dir, "batch_fill_test.docx")
        await officecli_create_docx(file_path=test_file)

        # 1. 预先加入 2 个段落供替换写盘
        add_cmds = [
            {"command": "add", "parent": "/", "type": "paragraph", "props": {"text": "投标人名称：______"}},
            {"command": "add", "parent": "/", "type": "paragraph", "props": {"text": "法定代表人：______"}},
        ]
        await office_cli_service.apply_batch(test_file, add_cmds)

        # 2. 批量原子更新
        updates = [
            {"path": "/body/p[1]", "value": "投标人名称：聚猫科技股份有限公司"},
            {"path": "/body/p[2]", "value": "法定代表人：张三"},
        ]
        updates_json = json.dumps(updates, ensure_ascii=False)

        res = await officecli_batch_fill_sentence_tool.ainvoke({
            "file_path": test_file,
            "updates_json_str": updates_json
        })
        assert "成功" in res

        await office_cli_service.save_and_close(test_file)


@pytest.mark.asyncio
async def test_officecli_fill_table_rows_tool_auto_index_should_succeed():
    """测试表格全量追加填充工具（自动生成 1..N 递增序号与表头保护）"""
    with tempfile.TemporaryDirectory() as tmp_dir:
        test_file = os.path.join(tmp_dir, "table_fill_test.docx")
        await officecli_create_docx(file_path=test_file)

        # 1. 预先加入 1 个表格
        add_tbl_cmds = [
            {"command": "add", "parent": "/", "type": "table"}
        ]
        await office_cli_service.apply_batch(test_file, add_tbl_cmds)

        # 2. 模拟插入 2 行数据，auto_index=True 自动在第1列注入序号 1, 2
        rows = [
            ["项目经理", "张三", "PMP"],
            ["技术负责人", "李四", "高级工程师"]
        ]
        rows_json = json.dumps(rows, ensure_ascii=False)

        res = await officecli_fill_table_rows_tool.ainvoke({
            "file_path": test_file,
            "table_path": "/body/tbl[1]",
            "rows_json_str": rows_json,
            "auto_index": True
        })
        assert "成功" in res

        await office_cli_service.save_and_close(test_file)


def test_get_all_office_cli_agent_tools_count_should_be_six():
    """测试 Office CLI Agent 工具导出列表数量 (共 6 个工具)"""
    tools = get_all_office_cli_agent_tools()
    assert len(tools) == 6


def test_all_bid_writing_worker_roles_should_receive_style_extractor_tool():
    """测试标书撰写 Worker 的所有角色工具包均包含格式识别工具"""
    role_cases = [
        ("报价表", "pricing", "needs_data"),
        ("实质性要求响应对照表", "deviation", "needs_data"),
        ("资格审查表", "qualification", "needs_data"),
        ("投标函", "bid_letter", "needs_fill"),
        ("其他固定格式表单", "_unknown", "needs_fill"),
    ]

    for chapter_title, mapping_hint, category in role_cases:
        tools = _build_worker_tools(
            docx_temp_path="/tmp/bid-writing-tool-test.docx",
            chapter_title=chapter_title,
            mapping_hint=mapping_hint,
            category=category,
        )
        assert any(tool.name == "extract_text_by_style" for tool in tools), (
            f"章节 [{chapter_title}] 的 Worker 未分配 extract_text_by_style"
        )


def test_writer_tools_should_expose_style_extractor_tool():
    """测试通用标书写作工具集暴露格式识别工具"""
    assert any(tool.name == "extract_text_by_style" for tool in WRITER_TOOLS)


def test_supervisor_audit_node_should_block_when_artifact_is_missing():
    """测试 Supervisor 质量审查：缺少工作副本时必须阻断发布"""
    state: BidFillerState = {
        "document_id": "doc123",
        "original_context": "",
        "slot_analysis": None,
        "worker_proposals": None,
        "db_session": None,
        "company_profile": MagicMock(),
        "original_docx": None,
        "docx_temp_path": "/invalid/non_existent.docx",
        "custom_instructions": None,
        "category_hints": None,
        "repair_count": 0,
        "max_repair_rounds": 2,
        "repair_instructions_map": None,
        "audit_passed": None,
        "audit_items": [],
        "review_findings": [],
        "audit_report": None,
        "filled_docx_bytes": None,
    }

    result = supervisor_audit_node(state)
    assert result.get("audit_passed") is False
    assert result.get("audit_blocked") is True


def test_should_repair_routing_logic():
    """测试 Supervisor 审核条件边路由跳转"""
    state_pass: BidFillerState = {"audit_passed": True}  # type: ignore
    assert should_repair(state_pass) == "write_docx_node"

    state_fail: BidFillerState = {"audit_passed": False}  # type: ignore
    assert should_repair(state_fail) == "agent_fill_node"

    state_blocked: BidFillerState = {"audit_passed": False, "audit_blocked": True}  # type: ignore
    assert should_repair(state_blocked) == "blocked_docx_node"


def test_parse_proposals_should_skip_info_tables_and_extract_valid_table_rows():
    """测试 _parse_proposals 智能跳过前置数据来源说明表，准确提炼 /body/tbl[2] 等写盘明细表格"""
    from app.agents.bid_filler_workers import _parse_proposals

    raw_markdown = """
### 【投标配置及分项报价表（投标报价分析表）】— 数据来源与核验

| 数据项 | 检索结果 | 说明 |
|---|---|---|
| 项目名称 | 和烁热能公司屋顶（400kW）分布式光伏发电项目 | 项目元数据库直查 |
| 招标编号 | SZDZ-2026-NG008号 | 项目元数据库直查 |
| 投标总报价（小写） | 1017934.21元 | 财务库直查 |

### ✅ 写盘明细表

| 序号 | DOM 节点路径 | 替换前模板原文 | 实际填入/扩写结果 | 提议类型 | 写盘状态 |
|---|---|---|---|---|---|
| 1 | /body/p[@paraId=31FC6AF8] | 招标编号：号 项目名称： | 招标编号：SZDZ-2026-NG008号 项目名称：某某项目 | sentence_batch | ✅ 已提交刷盘 |
| 2 | /body/tbl[2]（16 行数据明细） | 18 行空白数据行 | [["光伏组件", "630Wp", "国内一线", "块", "763", "882.69", "673492.47", "合规"]] | table_rows | ✅ 已提交刷盘 |
"""
    proposals = _parse_proposals(raw_markdown)
    assert len(proposals) == 2
    assert proposals[0]["path"] == "/body/p[@paraId=31FC6AF8]"
    assert proposals[0]["type"] == "sentence_batch"
    assert proposals[1]["path"] == "/body/tbl[2]"
    assert proposals[1]["type"] == "table_rows"
    assert "光伏组件" in proposals[1]["proposed_text"]


def test_fill_docx_proposals_in_dom_table_overwrite_and_auto_index():
    """测试 DOM 引擎表格填报：原位覆盖已有空白行、序号列智能对齐与末尾汇总行（投标总报价/交货期限）保护"""
    from docx import Document
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    doc = Document()
    tbl = doc.add_table(rows=6, cols=9)
    # Row 0: 表头
    headers = ["序号", "标的物名称", "品牌、规格、型号", "生产厂家", "单位", "数量", "单价", "合价", "备注"]
    for c_idx, h in enumerate(headers):
        tbl.rows[0].cells[c_idx].text = h

    # Rows 1-3: 空白数据行
    for r_idx in range(1, 4):
        for c_idx in range(9):
            tbl.rows[r_idx].cells[c_idx].text = ""

    # Rows 4-5: 底部汇总与落款行
    tbl.rows[4].cells[0].text = "投标总报价：大写：壹佰万元整"
    tbl.rows[5].cells[0].text = "交货期限：60天内完成"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)

        # 传入未提供序号的 8 列业务数据
        matrix = [
            ["光伏组件", "630Wp", "国产", "块", "763", "882.69", "673492.47", "符合要求"],
            ["逆变器", "110kW", "国产", "台", "6", "11555", "69330", "符合要求"],
        ]
        proposals = [
            {
                "path": "/body/tbl[1]",
                "original_context": "空白表格",
                "proposed_text": json.dumps(matrix, ensure_ascii=False),
                "type": "table_rows"
            }
        ]

        count = fill_docx_proposals_in_dom(temp_path, proposals)
        assert count >= 2

        res_doc = Document(temp_path)
        res_tbl = res_doc.tables[0]

        # 验证 Row 1: 第一列自动注入序号 1，且后续字段完全对齐
        r1_cells = [c.text.strip() for c in res_tbl.rows[1].cells]
        assert r1_cells[0] == "1"
        assert r1_cells[1] == "光伏组件"
        assert r1_cells[2] == "630Wp"
        assert r1_cells[3] == "国产"
        assert r1_cells[4] == "块"
        assert r1_cells[5] == "763"
        assert r1_cells[6] == "882.69"
        assert r1_cells[7] == "673492.47"
        assert r1_cells[8] == "符合要求"

        # 验证 Row 2: 第一列自动注入序号 2
        r2_cells = [c.text.strip() for c in res_tbl.rows[2].cells]
        assert r2_cells[0] == "2"
        assert r2_cells[1] == "逆变器"

        # 验证底部汇总行完好无损未被覆盖
        assert "投标总报价" in res_tbl.rows[4].cells[0].text
        assert "交货期限" in res_tbl.rows[5].cells[0].text

    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_parse_proposals_should_auto_expand_markdown_range_paths():
    """测试 Markdown 表格解析器：自动将连续的范围概括路径 (如 /body/tbl[6]/tr[2]~tr[19]) 展开为递增行号独立提案"""
    from app.agents.bid_filler_workers import _parse_proposals

    raw_markdown = """
### 写盘明细表
| 序号 | DOM 节点路径 | 替换前模板原文 | 实际填入/扩写结果 | 提议类型 | 写盘状态 |
|---|---|---|---|---|---|
| 1 | /body/tbl[6] | 表头... | 表头原样保留 | — | 已保护 |
| 2 | /body/tbl[6]/tr[2]~tr[19] | 空白 | 标的物A技术响应承诺详情 | sentence_batch | 已入刷盘队列 |
| 3 | /body/tbl[6]/tr[2]~tr[19] | 同上 | 标的物B技术响应承诺详情 | sentence_batch | 已入刷盘队列 |
| 4 | /body/tbl[6]/tr[2]~tr[19] | 同上 | 标的物C技术响应承诺详情 | sentence_batch | 已入刷盘队列 |
| 5 | /body/tbl[6]/tr[2]~tr[19] | 同上 | 标的物D技术响应承诺详情 | sentence_batch | 已入刷盘队列 |
"""

    proposals = _parse_proposals(raw_markdown)
    assert len(proposals) == 4
    # 验证 4 条概括路径已被成功自动展开为 tr[2] 到 tr[5] 的连续物理行
    assert proposals[0]["path"] == "/body/tbl[6]/tr[2]"
    assert "标的物A" in proposals[0]["proposed_text"]
    assert proposals[1]["path"] == "/body/tbl[6]/tr[3]"
    assert "标的物B" in proposals[1]["proposed_text"]
def test_parse_proposals_should_filter_pseudo_plus_summary_syntax():
    """测试 _parse_proposals 能够纯通用地拦截过滤大模型偷懒输出的加号运算符伪拼接表达式与行级模糊路径"""
    from app.agents.bid_filler_workers import _parse_proposals

    raw_markdown = """
### 写盘明细表
| 序号 | DOM 节点路径 | 替换前模板原文 | 实际填入/扩写结果 | 提议类型 | 写盘状态 |
|---|---|---|---|---|---|
| 1 | /body/tbl[2]/tr[2] | （空单元格，序号"1"） | 条款A原文描述 + 对应技术服务承诺 + "无" + "完全响应招标文件要求，无偏离。" | text | 已提交刷盘队列 |
| 2 | /body/tbl[2]/tr[11] | "..."（扩展占位行） | 条款B要求详情 + 承诺描述 + "无" | text | 已提交刷盘队列 |
| 3 | /body/tbl[2] | 表头原样保留 | 表头原样保留 | — | 已保护 |
| 4 | /body/p[12] | 授权代表：______ | 授权代表：张三 | sentence_batch | 已入刷盘队列 |
"""
    proposals = _parse_proposals(raw_markdown)
    # 前 3 条未指定单元格的行级模糊路径与伪语法概括全部被纯通用规则拦截，仅保留真实合法的段落提案
    assert len(proposals) == 1
    assert proposals[0]["path"] == "/body/p[12]"
    assert "张三" in proposals[0]["proposed_text"]


def test_fill_docx_proposals_in_dom_formatted_json_matrix_deviation_table_should_succeed():
    """测试技术偏离表等带有换行缩进的格式化 2D 矩阵 JSON 提案能否 100% 成功刷盘到 Word 表格中"""
    from docx import Document
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    doc = Document()
    tbl = doc.add_table(rows=4, cols=5)
    # Row 0: 表头（序号 / 招标文件要求 / 投标文件服务承诺 / 有无偏离 / 偏离内容及原因）
    headers = ["序号", "招标文件技术要求", "投标文件对应要求的服务承诺", "有无偏离", "偏离内容及原因"]
    for c_idx, h in enumerate(headers):
        tbl.rows[0].cells[c_idx].text = h

    # 预设 3 行空白模板行
    for r_idx in range(1, 4):
        for c_idx in range(5):
            tbl.rows[r_idx].cells[c_idx].text = ""

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)

        # 模拟大模型输出的带换行符与 4 列（未带序号）的偏离表 JSON 字符串
        formatted_json_matrix = (
            "[\n"
            '  ["光伏组件：规格型号≥630Wp，数量763块...", "我方完全响应并承诺：投报天合光能...", "无", "完全响应招标文件要求，无偏离。"],\n'
            '  ["逆变器：110kW逆变器2台...", "我方完全响应并承诺：投报华为组串式逆变器...", "无", "完全响应招标文件要求，无偏离。"],\n'
            '  ["并网柜：200kW室外柜2台...", "我方完全响应并承诺：投报诺电品牌...", "无", "完全响应招标文件要求，无偏离。"],\n'
            '  ["直流电缆：PV1-F 4mm2...", "我方完全响应并承诺：投报江南品牌...", "无", "完全响应招标文件要求，无偏离。"]\n'
            "]\n"
        )

        proposals = [
            {
                "path": "/body/tbl[1]",
                "original_context": "空白偏离表格",
                "proposed_text": formatted_json_matrix,
                "value": formatted_json_matrix,
                "type": "table_rows",
                "chapter_title": "技术要求响应及偏离表"
            }
        ]

        count = fill_docx_proposals_in_dom(temp_path, proposals)
        assert count == 4

        res_doc = Document(temp_path)
        res_tbl = res_doc.tables[0]

        # 验证总行数为 5 行（1 行表头 + 4 行数据）
        assert len(res_tbl.rows) == 5

        # 验证第 1 行数据：序号自动注入为 1，4 列内容精准对齐
        r1_cells = [c.text.strip() for c in res_tbl.rows[1].cells]
        assert r1_cells[0] == "1"
        assert "光伏组件" in r1_cells[1]
        assert "天合光能" in r1_cells[2]
        assert r1_cells[3] == "无"
        assert "完全响应" in r1_cells[4]

        # 验证第 4 行数据：自动追加新行，序号为 4
        r4_cells = [c.text.strip() for c in res_tbl.rows[4].cells]
        assert r4_cells[0] == "4"
        assert "直流电缆" in r4_cells[1]
        assert "江南品牌" in r4_cells[2]

    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_partial_deviation_matrix_should_not_erase_later_cell_proposals():
    """局部整表矩阵与单元格提案并存时，不能清空矩阵之后的已有数据行。"""
    import tempfile
    from docx import Document
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    doc = Document()
    table = doc.add_table(rows=23, cols=5)
    headers = ["序号", "招标文件商务条款中的要求", "投标文件对应要求的服务承诺", "有无偏离", "偏离内容及原因"]
    for c_idx, header in enumerate(headers):
        table.rows[0].cells[c_idx].text = header

    # 预置 22 行完整内容，模拟前一轮已经写入的结果。
    for r_idx in range(1, 23):
        values = [str(r_idx), f"原条款-{r_idx}", f"原承诺-{r_idx}", "无", "原偏离说明"]
        for c_idx, value in enumerate(values):
            table.rows[r_idx].cells[c_idx].text = value

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)
        partial_matrix = [
            [str(i), f"新条款-{i}", f"新承诺-{i}", "无", "完全响应招标文件要求，无偏离。"]
            for i in range(6, 16)
        ]
        proposals = [
            {
                "path": "/body/tbl[1]",
                "proposed_text": json.dumps(partial_matrix, ensure_ascii=False),
                "type": "table_rows",
                "chapter_title": "商务条款响应及偏离表",
            },
            {
                "path": "/body/tbl[1]/tr[13]/tc[3]",
                "proposed_text": "补写第12行承诺",
                "type": "text",
                "chapter_title": "商务条款响应及偏离表",
            },
        ]

        count = fill_docx_proposals_in_dom(temp_path, proposals)
        assert count >= 11

        result = Document(temp_path).tables[0]
        assert len(result.rows) == 23
        # 矩阵从序号 6 开始，必须按序号写入对应行；矩阵之后的行不能被清理。
        assert result.rows[1].cells[2].text.strip() == "原承诺-1"
        assert result.rows[6].cells[1].text.strip() == "新条款-6"
        assert result.rows[10].cells[2].text.strip() == "新承诺-10"
        assert result.rows[11].cells[1].text.strip() == "新条款-11"
        assert result.rows[11].cells[2].text.strip() == "新承诺-11"
        assert result.rows[12].cells[1].text.strip() == "新条款-12"
        assert result.rows[12].cells[2].text.strip() == "补写第12行承诺"
        assert result.rows[16].cells[2].text.strip() == "原承诺-16"
        assert result.rows[19].cells[2].text.strip() == "原承诺-20"
        assert result.rows[19].cells[3].text.strip() == "无"
        assert result.rows[19].cells[4].text.strip() == "原偏离说明"
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_full_deviation_matrix_should_trim_stale_rows_and_fill_all_columns():
    """完整的无序号商务矩阵应覆盖旧表，不保留历史追加行或空列。"""
    import tempfile
    from docx import Document
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    doc = Document()
    table = doc.add_table(rows=43, cols=5)
    headers = ["序号", "招标文件商务条款中的要求", "投标文件对应要求的服务承诺", "有无偏离", "偏离内容及原因"]
    for c_idx, header in enumerate(headers):
        table.rows[0].cells[c_idx].text = header
    for r_idx in range(1, 43):
        for c_idx, value in enumerate([str(r_idx), f"旧条款-{r_idx}", f"旧承诺-{r_idx}", "无", "旧说明"]):
            table.rows[r_idx].cells[c_idx].text = value

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)
        matrix = [
            [f"对应商务条款-{i}", f"我方承诺-{i}", "无", "完全响应招标文件要求，无偏离。"]
            for i in range(1, 33)
        ]
        matrix[-1][0] = "履约保证金：合同总价款的10%"
        proposals = [{
            "path": "/body/tbl[1]",
            "proposed_text": json.dumps(matrix, ensure_ascii=False),
            "type": "table_rows",
            "chapter_title": "商务条款响应及偏离表",
        }]

        fill_docx_proposals_in_dom(temp_path, proposals)
        result = Document(temp_path).tables[0]
        assert len(result.rows) == 33
        for r_idx in range(1, 33):
            values = [cell.text.strip() for cell in result.rows[r_idx].cells]
            assert values[0] == str(r_idx)
            assert all(values)
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_dual_path_merge_substantive_table_cell_proposals_should_preserve_all_cells():
    """测试实质性要求对照表等单单元格提案在双路融合时，绝不被误判为整表提案而误删其他单元格"""
    import re
    from docx import Document
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    # 1. 模拟 LLM 文本总结提取出的 20 个单元格提案 (tr[2]~tr[11] 的 tc[2] 与 tc[3])
    text_proposals = []
    for r in range(2, 12):
        text_proposals.append({
            "path": f"/body/tbl[1]/tr[{r}]/tc[2]/p[1]",
            "proposed_text": f"实质性要求条款内容_{r}",
            "value": f"实质性要求条款内容_{r}",
            "type": "text",
            "status": "success"
        })
        text_proposals.append({
            "path": f"/body/tbl[1]/tr[{r}]/tc[3]/p[1]",
            "proposed_text": "是",
            "value": "是",
            "type": "text",
            "status": "success"
        })

    # 2. 模拟工具调用捕获的单格权威提案（如最后一行单元格）
    chapter_collected_proposals = [
        {
            "path": "/body/tbl[1]/tr[11]/tc[3]/p[1]",
            "proposed_text": "是",
            "value": "是",
            "type": "text",
            "status": "success"
        }
    ]

    # 3. 运行双路融合逻辑
    proposals_dict = {}
    for p in text_proposals:
        p_path = str(p.get("path", "")).strip()
        if p_path:
            proposals_dict[p_path] = p

    for p in chapter_collected_proposals:
        p_path = str(p.get("path", "")).strip()
        if p_path:
            proposals_dict[p_path] = p
            if p.get("type") == "table_rows" or re.match(r'^/body/tbl\[\d+\]$', p_path):
                m_tbl = re.match(r'^(/body/tbl\[\d+\])$', p_path)
                if m_tbl:
                    tbl_prefix = m_tbl.group(1)
                    to_del = [k for k in proposals_dict.keys() if k.startswith(tbl_prefix + "/") and k != p_path]
                    for k in to_del:
                        proposals_dict.pop(k, None)

    merged_proposals = list(proposals_dict.values())
    # 验证全部 20 个单元格提案 100% 完整保留，未被误删
    assert len(merged_proposals) == 20

    # 4. 验证写盘引擎能将这 20 个单元格提案完整刷盘到 Word 文档的表格中
    doc = Document()
    tbl = doc.add_table(rows=11, cols=3)
    tbl.rows[0].cells[0].text = "序号"
    tbl.rows[0].cells[1].text = "第四章中项目需求中实质性要求"
    tbl.rows[0].cells[2].text = "是否响应"
    for r in range(1, 11):
        tbl.rows[r].cells[0].text = str(r)
        tbl.rows[r].cells[1].text = ""
        tbl.rows[r].cells[2].text = ""

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)
        count = fill_docx_proposals_in_dom(temp_path, merged_proposals)
        assert count == 20

        res_doc = Document(temp_path)
        res_tbl = res_doc.tables[0]
        # 验证 10 行数据均已成功填入实质性要求与“是”
        for r in range(1, 11):
            assert f"实质性要求条款内容_{r+1}" in res_tbl.rows[r].cells[1].text
            assert res_tbl.rows[r].cells[2].text.strip() == "是"
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_fill_docx_proposals_in_dom_prefilled_items_table_should_overwrite_in_place_without_duplication():
    """测试供货一览表等已预置货物名称/序号/斜杠的模板表，2D 矩阵写盘时 100% 原位覆盖填报，绝不重复生成多余行把原模板挤到末尾"""
    import json
    from docx import Document
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    doc = Document()
    # 模拟供货一览表：1 行表头 + 14 行预置标的物数据行
    tbl = doc.add_table(rows=15, cols=7)
    headers = ["编号", "货物名称", "品牌、型号（或规格）", "数量", "产地", "制造厂商", "备注"]
    for c_idx, h in enumerate(headers):
        tbl.rows[0].cells[c_idx].text = h

    item_names = [
        "光伏组件", "逆变器 1", "逆变器 2", "并网柜", "直流电缆",
        "交流电缆（逆变器）", "交流电缆（并网 1）", "交流电缆（并网 2）", "电缆桥架", "光伏支架",
        "接地系统", "其他辅材", "防水", "加固"
    ]
    for idx, name in enumerate(item_names, 1):
        tbl.rows[idx].cells[0].text = str(idx)
        tbl.rows[idx].cells[1].text = name
        tbl.rows[idx].cells[2].text = ""
        tbl.rows[idx].cells[3].text = "/"  # 预置斜杠
        tbl.rows[idx].cells[4].text = ""
        tbl.rows[idx].cells[5].text = ""
        tbl.rows[idx].cells[6].text = ""

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)

        # 模拟 Worker 产出的 14 行 6 列（未含首列序号，底层自动补序号）的 2D 矩阵提案
        matrix = []
        for idx, name in enumerate(item_names, 1):
            matrix.append([
                name,
                f"投报品牌型号_{idx}",
                f"{idx * 100}台" if idx <= 3 else "1批",
                "中国",
                f"制造厂商_{idx}",
                "满足技术要求"
            ])

        proposals = [
            {
                "path": "/body/tbl[1]",
                "proposed_text": json.dumps(matrix, ensure_ascii=False),
                "value": json.dumps(matrix, ensure_ascii=False),
                "type": "table_rows",
                "chapter_title": "供货一览表（主要标的物）"
            }
        ]

        count = fill_docx_proposals_in_dom(temp_path, proposals)
        assert count == 14

        res_doc = Document(temp_path)
        res_tbl = res_doc.tables[0]

        # 核心断言 1：总行数严格保持为 15 行（1 表头 + 14 数据行），严禁出现 29 行（14 行新增 + 14 行残留）！
        assert len(res_tbl.rows) == 15

        # 核心断言 2：每一行均在原有位置被精准更新填满，原标的物名称与新填入的品牌厂商完美对齐
        for idx, name in enumerate(item_names, 1):
            r_cells = [c.text.strip() for c in res_tbl.rows[idx].cells]
            assert r_cells[0] == str(idx)
            assert r_cells[1] == name
            assert r_cells[2] == f"投报品牌型号_{idx}"
            assert r_cells[4] == "中国"
            assert r_cells[5] == f"制造厂商_{idx}"
            assert r_cells[6] == "满足技术要求"

    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_fill_docx_proposals_in_dom_hierarchical_sub_items_table_should_overwrite_and_expand_cleanly():
    """测试投标配置及分项报价表：大类 1/2 与二级分项 2.1~2.14 展开，原位覆盖占位行并消除 ...... 占位符"""
    import json
    from docx import Document
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    doc = Document()
    # 模拟模板：1 表头 + 5 模板行（设计费/建设费/空行/空行/......）+ 1 表尾合并合计行
    tbl = doc.add_table(rows=7, cols=5)
    headers = ["序号", "项目/费用名称", "单价（元）", "分项总价（元）", "备注"]
    for c_idx, h in enumerate(headers):
        tbl.rows[0].cells[c_idx].text = h

    tbl.rows[1].cells[0].text = "1"
    tbl.rows[1].cells[1].text = "设计费（含电力设计、建筑设计费用）"
    tbl.rows[2].cells[0].text = "2"
    tbl.rows[2].cells[1].text = "建设费（含设备购置费）"
    tbl.rows[3].cells[0].text = ""
    tbl.rows[4].cells[0].text = ""
    tbl.rows[5].cells[0].text = "......"

    # 表尾合并合计行
    tbl.rows[6].cells[0].merge(tbl.rows[6].cells[4])
    tbl.rows[6].cells[0].text = "合计总价（合计总价=设计费+建设费合计）：与开标一览表一致"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)

        # 模拟 Worker 提交的 16 行二维矩阵（1个设计费 + 1个建设费汇总 + 14个 2.1~2.14 二级细项）
        matrix = [
            ["1", "设计费（含电力设计、建筑设计费用）", "0.00", "0.00", "深化设计费用已包含于建设费报价内，不再单独计列"],
            ["2", "建设费（含设备购置费）", "2,235,211.56", "2,235,211.56", "包含全部设备购置费、材料费、安装调试费及防水加固等费用，下含2.1~2.14分项明细"],
            ["2.1", "光伏组件", "882.69", "1,167,798.87", "天合光能 635Wp 单晶硅组件 1323台"],
            ["2.2", "逆变器 1", "11,555.00", "46,220.00", "华为 100kW 组串式逆变器 4台"],
            ["2.3", "逆变器 2", "11,555.00", "46,220.00", "华为 110kW 组串式逆变器 4台"],
            ["2.4", "并网柜", "9,180.35", "9,180.35", "诺电 200kW 室外并网柜 1台"],
            ["2.5", "直流电缆", "25,720.56", "25,720.56", "江南 PV1-F 4mm² 光伏专用直流电缆 1批"],
            ["2.6", "交流电缆（逆变器）", "17,859.45", "17,859.45", "江南 ZRC-YJLHV 3*120+1*70 1批"],
            ["2.7", "交流电缆（并网 1）", "11,882.62", "11,882.62", "江南 ZRC-YJV 3*240+1*120 1批"],
            ["2.8", "交流电缆（并网 2）", "33,600.00", "33,600.00", "江南 ZRC-YJV 3*150+1*70 1批"],
            ["2.9", "电缆桥架", "9,233.06", "9,233.06", "热镀锌钢制梯式桥架 1批"],
            ["2.10", "光伏支架", "117,600.00", "117,600.00", "彩钢瓦屋面专用铝合金支架 1套"],
            ["2.11", "接地系统", "6,580.98", "6,580.98", "光伏专用接地防雷系统 1批"],
            ["2.12", "其他辅材", "15,035.67", "15,035.67", "穿线管、接线盒、连接器等 1批"],
            ["2.13", "防水", "176,400.00", "176,400.00", "屋面全面防水处理及渗漏排查 1项"],
            ["2.14", "加固", "551,880.00", "551,880.00", "屋面加固处理及承载力验算 1项"]
        ]

        proposals = [
            {
                "path": "/body/tbl[1]",
                "proposed_text": json.dumps(matrix, ensure_ascii=False),
                "value": json.dumps(matrix, ensure_ascii=False),
                "type": "table_rows",
                "chapter_title": "投标配置及分项报价表"
            }
        ]

        count = fill_docx_proposals_in_dom(temp_path, proposals)
        assert count == 16

        res_doc = Document(temp_path)
        res_tbl = res_doc.tables[0]

        # 核心断言 1：总行数严格为 18 行（1 表头 + 16 细项数据行 + 1 表尾合并合计行）
        assert len(res_tbl.rows) == 18

        # 核心断言 2：前 5 行模板行（包含原来的 ......）已被 100% 干净覆盖，没有任何 ...... 残留
        row_texts = ["".join([c.text.strip() for c in r.cells]) for r in res_tbl.rows]
        assert not any("......" in t for t in row_texts)

        # 核心断言 3：各行序号与二级编号精准对齐
        assert res_tbl.rows[1].cells[0].text.strip() == "1"
        assert res_tbl.rows[2].cells[0].text.strip() == "2"
        assert res_tbl.rows[3].cells[0].text.strip() == "2.1"
        assert "光伏组件" in res_tbl.rows[3].cells[1].text
        assert res_tbl.rows[4].cells[0].text.strip() == "2.2"
        assert "逆变器 1" in res_tbl.rows[4].cells[1].text
        assert res_tbl.rows[16].cells[0].text.strip() == "2.14"
        assert "加固" in res_tbl.rows[16].cells[1].text

        # 核心断言 4：表尾合计行仍完好保留在最后一行
        assert "合计总价" in res_tbl.rows[17].cells[0].text

    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_fill_docx_proposals_in_dom_auto_split_and_deduplicate_hierarchical_prefix():
    """测试 4 列数据传入时自动将 '2.1 光伏组件' 拆分为序号 '2.1' 与名称 '光伏组件'，以及 5 列数据重复前缀去重"""
    import json
    from docx import Document
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    doc = Document()
    # 模拟模板：1 表头 + 2 模板行 + 1 合并合计行
    tbl = doc.add_table(rows=4, cols=5)
    headers = ["序号", "项目/费用名称", "单价（元）", "分项总价（元）", "备注"]
    for c_idx, h in enumerate(headers):
        tbl.rows[0].cells[c_idx].text = h

    # 表尾合并合计行
    tbl.rows[3].cells[0].merge(tbl.rows[3].cells[4])
    tbl.rows[3].cells[0].text = "合计总价：2235211.56"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)

        # 4 列数据测试（首项包含层级编号 1、2、2.1、2.2）
        matrix_4cols = [
            ["1 设计费（含电力设计、建筑设计费用）", "0.00", "0.00", ""],
            ["2 建设费（含设备购置费）", "2235211.56", "2235211.56", ""],
            ["2.1 光伏组件", "882.69", "1167798.87", ""],
            ["2.2、逆变器 1", "11555.00", "46220.00", ""],
            ["合计总价（合计总价=设计费+建设费合计）", "", "2235211.56", ""]
        ]

        proposals = [
            {
                "path": "/body/tbl[1]",
                "proposed_text": json.dumps(matrix_4cols, ensure_ascii=False),
                "value": json.dumps(matrix_4cols, ensure_ascii=False),
                "type": "table_rows",
                "chapter_title": "投标配置及分项报价表"
            }
        ]

        count = fill_docx_proposals_in_dom(temp_path, proposals)
        assert count == 5

        res_doc = Document(temp_path)
        res_tbl = res_doc.tables[0]

        # 核心断言 1：第 1 列为纯层级序号，第 2 列为纯名称（无序号混入）
        assert res_tbl.rows[1].cells[0].text.strip() == "1"
        assert res_tbl.rows[1].cells[1].text.strip() == "设计费（含电力设计、建筑设计费用）"

        assert res_tbl.rows[2].cells[0].text.strip() == "2"
        assert res_tbl.rows[2].cells[1].text.strip() == "建设费（含设备购置费）"

        assert res_tbl.rows[3].cells[0].text.strip() == "2.1"
        assert res_tbl.rows[3].cells[1].text.strip() == "光伏组件"

        assert res_tbl.rows[4].cells[0].text.strip() == "2.2"
        # 核心断言 2：合计行包含合计总价或合计金额
        assert "合计总价" in res_tbl.rows[5].cells[0].text or "2235211.56" in res_tbl.rows[5].cells[0].text

    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_fill_docx_proposals_in_dom_full_paragraph_should_not_duplicate_text():
    """
    测试针对包含多个空白/下划线占位符的段落（如《投标函》主句），
    当 Worker 提交完整覆盖句子时，写盘引擎能够智能识别整段替换并生成精确的 Diff Run，
    彻底消除文本重复与未填占位符残留，并为填入的数据自动添加下划线。
    """
    from docx import Document
    from docx.oxml.ns import qn
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    doc = Document()
    orig_para_text = "根据贵方的                   号招标文件，正式授权下述签字人       代表我方                           （投标人的名称），全权处理本次项目投标的有关事宜。"
    p = doc.add_paragraph(orig_para_text)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)

        proposed_full_text = "根据贵方的RXCG2026-G003号招标文件，正式授权下述签字人李四（授权代表）代表我方四川石楠建设工程有限公司，全权处理本次项目投标的有关事宜。"

        proposals = [
            {
                "path": "/body/p[1]",
                "original_context": orig_para_text,
                "proposed_text": proposed_full_text,
                "value": proposed_full_text,
                "type": "text",  # 模拟大模型在 JSON 中写了普通 text 类型
                "chapter_title": "投标函格式"
            }
        ]

        count = fill_docx_proposals_in_dom(temp_path, proposals)
        assert count >= 1

        res_doc = Document(temp_path)
        res_p = res_doc.paragraphs[0]
        res_text = res_p.text.strip()

        # 核心断言 1：绝不能出现重复前缀或重复后缀
        assert res_text.count("根据贵方") == 1
        assert res_text.count("全权处理本次项目投标的有关事宜") == 1
        assert res_text.count("正式授权下述签字人") == 1
        assert res_text.count("代表我方") == 1

        # 核心断言 2：结果文本完整且无残留空格/下划线占位符
        assert res_text == proposed_full_text
        assert "                   " not in res_text
        assert "（投标人的名称）" not in res_text

        # 核心断言 3：验证填入的数据携带下划线，模板固定文字不带下划线
        underlined_texts = []
        normal_texts = []
        for r in res_p.runs:
            rPr = r._element.find(qn('w:rPr'))
            has_u = (rPr is not None and rPr.find(qn('w:u')) is not None)
            if has_u:
                underlined_texts.append(r.text)
            else:
                normal_texts.append(r.text)

        # 验证填入的核心业务数据有下划线
        full_underlined = "".join(underlined_texts)
        assert "RXCG2026-G003" in full_underlined
        assert "李四" in full_underlined
        assert "四川石楠建设工程有限公司" in full_underlined

        # 验证固定模板引导词无下划线
        full_normal = "".join(normal_texts)
        assert "根据贵方的" in full_normal
        assert "号招标文件" in full_normal
        assert "全权处理本次项目投标的有关事宜" in full_normal

    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def test_is_full_paragraph_replacement_detection():
    """测试 _is_full_paragraph_replacement 智能判定器的多维准确性"""
    from app.agents.bid_filler_agent import _is_full_paragraph_replacement

    # 1. 显式类型
    assert _is_full_paragraph_replacement("某段落原文", "某段落修改后文本", prop_type="sentence_batch") is True
    assert _is_full_paragraph_replacement("某段落原文", "某段落修改后文本", prop_type="paragraph") is True

    # 2. 投标函主句：首尾双锚点匹配
    orig_bid = "根据贵方的                   号招标文件，正式授权下述签字人       代表我方                           （投标人的名称），全权处理本次项目投标的有关事宜。"
    prop_bid = "根据贵方的RXCG2026-G003号招标文件，正式授权下述签字人李四（授权代表）代表我方四川石楠建设工程有限公司，全权处理本次项目投标的有关事宜。"
    assert _is_full_paragraph_replacement(orig_bid, prop_bid, prop_type="text") is True

    # 3. 授权委托书：首尾双锚点与骨架重合
    orig_auth = "兹委托受托人______合法地代表我单位参加______组织的______项目的采购活动，受托人有权在该投标活动中，以我单位的名义签署投标书和投标文件，与采购人协商、澄清、解释，签订合同书并执行一切与此有关的事项。"
    prop_auth = "兹委托受托人李四合法地代表我单位参加张家港润信项目咨询有限公司组织的窑厂工业区840KW分布式光伏发电项目的采购活动，受托人有权在该投标活动中，以我单位的名义签署投标书和投标文件，与采购人协商、澄清、解释，签订合同书并执行一切与此有关的事项。"
    assert _is_full_paragraph_replacement(orig_auth, prop_auth, prop_type="text") is True

    # 4. 普通单槽位填空值：必须判定为 False
    assert _is_full_paragraph_replacement("地    址：________________________", "四川省成都市高新区天府大道北段128号", prop_type="text") is False
    assert _is_full_paragraph_replacement("投标单位代表姓名（签字）：______", "李四", prop_type="text") is False


def test_parse_proposals_fusion_should_preserve_specific_type_from_markdown():
    """测试 _parse_proposals 在 JSON 中 type 为 text 时，能从 Markdown 表格中继承 sentence_batch 类型"""
    from app.agents.bid_filler_workers import _parse_proposals

    raw_text = """
### 操作总结
| 序号 | DOM 节点路径 | 替换前模板原文 | 实际填入/扩写结果 | 提议类型 | 写盘状态 |
|---|---|---|---|---|---|
| 1 | /body/p[@paraId=32595626] | 根据贵方的______号招标文件... | 根据贵方的RXCG2026-G003号招标文件，正式授权下述签字人李四代表我方四川石楠建设工程有限公司，全权处理本次项目投标的有关事宜。 | sentence_batch | 已提交刷盘队列 |

```json
[
  {
    "path": "/body/p[@paraId=32595626]",
    "proposed_text": "根据贵方的RXCG2026-G003号招标文件，正式授权下述签字人李四代表我方四川石楠建设工程有限公司，全权处理本次项目投标的有关事宜。",
    "value": "根据贵方的RXCG2026-G003号招标文件，正式授权下述签字人李四代表我方四川石楠建设工程有限公司，全权处理本次项目投标的有关事宜。",
    "type": "text",
    "status": "success"
  }
]
```
"""
    proposals = _parse_proposals(raw_text)
    assert len(proposals) == 1
    assert proposals[0]["path"] == "/body/p[@paraId=32595626]"
    # 验证提议类型已成功提升为更精确的 sentence_batch
    assert proposals[0]["type"] == "sentence_batch"


def test_fill_docx_proposals_in_dom_hierarchical_pricing_table_exact_alignment_and_footer_reuse():
    """
    测试《投标配置及分项报价表》全流程原位写入：
    1. 验证 4 列数据（省略空备注）能够 100% 精准对齐 5 列表格，杜绝名称重复占用单价列与向右错位；
    2. 验证原模板预置的 1 设计费、2 建设费、空白行及 ...... 占位行 100% 被二级细项原位覆盖；
    3. 验证原模板最后一行的合计总价行被原位复用刷新总金额，绝不在下方残留旧行或产生重复合计行。
    """
    from docx import Document
    from app.agents.bid_filler_agent import fill_docx_proposals_in_dom

    doc = Document()
    tbl = doc.add_table(rows=7, cols=5)

    # Row 0: 表头
    headers = ["序号", "项目/费用名称", "单价（元）", "分项总价（元）", "备注"]
    for c_idx, h in enumerate(headers):
        tbl.rows[0].cells[c_idx].text = h

    # Row 1: 1 设计费（模拟原模板中的单元格合并）
    tbl.rows[1].cells[1].merge(tbl.rows[1].cells[2])
    tbl.rows[1].cells[0].text = "1"
    tbl.rows[1].cells[1].text = "设计费（含电力设计、建筑设计费用）"

    # Row 2: 2 建设费（模拟原模板中的单元格合并）
    tbl.rows[2].cells[1].merge(tbl.rows[2].cells[2])
    tbl.rows[2].cells[0].text = "2"
    tbl.rows[2].cells[1].text = "建设费（含设备购置费）"

    # Row 3 & 4: 空白数据行
    # Row 5: ...... 占位行（整行合并）
    tbl.rows[5].cells[0].merge(tbl.rows[5].cells[4])
    tbl.rows[5].cells[0].text = "......"

    # Row 6: 原模板表尾合计总价行（前两列合并）
    tbl.rows[6].cells[0].merge(tbl.rows[6].cells[1])
    tbl.rows[6].cells[0].text = "合计总价（合计总价=设计费+建设费合计）"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)

        # 模拟大模型输出的 17 行 4 列数据矩阵（包含明细与合计，省略了末尾空的备注列）
        pricing_matrix = [
            ["1", "设计费（含电力设计、建筑设计费用）", "0.00", "0.00"],
            ["2", "建设费（含设备购置费）", "", "2235211.56"],
            ["2.1", "光伏组件", "882.69", "1167798.87"],
            ["2.2", "逆变器 1", "11555.00", "46220.00"],
            ["2.3", "逆变器 2", "11555.00", "46220.00"],
            ["2.4", "并网柜", "9180.35", "9180.35"],
            ["2.5", "直流电缆", "25720.56", "25720.56"],
            ["2.6", "交流电缆（逆变器）", "17859.45", "17859.45"],
            ["2.7", "交流电缆（并网1）", "11882.62", "11882.62"],
            ["2.8", "交流电缆（并网2）", "33600.00", "33600.00"],
            ["2.9", "电缆桥架", "9233.06", "9233.06"],
            ["2.10", "光伏支架", "117600.00", "117600.00"],
            ["2.11", "接地系统", "6580.98", "6580.98"],
            ["2.12", "其他辅材", "15035.67", "15035.67"],
            ["2.13", "防水", "176400.00", "176400.00"],
            ["2.14", "加固", "551880.00", "551880.00"],
            ["合计总价（合计总价=设计费+建设费合计）", "", "2235211.56", ""]
        ]

        proposals = [
            {
                "path": "/body/tbl[1]",
                "proposed_text": json.dumps(pricing_matrix, ensure_ascii=False),
                "value": json.dumps(pricing_matrix, ensure_ascii=False),
                "type": "table_rows",
                "chapter_title": "投标配置及分项报价表"
            }
        ]

        count = fill_docx_proposals_in_dom(temp_path, proposals)
        assert count >= 16

        res_doc = Document(temp_path)
        res_tbl = res_doc.tables[0]

        # 核心断言 1：总行数严格为 18 行（1 行表头 + 16 行明细 + 1 行合计），绝无任何多余旧行残留
        assert len(res_tbl.rows) == 18

        # 核心断言 2：Row 1 为大类设计费，单价为 0.00，分项总价为 0.00，备注为空
        assert res_tbl.rows[1].cells[0].text.strip() == "1"
        assert res_tbl.rows[1].cells[1].text.strip() == "设计费（含电力设计、建筑设计费用）"
        assert res_tbl.rows[1].cells[2].text.strip() == "0.00"
        assert res_tbl.rows[1].cells[3].text.strip() == "0.00"
        assert res_tbl.rows[1].cells[4].text.strip() == ""

        # 核心断言 3：Row 2 为大类建设费
        assert res_tbl.rows[2].cells[0].text.strip() == "2"
        assert res_tbl.rows[2].cells[1].text.strip() == "建设费（含设备购置费）"
        assert res_tbl.rows[2].cells[3].text.strip() == "2235211.56"

        # 核心断言 4：Row 3 为 2.1 光伏组件，列数据 100% 精准对齐（绝不发生名称写进单价列！）
        assert res_tbl.rows[3].cells[0].text.strip() == "2.1"
        assert res_tbl.rows[3].cells[1].text.strip() == "光伏组件"
        assert res_tbl.rows[3].cells[2].text.strip() == "882.69"  # 单价必须在第 2 列
        assert res_tbl.rows[3].cells[3].text.strip() == "1167798.87"  # 分项总价必须在第 3 列
        assert res_tbl.rows[3].cells[4].text.strip() == ""  # 备注必须在第 4 列

        # 核心断言 5：Row 4 为 2.2 逆变器 1
        assert res_tbl.rows[4].cells[0].text.strip() == "2.2"
        assert res_tbl.rows[4].cells[1].text.strip() == "逆变器 1"
        assert res_tbl.rows[4].cells[2].text.strip() == "11555.00"
        assert res_tbl.rows[4].cells[3].text.strip() == "46220.00"

        # 核心断言 6：最后一行合计总价行已原位刷新，且表格仅有唯一个合计总价行
        last_row = res_tbl.rows[-1]
        assert "合计总价" in last_row.cells[0].text
        assert "2235211.56" in "".join(c.text for c in last_row.cells)

    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)
