"""
章节-表格所有权拓扑映射与跨章节表格冲突守卫测试 (test_chapter_table_ownership_and_collision_guard.py)
"""

import os
import json
import tempfile
import pytest
from docx import Document
from app.utils.table_utils import get_chapter_specific_table_indices, get_doc_chapter_tables_mapping
from app.agents.bid_filler_agent import fill_docx_proposals_in_dom


def test_get_chapter_specific_table_indices_fuzzy_and_noisy_titles():
    """验证即使章节标题含有 MinerU 提取残留的重复字符噪声，依然能精准绑定专属表格。"""
    doc = Document()
    
    # 章节 1
    doc.add_paragraph("一、投标函")
    doc.add_paragraph("致：招标人...")
    
    # 章节 2：投标配置及分项报价表
    doc.add_paragraph("五、投标配置及分项报价表五五、、投标配置及分项报价表")
    t1 = doc.add_table(rows=3, cols=5)
    t1.rows[0].cells[0].text = "序号"
    t1.rows[0].cells[1].text = "项目/费用名称"
    t1.rows[0].cells[2].text = "单价（元）"
    t1.rows[0].cells[3].text = "分项总价（元）"
    t1.rows[0].cells[4].text = "备注"

    # 章节 3：常用零件及耗材报价表
    doc.add_paragraph("七、常用零件及耗材报价表七七、常用零件及耗材报价表、常用零件及耗材报价表")
    t2 = doc.add_table(rows=3, cols=6)
    t2.rows[0].cells[0].text = "序号"
    t2.rows[0].cells[1].text = "产品名称"
    t2.rows[0].cells[2].text = "规格型号"
    t2.rows[0].cells[3].text = "计量单位"
    t2.rows[0].cells[4].text = "数量"
    t2.rows[0].cells[5].text = "单价（元）"

    # 1. 验证投标配置及分项报价表精准命中表格 0 (即 /body/tbl[1])
    indices_pricing = get_chapter_specific_table_indices(doc, "投标配置及分项报价表")
    assert indices_pricing == [0]

    # 2. 验证常用零件及耗材报价表精准命中表格 1 (即 /body/tbl[2])
    indices_parts = get_chapter_specific_table_indices(doc, "常用零件及耗材报价表")
    assert indices_parts == [1]


def test_fill_docx_proposals_in_dom_table_collision_guard_should_redirect_and_protect():
    """
    核心冲突测试：
    当 Worker [8]《常用零件及耗材报价表》误将 Path 也填报为 `/body/tbl[1]`（与 Worker [7] 冲突）时，
    冲突守卫必须：
    1. 保护 `/body/tbl[1]` 保持 Worker [7] 的完整报价数据不被覆盖；
    2. 自动将 Worker [8] 的提案智能重定向到其真实的 `/body/tbl[2]`！
    """
    doc = Document()

    # 章节 1：分项报价表
    doc.add_paragraph("五、投标配置及分项报价表")
    t1 = doc.add_table(rows=4, cols=5)
    t1.rows[0].cells[0].text = "序号"
    t1.rows[0].cells[1].text = "项目/费用名称"
    t1.rows[0].cells[2].text = "单价（元）"
    t1.rows[0].cells[3].text = "分项总价（元）"
    t1.rows[0].cells[4].text = "备注"

    # 章节 2：常用零件及耗材报价表
    doc.add_paragraph("七、常用零件及耗材报价表")
    t2 = doc.add_table(rows=4, cols=6)
    t2.rows[0].cells[0].text = "序号"
    t2.rows[0].cells[1].text = "产品名称"
    t2.rows[0].cells[2].text = "规格型号"
    t2.rows[0].cells[3].text = "计量单位"
    t2.rows[0].cells[4].text = "数量"
    t2.rows[0].cells[5].text = "单价（元）"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        temp_path = tf.name

    try:
        doc.save(temp_path)

        # Worker [7] 正确生成的 17 行分项报价数据
        pricing_matrix = [
            ["1", "设计费", "0.00", "0.00", ""],
            ["2", "建设费", "——", "2235211.56", ""],
            ["2.1", "光伏组件", "882.69", "1167798.87", ""],
            ["2.2", "逆变器 1", "11555.00", "46220.00", ""],
            ["合计总价（合计总价=设计费+建设费合计）", "", "", "2235211.56", ""]
        ]

        # Worker [8] 误报了同样的 /body/tbl[1] 路径的 6 行耗材数据
        parts_matrix = [
            ["直流光伏电缆", "PV1-F 4mm²", "批", "1", "25720.56"],
            ["交流电缆（逆变器）", "ZRC-YJLHV", "批", "1", "17859.45"],
            ["电缆桥架", "钢制梯式", "批", "1", "9233.06"]
        ]

        proposals = [
            {
                "path": "/body/tbl[1]",
                "chapter_title": "投标配置及分项报价表",
                "proposed_text": json.dumps(pricing_matrix, ensure_ascii=False),
                "type": "table_rows"
            },
            {
                # 故意设置冲突的错误路径！
                "path": "/body/tbl[1]",
                "chapter_title": "常用零件及耗材报价表",
                "proposed_text": json.dumps(parts_matrix, ensure_ascii=False),
                "type": "table_rows"
            }
        ]

        # 执行写盘
        count = fill_docx_proposals_in_dom(temp_path, proposals)
        assert count > 0

        # 重新读取文档进行断言校验
        res_doc = Document(temp_path)
        res_t1 = res_doc.tables[0] # /body/tbl[1]
        res_t2 = res_doc.tables[1] # /body/tbl[2]

        # 核心断言 1：表格 1 (分项报价表) 必须完好保留 Worker [7] 的数据，绝无被耗材数据篡改覆盖！
        assert res_t1.rows[1].cells[1].text.strip() == "设计费"
        assert res_t1.rows[2].cells[1].text.strip() == "建设费"
        assert res_t1.rows[3].cells[1].text.strip() == "光伏组件"
        assert res_t1.rows[4].cells[1].text.strip() == "逆变器 1"

        # 核心断言 2：表格 2 (常用零件表) 必须成功接收被智能重定向过来的 Worker [8] 耗材数据！
        assert res_t2.rows[1].cells[1].text.strip() == "直流光伏电缆"
        assert res_t2.rows[1].cells[2].text.strip() == "PV1-F 4mm²"
        assert res_t2.rows[2].cells[1].text.strip() == "交流电缆（逆变器）"
        assert res_t2.rows[3].cells[1].text.strip() == "电缆桥架"

    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)
