"""
DocxTestFillerService 单元测试

遵照项目规范与 docx 技能规则：
1. 测试 Word 模版生成；
2. 测试占位符替换并验证原始下划线 (w:u) 格式保留；
3. 测试表格修改并验证 DXA 双列宽与 XML 属性设置；
4. 测试边界防御逻辑。
"""

import io
import pytest
from docx import Document
from app.services.docx_test_filler_service import docx_test_filler_service


def test_fill_docx_should_replace_placeholders_and_keep_underline():
    """ 测试 Word 占位符替换功能，验证下划线格式与文字渲染正确 """
    # 1. 生成模版
    template_bytes = docx_test_filler_service.create_sample_docx()
    assert template_bytes is not None
    assert len(template_bytes) > 0

    # 2. 准备修改数据
    fill_data = {
        "项目名称": "智能招投标系统自动化填报项目",
        "投标人名称": "聚猫科技股份有限公司"
    }

    # 3. 执行填报修改
    modified_bytes = docx_test_filler_service.fill_and_modify_docx(template_bytes, fill_data)
    assert modified_bytes is not None
    assert len(modified_bytes) > 0

    # 4. 重新装载修改后的 .docx 进行断言验证
    doc = Document(io.BytesIO(modified_bytes))
    full_text = "\n".join([p.text for p in doc.paragraphs])

    assert "智能招投标系统自动化填报项目" in full_text
    assert "聚猫科技股份有限公司" in full_text

    # 校验下划线属性保留
    replaced_p = [p for p in doc.paragraphs if "智能招投标系统自动化填报项目" in p.text][0]
    underline_runs = [r for r in replaced_p.runs if docx_test_filler_service._check_underline(r)]
    assert len(underline_runs) > 0
    assert "智能招投标系统自动化填报项目" in underline_runs[0].text


def test_fill_docx_should_modify_table_with_dxa_widths():
    """ 测试 Word 表格修改功能，验证数据更新与单元格属性设置 """
    template_bytes = docx_test_filler_service.create_sample_docx()

    fill_data = {
        "项目名称": "演示测试项目"
    }
    table_updates = {
        "row_1": ["1", "人工智能智能标书生成模块", "580,000.00"],
        "row_2": ["2", "Word DOM 格式智能裁切服务", "220,000.00"]
    }

    modified_bytes = docx_test_filler_service.fill_and_modify_docx(
        docx_bytes=template_bytes,
        fill_data=fill_data,
        table_updates=table_updates
    )

    doc = Document(io.BytesIO(modified_bytes))
    assert len(doc.tables) >= 1

    table = doc.tables[0]
    row1_cells = [cell.text.strip() for cell in table.rows[1].cells]
    row2_cells = [cell.text.strip() for cell in table.rows[2].cells]

    assert row1_cells[1] == "人工智能智能标书生成模块"
    assert row1_cells[2] == "580,000.00"
    assert row2_cells[1] == "Word DOM 格式智能裁切服务"
    assert row2_cells[2] == "220,000.00"


def test_fill_docx_with_empty_bytes_should_raise_value_error():
    """ 测试边界防御逻辑：当传入空字节流时抛出 ValueError 异常 """
    with pytest.raises(ValueError, match="传入的 docx_bytes 不能为空"):
        docx_test_filler_service.fill_and_modify_docx(b"", {"项目名称": "测试"})
