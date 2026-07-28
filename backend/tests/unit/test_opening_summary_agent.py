"""
单元测试：OpeningSummaryAgent 开标一览表专项 Agent (test_opening_summary_agent.py)
"""

import os
import tempfile
import pytest
from unittest.mock import MagicMock, patch
from docx import Document as DocxDocument

from app.utils.rmb_formatter import number_to_chinese_rmb
from app.agents.nodes.opening_summary_agent import (
    build_standard_opening_summary_docx,
    extract_and_fill_opening_summary_docx,
    generate_opening_summary_node
)


def test_build_standard_opening_summary_docx():
    """测试标准 GB/T 国标开标一览表 Word 文档在线合成与落盘"""
    with tempfile.TemporaryDirectory() as tmp_dir:
        output_docx = os.path.join(tmp_dir, "test_opening_summary.docx")
        summary_data = {
            "project_name": "测试智慧政务系统项目",
            "project_code": "TEST-2026-001",
            "bidder_name": "测试工程软件有限公司",
            "total_cost": 1234567.89,
            "construction_period": "60日历天",
            "validity_period": "90日历天",
            "quality_standard": "合格",
            "pm_name": "张经理"
        }

        res_path = build_standard_opening_summary_docx(output_docx, summary_data)
        assert os.path.exists(res_path)

        # 验证生成的 docx 文件可被正常解析并包含关键数据
        doc = DocxDocument(res_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        for t in doc.tables:
            for row in t.rows:
                full_text += "\n" + " ".join([c.text for c in row.cells])

        assert "开 标 一 览 表" in full_text
        assert "测试智慧政务系统项目" in full_text
        assert number_to_chinese_rmb(1234567.89) in full_text
        assert "¥1,234,567.89" in full_text
        assert "测试工程软件有限公司" in full_text


def test_extract_and_fill_opening_summary_docx_non_existent_file():
    """测试当原始 Word 文件不存在时优雅返回 False"""
    res = extract_and_fill_opening_summary_docx(
        original_docx_path="/non/existent/path.docx",
        output_docx_path="/tmp/output.docx",
        summary_data={}
    )
    assert res is False


def test_extract_and_fill_opening_summary_docx_with_mock_doc():
    """测试从模拟含开标一览表的原始 Word 文件中定位、提取并原位修改数据"""
    with tempfile.TemporaryDirectory() as tmp_dir:
        input_docx = os.path.join(tmp_dir, "mock_tender.docx")
        output_docx = os.path.join(tmp_dir, "output_opening_summary.docx")

        # 创建一个简单模拟的原招标文件 docx
        doc = DocxDocument()
        doc.add_paragraph("第一章 招标公告")
        doc.add_paragraph("第二章 投标人须知")
        doc.add_paragraph("三、开标一览表")
        doc.add_paragraph("项目名称：原项目名称占位")
        doc.add_paragraph("招标编号：原编号占位")
        
        table = doc.add_table(rows=2, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "项目名称"
        hdr_cells[1].text = "技术与工期要求"
        hdr_cells[2].text = "投标总价 (小写)"

        data_cells = table.rows[1].cells
        data_cells[0].text = "待填项目"
        data_cells[1].text = "按要求"
        data_cells[2].text = "¥0.00"

        doc.add_paragraph("投标人名称：原公司占位")
        doc.add_paragraph("日期：2026年1月1日")
        doc.add_paragraph("四、法定代表人授权书")

        doc.save(input_docx)

        summary_data = {
            "project_name": "全自动数字化标书系统项目",
            "project_code": "BID-2026-888",
            "bidder_name": "人工智能科技股份有限公司",
            "total_cost": 888888.00,
            "construction_period": "30日历天",
            "validity_period": "90日历天"
        }

        success = extract_and_fill_opening_summary_docx(
            original_docx_path=input_docx,
            output_docx_path=output_docx,
            summary_data=summary_data
        )

        assert success is True
        assert os.path.exists(output_docx)

        # 检查输出 docx 文件内容
        res_doc = DocxDocument(output_docx)
        res_text = "\n".join([p.text for p in res_doc.paragraphs])
        for t in res_doc.tables:
            for row in t.rows:
                res_text += "\n" + " ".join([c.text for c in row.cells])

        assert "全自动数字化标书系统项目" in res_text
        assert "BID-2026-888" in res_text
        assert "人工智能科技股份有限公司" in res_text
        assert "888,888.00" in res_text
        # 确认下一个无关的大章“四、法定代表人授权书”被截断剔除
        assert "四、法定代表人授权书" not in res_text


def test_generate_opening_summary_node_db_flow():
    """测试 Node 入口节点读取 DB 并完整运行流程"""
    mock_doc = MagicMock()
    mock_doc.id = "test-doc-123"
    mock_doc.file_path = ""
    mock_doc.parsed_metadata = {
        "project_name": "测试政务云项目",
        "project_code": "GOV-2026-009",
        "cost_analysis": {"total_cost": 500000.0},
        "company_quals": "云计算科技有限公司"
    }

    mock_db = MagicMock()
    mock_db.query.return_value.filter.return_value.first.return_value = mock_doc

    with patch("app.agents.nodes.opening_summary_agent.SessionLocal", return_value=mock_db), \
         patch("app.agents.nodes.opening_summary_agent.emit_agent_log"):

        state = {
            "document_id": "test-doc-123",
            "user_id": "user-1",
            "tenant_id": "default-tenant"
        }

        res = generate_opening_summary_node(state)
        assert res["status"] == "success"
        assert "opening_summary_path" in res
        assert res["summary_data"]["total_cost"] == 500000.0
        assert res["summary_data"]["project_name"] == "测试政务云项目"
