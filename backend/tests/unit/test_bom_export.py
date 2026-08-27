"""
单元测试：BOM 成本测算 Word 导出服务与大写金额汇总 (test_bom_export.py)
"""

import io
import pytest
from docx import Document

from app.services.bom_export_service import generate_bom_docx
from app.utils.rmb_formatter import number_to_chinese_rmb


def test_bom_export_docx_generation_with_chinese_rmb():
    """测试生成 BOM Word 文档，验证表格结构与表尾大小写总价"""
    test_items = [
        {
            "name": "单晶光伏组件",
            "section_name": "光伏发电设备",
            "spec_requirement": "≥630Wp，组件效率≥20.2%",
            "matched_name": "单晶光伏组件",
            "matched_brand": "天合光能",
            "matched_model": "635Wp",
            "matched_manufacturer": "天合光能股份有限公司",
            "match_quality": "精准匹配",
            "qty": 763,
            "unit": "块",
            "ref_price": 882.69,
            "subtotal": 673492.47,
        },
        {
            "name": "组串式逆变器",
            "section_name": "逆变与配电系统",
            "spec_requirement": "110kW 组串式逆变器",
            "matched_brand": "华为",
            "matched_model": "SUN2000-110KTL",
            "matched_manufacturer": "华为技术有限公司",
            "match_quality": "精准匹配",
            "qty": 5,
            "unit": "台",
            "ref_price": 25000.0,
            "subtotal": 125000.0,
        },
    ]

    total_cost = 798492.47
    expected_chinese_rmb = number_to_chinese_rmb(total_cost)
    assert expected_chinese_rmb == "柒拾玖万捌仟肆佰玖拾贰元肆角柒分"

    doc_io = generate_bom_docx(
        document_title="某新能源光伏电站标书",
        items=test_items,
        total_cost=total_cost,
        budget_limit="¥1,000,000.00",
        status_text="在最高投标限价内可控（使用率 79.8%）",
        analysis_summary="本批次2项设备均与参考库精准匹配，建议按指导价起草报价清单。"
    )

    assert isinstance(doc_io, io.BytesIO)
    doc_bytes = doc_io.getvalue()
    assert len(doc_bytes) > 0

    # 重新读取并验证 docx 内部内容
    doc = Document(io.BytesIO(doc_bytes))
    
    # 验证标题
    assert "拟投入设备及 BOM 成本测算清单" in doc.paragraphs[0].text
    
    # 验证项目信息
    info_text = doc.paragraphs[1].text
    assert "某新能源光伏电站标书" in info_text
    assert "¥1,000,000.00" in info_text
    assert "在最高投标限价内可控" in info_text

    # 验证表格行数与内容
    assert len(doc.tables) >= 1
    table = doc.tables[0]
    # 表头(1) + 数据行(2) + 表尾合计(1) = 4
    assert len(table.rows) == 4

    # 验证第一行数据
    row_1_text = " ".join(c.text for c in table.rows[1].cells)
    assert "单晶光伏组件" in row_1_text
    assert "天合光能" in row_1_text
    assert "673,492.47" in row_1_text

    # 验证表尾合计行包含大写总价与小写数值
    footer_row_text = " ".join(c.text for c in table.rows[3].cells)
    assert "【合计】预估总成本" in footer_row_text
    assert "柒拾玖万捌仟肆佰玖拾贰元肆角柒分" in footer_row_text
    assert "798,492.47" in footer_row_text

    # 验证专家评估意见
    full_doc_text = "\n".join(p.text for p in doc.paragraphs)
    assert "专家评估指导意见" in full_doc_text
    assert "建议按指导价起草报价清单" in full_doc_text


def test_bom_export_xlsx_generation_with_chinese_rmb():
    """测试生成 BOM Excel (.xlsx) 工作簿，验证 9 列格式与表尾大小写合计"""
    from openpyxl import load_workbook
    from app.services.bom_export_service import generate_bom_xlsx

    test_items = [
        {
            "name": "单晶光伏组件",
            "spec_requirement": "≥630Wp 单晶硅组件",
            "matched_brand": "天合光能",
            "matched_model": "635Wp",
            "matched_manufacturer": "天合光能股份有限公司",
            "qty": 763,
            "unit": "块",
            "ref_price": 882.69,
            "subtotal": 673492.47,
        },
        {
            "name": "组串式逆变器",
            "spec_requirement": "110kW 组串式逆变器",
            "matched_brand": "华为",
            "matched_model": "SUN2000-110KTL",
            "matched_manufacturer": "华为技术有限公司",
            "qty": 5,
            "unit": "台",
            "ref_price": 25000.0,
            "subtotal": 125000.0,
        },
    ]

    total_cost = 798492.47
    excel_io = generate_bom_xlsx(
        document_title="某新能源光伏电站标书",
        items=test_items,
        total_cost=total_cost,
        budget_limit="¥1,000,000.00",
        status_text="在最高投标限价内可控",
        analysis_summary="本批次设备均与参考库精准对标。"
    )

    assert isinstance(excel_io, io.BytesIO)
    excel_bytes = excel_io.getvalue()
    assert len(excel_bytes) > 0

    wb = load_workbook(io.BytesIO(excel_bytes))
    ws = wb.active
    assert ws.title == "BOM成本测算清单"

    # 验证标题与信息
    assert "拟投入设备及 BOM 成本测算清单" in ws['A1'].value
    assert "某新能源光伏电站标书" in ws['A2'].value

    # 验证表头（9 列）
    headers = [ws.cell(row=5, column=c).value for c in range(1, 10)]
    assert headers == ["序号", "标的物名称", "品牌、规格、型号", "生产厂家", "单位", "数量", "单价(元)", "总价(元)", "备注"]

    # 验证数据行
    assert ws['B6'].value == "单晶光伏组件"
    assert ws['C6'].value == "天合光能 635Wp"
    assert ws['D6'].value == "天合光能股份有限公司"
    assert ws['G6'].value == 882.69
    assert ws['H6'].value == 673492.47

    # 验证表尾合计行 (Row 8)
    footer_row_val = ws['A8'].value
    assert "【合计】预估总成本" in footer_row_val
    assert "柒拾玖万捌仟肆佰玖拾贰元肆角柒分" in footer_row_val
    assert ws['H8'].value == 798492.47

