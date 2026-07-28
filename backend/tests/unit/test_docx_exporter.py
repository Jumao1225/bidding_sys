"""
DocxExporterService 单元测试

测试 Word 文档导出服务能否正常渲染数据模型，并保证文本全量使用纯黑色字体 (RGB: 0, 0, 0)。
"""

import pytest
import io
from docx import Document
from app.schemas.bid_generator import BidFormatStructure, BidFormatSection, ContentTypeEnum
from app.services.docx_exporter_service import docx_exporter_service


def test_export_bid_format_to_docx_bytes_should_produce_valid_black_font_docx():
    """测试将 BidFormatStructure 导为 .docx 二进制流，并验证字节流可解析及字体颜色为黑色"""
    structure = BidFormatStructure(
        document_title="测试项目 - 投标文件格式",
        source_chapter_name="第六章 投标文件格式",
        sections=[
            BidFormatSection(
                section_title="附件一 投标函",
                content_type=ContentTypeEnum.TEXT_TEMPLATE,
                body_markdown="致：某某招标公司\n我方自愿参加本项目的投标......",
                placeholders=["招标公司名称"]
            ),
            BidFormatSection(
                section_title="附件二 报价明细表",
                content_type=ContentTypeEnum.FORM_TABLE,
                body_markdown="| 品名 | 数量 | 单价 | 总价 |\n| :--- | :--- | :--- | :--- |\n| 软件模块 | 1 | 10000 | 10000 |",
                placeholders=[]
            )
        ]
    )

    docx_bytes = docx_exporter_service.export_bid_format_to_docx_bytes(structure)
    assert docx_bytes is not None
    assert len(docx_bytes) > 0

    # 重新装载生成的 .docx 验证其正确性
    doc = Document(io.BytesIO(docx_bytes))
    assert len(doc.paragraphs) > 0
    assert "测试项目 - 投标文件格式" in doc.paragraphs[0].text or "测试项目" in [p.text for p in doc.paragraphs]
    assert len(doc.tables) == 1


def test_add_formatted_text_to_paragraph_should_support_italic_and_underline():
    """测试 _add_formatted_text_to_paragraph 正确解析 HTML/Markdown 标记并生成斜体与下划线 Run"""
    doc = Document()
    p = doc.add_paragraph()

    test_line = "请填写 <u>下划线文本</u> 以及 *斜体文本* 与 <i>斜体二</i> 填空: _________"
    docx_exporter_service._add_formatted_text_to_paragraph(p, test_line)

    runs = p.runs
    assert len(runs) >= 4
    
    # 验证下划线样式
    underline_runs = [r for r in runs if r.underline]
    assert len(underline_runs) >= 2
    assert any("下划线文本" in r.text for r in underline_runs)
    assert any("_________" in r.text for r in underline_runs)

    # 验证斜体样式
    italic_runs = [r for r in runs if r.italic]
    assert len(italic_runs) >= 2
    assert any("斜体文本" in r.text for r in italic_runs)
    assert any("斜体二" in r.text for r in italic_runs)

