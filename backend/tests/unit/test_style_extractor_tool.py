"""
样式定向提取 Agent 工具单元测试 (Style Extractor Tool Unit Tests)
"""

import os
import pytest
import docx
from unittest.mock import patch, MagicMock
from app.agents.tools.style_extractor_tool import extract_text_by_style, resolve_document_file_path


def test_extract_text_by_style_non_existent_file():
    """测试当文件或 document_id 不存在时，工具返回清晰说明"""
    res = extract_text_by_style.invoke({
        "file_path": "non_existent_file.docx",
        "chapter_keyword": "第四章",
        "style_type": "italic_underline"
    })
    assert "无法定位文档对应的物理磁盘文件" in res or "文件不存在" in res


def test_extract_text_by_style_missing_all_params():
    """测试当无 document_id 且无 file_path 时提示错误"""
    res = extract_text_by_style.invoke({
        "chapter_keyword": "第四章",
        "style_type": "italic_underline"
    })
    assert "必须至少提供 document_id 或 file_path" in res


def test_extract_text_by_style_non_docx_file(tmp_path):
    """测试非 docx 文件调用时提示错误"""
    txt_file = tmp_path / "test.txt"
    txt_file.write_text("hello", encoding="utf-8")

    res = extract_text_by_style.invoke({
        "file_path": str(txt_file),
        "chapter_keyword": "第四章",
        "style_type": "italic_underline"
    })
    assert "仅支持 Word (.docx) 格式" in res


def test_extract_text_by_style_document_id_db_resolution(tmp_path):
    """测试传入 document_id 时从 DB 成功智能解析文件磁盘路径并完成样式提取"""
    docx_file = tmp_path / "sample_doc_id.docx"
    doc = docx.Document()
    doc.add_heading("第四章 商务要求", level=1)
    p = doc.add_paragraph()
    r = p.add_run("付款比例为 90%")
    r.italic = True
    r.underline = True
    doc.save(str(docx_file))

    # Mock DB 查询返回 docx_file
    mock_doc = MagicMock()
    mock_doc.file_path = str(docx_file)

    mock_db = MagicMock()
    mock_db.query.return_value.filter.return_value.first.return_value = mock_doc

    with patch("app.db.session.SessionLocal", return_value=mock_db):
        res = extract_text_by_style.invoke({
            "document_id": "test-doc-uuid-9999",
            "chapter_keyword": "第四章",
            "style_type": "italic_underline"
        })
        assert "成功找到样式类型 [italic_underline] 匹配结果" in res
        assert "付款比例为 90%" in res


def test_extract_text_by_style_integration(tmp_path):
    """集成测试：使用 python-docx 动态构造具有特定样式的 docx 并测试提取"""
    docx_file = tmp_path / "sample_spec.docx"
    doc = docx.Document()

    # 第一章
    doc.add_heading("第一章 总则", level=1)
    p1 = doc.add_paragraph()
    r1 = p1.add_run("本段为普通说明文字。")

    # 第四章
    doc.add_heading("第四章 技术规范与要求", level=1)
    p4 = doc.add_paragraph()
    r4_1 = p4.add_run("系统响应要求：")
    r4_1.bold = True

    r4_2 = p4.add_run("峰值并发下响应延时必须小于50ms")
    r4_2.italic = True
    r4_2.underline = True

    doc.save(str(docx_file))

    # 1. 查找第四章中斜体+下划线文本
    res_iu = extract_text_by_style.invoke({
        "file_path": str(docx_file),
        "chapter_keyword": "第四章",
        "style_type": "italic_underline"
    })
    assert "成功找到样式类型 [italic_underline] 匹配结果" in res_iu
    assert "第四章 技术规范与要求" in res_iu
    assert "峰值并发下响应延时必须小于50ms" in res_iu

    # 2. 查找不存在匹配项的章节
    res_none = extract_text_by_style.invoke({
        "file_path": str(docx_file),
        "chapter_keyword": "第一章",
        "style_type": "italic_underline"
    })
    assert "未找到样式类型为 [italic_underline] 的文本" in res_none
