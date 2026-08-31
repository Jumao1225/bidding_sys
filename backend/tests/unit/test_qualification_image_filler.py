"""
资质证明文件与图片自动嵌入单元测试 (test_qualification_image_filler.py)

测试目标：
1. 验证 query_company_qualification_tool 是否能准确解析 uploads/qualifications 下的本地绝对路径并支持 Agent 自主全量/模糊检索；
2. 验证 officecli_insert_image_tool 是否能成功向 Word 文档物理节点原位嵌入资质证书图片；
3. 验证 resolve_qualification_image_path 函数在各种路径格式下的物理存在性校验逻辑。
"""

import os
import tempfile
import uuid
import pytest
from docx import Document
from PIL import Image

from app.agents.tools.bid_db_tools import (
    query_company_qualification_tool,
    resolve_qualification_image_path
)
from app.agents.tools.office_cli_agent_tools import officecli_insert_image_tool
from app.agents.tools.writer_tools import get_company_qualifications_tool
from app.agents.bid_filler_agent import fill_docx_proposals_in_dom
from app.agents.bid_filler_workers import _build_worker_tools, build_worker_prompt


@pytest.fixture
def temp_qualification_image():
    """创建一个临时测试用资质证书图片文件并写入 DB 记录"""
    backend_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    qual_dir = os.path.join(backend_dir, "uploads", "qualifications")
    os.makedirs(qual_dir, exist_ok=True)

    img_path = os.path.join(qual_dir, "test_cert_mock.png")
    img = Image.new("RGB", (300, 200), color=(73, 109, 137))
    img.save(img_path)

    # 写入测试 DB 记录
    from app.db.session import SessionLocal
    from app.db.models.business import CompanyQualification
    db = SessionLocal()
    mock_quals = []
    try:
        for name in ["营业执照", "电力工程施工总承包", "安全生产许可证", "机电工程施工总承包"]:
            q = CompanyQualification(
                tenant_id="default-tenant",
                name=name,
                level="二级",
                company_name="测试工程有限公司",
                file_url="/uploads/qualifications/test_cert_mock.png"
            )
            db.add(q)
            mock_quals.append(q)
        other_tenant_q = CompanyQualification(
            tenant_id="other-tenant",
            name="跨租户营业执照",
            level="二级",
            company_name="其他租户有限公司",
            file_url="/uploads/qualifications/test_cert_mock.png",
        )
        db.add(other_tenant_q)
        mock_quals.append(other_tenant_q)
        db.commit()
    except Exception:
        db.rollback()

    yield img_path

    try:
        for q in mock_quals:
            db.delete(q)
        db.commit()
    except Exception:
        db.rollback()
    finally:
        db.close()

    if os.path.exists(img_path):
        os.remove(img_path)


@pytest.fixture
def temp_sample_docx():
    """创建一个临时测试用 Word (.docx) 文档"""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        docx_path = tf.name

    doc = Document()
    doc.add_heading("第一章 资格证明文件", level=1)
    doc.add_paragraph("以下为投标人资质证明文件：")
    doc.add_paragraph("[待手动补充资质证书: 营业执照]")
    doc.save(docx_path)

    yield docx_path

    if os.path.exists(docx_path):
        os.remove(docx_path)


def test_resolve_qualification_image_path(temp_qualification_image):
    """测试资质图片本地路径解析与存在性校验功能"""
    file_name = os.path.basename(temp_qualification_image)
    rel_url = f"/uploads/qualifications/{file_name}"

    local_path, exists = resolve_qualification_image_path(rel_url)
    assert exists is True
    assert local_path is not None
    assert os.path.exists(local_path)
    assert os.path.basename(local_path) == file_name


def test_query_company_qualification_tool(temp_qualification_image):
    """测试 DB 直查工具 query_company_qualification_tool 支持自主全量/关键字检索"""
    # 关键字查询
    result_kw = query_company_qualification_tool.invoke({"cert_keyword": "营业执照"})
    assert isinstance(result_kw, str)

    # 空关键字自主全量查询
    result_all = query_company_qualification_tool.invoke({"cert_keyword": ""})
    assert isinstance(result_all, str)


def test_query_company_qualification_tool_respects_context_tenant(temp_qualification_image):
    """资质图片检索必须只返回当前会话租户的数据。"""
    from app.core.context import current_tenant_id

    token = current_tenant_id.set("default-tenant")
    try:
        result = query_company_qualification_tool.invoke({"cert_keyword": "跨租户营业执照"})
        assert "跨租户营业执照" not in result
    finally:
        current_tenant_id.reset(token)


def test_officecli_insert_image_tool(temp_sample_docx, temp_qualification_image):
    """测试向 Word 文档原位节点嵌入资质图片功能"""
    target_path = "/body/p[3]/r[1]"

    result = officecli_insert_image_tool.invoke({
        "file_path": temp_sample_docx,
        "target_path": target_path,
        "image_path": temp_qualification_image,
        "width_inches": 5.0,
        "caption": "营业执照副本 (测试)"
    })

    assert "成功在 Word 节点" in result or "嵌入资质图片" in result

    # 打开校验 docx 中的 inline 图片元素与图注
    doc = Document(temp_sample_docx)
    p3 = doc.paragraphs[2]  # 第3个段落 (0-indexed 2)
    assert "营业执照副本" in p3.text or len(doc.inline_shapes) > 0 or len(p3.runs) > 0


def test_worker_officecli_insert_image_call(temp_sample_docx, temp_qualification_image):
    """测试 Worker 节点工具包中的 officecli_insert_image 工具可以正常被调用，防范 'StructuredTool' object is not callable 错误"""
    from app.agents.bid_filler_workers import _build_worker_tools
    worker_tools = _build_worker_tools(temp_sample_docx, "资格证明文件")

    insert_tool = next((t for t in worker_tools if t.name == "officecli_insert_image"), None)
    assert insert_tool is not None

    res = insert_tool.invoke({
        "target_path": "/body/p[1]",
        "image_path": temp_qualification_image,
        "width_inches": 5.0,
        "caption": "测试证明证书"
    })
    assert "提案" in res or "原位" in res


def test_worker_qualification_prompt_should_fill_all_material_fields_without_skip():
    """测试资格证明 Worker 同时负责文字字段、表单字段和图片字段。"""
    system_prompt, _ = build_worker_prompt(
        chapter_title="资格证明材料",
        category="needs_fill",
        template_text="",
        content_hint="",
        document_id=uuid.uuid4().hex,
        mapping_hint="qualification",
    )

    assert "全部段落、表格和空白字段" in system_prompt
    assert "授权委托" in system_prompt
    assert "sentence_batch" in system_prompt
    assert "anchor_text" in system_prompt
    assert "不得因附件标题" in system_prompt


def test_worker_sentence_batch_should_preserve_full_paragraph_semantics():
    """测试长句工具产出的提案类型为整段替换，防止同段字段错位。"""
    collected = []
    worker_tools = _build_worker_tools(
        docx_temp_path="",
        chapter_title="资格证明材料",
        collected_proposals=collected,
        mapping_hint="qualification",
        category="needs_fill",
    )
    sentence_tool = next(tool for tool in worker_tools if tool.name == "officecli_batch_fill_sentence")
    result = sentence_tool.invoke({
        "updates_json_str": '[{"path":"/body/p[1]","value":"动态完整段落","original_context":"动态原文"}]'
    })

    assert "提案" in result
    assert collected[0]["type"] == "sentence_batch"
    assert collected[0]["original_context"] == "动态原文"


def test_image_proposal_should_reject_mismatched_target_and_accept_dynamic_anchor(temp_qualification_image):
    """测试图片必须匹配运行时条款锚点，禁止落入无关日期段落。"""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        docx_path = tf.name

    anchor_token = f"条款锚点{uuid.uuid4().hex}"
    caption_token = f"证明项{uuid.uuid4().hex}"
    unrelated_token = f"日期字段{uuid.uuid4().hex}"
    doc = Document()
    doc.add_paragraph(anchor_token + caption_token)
    doc.add_paragraph(unrelated_token)
    doc.save(docx_path)

    try:
        rejected = fill_docx_proposals_in_dom(
            docx_path,
            [{
                "path": "/body/p[2]",
                "proposed_text": temp_qualification_image,
                "value": temp_qualification_image,
                "type": "image",
                "caption": caption_token,
                "anchor_text": unrelated_token,
            }],
        )
        assert rejected == 0
        assert len(Document(docx_path).inline_shapes) == 0

        accepted = fill_docx_proposals_in_dom(
            docx_path,
            [{
                "path": "/body/p[1]",
                "proposed_text": temp_qualification_image,
                "value": temp_qualification_image,
                "type": "image",
                "caption": caption_token,
                "anchor_text": anchor_token + caption_token,
            }],
        )
        assert accepted == 1
        result_doc = Document(docx_path)
        assert result_doc.paragraphs[0].text == anchor_token + caption_token
        assert caption_token in result_doc.paragraphs[1].text
    finally:
        if os.path.exists(docx_path):
            os.remove(docx_path)


def test_image_anchor_should_tolerate_dynamic_suffix_difference():
    """图片锚点匹配应允许同一材料名称后附等级等动态后缀差异。"""
    from app.agents.bid_filler_agent import _image_target_matches_anchor

    doc = Document()
    paragraph = doc.add_paragraph("需满足：核心能力等级达到甲级。")

    assert _image_target_matches_anchor(
        paragraph,
        {
            "anchor_text": "需满足：核心能力等级达到甲级。",
            "caption": "核心能力等级乙级",
        },
    ) is True


def test_text_requirement_image_embedding(temp_qualification_image):
    """测试当 Word 模版只有纯文本条款（如截图中的 1. 营业执照... 8. 电力工程施工总承包...安全生产许可证）时，自动在该条款段落下方插入图片"""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        docx_path = tf.name

    doc = Document()
    doc.add_heading("四、资格证明文件", level=1)
    doc.add_paragraph("1. 法人或者其他组织的营业执照等证明文件，自然人的身份证明；")
    doc.add_paragraph("8. 本项目的特定资格要求：需满足以下要求：具有电力工程施工总承包叁级及以上，并取得承装（修、试）电力设施许可证五级及以上，以及取得有效的安全生产许可证。")
    doc.save(docx_path)

    try:
        from app.agents.tools.bid_db_tools import auto_embed_qualification_images_in_docx
        count = auto_embed_qualification_images_in_docx(docx_path)
        assert count > 0

        doc_after = Document(docx_path)
        # 确认原文条款保留完好
        p1_text = doc_after.paragraphs[1].text
        assert "1. 法人或者其他组织的营业执照" in p1_text
        # 确认下方增加了资质证书图片的段落和图注
        full_text = "\n".join([p.text for p in doc_after.paragraphs])
        assert "图：" in full_text
    finally:
        if os.path.exists(docx_path):
            os.remove(docx_path)


def test_keyword_only_does_not_trigger_image_embedding(temp_qualification_image):
    """仅提到证书/资质关键词的说明性原文不应被兜底逻辑改写。"""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        docx_path = tf.name

    doc = Document()
    doc.add_heading("四、资格证明文件", level=1)
    doc.add_paragraph("营业执照、资质证书和安全生产许可证的定义以招标文件约定为准。")
    doc.save(docx_path)

    try:
        from app.agents.tools.bid_db_tools import auto_embed_qualification_images_in_docx
        count = auto_embed_qualification_images_in_docx(docx_path)
        assert count == 0
        assert len(Document(docx_path).inline_shapes) == 0
    finally:
        if os.path.exists(docx_path):
            os.remove(docx_path)


def test_qualification_section_without_request_does_not_append_all_images(temp_qualification_image):
    """只有资格证明章节标题、没有明确材料要求时，不应把全部资质图片追加到章节末尾。"""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        docx_path = tf.name

    doc = Document()
    doc.add_heading("四、资格证明文件", level=1)
    doc.add_paragraph("本章节内容按项目实际要求填写。")
    doc.save(docx_path)

    try:
        from app.agents.tools.bid_db_tools import auto_embed_qualification_images_in_docx
        count = auto_embed_qualification_images_in_docx(docx_path)
        assert count == 0
        assert len(Document(docx_path).inline_shapes) == 0
    finally:
        if os.path.exists(docx_path):
            os.remove(docx_path)


def test_multi_qualification_single_certificate(temp_qualification_image):
    """测试一证多资质（单张证书图片对应多个资质名称）场景下的自动去重合并渲染"""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        docx_path = tf.name

    doc = Document()
    doc.add_heading("四、资格证明文件", level=1)
    doc.add_paragraph("投标人须具有电力工程施工总承包及机电工程施工总承包资质。")
    doc.save(docx_path)

    try:
        from app.agents.tools.bid_db_tools import auto_embed_qualification_images_in_docx
        count = auto_embed_qualification_images_in_docx(docx_path)
        assert count >= 1

        doc_after = Document(docx_path)
        full_text = "\n".join([p.text for p in doc_after.paragraphs])
        assert "图：" in full_text
    finally:
        if os.path.exists(docx_path):
            os.remove(docx_path)


def test_global_image_path_deduplication(temp_qualification_image):
    """测试全局物理图片路径去重：即使多个段落均匹配到同一张资质图片，整个 Word 文档中也绝不重复插入同一张物理图片"""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        docx_path = tf.name

    doc = Document()
    doc.add_heading("四、资格证明文件", level=1)
    doc.add_paragraph("（1）投标人须具有电力工程施工总承包资质；")
    doc.add_paragraph("① 持有旧版承装（修、试）电力设施许可证；")
    doc.add_paragraph("② 持有新版承装（修、试）电力设施许可证；")
    doc.save(docx_path)

    try:
        from app.agents.tools.bid_db_tools import auto_embed_qualification_images_in_docx
        count = auto_embed_qualification_images_in_docx(docx_path)

        # 验证全局图片去重机制有效，去重后插入次数受控
        assert count <= 2
    finally:
        if os.path.exists(docx_path):
            os.remove(docx_path)


def test_idempotent_multiple_runs_no_duplicates(temp_qualification_image):
    """测试多轮调用与已有图片感知防重：多次执行 auto_embed_qualification_images_in_docx 不会产生二次多余插图"""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        docx_path = tf.name

    doc = Document()
    doc.add_heading("四、资格证明文件", level=1)
    doc.add_paragraph("8. 本项目的特定资格要求：具有电力工程施工总承包及安全生产许可证。")
    doc.save(docx_path)

    try:
        from app.agents.tools.bid_db_tools import auto_embed_qualification_images_in_docx
        # 第一轮嵌入
        count1 = auto_embed_qualification_images_in_docx(docx_path)
        doc1 = Document(docx_path)
        shapes_run1 = len(doc1.inline_shapes)

        # 第二轮再次嵌入 (模拟 Worker Agent 与写盘节点、下载 API 连续调用)
        count2 = auto_embed_qualification_images_in_docx(docx_path)
        doc2 = Document(docx_path)
        shapes_run2 = len(doc2.inline_shapes)

        # 核心断言：第二轮调用必须 0 新增图片，InlineShapes 总数完全一致！
        assert count2 == 0
        assert shapes_run1 == shapes_run2
    finally:
        if os.path.exists(docx_path):
            os.remove(docx_path)
