"""
Docx 样式转义解析器单元测试 (Docx Style Transpilation Unit Tests)
验证 DocxParser 是否能够精准识别并转义加粗、斜体、下划线、斜体+下划线组合以及字体颜色。
"""

import os
import pytest
import docx
from unittest.mock import MagicMock
from app.services.parsers.docx_parser import DocxParser, docx_parser


def test_render_styled_paragraph_bold_italic_underline_combination():
    """测试 _render_styled_paragraph 处理加粗、斜体、下划线及组合样式的转义输出"""
    # 模拟单个 Run: 斜体 + 下划线
    run_iu = MagicMock()
    run_iu.text = "斜体带下划线文本"
    run_iu.bold = False
    run_iu.italic = True
    run_iu.underline = True
    run_iu.font.color.rgb = None

    p_iu = MagicMock()
    p_iu.runs = [run_iu]
    p_iu.text = "斜体带下划线文本"

    rendered_iu = DocxParser._render_styled_paragraph(p_iu)
    assert '<span class="style-italic-underline"><u>*斜体带下划线文本*</u></span>' in rendered_iu

    # 模拟单个 Run: 纯加粗
    run_b = MagicMock()
    run_b.text = "加粗文本"
    run_b.bold = True
    run_b.italic = False
    run_b.underline = False
    run_b.font.color.rgb = None

    p_b = MagicMock()
    p_b.runs = [run_b]
    p_b.text = "加粗文本"

    rendered_b = DocxParser._render_styled_paragraph(p_b)
    assert "**加粗文本**" in rendered_b

    # 模拟混合段落: 普通 + (斜体+下划线)
    run_normal = MagicMock()
    run_normal.text = "普通说明："
    run_normal.bold = False
    run_normal.italic = False
    run_normal.underline = False
    run_normal.font.color.rgb = None

    p_mix = MagicMock()
    p_mix.runs = [run_normal, run_iu]
    p_mix.text = "普通说明：斜体带下划线文本"

    rendered_mix = DocxParser._render_styled_paragraph(p_mix)
    assert rendered_mix == '普通说明：<span class="style-italic-underline"><u>*斜体带下划线文本*</u></span>'


def test_render_styled_paragraph_font_color():
    """测试带颜色属性的字体转义输出"""
    run_red = MagicMock()
    run_red.text = "废标红线"
    run_red.bold = True
    run_red.italic = False
    run_red.underline = False
    run_red.font.color.rgb = "FF0000"

    p_red = MagicMock()
    p_red.runs = [run_red]
    p_red.text = "废标红线"

    rendered_red = DocxParser._render_styled_paragraph(p_red)
    assert '<span style="color:#FF0000">**废标红线**</span>' in rendered_red


def test_docx_parser_integration_file_parse(tmp_path):
    """集成测试：使用 python-docx 动态生成包含样式的 real .docx 文件，验证 DocxParser.parse 输出"""
    docx_file = tmp_path / "style_test.docx"
    doc = docx.Document()

    # 1. 标题
    doc.add_heading("第四章 技术规范", level=1)

    # 2. 包含斜体且带下划线的特殊段落
    p = doc.add_paragraph()
    r1 = p.add_run("注意：")
    r1.bold = True

    r2 = p.add_run("响应时间不得高于 50ms。")
    r2.italic = True
    r2.underline = True

    doc.save(str(docx_file))

    # 执行解析
    res = docx_parser.parse(str(docx_file))
    assert res is not None
    assert "markdown_content" in res

    md_content = res["markdown_content"]
    assert "# 第四章 技术规范" in md_content
    assert "**注意：**" in md_content
    assert '<span class="style-italic-underline"><u>*响应时间不得高于 50ms。*</u></span>' in md_content
