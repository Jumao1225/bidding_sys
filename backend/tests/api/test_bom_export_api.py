"""
BOM 成本测算 Word 文档导出 API 测试 (test_bom_export_api.py)
"""

import io
import pytest
import httpx
from docx import Document as DocxDocument

from app.main import app
from app.db.session import SessionLocal
from app.db.models.user import User, Tenant
from app.db.models.project import Project, Document
from app.db.models.ai_analysis import CostEstimate
from app.core.security import create_access_token


@pytest.fixture
def bom_export_env():
    """创建测试用的租户、用户与测试文档"""
    db = SessionLocal()
    tenant = None
    user = None
    project = None
    doc = None
    cost_item = None

    try:
        tenant = Tenant(name="BomExportTenant", is_active=True)
        db.add(tenant)
        db.flush()

        user = User(
            email="bom_user@example.com",
            hashed_password="fakehashedpassword",
            tenant_id=tenant.id,
            role="user",
            is_active=True
        )
        db.add(user)
        db.flush()

        project = Project(name="BOM测算工程项目", tenant_id=tenant.id, status="created")
        db.add(project)
        db.flush()

        doc = Document(
            tenant_id=tenant.id,
            user_id=user.id,
            project_id=project.id,
            filename="某工程光伏招标文件.docx",
            file_path="dummy_path.docx",
            parse_status="completed",
            parsed_metadata={
                "budget_limit": "¥1,000,000.00",
                "cost_analysis": {
                    "analysis_summary": "测试汇总说明",
                    "items": []
                }
            }
        )
        db.add(doc)
        db.flush()

        cost_item = CostEstimate(
            tenant_id=tenant.id,
            user_id=user.id,
            document_id=doc.id,
            project_id=project.id,
            item_code="A-01",
            item_name="光伏组件",
            spec_requirement="650W",
            quantity=100.0,
            unit="块",
            unit_price=800.0,
            calculated_total=80000.0,
            matched_name="光伏组件",
            matched_brand="天合",
            matched_model="650W",
            section_name="发电设备"
        )
        db.add(cost_item)
        db.commit()

        token = create_access_token(user.id)

        yield {
            "tenant": tenant,
            "user": user,
            "doc": doc,
            "token": token
        }
    finally:
        if cost_item:
            db.delete(cost_item)
        if doc:
            db.delete(doc)
        if project:
            db.delete(project)
        if user:
            db.delete(user)
        if tenant:
            db.delete(tenant)
        db.commit()
        db.close()


@pytest.mark.asyncio
async def test_export_bom_docx_with_payload_success(bom_export_env):
    """测试通过 POST payload 传递定制 BOM 列表成功导出 docx 文件"""
    env = bom_export_env
    doc_id = env["doc"].id
    token = env["token"]

    payload = {
        "document_title": "测试项目定制标书",
        "total_cost": 500000.0,
        "budget_limit": "¥600,000.00",
        "status_text": "在最高投标限价内可控",
        "analysis_summary": "前端实时编辑的评估建议",
        "items": [
            {
                "item_code": "01",
                "name": "高压逆变柜",
                "spec_requirement": "10kV / 500kW",
                "qty": 2,
                "unit": "台",
                "ref_price": 250000.0,
                "subtotal": 500000.0,
                "matched_name": "高压逆变柜",
                "matched_brand": "阳光电源",
                "matched_model": "SG500",
                "section_name": "变电系统"
            }
        ]
    }

    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
        response = await client.post(
            f"/api/v1/analysis/{doc_id}/export-bom-docx",
            json=payload,
            headers={"Authorization": f"Bearer {token}"}
        )

    assert response.status_code == 200
    assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in response.headers.get("content-type", "")
    assert len(response.content) > 0

    # 验证导出的 docx 内容
    doc_file = DocxDocument(io.BytesIO(response.content))
    full_text = "\n".join(p.text for p in doc_file.paragraphs)
    assert "拟投入设备及 BOM 成本测算清单" in full_text
    assert "测试项目定制标书" in full_text
    assert "前端实时编辑的评估建议" in full_text


@pytest.mark.asyncio
async def test_export_bom_docx_fallback_to_db_success(bom_export_env):
    """测试不传 payload items 时回退到数据库 CostEstimate 记录生成 docx"""
    env = bom_export_env
    doc_id = env["doc"].id
    token = env["token"]

    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
        response = await client.post(
            f"/api/v1/analysis/{doc_id}/export-bom-docx",
            json={},
            headers={"Authorization": f"Bearer {token}"}
        )

    assert response.status_code == 200
    assert len(response.content) > 0

    doc_file = DocxDocument(io.BytesIO(response.content))
    table = doc_file.tables[0]
    row_text = " ".join(c.text for c in table.rows[1].cells)
    assert "光伏组件" in row_text
    assert "天合" in row_text


@pytest.mark.asyncio
async def test_export_bom_xlsx_api_success(bom_export_env):
    """测试通过 POST payload 导出标准 Excel (.xlsx) 工作簿接口"""
    from openpyxl import load_workbook

    env = bom_export_env
    doc_id = env["doc"].id
    token = env["token"]

    payload = {
        "document_title": "测试光伏标书Excel",
        "total_cost": 300000.0,
        "budget_limit": "¥400,000.00",
        "status_text": "在最高投标限价内可控",
        "analysis_summary": "通过Excel导出的专家建议",
        "items": [
            {
                "name": "集中式逆变器",
                "spec_requirement": "1500V / 3.125MW",
                "matched_brand": "特变电工",
                "matched_model": "TC3125KF",
                "matched_manufacturer": "特变电工新能源",
                "qty": 1,
                "unit": "台",
                "ref_price": 300000.0,
                "subtotal": 300000.0,
            }
        ]
    }

    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
        response = await client.post(
            f"/api/v1/analysis/{doc_id}/export-bom-xlsx",
            json=payload,
            headers={"Authorization": f"Bearer {token}"}
        )

    assert response.status_code == 200
    assert "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet" in response.headers.get("content-type", "")
    assert len(response.content) > 0

    # 验证生成的 xlsx 内容
    wb = load_workbook(io.BytesIO(response.content))
    ws = wb.active
    assert "拟投入设备及 BOM 成本测算清单" in ws['A1'].value
    assert "测试光伏标书Excel" in ws['A2'].value
    assert ws['B6'].value == "集中式逆变器"
    assert ws['C6'].value == "特变电工 TC3125KF"
    assert ws['D6'].value == "特变电工新能源"
    assert ws['H6'].value == 300000.0
    assert "叁拾万元整" in ws['A7'].value

