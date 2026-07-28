"""
集成测试：拟人化标书自动填报 Agent 全流程 (test_human_like_bid_agent.py)
"""

import pytest
import os
import tempfile
from unittest.mock import patch, AsyncMock, MagicMock
from docx import Document

from app.agents.human_like_bid_agent import HumanLikeBidFillerAgent, human_like_bid_filler_agent
from app.services.llm_slot_analyzer import SlotAnalysisReport, SlotItem


@pytest.mark.asyncio
async def test_human_like_bid_agent_pipeline():
    """测试拟人化标书自动填报 Agent 的完整 Task 执行与 Office CLI / SQL 工具编排流程"""
    # 1. 创建临时的真实测试 .docx
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
        doc = Document()
        doc.add_paragraph("一、投标函")
        doc.add_paragraph("投标人名称：______")
        doc.add_paragraph("统一社会信用代码：【 】")
        doc.save(tmp_file.name)
        tmp_doc_path = tmp_file.name

    tmp_out_path = tmp_doc_path.replace(".docx", "_filled.docx")

    try:
        # Mock LLM Slot Analyzer 返回 2 个槽位
        mock_slot_report = SlotAnalysisReport(
            total_slots_found=2,
            slots=[
                SlotItem(
                    path="/body/p[2]",
                    run_index=1,
                    label="投标人名称：",
                    raw_placeholder="______",
                    target_field_intent="company_name",
                    confidence_score=0.98,
                    reasoning="识别为公司名称"
                ),
                SlotItem(
                    path="/body/p[3]",
                    run_index=0,
                    label="统一社会信用代码：",
                    raw_placeholder="【 】",
                    target_field_intent="credit_code",
                    confidence_score=0.95,
                    reasoning="识别为信用代码"
                )
            ],
            summary="测试样本分析"
        )

        with patch("app.agents.human_like_bid_agent.analyze_slots_with_llm", return_value=mock_slot_report), \
             patch("app.services.office_cli_service.office_cli_service.query_structure", new_callable=AsyncMock) as mock_query, \
             patch("app.services.office_cli_service.office_cli_service.batch_update", new_callable=AsyncMock) as mock_batch:

            mock_query.return_value = "/body/p[2]: 投标人名称：______\n/body/p[3]: 统一社会信用代码：【 】"
            mock_batch.return_value = {"success": True, "executed_commands_count": 2}

            agent = HumanLikeBidFillerAgent(task_id="test_task_123")
            result = await agent.execute_fill_pipeline(
                document_id="test_doc_001",
                template_doc_path=tmp_doc_path,
                output_doc_path=tmp_out_path
            )

            assert result["success"] is True
            assert result["total_slots_detected"] == 2
            assert result["total_slots_filled"] == 2
            assert len(result["filled_items"]) == 2

            # 验证查库结果正确性 (兼容真实数据库与默认演示库)
            company_item = next(item for item in result["filled_items"] if item["intent"] == "company_name")
            assert company_item["looked_up_value"] in ["四川石楠建设工程有限公司", "聚猫科技股份有限公司"]

            credit_item = next(item for item in result["filled_items"] if item["intent"] == "credit_code")
            assert credit_item["looked_up_value"] in ["91510000MA6X12345X", "91110108MA01988888X"]

    finally:
        if os.path.exists(tmp_doc_path):
            os.remove(tmp_doc_path)
        if os.path.exists(tmp_out_path):
            os.remove(tmp_out_path)
