from docx import Document

def create_test_doc():
    doc = Document()
    doc.add_heading('招标文件 (Bidding Document)', 0)
    
    doc.add_heading('第一章 招标公告', level=1)
    doc.add_paragraph('本项目要求投标人必须具备建筑工程施工总承包一级资质。')
    doc.add_paragraph('投标保证金：人民币50万元。')
    
    doc.add_heading('第二章 评标标准', level=1)
    doc.add_paragraph('评分项1：企业资质（满分20分）。具有特级资质得20分，一级得10分，其他不得分。')
    doc.add_paragraph('评分项2：过往业绩（满分30分）。每提供一个类似项目合同得5分，最多30分。')
    
    doc.add_heading('第三章 合同条款', level=1)
    doc.add_paragraph('1. 违约责任：如工期延误，违约金为合同总价的 5%。')
    doc.add_paragraph('2. 必须承诺在中标后 10 日内签订合同。')

    doc.save('test_bidding.docx')
    print("生成测试文档: test_bidding.docx 成功")

if __name__ == "__main__":
    create_test_doc()
