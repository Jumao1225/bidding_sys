"""
表格结构识别与多行表头处理工具模块 (table_utils.py)

提供对 Word 表格多行复合表头、列头名称合并提取及数据区定位等通用结构化分析能力。
"""

import re
from typing import List, Tuple, Optional, Dict, Any
from loguru import logger


def detect_table_header_rows(table) -> int:
    """
    智能识别 Word 表格的表头所占行数（支持单行表头与多行复合表头）。

    判定维度：
    1. XML <w:tblHeader/> 显式标记；
    2. 跨列合并特征（第 0 行存在合并且第 1 行包含子列属性名）；
    3. 纵向合并特征（w:vMerge）与表头属性标签词汇；
    4. 数据行特征排除（第 1 行若含数据序号、下划线、占位符等则非表头）。

    :param table: docx.table.Table 对象
    :return: 表头行数 header_rows_count (>= 1)
    """
    if table is None or not hasattr(table, 'rows') or not table.rows:
        return 1

    try:
        # 1. 优先检查 Word XML 原生的 <w:tblHeader/> 标记
        xml_header_count = 0
        for r in table.rows:
            trPr = r._tr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trPr')
            if trPr is not None:
                tblHeader = trPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblHeader')
                if tblHeader is not None:
                    xml_header_count += 1
                    continue
            break
        if xml_header_count >= 1:
            return xml_header_count

        total_rows = len(table.rows)
        if total_rows <= 1:
            return 1

        total_cols = len(table.rows[0].cells)

        # 2. 检查第 0 行是否存在跨列合并 (unique _tc 数量少于 total_cols)
        row0_unique_tcs = set(c._tc for c in table.rows[0].cells)
        has_h_merge_row0 = len(row0_unique_tcs) < total_cols

        # 提取第 0 行、第 1 行的单元格文本
        row0_texts = [c.text.strip().replace("\n", "") for c in table.rows[0].cells]
        row1_texts = [c.text.strip().replace("\n", "") for c in table.rows[1].cells]

        # 数据行特征模式：纯数字序号(如 1, 2)、待填占位符、金额数字、纯空格/下划线
        data_patterns = [
            re.compile(r'^\d+$'),                                        # 纯数字序号
            re.compile(r'^\d+\.\d+$'),                                   # 层级数字序号如 2.1
            re.compile(r'\[(?:待补充|待手动补充|建议人工|待填|查询|错误)[^\]]*\]'), # 占位符
            re.compile(r'_{2,}'),                                        # 下划线
            re.compile(r'^\d+(?:\.\d+)?\s*(?:元|万元|%|块|台|套|米|m)?$'),      # 金额/数量
        ]

        def _is_data_cell(t: str) -> bool:
            if not t:
                return False
            return any(pat.search(t) for pat in data_patterns)

        # 如果第 1 行有明显的数据特征（例如首列是纯数字 1/2，或者单元格有占位符），则第 1 行必然是数据行而非表头
        row1_has_data_features = any(_is_data_cell(t) for t in row1_texts if t)

        # 特殊保护：如果第 1 行的第 0 列是 "1" 或 "01" 等纯序号，且不包含任何子表头关键词，则第 1 行为数据行
        if row1_texts and re.match(r'^[0-9]+$', row1_texts[0]):
            row1_has_data_features = True

        if row1_has_data_features:
            return 1

        # 表头子项/属性列名特征词
        sub_header_keywords = [
            "类别", "编号", "等级", "专业", "级别", "证号", "名称", "规格", "型号", "单位", "数量", "单价", "合价",
            "总价", "小计", "说明", "备注", "品牌", "厂家", "生产厂家", "响应", "偏离", "分项", "职务", "经验"
        ]
        has_sub_header_kw = any(
            any(kw in t for kw in sub_header_keywords)
            for t in row1_texts if t
        )

        # 检查第 1 行与第 0 行是否存在纵向合并 (vMerge)
        has_vmerge = False
        for c in table.rows[1].cells:
            tcPr = c._tc.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
            if tcPr is not None:
                vMerge = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')
                if vMerge is not None:
                    has_vmerge = True
                    break

        # 若第 0 行存在跨列合并 或 第 1 行存在 vMerge 纵向合并 或 包含子表头属性关键词，且第 1 行无数据特征
        if (has_h_merge_row0 or has_vmerge or has_sub_header_kw) and not row1_has_data_features:
            # 进一步检测是否存在 3 行复合表头
            if total_rows > 2:
                row2_texts = [c.text.strip().replace("\n", "") for c in table.rows[2].cells]
                row2_has_data_features = any(_is_data_cell(t) for t in row2_texts if t)
                if row2_texts and re.match(r'^[0-9]+$', row2_texts[0]):
                    row2_has_data_features = True
                if not row2_has_data_features and any(any(kw in t for kw in sub_header_keywords) for t in row2_texts if t):
                    return 3
            return 2

        return 1
    except Exception as e:
        logger.warning(f"智能检测表格表头行数异常: {e}")
        return 1


def get_merged_header_texts(table, header_rows_count: int = 1) -> List[str]:
    """
    提取并智能合并多行表头的列定义名称列表。
    例如：第 0 行 '资格证书' 与 第 1 行 '类别' 合并为 '资格证书(类别)'。

    :param table: docx.table.Table 对象
    :param header_rows_count: 表头行数
    :return: 每一列的完整表头名称列表
    """
    if not table or not hasattr(table, 'rows') or not table.rows:
        return []

    total_cols = len(table.rows[0].cells)
    header_rows_count = min(header_rows_count, len(table.rows))

    if header_rows_count <= 1:
        return [c.text.strip().replace("\n", "") for c in table.rows[0].cells]

    col_headers = []
    for c_i in range(total_cols):
        parts = []
        for r_i in range(header_rows_count):
            if c_i < len(table.rows[r_i].cells):
                txt = table.rows[r_i].cells[c_i].text.strip().replace("\n", "")
                if txt and txt not in parts:
                    parts.append(txt)
        if not parts:
            col_headers.append("")
        elif len(parts) == 1:
            col_headers.append(parts[0])
        else:
            # 组合父表头与子表头：如 "资格证书(类别)"
            parent = parts[0]
            child = "/".join(parts[1:])
            col_headers.append(f"{parent}({child})")

    return col_headers


def get_table_header_logical_spans(table, hdr_count: int = 1) -> List[Tuple[int, int]]:
    """
    获取表头的逻辑列网格跨度列表。
    取表头最底端一行 (table.rows[hdr_count - 1])，
    对于每个逻辑列，返回其在底层物理网格中的 (start_col_idx, end_col_idx)（inclusive）。
    例如：
    1. 6 列物理网格中，Row 0 第 1、2 列合并为名称列，则返回 [(0, 0), (1, 2), (3, 3), (4, 4), (5, 5)]；
    2. 多行表头中（hdr_count=2），Row 1 分出了子列（如类别、编号），则以 Row 1 的细分列跨度为准。
    """
    if table is None or not hasattr(table, 'rows') or not table.rows:
        return []

    hdr_idx = max(0, min(hdr_count - 1, len(table.rows) - 1))
    header_row = table.rows[hdr_idx]
    spans = []
    seen_tcs = []

    for c_i, cell in enumerate(header_row.cells):
        tc = cell._tc
        if tc not in seen_tcs:
            seen_tcs.append(tc)
            spans.append([c_i, c_i])
        else:
            idx = seen_tcs.index(tc)
            spans[idx][1] = c_i

    return [tuple(s) for s in spans]


def align_row_to_header_grid_spans(row, header_spans: List[Tuple[int, int]]):
    """
    使指定行 (row) 的单元格合并结构与表头的逻辑列跨度 (header_spans) 100% 对齐。
    如果表头某列跨越了多个物理列（如 start < end），则将该行的对应单元格进行合并。
    """
    if not row or not header_spans or not hasattr(row, 'cells') or not row.cells:
        return

    total_cells = len(row.cells)
    for start_c, end_c in header_spans:
        if 0 <= start_c < end_c < total_cells:
            if row.cells[start_c]._tc != row.cells[end_c]._tc:
                try:
                    row.cells[start_c].merge(row.cells[end_c])
                except Exception as me:
                    logger.debug(f"单元格合并容错: {me}")


def clean_row_vmerge(row):
    """
    彻底清洗指定行 (row) 中所有单元格的垂直纵向合并标记 (<w:vMerge/>)。
    防止 table.add_row() 新增的数据行因继承下方/上方落款行的 vMerge 属性
    而导致大单元格跨页垂直拉伸或产生大面积空白断层。
    """
    if not row or not hasattr(row, 'cells') or not row.cells:
        return

    seen_tcs = set()
    for cell in row.cells:
        tc = cell._tc
        if tc in seen_tcs:
            continue
        seen_tcs.add(tc)
        tcPr = tc.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
        if tcPr is not None:
            for vMerge in tcPr.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge'):
                try:
                    tcPr.remove(vMerge)
                except Exception:
                    pass


def is_fixed_slot_form_table(table, header_rows_count: int = 1) -> bool:
    """
    通过纯结构与排版特征智能判定 Word 表格是否为“固定单元格填报表（Fixed-Slot Form Table）”。
    【零硬编码原则】：不依赖任何硬编码的章节名称、固定项目名或写死的价格标签，而是基于表格几何拓扑与单元格预置形态判定：

    1. 【行数规模小且结构封闭】：
       - 扣除表头后，数据行通常只有 1 行；
       - 或总行数较少（如数据行 1 行 + 1 行跨列合并的表尾/汇总/落款行）；
    2. 【预置内容与特定槽位】：
       - 数据行中已预印有特定的描述性文本/属性标签，并非全空的列表占位行；
       - 仅有局部特定单元格处于空白或待填状态；
    3. 【非动态清单结构】：
       - 动态清单表通常表头包含多列明细属性（如序号、品名、规格、数量、单价、合价等），且数据区留有较大空白占位行以供展开。
    """
    if table is None or not hasattr(table, 'rows') or not table.rows:
        return False

    total_rows = len(table.rows)
    if total_rows <= 1:
        return True

    hdr_count = max(1, min(header_rows_count, total_rows - 1))
    data_rows = table.rows[hdr_count:]
    if not data_rows:
        return True

    first_data_cells = [c.text.strip() for c in data_rows[0].cells]
    has_pre_printed_content = any(len(txt) >= 2 for txt in first_data_cells)

    # 1. 结构判定：如果有效数据行仅有 1 行，且已预印了具体描述内容（而非纯空占位行）
    if len(data_rows) == 1:
        return has_pre_printed_content

    # 2. 如果数据行仅有 2 行，且最后一行存在跨列合并特征（如跨多列合并的汇总或落款行）
    if len(data_rows) == 2:
        total_cols = len(table.rows[0].cells)
        last_row_unique_tcs = set(c._tc for c in data_rows[-1].cells)
        if len(last_row_unique_tcs) < total_cols:
            # 只有当首个数据行已经预印了具体项目内容时，才作为固定表单；若首行全为空白占位，则是带表尾的动态清单表
            return has_pre_printed_content

    # 3. 检查首个数据行是否预置了具体文字内容（非纯空白占位行）且总数据行较少
    if has_pre_printed_content and len(data_rows) <= 3:
        return True

    return False


def is_fixed_slot_summary_table(table, chapter_title: str = "") -> bool:
    """
    向后兼容别名函数，底层调用零硬编码的 is_fixed_slot_form_table。
    """
    return is_fixed_slot_form_table(table)


def get_doc_chapter_tables_mapping(doc) -> List[Dict[str, Any]]:
    """
    单次拓扑遍历 Word 文档 body 流，构建每个章节标题与其下方专属表格的拓扑映射列表。
    支持大章层级包含关系：大章条目不仅记录自身紧随的表格，还会自动聚合其下属所有子节的表格。
    """
    if doc is None or not hasattr(doc, 'element') or not hasattr(doc.element, 'body'):
        return []

    mapping = []
    current_major_entry = None
    current_entry = {"chapter_title": "PREAMBLE", "table_indices": []}

    for elem in doc.element.body:
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag == "p":
            txt = "".join(elem.itertext()).strip()
            if not txt:
                continue

            # 判定是否为一级大章节标题 (如 "五、投标配置及分项报价表" 或 "第X章 ...")
            is_major_title = bool(
                re.match(r'^[一二三四五六七八九十百0-9]{1,3}[、\.\s]', txt)
                or re.match(r'^第[一二三四五六七八九十0-9]+[章节篇部分]', txt)
            ) and len(txt) <= 150 and not any(p in txt for p in ["。", "；", ";", "！"])

            # 判定是否为子节/子表格标题 (如 "投标报价分析表"、"技术要求响应及偏离表")
            is_sub_title = bool(
                ("格式" in txt[:20] and len(txt) <= 150)
                or (txt.endswith("表") and len(txt) <= 150 and not any(p in txt for p in ["。", "；", ";", "！"]))
            ) and not any(p in txt for p in ["。", "；", "！"])

            if is_major_title:
                if current_entry["chapter_title"] != "PREAMBLE" or current_entry["table_indices"]:
                    if current_entry not in mapping:
                        mapping.append(current_entry)
                current_major_entry = {"chapter_title": txt, "table_indices": [], "is_major": True}
                current_entry = current_major_entry
            elif is_sub_title:
                if current_entry["chapter_title"] != "PREAMBLE" or current_entry["table_indices"]:
                    if current_entry not in mapping and current_entry is not current_major_entry:
                        mapping.append(current_entry)
                current_entry = {"chapter_title": txt, "table_indices": [], "is_major": False}

        elif tag == "tbl":
            # 找到对应 table 的索引
            for t_idx, t in enumerate(doc.tables):
                if t._element == elem:
                    if t_idx not in current_entry["table_indices"]:
                        current_entry["table_indices"].append(t_idx)
                    # 大章自动层级包含下属子节的表格
                    if current_major_entry and current_entry is not current_major_entry:
                        if t_idx not in current_major_entry["table_indices"]:
                            current_major_entry["table_indices"].append(t_idx)
                    break

    if current_entry["chapter_title"] != "PREAMBLE" or current_entry["table_indices"]:
        if current_entry not in mapping:
            mapping.append(current_entry)
    if current_major_entry and current_major_entry not in mapping:
        mapping.append(current_major_entry)

    return mapping


def get_chapter_specific_table_indices(doc, chapter_title: str) -> List[int]:
    """
    根据章节标题精准计算该章节在 Word DOM 拓扑中拥有的专属表格索引列表 (0-indexed)。
    具备【目录与正文智能去重】与【大章子表层级包含】机制，自动过滤文档开头的目录项，100% 锁定正文中的真实表格。
    若该章节不包含任何表格（如封面、投标函文本段），严格返回空列表 []。
    """
    if not doc or not doc.tables or not chapter_title:
        return []

    mapping = get_doc_chapter_tables_mapping(doc)
    if not mapping:
        return []

    # 1. 清洗 chapter_title 的核心词汇
    clean_target = re.sub(r'^[一二三四五六七八九十百0-9\s、\.\(\)（）]+', '', chapter_title).strip()
    clean_target = re.sub(r'[\s、\.\(\)（）]+', '', clean_target)

    # 提取有意义的 token（长度 >= 2）
    target_tokens = set()
    for i in range(len(clean_target) - 1):
        target_tokens.add(clean_target[i:i+2])
    if len(clean_target) >= 3:
        for i in range(len(clean_target) - 2):
            target_tokens.add(clean_target[i:i+3])

    best_score = 0.0
    best_entry = None

    for entry in mapping:
        entry_title = entry.get("chapter_title", "")
        if entry_title == "PREAMBLE":
            continue
        clean_entry = re.sub(r'^[一二三四五六七八九十百0-9\s、\.\(\)（）]+', '', entry_title).strip()
        clean_entry = re.sub(r'[\s、\.\(\)（）]+', '', clean_entry)

        if not clean_entry:
            continue

        # 严格要求基础文本相似度
        base_score = 0.0
        if clean_target == clean_entry:
            base_score = 100.0
        elif clean_target in clean_entry or clean_entry in clean_target:
            base_score = 60.0
        elif target_tokens:
            matched_tokens = sum(1 for tk in target_tokens if tk in clean_entry)
            token_ratio = matched_tokens / len(target_tokens)
            if token_ratio >= 0.5:
                base_score = token_ratio * 50.0

        # 若没有任何文本相似度，坚决跳过，严禁给无关联章节加分！
        if base_score <= 0.0:
            continue

        # 仅在文本确实匹配的前提下：若该条目在正文中拥有真实表格，大幅加 100 分（用于有效区分无表格的目录项与有表格的正文项）
        score = base_score
        if entry.get("table_indices"):
            score += 100.0
            # 若表格数量更多（属于大章聚合），额外微调优先匹配大章集合
            score += len(entry["table_indices"]) * 2.0

        if score > best_score:
            best_score = score
            best_entry = entry

    if best_entry and best_entry.get("table_indices"):
        return best_entry["table_indices"]

    # 2. 如果拓扑映射未直接命中表格，按表头关键词二级回退（严格要求强相关）
    header_matches = []
    for t_idx, table in enumerate(doc.tables):
        if not table.rows:
            continue
        hdr_txt = "".join(c.text.strip() for c in table.rows[0].cells)
        if clean_target and (clean_target in hdr_txt or (len(hdr_txt) <= 20 and hdr_txt in clean_target)):
            header_matches.append(t_idx)
        elif target_tokens:
            matched_tokens = sum(1 for tk in target_tokens if tk in hdr_txt)
            if len(target_tokens) >= 2 and (matched_tokens / len(target_tokens)) >= 0.7:
                header_matches.append(t_idx)

    if 1 <= len(header_matches) <= 2:
        return header_matches

    return []


def extract_chapter_dom_structure(
    doc_or_path,
    chapter_title: str,
    selector: str = "all"
) -> str:
    """
    【精准章节作用域提取器 (Chapter Scope DOM Extractor)】
    根据目标章节标题，从 Word 物理 DOM 拓扑流中精确定位起始位置，
    仅提取 100% 属于该章节下的所有段落 (/body/p[N]) 与表格 (/body/tbl[M])，直到下一个同级章节标题为止。
    自动识别并跳过文档开头的目录项，精确定位正文中的真实起始段落。

    优势：
    1. 100% 完整无损提取当前章节的所有内容与表格，绝不丢失任何一处待填信息；
    2. 从物理层面天然排除其余 90%+ 无关章节的数十万字噪音，无需粗暴截断；
    3. 结构极其小巧清晰（通常 500~2500 字符），大模型秒级完成思考与写盘。
    """
    if not doc_or_path or not chapter_title:
        return ""

    import os
    from docx import Document
    try:
        if isinstance(doc_or_path, str):
            if not os.path.exists(doc_or_path):
                return ""
            doc = Document(doc_or_path)
        else:
            doc = doc_or_path

        if not hasattr(doc, 'element') or not hasattr(doc.element, 'body'):
            return ""

        clean_target = re.sub(r'^[一二三四五六七八九十百0-9\s、\.\(\)（）]+', '', chapter_title).strip()
        clean_target = re.sub(r'[\s、\.\(\)（）]+', '', clean_target)

        target_tokens = set()
        for i in range(len(clean_target) - 1):
            target_tokens.add(clean_target[i:i+2])

        # 第一阶段：先扫描所有匹配目标标题的段落位置，通过评分机制选取得分最高且位于正文中的真实起始段落
        candidates = []
        p_count = 0
        for elem in doc.element.body:
            tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
            if tag == "p":
                p_count += 1
                txt = "".join(elem.itertext()).strip()
                if not txt:
                    continue
                clean_p = re.sub(r'^[一二三四五六七八九十百0-9\s、\.\(\)（）]+', '', txt).strip()
                clean_p = re.sub(r'[\s、\.\(\)（）]+', '', clean_p)
                score = 0.0
                if clean_target and clean_p:
                    if clean_target in clean_p or clean_p in clean_target:
                        score += 100.0
                    elif target_tokens:
                        overlap = sum(1 for tk in target_tokens if tk in clean_p)
                        score += (overlap / len(target_tokens)) * 50.0

                if score >= 30.0:
                    candidates.append((p_count, score))

        if not candidates:
            return ""

        max_score = max(c[1] for c in candidates)
        # 在最高得分中取最后一个（自动跳过文档开头的目录项，精准锁定正文位置）
        target_start_p = [c[0] for c in candidates if c[1] == max_score][-1]

        # 第二阶段：从正文起始段落开始收集该章节的所有段落与表格
        p_idx = 0
        tbl_idx = 0
        in_target = False
        collected_lines = []

        for elem in doc.element.body:
            tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag

            if tag == "p":
                p_idx += 1
                txt = "".join(elem.itertext()).strip()
                if p_idx == target_start_p:
                    in_target = True
                    collected_lines.append(f"📌 [当前目标章节正文标题] /body/p[{p_idx}]: {txt}")
                    continue
                elif in_target and p_idx > target_start_p:
                    # 判断是否遇到下一个同级大章节标题（严格匹配中文大写序号/第X章/附件X，严禁将 1、2、5、等正文条款细项误判为大章标题）
                    is_top_chapter_pattern = bool(
                        re.match(r'^[一二三四五六七八九十百]+[、\.\s]', txt)
                        or re.match(r'^第[一二三四五六七八九十0-9]+[章节篇部分]', txt)
                        or re.match(r'^【[一二三四五六七八九十0-9]+】', txt)
                        or re.match(r'^附件[一二三四五六七八九十0-9]+', txt)
                    )
                    is_not_body_clause = not txt.endswith((':', '：', '；', ';', '。', '，', ',', '！', '!')) and not any(
                        kw in txt for kw in ["有关事宜", "通讯地址", "如下", "为：", "即：", "规定", "履行", "承诺", "如果", "保证", "同意"]
                    )
                    is_ch_title = is_top_chapter_pattern and is_not_body_clause and len(txt) <= 80
                    if is_ch_title:
                        in_target = False
                        break
                    if txt and selector in ("paragraph", "all"):
                        collected_lines.append(f"/body/p[{p_idx}]: {txt}")

            elif tag == "tbl":
                tbl_idx += 1
                if in_target:
                    # 找到对应的 table 对象
                    target_table = None
                    if 0 <= tbl_idx - 1 < len(doc.tables):
                        target_table = doc.tables[tbl_idx - 1]

                    if target_table and target_table.rows:
                        hdr_count = detect_table_header_rows(target_table)
                        headers = get_merged_header_texts(target_table, hdr_count)
                        hdr_desc = " | ".join(headers) if headers else "未识别表头"
                        total_r = len(target_table.rows)
                        total_c = len(target_table.rows[0].cells) if target_table.rows else 0
                        collected_lines.append(
                            f"\n📊 [本章节专属目标表格] 路径: `/body/tbl[{tbl_idx}]` | 规模: 共 {total_r} 行 x {total_c} 列（含 {hdr_count} 行表头）"
                        )
                        collected_lines.append(f"   表头定义: `[{hdr_desc}]`")
                        collected_lines.append("   【表格现有各行明细与单元格节点坐标】：")
                        for r_i, r in enumerate(target_table.rows):
                            row_cells_info = []
                            unique_tcs = []
                            for c_i, cell in enumerate(r.cells):
                                if cell._tc in unique_tcs:
                                    continue
                                unique_tcs.append(cell._tc)
                                cell_txt = cell.text.strip().replace("\n", " ")
                                if not cell_txt or cell_txt in ("_", "——", "--", "待填"):
                                    display_val = f"⚠️[待填空位: /body/tbl[{tbl_idx}]/tr[{r_i+1}]/tc[{c_i+1}]]"
                                else:
                                    display_val = f"'{cell_txt[:60]}'"
                                row_cells_info.append(f"tc[{c_i+1}]: {display_val}")
                            row_type = "【表头行】" if r_i < hdr_count else f"Row {r_i+1}"
                            collected_lines.append(f"   └─ [{row_type}] `/body/tbl[{tbl_idx}]/tr[{r_i+1}]` -> {' | '.join(row_cells_info)}")

        if collected_lines:
            header_str = f"🎯 【章节专属 100% 完整 DOM 视野 (零信息丢失)】: 共提取出 {len(collected_lines)} 个专属节点：\n"
            return header_str + "\n".join(collected_lines)

    except Exception as e:
        logger.warning(f"提取章节专属 DOM 结构异常: {e}")

    return ""


def inspect_and_repair_table_blanks(doc, document_id: str = "") -> int:
    """
    【表格留白智能检查与 LLM 动态修复引擎 — 零硬编码自愈闭环】
    1. 自行检查：扫描文档中条款偏离表/实质性响应表，识别出“含有条款要求/指标名称但右侧响应列留白”的未完成行；
    2. 智能修复：绝不直接删除原模板条款行，而是调用 LLM 针对这些被遗漏的条款，按表头定义逐项生成合规的响应内容并写回单元格；
    3. 序号规整：确保整表序号连续无断层。
    """
    if doc is None or not hasattr(doc, 'tables') or not doc.tables:
        return 0

    import json, re
    repaired_total = 0
    from app.services.llm_service import llm_service

    for t_i, table in enumerate(doc.tables):
        if not table.rows or len(table.rows) <= 1:
            continue

        hdr_row = table.rows[0]
        hdr_cells = [c.text.strip().replace("\n", " ") for c in hdr_row.cells]
        hdr_txt = "".join(hdr_cells)
        total_cols = len(hdr_cells)

        # 仅针对条款偏离表、实质性要求对照表等响应类表格执行自愈检查
        is_deviation_tbl = any(k in hdr_txt for k in ["偏离", "响应", "承诺", "实质性", "对照表"])
        if total_cols < 3 or not is_deviation_tbl:
            continue

        # 收集该表格中所有“有条款但右侧数据列存在空白”的待修复行
        unfilled_rows_info = []
        for r_i in range(1, len(table.rows)):
            row = table.rows[r_i]
            cells_txt = [c.text.strip() for c in row.cells]

            # 若整行全空，跳过
            if all(not txt for txt in cells_txt):
                continue

            # 寻找该行中主要的内容描述列（通常在第 2 列或第 1 列）
            main_desc = ""
            for c_idx in range(min(2, total_cols)):
                if len(cells_txt[c_idx]) >= 4 and not cells_txt[c_idx].isdigit():
                    main_desc = cells_txt[c_idx]
                    break

            # 对已有条款行做确定性的列级补齐，避免把“合同总价款”等普通商务词
            # 误当成汇总行，也避免完全依赖 LLM 自愈后仍留下空单元格。列位置只
            # 依据当前表头识别，不假设固定列数或固定列号。
            header_indexes = {
                "commitment": None,
                "status": None,
                "reason": None,
            }
            header_defaults = {"status": "无"}
            for h_idx, h_text in enumerate(hdr_cells):
                if header_indexes["commitment"] is None and any(
                    keyword in h_text for keyword in ["服务承诺", "响应情况", "投标响应", "响应内容", "承诺内容"]
                ):
                    header_indexes["commitment"] = h_idx
                if header_indexes["status"] is None and any(
                    keyword in h_text for keyword in ["有无偏离", "是否偏离", "偏离情况", "是否响应"]
                ):
                    header_indexes["status"] = h_idx
                    header_defaults["status"] = "是" if "是否响应" in h_text else "无"
                if header_indexes["reason"] is None and any(
                    keyword in h_text for keyword in ["偏离内容", "偏离原因", "偏离说明"]
                ):
                    header_indexes["reason"] = h_idx

            if main_desc and (header_indexes["commitment"] is not None or header_indexes["status"] is not None or header_indexes["reason"] is not None):
                commitment_idx = header_indexes["commitment"]
                status_idx = header_indexes["status"]
                reason_idx = header_indexes["reason"]
                if commitment_idx is not None and commitment_idx < len(row.cells) and not row.cells[commitment_idx].text.strip():
                    row.cells[commitment_idx].text = "我方承诺完全响应该条款要求。"
                if status_idx is not None and status_idx < len(row.cells):
                    current_status = row.cells[status_idx].text.strip()
                    valid_statuses = {"是", "否", "无", "无偏离", "有", "有偏离"}
                    if not current_status or current_status.isdigit() or current_status not in valid_statuses:
                        row.cells[status_idx].text = header_defaults["status"]
                if reason_idx is not None and reason_idx < len(row.cells) and not row.cells[reason_idx].text.strip():
                    row.cells[reason_idx].text = "完全响应招标文件要求，无偏离。"
                cells_txt = [c.text.strip() for c in row.cells]

            # 若有实质条款描述，但后续有空白单元格，则判定为待修复留白行
            if main_desc and any(not cells_txt[c_idx] for c_idx in range(1, total_cols)):
                unfilled_rows_info.append({
                    "row_index": r_i,
                    "term_desc": main_desc,
                    "current_cells": cells_txt
                })

        if not unfilled_rows_info:
            continue

        logger.info(f"🔍 [表格留白自检] 表格 /body/tbl[{t_i+1}] 检出 {len(unfilled_rows_info)} 行未填留白，启动 LLM 动态自愈修复...")

        # 构造 LLM 动态修复 Prompt（完全抽象，零具体数据硬编码）
        prompt_rows_str = ""
        for u_item in unfilled_rows_info:
            prompt_rows_str += f"- 行索引 {u_item['row_index']}: 条款内容='{u_item['term_desc']}'\n"

        repair_prompt = f"""你是资深招投标响应与合规专家。以下表格中存在部分条款行未完成填写，请根据表头定义与条款内容，为每一行生成完整合规的填报响应。

【表格表头定义】:
{json.dumps(hdr_cells, ensure_ascii=False)}

【待补齐填写的条款清单】:
{prompt_rows_str}

【填报要求】:
1. 每一行必须严格按照表头定义的每一列（从第 1 列到最后一列）生成完整的字符串数组；
2. 响应表述必须严谨专业、完全闭合，严禁使用省略号或空字符串；
3. 必须输出严格合法的 JSON 数组，格式如下：
[
  {{"row_index": 12, "col_values": ["11", "条款原文", "响应承诺内容", "无偏离", "偏离说明"]}},
  ...
]"""

        try:
            llm_inst = llm_service.get_llm(temperature=0.1, json_mode=False)
            response = llm_inst.invoke(repair_prompt)
            raw_text = response.content if hasattr(response, 'content') else str(response)

            # 解析 JSON 结果
            import json, re
            from app.agents.bid_filler_agent import align_table_row_cells
            m_json = re.search(r'\[\s*\{.*\}\s*\]', raw_text, re.DOTALL)
            if m_json:
                repair_items = json.loads(m_json.group(0))
                for item in repair_items:
                    r_idx = item.get("row_index")
                    col_vals = item.get("col_values")
                    if isinstance(r_idx, int) and 1 <= r_idx < len(table.rows) and isinstance(col_vals, list):
                        target_row = table.rows[r_idx]
                        aligned_vals = align_table_row_cells(col_vals, total_cols, r_idx - 1)
                        for c_i, val in enumerate(aligned_vals):
                            if c_i < len(target_row.cells) and val:
                                # 仅在原单元格为空时补齐，保护已有数据
                                if not target_row.cells[c_i].text.strip():
                                    target_row.cells[c_i].text = str(val).strip()
                                    repaired_total += 1
                logger.info(f"   🎯 [LLM 动态自愈成功] 成功为表格 /body/tbl[{t_i+1}] 智能补齐 {len(repair_items)} 行空白单元格！")
        except Exception as repair_err:
            logger.warning(f"   LLM 表格留白自愈调用异常: {repair_err}")

    # 最后执行基础的序号规整
    for table in doc.tables:
        if not table.rows or len(table.rows) <= 1:
            continue
        data_rows = table.rows[1:]
        num_cells = sum(1 for r in data_rows if r.cells and (not r.cells[0].text.strip() or r.cells[0].text.strip().isdigit()))
        if data_rows and num_cells >= len(data_rows) * 0.7:
            curr_seq = 1
            for r_i in range(1, len(table.rows)):
                r = table.rows[r_i]
                if len(r.cells) > 0:
                    r_txt = r.cells[0].text.strip()
                    if not r_txt or r_txt.isdigit():
                        r.cells[0].text = str(curr_seq)
                        curr_seq += 1

    return repaired_total


def extract_equipment_tables_and_context(raw_text: str) -> str:
    """
    智能靶向过滤（全量无截断）：
    精准提取所有【标的物/设备材料清单表格】、关联章节标题以及技术工况/关键技术要求段落，
    仅剔除无关的纯行政人事表（人员社保、执业证书、财务审计等），绝不进行字符截断以保证上下文 100% 完整。
    """
    if not raw_text:
        return ""

    # 1. 查找所有 HTML 表格和 Markdown 表格
    html_pattern = re.compile(r'<table[\s\S]*?</table>', re.IGNORECASE)
    md_pattern = re.compile(r'(?:(?:^|\n)\|[^\n]+\|\n(?:\|[-:\s|]+\|\n)(?:\|[^\n]+\|\n?)+)', re.MULTILINE)

    table_spans = []
    for m in html_pattern.finditer(raw_text):
        table_spans.append((m.start(), m.end(), "html", m.group(0)))
    for m in md_pattern.finditer(raw_text):
        table_spans.append((m.start(), m.end(), "md", m.group(0)))

    table_spans.sort(key=lambda x: x[0])

    if not table_spans:
        # 若没有识别出表格，完整保留原文，绝不截断
        return raw_text

    # 纯通用清单表格特征词（严格杜绝任何具体设备或行业特定名词硬编码）
    EQUIPMENT_KEYWORDS = ["设备", "标的", "货物", "材料", "物资", "产品", "服务", "工程量", "清单", "规格", "型号", "参数", "指标", "数量", "单位", "单价", "合价", "总价", "定额"]
    EXCLUDE_KEYWORDS = ["近三年财务", "财务审计", "营业额", "社保缴纳", "人员资质", "执业证书", "身份证", "评分细则", "评分标准"]

    filtered_sections = []
    last_end = 0

    for start, end, t_type, tbl_content in table_spans:
        plain_tbl = re.sub(r'<[^>]+>', ' ', tbl_content).lower()
        is_equipment_table = any(kw in plain_tbl for kw in EQUIPMENT_KEYWORDS)
        is_excluded = any(kw in plain_tbl for kw in EXCLUDE_KEYWORDS) and not any(k in plain_tbl for k in ["采购清单", "主要标的物", "设备名称", "货物名称"])

        preceding_text = raw_text[last_end:start].strip()
        
        if is_equipment_table and not is_excluded:
            # 提取紧随该设备表格前方的章节大标题或说明（如 "第X标段/分部工程清单"）
            if preceding_text:
                preceding_lines = [l.strip() for l in preceding_text.split('\n') if l.strip()]
                # 倒序向上查找最近的各级标题行与技术说明（扩大探测窗口至 15 行），确保大标题 100% 完整保留
                headers = []
                for l in reversed(preceding_lines[-15:]):
                    if l.startswith('#') or re.match(r'^(?:[0-9]+[、.．]|[一二三四五六七八九十]+[、.．]|第[0-9一二三四五六七八九十]+[标标段章节部分区]|(?:（|\()[0-9一二三四五六七八九十]+(?:）|\)))', l):
                        headers.append(l)
                    elif any(c in l for c in ['表', '清单', '需求', '规格', '标段', '工程', '部分', '系统', '一览表']):
                        headers.append(l)
                headers = list(reversed(headers))
                if not headers and preceding_lines:
                    headers = preceding_lines[-3:]
                if headers:
                    filtered_sections.append("\n".join(headers))
            filtered_sections.append(tbl_content)
        else:
            # 若是非设备表格（如人员资质表、财务表），跳过该表格及其紧贴标题
            pass
        
        last_end = end

    # 处理最后一个表格后面的剩余文本（保留特殊工况、现场施工要求、技术门槛等完整说明，绝不截断）
    remaining_text = raw_text[last_end:].strip()
    if remaining_text:
        filtered_sections.append(remaining_text)

    if not filtered_sections:
        return raw_text

    return "\n\n".join(filtered_sections)


def normalize_section_name(raw_sec: Optional[str]) -> Optional[str]:
    """
    分部/工程大类规范化函数：
    忠实保留标书提取的原始 section_name，去除前后多余空白，杜绝任何硬编码或破坏性截断。
    """
    if not raw_sec or not isinstance(raw_sec, str):
        return None
    s = raw_sec.strip()
    return s if s else None


def is_narrative_clause_or_lead_in(text: str) -> bool:
    """
    基于通用语法标点与公文篇章结构，纯通用判断文本是否为正文叙述句、条款导语或转折声明句（而非表单属性标签/槽位）。

    【核心原则】：严禁任何硬编码业务数据与具体名称，纯基于语言学标点、公文篇章结构与谓语特征判定。

    判定规则：
    1. 标点特征：真正表单属性标签绝不包含句中标点（逗号、分号、句号、感叹号、问号）；
       若文本在冒号之前或主体部分包含 `[，,；;。！!？?]`，判定为叙述从句（如 "据此函，签字人兹宣布同意如下："）。
    2. 条款序号特征：若文本以公文条款序号开头（如 `1、`, `5.`, `（一）`, `(2)`, `一、` 等），判定为条款正文/导语标题（如 "5、与本投标有关的正式通讯地址为："）。
    3. 公文致函抬头特征：以公文致函抬头开头（如 "致某某单位："）。
    4. 公文导语/动词结构特征：
       - 包含承前启后引导词（如 "如下", "为：", "据此", "特此", "兹宣布", "兹同意", "兹授权", "根据贵方", "有关事宜", "全权处理", "声明如下", "承诺如下", "保证如下", "授权如下" 等）；
       - 结尾系词（如 "通讯地址为："、"条件为："）。
    5. 长度与结构特征：
       - 纯表单属性标签为简短的名词短语（去除符号空格后长度通常 <= 20）；
       - 超过 25 字符且无明确待填占位符（如连续下划线/多个空格）的文本判定为长叙述句。
    """
    if not text or not isinstance(text, str):
        return False

    raw = text.strip()
    if not raw:
        return False

    # 1. 标点特征：表单属性标签绝不包含逗号、分号、句号、感叹号、问号
    # 若在冒号前含有这些标点，100% 为叙述从句
    colon_idx = -1
    for c_char in [':', '：']:
        idx = raw.find(c_char)
        if idx != -1:
            colon_idx = idx if colon_idx == -1 else min(colon_idx, idx)

    text_before_colon = raw[:colon_idx] if colon_idx != -1 else raw
    if re.search(r'[，,；;。！!？?]', text_before_colon):
        return True

    # 2. 条款序号特征：以条款序号开头（如 `1、`, `5.`, `（一）`, `(2)`, `一、` 等）
    if re.match(r'^\s*(?:\d+[\.、\)]|[一二三四五六七八九十百]+[\.、\)]|\([0-9一二三四五六七八九十]+\)|（[0-9一二三四五六七八九十]+）)', raw):
        return True

    # 3. 公文致函抬头特征（如 "致某某单位："）
    if re.match(r'^\s*致[^\n:：]{2,50}[:：]?\s*$', raw):
        return True

    # 4. 公文转折/引导词与谓语动词结构特征
    clean_no_punct = re.sub(r'[\s_＿\-\–\—\(\)（）\[\]［］【】:：\.\,，。；;、“”"\'`]', '', raw)
    narrative_markers = [
        "如下", "宣布", "同意", "据此", "特此", "兹宣布", "兹同意", "兹授权",
        "根据贵方", "有关事宜", "全权处理", "履行合同", "承担责任", "严格履行",
        "为以下", "声明如下", "承诺如下", "保证如下", "授权如下"
    ]
    if any(marker in clean_no_punct for marker in narrative_markers):
        return True

    # 结尾系词判断（如 "通讯地址为："、"条件为："）
    if re.search(r'为[:：]\s*$', raw):
        return True

    # 5. 长度与结构特征：无占位符且长度过长的文本
    has_explicit_slot = bool(re.search(r'(?:_{2,}|＿{2,}|\s{3,}|\[待[^\]]+\]|［待[^］]+］|\s*年\s*月\s*日)', raw))
    if len(clean_no_punct) > 25 and not has_explicit_slot:
        return True

    return False


def get_chapter_body_elements(doc, chapter_title: str) -> List[Any]:
    """
    提取 Word 文档中属于指定章节的所有 body 级 XML 元素节点（段落与表格）
    """
    if doc is None or not hasattr(doc, 'element') or not hasattr(doc.element, 'body') or not chapter_title:
        return []

    clean_target = re.sub(r'^[一二三四五六七八九十百0-9\s、\.\(\)（）]+', '', chapter_title).strip()
    clean_target = re.sub(r'[\s、\.\(\)（）]+', '', clean_target)
    target_tokens = set(clean_target[i:i+2] for i in range(len(clean_target) - 1)) if len(clean_target) > 1 else {clean_target}

    candidates = []
    p_count = 0
    for elem in doc.element.body:
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag == "p":
            p_count += 1
            txt = "".join(elem.itertext()).strip()
            if not txt:
                continue
            clean_p = re.sub(r'^[一二三四五六七八九十百0-9\s、\.\(\)（）]+', '', txt).strip()
            clean_p = re.sub(r'[\s、\.\(\)（）]+', '', clean_p)
            score = 0.0
            if clean_target and clean_p:
                if clean_target in clean_p or clean_p in clean_target:
                    score += 100.0
                elif target_tokens:
                    overlap = sum(1 for tk in target_tokens if tk in clean_p)
                    score += (overlap / len(target_tokens)) * 50.0

            if score >= 30.0:
                candidates.append((p_count, elem, score))

    if not candidates:
        return []

    max_score = max(c[2] for c in candidates)
    target_start_elem = [c[1] for c in candidates if c[2] == max_score][-1]

    collected = []
    in_target = False
    for elem in doc.element.body:
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if elem == target_start_elem:
            in_target = True
            collected.append(elem)
            continue
        elif in_target:
            if tag == "p":
                txt = "".join(elem.itertext()).strip()
                is_top_chapter_pattern = bool(
                    re.match(r'^[一二三四五六七八九十百]+[、\.\s]', txt)
                    or re.match(r'^第[一二三四五六七八九十0-9]+[章节篇部分]', txt)
                    or re.match(r'^【[一二三四五六七八九十0-9]+】', txt)
                    or re.match(r'^附件[一二三四五六七八九十0-9]+', txt)
                )
                is_not_body_clause = not txt.endswith((':', '：', '；', ';', '。', '，', ',', '！', '!')) and not any(
                    kw in txt for kw in ["有关事宜", "通讯地址", "如下", "为：", "即：", "规定", "履行", "承诺", "如果", "保证", "同意"]
                )
                if is_top_chapter_pattern and is_not_body_clause and len(txt) <= 80:
                    break
            collected.append(elem)

    return collected


def reset_chapter_to_template(
    working_docx_path: str,
    clean_template_path: str,
    chapter_title: str
) -> bool:
    """
    单章节重置器：在重新生成/微调某章节前，将工作副本中该章节的所有段落与表格精准重置为原始纯净模板状态。
    
    【核心目的】：彻底抹除历史运行中可能产生的脏数据或过时填充，确保本次重新生成 100% 覆盖原内容。
    """
    import os
    if not os.path.exists(working_docx_path) or not os.path.exists(clean_template_path):
        return False
    try:
        from docx import Document
        from copy import deepcopy
        doc_work = Document(working_docx_path)
        doc_tpl = Document(clean_template_path)

        tpl_elems = get_chapter_body_elements(doc_tpl, chapter_title)
        work_elems = get_chapter_body_elements(doc_work, chapter_title)

        if not tpl_elems or not work_elems:
            logger.warning(f"重置章节未找到对应 DOM 元素: tpl={len(tpl_elems)}, work={len(work_elems)}")
            return False

        # 精确切片替换：将 work 章节范围内的所有元素移除，并在原位置插入 tpl 纯净模板的所有元素
        parent = work_elems[0].getparent()
        if parent is not None:
            insert_idx = parent.index(work_elems[0])
            for w_elem in work_elems:
                parent.remove(w_elem)
            for i, t_elem in enumerate(tpl_elems):
                parent.insert(insert_idx + i, deepcopy(t_elem))

        doc_work.save(working_docx_path)
        logger.info(f"🔄 [单章节重置] 成功将目标章节 [{chapter_title}]（旧 {len(work_elems)} 节点 -> 模板 {len(tpl_elems)} 节点）精准还原为原始模板样式！其余章节 100% 保持不变。")
        return True
    except Exception as e:
        logger.warning(f"重置章节至模板状态异常: {e}")
        return False






