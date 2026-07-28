"""
投标书 Word 文档生成器 (WordGenerator)

根据从招标文件「投标文件格式」章节提取的目录结构，
动态编排并生成符合甲方格式要求的投标书 .docx 草稿。

三层样式合并策略：
  优先级1: LLM 从格式章节文字中提取的排版规范
  优先级2: 原始 .docx 文件的内置样式
  优先级3: 中文招标文档通用默认值
"""

import io
import os
import copy
from typing import Optional, Any
from loguru import logger
from pydantic import BaseModel, Field
from docx import Document
from docx import Document as DocxDocument
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from app.utils.rmb_formatter import number_to_chinese_rmb


# ============================================================
# 原生 .docx 深拷贝与数据位置换引擎 (100% 还原原标书视觉外观)
# ============================================================



def _set_tc_text_cleanly(tc, val: str):
    """
    干净原位替换单元格文本：将 val 放入首个 <w:t> 节点，并将其余 <w:t> 节点清空。
    若单元格全空（无 <w:t> 节点），自动在段落中追加 <w:r><w:t> 节点并写入。
    """
    t_nodes = [t for t in tc.iter() if t.tag.endswith('t')]
    if t_nodes:
        t_nodes[0].text = str(val)
        for t in t_nodes[1:]:
            t.text = ""
    else:
        p_nodes = [p for p in tc.iter() if p.tag.endswith('p')]
        if p_nodes:
            p = p_nodes[0]
        else:
            p = parse_xml(f'<w:p {nsdecls("w")}/>')
            tc.append(p)
        r = parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{str(val)}</w:t></w:r>')
        p.append(r)


def _fill_table_with_mapping(
    table_elem, 
    tr_elems: list, 
    data_items: list, 
    column_mappings: list, 
    total_cost: float = 0.0
):
    """
    通用矩阵填充器 (Zero Hardcoding):
    根据 TableAgent 决定的 column_mappings，零硬编码将 data_items 动态塞入 Word 单元格！
    """
    if not tr_elems or not data_items:
        return

    header_tr = tr_elems[0]
    # 移除模板展示行
    for tr in tr_elems[1:]:
        table_elem.remove(tr)

    # 建立 col_index -> ColumnMapping 快捷映射
    mapping_by_col = {m.col_index: m for m in column_mappings}

    for i, item in enumerate(data_items):
        row_tr = copy.deepcopy(header_tr)
        tc_elems = [e for e in row_tr if e.tag.endswith('tc')]

        for col_idx, tc in enumerate(tc_elems):
            mapping = mapping_by_col.get(col_idx)
            val = ""
            if mapping:
                fkey = mapping.field_key
                default_v = mapping.default_val or ""

                if fkey == "seq":
                    val = str(i + 1)
                elif fkey == "name":
                    val = item.get("name") or item.get("item_name") or item.get("req") or default_v
                elif fkey == "spec":
                    val = str(item.get("spec_requirement") or item.get("spec") or item.get("model") or item.get("level") or default_v)[:40]
                elif fkey == "brand":
                    val = item.get("brand") or item.get("matched_brand") or default_v or "一线品牌"
                elif fkey == "manufacturer":
                    val = item.get("manufacturer") or item.get("matched_manufacturer") or item.get("factory") or item.get("company") or default_v or "知名品牌/合规厂家"
                elif fkey == "unit":
                    val = item.get("unit") or "项"
                elif fkey == "qty":
                    val = str(item.get("qty") or item.get("quantity") or "1")
                elif fkey == "price":
                    p = item.get("ref_price") or item.get("price") or 0.0
                    val = f"{float(p):,.2f}"
                elif fkey == "subtotal":
                    sub = item.get("subtotal") or 0.0
                    val = f"{float(sub):,.2f}"
                elif fkey == "remark":
                    val = item.get("remark") or default_v or "满足招标文件要求"
                elif fkey == "status":
                    val = item.get("status") or default_v or "完全响应"
                elif fkey == "expiry":
                    val = item.get("expiry") or item.get("expiry_date") or default_v or "长期有效"
                elif fkey == "role":
                    val = item.get("role") or default_v or "核心人员"
                else:
                    val = default_v

            _set_tc_text_cleanly(tc, val)

        table_elem.append(row_tr)


def _populate_native_table(table_elem, current_section: str, metadata: dict, analysis: dict, chapter_results: Optional[dict] = None):
    """
    针对原生的 OpenXML Table 节点 (w:tbl)，利用 TableAgent 自主识别表格表头与列语义，
    零硬编码地将 BOM 采购清单、资质证书、团队人员或响应条款填充至 Word 单元格中！
    """
    tr_elems = [e for e in table_elem if e.tag.endswith('tr')]
    if not tr_elems:
        return

    # 1. 提取表格第一行作为表头列名
    header_tr = tr_elems[0]
    tc_elems_header = [e for e in header_tr if e.tag.endswith('tc')]
    if not tc_elems_header:
        return

    header_texts = ["".join(tc.itertext()).strip() for tc in tc_elems_header]

    # 2. 调度 TableAgent 进行智能表格识别与列语义映射
    from app.agents.nodes.table_agent import analyze_table_structure_and_map
    decision = analyze_table_structure_and_map(header_texts, current_section)

    cost = analysis.get("cost_analysis", {})
    cost_items = cost.get("items", [])
    total_cost = cost.get("total_cost", 0.0)

    timeline = metadata.get("timeline") or {}
    project_name = timeline.get("project_name") or "分布式光伏发电项目"

    # 3. 兜底获取 chapter_results 中的数据
    if not cost_items and chapter_results:
        for tid, cres in chapter_results.items():
            if cres.get("mapping_hint") in ["pricing", "cost"] and cres.get("table_rows"):
                cost_items = cres.get("table_rows")
                break

    qual_rows = []
    if chapter_results:
        for tid, cres in chapter_results.items():
            if cres.get("mapping_hint") == "qualification" and cres.get("table_rows"):
                qual_rows = cres.get("table_rows")
                break

    eval_data = metadata.get("evaluation", {})
    qual_md = metadata.get("qualification", {})

    # 4. 依据 TableAgent 决策结果调度通用矩阵填充器
    if decision.table_type == "pricing_bom" and cost_items:
        _fill_table_with_mapping(table_elem, tr_elems, cost_items, decision.column_mappings, total_cost)

    elif decision.table_type == "qualification_certs" and qual_rows:
        _fill_table_with_mapping(table_elem, tr_elems, qual_rows, decision.column_mappings, total_cost)

    elif decision.table_type == "clause_compliance":
        hard_service = eval_data.get("hard_service_requirements") or {}
        hard_quals = qual_md.get("mandatory_qualifications") or []
        combined_requirements = []
        if isinstance(hard_service, dict):
            for k, v in hard_service.items():
                combined_requirements.append({"name": f"{k}: {v}", "spec": f"{k}: {v}", "status": "完全响应"})
        if isinstance(hard_quals, list):
            for q in hard_quals:
                combined_requirements.append({"name": str(q), "spec": str(q), "status": "完全响应"})

        if combined_requirements:
            _fill_table_with_mapping(table_elem, tr_elems, combined_requirements, decision.column_mappings, total_cost)

    elif decision.table_type == "team_personnel":
        personnel = qual_md.get("personnel_requirements") or []
        if personnel:
            personnel_rows = []
            for p in personnel:
                if isinstance(p, dict):
                    personnel_rows.append({"name": p.get("role", "核心人员"), "spec": p.get("requirement", "具备相应资格")})
                else:
                    personnel_rows.append({"name": str(p), "spec": "具备相应资格"})
            _fill_table_with_mapping(table_elem, tr_elems, personnel_rows, decision.column_mappings, total_cost)

    elif decision.table_type == "opening_summary" or "开标一览表" in current_section:
        # 原位填充开标一览表 (In-Place Fill)，100% 保持原生表格列宽、合并单元格与模板行
        rmb_str = number_to_chinese_rmb(total_cost) if total_cost > 0 else ""
        formatted_cost = f"¥{total_cost:,.2f}" if total_cost > 0 else ""

        for tr in tr_elems[1:]:
            row_text = "".join(tr.itertext()).strip()

            # 1. 大写总报价行
            if "大写" in row_text or "人民币" in row_text:
                tc_nodes = [e for e in tr if e.tag.endswith('tc')]
                for tc in tc_nodes:
                    tc_text = "".join(tc.itertext()).strip()
                    if ("大写" not in tc_text or tc_text == "") and rmb_str:
                        _set_tc_text_cleanly(tc, f"人民币 {rmb_str}")
                        break

            # 2. 项目概况与小写总价行
            else:
                tc_nodes = [e for e in tr if e.tag.endswith('tc')]
                if len(tc_nodes) >= 3:
                    if project_name:
                        _set_tc_text_cleanly(tc_nodes[0], project_name)
                    # 单元格 1: 完美保留原模板自带的技术要求描述（绝不写死）
                    c1_text = "".join(tc_nodes[1].itertext()).strip()
                    if not c1_text:
                        _set_tc_text_cleanly(tc_nodes[1], "完全响应招标文件技术要求")
                    if formatted_cost:
                        _set_tc_text_cleanly(tc_nodes[2], formatted_cost)
    else:
        logger.info(f"章节 [{current_section}] 的表格 TableAgent 判定保持原生模版框架")




def replace_paragraph_text_smartly(elem, old_pattern: str, new_text: str) -> bool:
    """
    智能跨 Run / 跨 <w:t> 节点替换段落文本。
    对段落拼接全文本，在第一个匹配的 <w:t> 节点上更新为 new_text，清空段落内后续匹配到的划线节点。
    """
    import re
    t_nodes = [t for t in elem.iter() if t.tag.endswith('t') and t.text]
    if not t_nodes:
        return False

    full_p_text = "".join(t.text for t in t_nodes)
    if not re.search(old_pattern, full_p_text):
        return False

    # 使用正则替换全段文本
    replaced_p_text = re.sub(old_pattern, new_text, full_p_text)

    # 替换后的文本填入首个节点，后续节点置空
    t_nodes[0].text = replaced_p_text
    for t in t_nodes[1:]:
        t.text = ""
    return True



def _replace_xml_element_variables(
    elem, 
    current_section: str, 
    metadata: dict, 
    analysis: dict,
    chapter_results: Optional[dict] = None
):
    """
    在原生的 OpenXML Element 节点（段落或表格）中智能检索并按语义置换数据。
    支持段落跨 Run / 跨 <w:t> 节点的划线占位符与标书字段语义置换。
    """
    if elem.tag.endswith('tbl'):
        _populate_native_table(elem, current_section, metadata, analysis, chapter_results)
        return

    timeline = metadata.get("timeline") or {}
    project_name = timeline.get("project_name") or ""
    project_id = timeline.get("project_id_code") or timeline.get("project_id") or ""

    contacts = timeline.get("contacts") or []
    tenderer = ""
    for c in contacts:
        if isinstance(c, dict):
            role = c.get("role_type", "")
            unit = c.get("unit_name", "")
            if any(k in role for k in ["招标", "买方", "甲方", "业主", "建设"]):
                tenderer = unit
                if tenderer:
                    break
    if not tenderer:
        tenderer = (
            timeline.get("tenderer_name") or 
            timeline.get("buyer_name") or 
            timeline.get("unit_name") or 
            timeline.get("project_owner") or ""
        )

    company_name = (
        analysis.get("company_quals") or 
        metadata.get("qualification", {}).get("company_name") or 
        timeline.get("supplier_name") or 
        timeline.get("company_name") or 
        timeline.get("bidder_name") or ""
    )

    if not company_name:
        try:
            from app.agents.tools.writer_tools import get_company_qualifications_tool
            quals = get_company_qualifications_tool()
            if quals and isinstance(quals, list) and len(quals) > 0:
                company_name = quals[0].get("company_name", "")
        except Exception:
            pass

    cost = analysis.get("cost_analysis") or {}
    total_cost = cost.get("total_cost", 0.0)
    total_cost_rmb = number_to_chinese_rmb(total_cost) if total_cost > 0 else ""

    p_text = "".join(elem.itertext()).strip()
    import re
    import datetime

    # 仅在【封面】或【投标函】及其格式小节处理
    is_target_section = any(k in current_section for k in ["封面", "投标函", "开标一览表"])

    if is_target_section and p_text:
        # 1. 项目名称替换
        if ("项目名称：" in p_text or "项目名称:" in p_text) and not any(k in p_text for k in ["历史", "业绩", "拟派", "人员", "以往"]):
            if project_name:
                prefix = "项目名称：" if "项目名称：" in p_text else "项目名称:"
                replace_paragraph_text_smartly(elem, r'项目名称[：:]\s*[_＿\s—]*.*$', f"{prefix}{project_name}")

        # 2. 招标编号 / 项目编号替换
        if ("招标编号：" in p_text or "招标编号:" in p_text or "项目编号：" in p_text or "项目编号:" in p_text):
            if project_id:
                prefix = "招标编号：" if "招标编号" in p_text else "项目编号："
                replace_paragraph_text_smartly(elem, r'(招标编号|项目编号)[：:]\s*[_＿\s—]*.*$', f"{prefix}{project_id}")

        # 3. 投标函 -> 致：XXX（招标人/买方单位）
        if "投标函" in current_section and ("致：" in p_text or "致:" in p_text):
            if tenderer:
                prefix = "致：" if "致：" in p_text else "致:"
                replace_paragraph_text_smartly(elem, r'致[：:]\s*[_＿\s—]*.*$', f"{prefix}{tenderer}：")

        # 4. 投标函 -> 根据贵方的______号招标文件
        if "投标函" in current_section and ("号招标文件" in p_text or "招标文件" in p_text):
            if project_id:
                replace_paragraph_text_smartly(elem, r'根据贵方的?\s*[_＿\s—]*号招标文件', f"根据贵方的 {project_id} 号招标文件")
                replace_paragraph_text_smartly(elem, r'[_＿]{2,}\s*号招标文件', f"{project_id}号招标文件")

        # 5. 投标函 -> 代表投标人______ (投标人的名称)
        if "投标函" in current_section and ("代表投标人" in p_text or "投标人名称" in p_text):
            if company_name:
                replace_paragraph_text_smartly(elem, r'代表投标人\s*[_＿\s—]*（?投标人的名称）?', f"代表投标人：{company_name}")

        # 6. 封面 & 投标函 -> 投标人名称：______
        if ("封面" in current_section or "投标函" in current_section) and ("投标人名称：" in p_text or "投标人名称:" in p_text):
            if company_name:
                prefix = "投标人名称：" if "投标人名称：" in p_text else "投标人名称:"
                replace_paragraph_text_smartly(elem, r'投标人名称[：:]\s*[_＿\s—]*.*$', f"{prefix}{company_name}")

        # 7. 投标函 -> 授权下述签字人____ (姓名和职务)
        if "投标函" in current_section and ("签字人" in p_text or "法定代表人" in p_text):
            replace_paragraph_text_smartly(elem, r'签字人\s*[_＿\s—]*（?姓名和职务）?', "签字人 (法定代表人/授权代理人)")

        # 8. 投标函 -> 投标总价与大写金额
        if "投标函" in current_section and any(k in p_text for k in ["投标总价", "投标报价", "报价总额", "人民币"]):
            if total_cost > 0:
                replace_paragraph_text_smartly(elem, r'[_＿]{2,}', f" 人民币{total_cost_rmb}（¥{total_cost:,.2f}元） ")

        # 9. 投标函 -> 日期：____年__月__日
        if ("封面" in current_section or "投标函" in current_section) and ("日期：" in p_text or "日期:" in p_text or "年" in p_text and "月" in p_text and "日" in p_text):
            curr_date_str = datetime.date.today().strftime("%Y 年 %m 月 %d 日")
            replace_paragraph_text_smartly(elem, r'日期[：:]\s*[_＿\s—]*年[_＿\s—]*月[_＿\s—]*日', f"日期：{curr_date_str}")
            replace_paragraph_text_smartly(elem, r'[_＿\s—]*年[_＿\s—]*月[_＿\s—]*日', curr_date_str)




def clone_format_section_from_original_docx(
    file_path: str,
    metadata: dict,
    analysis: dict,
    chapter_results: Optional[dict] = None
) -> Optional[bytes]:
    """
    核心克隆引擎：直接从原始 .docx 招标文件中定位「投标文件格式」大章，
    原汁原味地克隆该章节及其后所有原生 Word 段落与表格结构（100% 保持原生样式、边框、缩进、字号），
    并在原段落与表格中执行智能上下文数据置换与分类表单填充。
    """
    if not file_path or not os.path.exists(file_path) or not file_path.lower().endswith(".docx"):
        return None

    try:
        orig_doc = DocxDocument(file_path)
        body = orig_doc.element.body

        elements = list(body)
        total_elements = len(elements)
        keywords = ["投标文件格式", "投标主要文件目录", "一、封面", "投标文件封面", "投标文件组成"]

        candidates = []
        import re

        for idx, elem in enumerate(elements):
            text = "".join(elem.itertext()).strip()
            if not text:
                continue

            if any(k in text for k in keywords):
                # 排除目录行（包含多点、领导点、或末尾跟页码数字如 '格式 40'）
                is_toc = bool(re.search(r'(\.|\u2026|_|-){2,}\s*\d+|\b\d{1,3}$', text))
                # 排除引用说明句（如 "详见第七章"）
                is_ref = any(ref in text for ref in ["详见", "参见", "参照", "按第"])

                if not is_toc and not is_ref:
                    score = 0
                    # 如果位于文档中后半段（通常 > 20% 节点处），大幅加分，避开前几页的目录
                    if idx > total_elements * 0.20:
                        score += 50
                    if len(text) < 50:
                        score += 30
                    if any(h in text for h in ["第", "章", "格式", "封面"]):
                        score += 20
                    candidates.append((score, idx, text, elem))

        start_element = None
        start_idx = -1
        if candidates:
            # 排序：优先按得分降序，得分相同时取位置靠后的（正文大章通常在文档后半部分）
            candidates.sort(key=lambda x: (x[0], x[1]), reverse=True)
            best_score, start_idx, best_text, start_element = candidates[0]
            logger.info(f"在原 .docx 中精确定位到格式大章真实起点 (索引 {start_idx}/{total_elements}, 得分 {best_score}, 节点: '{best_text[:30]}')")

        if start_element is None or start_idx < 0:
            logger.warning("原 .docx 中未搜寻到 '投标文件格式' 正文大章节点")
            return None

        target_elements = elements[start_idx:]
        logger.info(f"成功截取格式大章正文节点，丢弃前 {start_idx} 个无关节点，克隆后 {len(target_elements)} 个原生 Word 节点")


        new_doc = DocxDocument(file_path)
        new_body = new_doc.element.body

        for child in list(new_body):
            if child.tag.endswith('sectPr'):
                continue
            new_body.remove(child)

        current_section = "封面"
        known_section_rules = [
            ("封面", ["封面"]),
            ("投标函", ["投标函"]),
            ("开标一览表", ["开标一览表", "开标表"]),
            ("资格证明", ["资格证明", "资格审查", "资格文件", "书面承诺"]),
            ("投标配置及分项报价表", ["分项报价", "投标配置", "报价清单", "采购清单", "报价分析", "报价表"]),
            ("常用零件及耗材", ["常用零件", "耗材报价"]),
            ("实质性条款", ["实质性", "响应对照表", "硬性要求"]),
            ("设计方案", ["设计方案", "技术方案", "施工组织"]),
            ("项目负责人", ["项目负责人", "人员介绍", "团队人员", "人员配备"]),
            ("投标人情况", ["投标人情况", "投标人简介"]),
            ("技术要求偏离", ["技术要求偏离", "技术偏离"]),
            ("商务要求偏离", ["商务条款偏离", "商务偏离"]),
            ("其他材料", ["其他材料", "其它材料"])
        ]

        for elem in target_elements:
            text_content = "".join(elem.itertext()).strip()
            if elem.tag.endswith('p') and text_content:
                clean_text = re.sub(r'\s+', '', text_content)
                if len(clean_text) < 60:
                    for sec_name, keywords in known_section_rules:
                        if any(kw in clean_text for kw in keywords):
                            current_section = sec_name
                            break


            copied_elem = copy.deepcopy(elem)
            _replace_xml_element_variables(copied_elem, current_section, metadata, analysis, chapter_results)
            new_body.append(copied_elem)

        stream = io.BytesIO()
        new_doc.save(stream)
        stream.seek(0)

        logger.info(f"成功直接从原 .docx 克隆格式大章，生成字节大小: {stream.getbuffer().nbytes}")
        return stream.getvalue()

    except Exception as e:
        logger.error(f"从原 .docx 克隆格式大章过程发生异常: {e}")
        return None


# ============================================================
# 一、Pydantic Schema 定义
# ============================================================

class FormattingSpec(BaseModel):
    """从格式章节文字中提取的排版规范"""
    paper_size: Optional[str] = Field("A4", description="纸张规格")
    body_font: Optional[str] = Field(None, description="正文字体（如 '仿宋_GB2312'）")
    body_font_size: Optional[str] = Field(None, description="正文字号（如 '小四'）")
    heading_font: Optional[str] = Field(None, description="标题字体（如 '黑体'）")
    heading_font_size: Optional[str] = Field(None, description="标题字号（如 '三号'）")
    line_spacing: Optional[str] = Field(None, description="行距（如 '28磅' / '1.5倍'）")
    margins: Optional[str] = Field(None, description="页边距说明")
    other_notes: Optional[str] = Field(None, description="其他排版说明")


class OutlineItem(BaseModel):
    """投标文件目录中的一个条目"""
    number: Optional[str] = Field(None, description="编号（如 '一'、'（三）'、'1.1'）")
    title: Optional[str] = Field(None, description="章节标题")
    sub_items: list["OutlineItem"] = Field(default_factory=list, description="子条目")
    content_hint: Optional[str] = Field(None, description="甲方对该章节的说明")
    mapping_hint: Optional[str] = Field(None, description="数据映射标识")


class BidDocOutline(BaseModel):
    """从招标文件格式章节提取的完整结构"""
    source_chapter: str = Field(..., description="来源章节名（如'第七章 投标文件格式'）")
    outline: list[OutlineItem] = Field(default_factory=list, description="完整目录树")
    formatting: Optional[FormattingSpec] = Field(None, description="排版规范")


# ============================================================
# 二、中文字号与磅值映射
# ============================================================

CHINESE_FONT_SIZE_MAP: dict[str, float] = {
    "初号": 42.0, "小初": 36.0,
    "一号": 26.0, "小一": 24.0,
    "二号": 22.0, "小二": 18.0,
    "三号": 16.0, "小三": 15.0,
    "四号": 14.0, "小四": 12.0,
    "五号": 10.5, "小五": 9.0,
    "六号": 7.5, "小六": 6.5,
    "七号": 5.5, "八号": 5.0,
}


def parse_font_size(size_str: Optional[str]) -> Optional[float]:
    """将中文字号或数字字号解析为磅值 (pt)"""
    if not size_str:
        return None
    size_str = size_str.strip()
    # 先查中文字号映射
    if size_str in CHINESE_FONT_SIZE_MAP:
        return CHINESE_FONT_SIZE_MAP[size_str]
    # 尝试直接解析数字（可能带 "pt" / "磅" 后缀）
    import re
    match = re.search(r'(\d+\.?\d*)', size_str)
    if match:
        return float(match.group(1))
    return None


# ============================================================
# 三、从原始 .docx 文件中读取样式
# ============================================================

def extract_styles_from_docx(file_path: str) -> dict:
    """
    从原始 .docx 文件中提取样式信息。
    仅在原文件为 .docx 格式时调用，作为样式合并的优先级2补充。
    """
    result: dict[str, Any] = {
        "default_font": None,
        "default_font_size_pt": None,
        "page_width": None,
        "page_height": None,
        "margins": {},
    }

    if not file_path or not os.path.exists(file_path):
        return result
    if not file_path.lower().endswith(".docx"):
        return result

    try:
        doc = Document(file_path)

        # 读取默认段落样式 (Normal)
        normal_style = doc.styles['Normal']
        if normal_style.font.name:
            result["default_font"] = normal_style.font.name
        if normal_style.font.size:
            result["default_font_size_pt"] = normal_style.font.size.pt

        # 读取页面设置（取第一个 Section）
        if doc.sections:
            section = doc.sections[0]
            result["page_width"] = section.page_width
            result["page_height"] = section.page_height
            result["margins"] = {
                "top": section.top_margin,
                "bottom": section.bottom_margin,
                "left": section.left_margin,
                "right": section.right_margin,
            }

        logger.info(f"成功从原始 .docx 读取样式: 字体={result['default_font']}, 字号={result['default_font_size_pt']}pt")
    except Exception as e:
        logger.warning(f"读取原始 .docx 样式失败 ({file_path}): {e}")

    return result


# ============================================================
# 四、样式合并器
# ============================================================

class MergedStyles:
    """合并后的最终样式配置"""

    # 通用默认值（优先级3兜底）
    DEFAULT_BODY_FONT = "宋体"
    DEFAULT_BODY_SIZE_PT = 12.0      # 小四
    DEFAULT_HEADING_FONT = "黑体"
    DEFAULT_H1_SIZE_PT = 16.0        # 三号
    DEFAULT_H2_SIZE_PT = 14.0        # 四号
    DEFAULT_LINE_SPACING_PT = None   # None 表示使用 1.5 倍行距
    DEFAULT_LINE_SPACING_MULTIPLE = 1.5

    def __init__(
        self,
        formatting_spec: Optional[FormattingSpec] = None,
        docx_styles: Optional[dict] = None
    ):
        """三层合并初始化"""
        spec = formatting_spec or FormattingSpec()
        ds = docx_styles or {}

        # 正文字体：优先级1(spec) > 优先级2(docx) > 优先级3(默认)
        self.body_font: str = (
            spec.body_font
            or ds.get("default_font")
            or self.DEFAULT_BODY_FONT
        )

        # 正文字号
        spec_body_size = parse_font_size(spec.body_font_size)
        self.body_size_pt: float = (
            spec_body_size
            or ds.get("default_font_size_pt")
            or self.DEFAULT_BODY_SIZE_PT
        )

        # 标题字体
        self.heading_font: str = spec.heading_font or self.DEFAULT_HEADING_FONT

        # 标题字号
        spec_heading_size = parse_font_size(spec.heading_font_size)
        self.h1_size_pt: float = spec_heading_size or self.DEFAULT_H1_SIZE_PT
        self.h2_size_pt: float = self.DEFAULT_H2_SIZE_PT

        # 行距
        self.line_spacing_pt: Optional[float] = None
        self.line_spacing_multiple: float = self.DEFAULT_LINE_SPACING_MULTIPLE
        if spec.line_spacing:
            import re
            # 解析 "28磅" 格式
            pt_match = re.search(r'(\d+\.?\d*)\s*磅', spec.line_spacing)
            if pt_match:
                self.line_spacing_pt = float(pt_match.group(1))
            # 解析 "1.5倍" 格式
            mul_match = re.search(r'(\d+\.?\d*)\s*倍', spec.line_spacing)
            if mul_match:
                self.line_spacing_multiple = float(mul_match.group(1))

        # 页边距（优先级2: docx 文件的值 或 优先级3: 默认值）
        margins = ds.get("margins", {})
        self.margin_top = margins.get("top") or Cm(2.54)
        self.margin_bottom = margins.get("bottom") or Cm(2.54)
        self.margin_left = margins.get("left") or Cm(3.17)
        self.margin_right = margins.get("right") or Cm(3.17)

        logger.info(
            f"样式合并完成: 正文={self.body_font} {self.body_size_pt}pt, "
            f"标题={self.heading_font} {self.h1_size_pt}pt"
        )


# ============================================================
# 五、WordGenerator 核心生成器
# ============================================================

class WordGenerator:
    """
    投标书 Word 文档生成器。
    按照从招标文件格式章节提取的目录结构，动态生成投标书草稿。
    """

    def __init__(self, styles: MergedStyles):
        self.styles = styles

    # ----------------------------------------------------------
    # 5.1 公共辅助：设置字体（兼容中文东亚字体）
    # ----------------------------------------------------------

    def _set_run_font(self, run, font_name: str, size_pt: float, bold: bool = False):
        """为 Run 设置字体，兼容中文东亚字体设置"""
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.name = font_name
        # 设置东亚字体（确保中文正确渲染）
        r_element = run._element
        r_pr = r_element.get_or_add_rPr()
        r_fonts = r_pr.find(qn('w:rFonts'))
        if r_fonts is None:
            r_fonts = r_pr.makeelement(qn('w:rFonts'), {})
            r_pr.insert(0, r_fonts)
        r_fonts.set(qn('w:eastAsia'), font_name)

    def _set_paragraph_spacing(self, paragraph):
        """为段落设置行距"""
        pf = paragraph.paragraph_format
        if self.styles.line_spacing_pt:
            pf.line_spacing = Pt(self.styles.line_spacing_pt)
        else:
            pf.line_spacing = self.styles.line_spacing_multiple

    # ----------------------------------------------------------
    # 5.2 公共辅助：添加段落、标题、表格
    # ----------------------------------------------------------

    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """添加标题段落"""
        para = doc.add_paragraph()
        run = para.add_run(text)
        size = self.styles.h1_size_pt if level == 1 else self.styles.h2_size_pt
        self._set_run_font(run, self.styles.heading_font, size, bold=True)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 标题前后间距
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        return para

    def _add_body_paragraph(self, doc: Document, text: str, indent: bool = False):
        """添加正文段落"""
        para = doc.add_paragraph()
        run = para.add_run(text)
        self._set_run_font(run, self.styles.body_font, self.styles.body_size_pt)
        self._set_paragraph_spacing(para)
        if indent:
            para.paragraph_format.first_line_indent = Pt(self.styles.body_size_pt * 2)
        return para

    def _add_placeholder(self, doc: Document, title: str):
        """添加占位符提示段落"""
        para = doc.add_paragraph()
        run = para.add_run(f"[此处需手动补充: {title}]")
        self._set_run_font(run, self.styles.body_font, self.styles.body_size_pt)
        run.font.color.rgb = None  # 使用默认黑色
        run.font.italic = True
        # 使用灰色标识占位
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        self._set_paragraph_spacing(para)
        return para

    def _create_table(self, doc: Document, headers: list[str], rows: list[list[str]]):
        """创建标准格式表格"""
        from docx.shared import RGBColor

        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # 表头行
        header_row = table.rows[0]
        for i, header_text in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(header_text)
            self._set_run_font(run, self.styles.heading_font, self.styles.body_size_pt, bold=True)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 标题行底色
            shading = cell._element.get_or_add_tcPr()
            shading_elem = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): 'D9E2F3',
                qn('w:val'): 'clear',
            })
            shading.append(shading_elem)

        # 数据行
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(str(cell_text))
                self._set_run_font(run, self.styles.body_font, self.styles.body_size_pt)

        return table

    # ----------------------------------------------------------
    # 5.3 各章节内容构建器
    # ----------------------------------------------------------

    def _build_bid_letter(self, doc: Document, metadata: dict, analysis: dict) -> bool:
        """构建投标函章节"""
        timeline = metadata.get("timeline", {})
        project_name = timeline.get("project_name") or "[项目名称]"
        project_id = timeline.get("project_id_code") or "[项目编号]"

        self._add_body_paragraph(doc, f"致：[招标人名称]", indent=True)
        self._add_body_paragraph(doc, "")

        lines = [
            f"1. 我方已仔细研究了「{project_name}」（编号：{project_id}）的招标文件，"
            f"我方愿意按照招标文件的规定和要求，提供本投标文件。",
            "",
            "2. 我方承诺所提交的投标文件内容真实、有效，且在投标有效期内不予撤销。",
            "",
            "3. 如果我方中标，我方承诺按照招标文件的要求和投标文件的承诺忠实履行合同义务。",
        ]
        for line in lines:
            if line:
                self._add_body_paragraph(doc, line, indent=True)
            else:
                self._add_body_paragraph(doc, "")

        self._add_body_paragraph(doc, "")
        self._add_body_paragraph(doc, "投标人：[公司全称]（盖章）")
        self._add_body_paragraph(doc, "法定代表人或其委托代理人：（签字）")
        self._add_body_paragraph(doc, "日期：    年    月    日")
        return True

    def _build_authorization(self, doc: Document, metadata: dict, analysis: dict) -> bool:
        """构建法定代表人授权书章节"""
        self._add_body_paragraph(doc, "本授权书声明：", indent=True)
        self._add_body_paragraph(doc, "")
        self._add_body_paragraph(
            doc,
            "我________（法定代表人姓名）系________（投标人名称）的法定代表人，"
            "现授权________（被授权人姓名）为我方代理人，"
            "代表我方全权处理本项目投标活动的一切事宜。",
            indent=True
        )
        self._add_body_paragraph(doc, "")
        self._add_body_paragraph(doc, "法定代表人（签字）：")
        self._add_body_paragraph(doc, "被授权人（签字）：")
        self._add_body_paragraph(doc, "职务：")
        self._add_body_paragraph(doc, "身份证号码：")
        self._add_body_paragraph(doc, "联系电话：")
        self._add_body_paragraph(doc, "投标人（盖章）：")
        self._add_body_paragraph(doc, "日期：    年    月    日")
        return True

    def _build_qualification_section(self, doc: Document, metadata: dict, analysis: dict) -> bool:
        """构建资质响应/资格审查章节"""
        qual_analysis = analysis.get("qualifications_analysis", {})
        items = qual_analysis.get("items", [])
        qual_md = metadata.get("qualification", {})

        # 如果有资质分析条目
        if items:
            match_score = qual_analysis.get("match_score", "N/A")
            self._add_body_paragraph(
                doc, f"综合匹配度评估：{match_score} 分", indent=True
            )
            self._add_body_paragraph(doc, "")

            headers = ["序号", "招标要求", "响应状态", "说明"]
            rows = []
            for i, item in enumerate(items):
                rows.append([
                    str(i + 1),
                    item.get("requirement", ""),
                    item.get("status", ""),
                    item.get("reason", ""),
                ])
            self._create_table(doc, headers, rows)
            return True

        # 若无资质分析结果，但有 SQL 提取的资质门槛
        hard_quals = qual_md.get("mandatory_qualifications") or []
        if hard_quals:
            self._add_body_paragraph(doc, "招标文件强制性资质门槛如下，我方承诺完全满足：", indent=True)
            for q in hard_quals:
                self._add_body_paragraph(doc, f"• {q}", indent=True)
            return True

        return False

    def _build_cost_section(self, doc: Document, metadata: dict, analysis: dict) -> bool:
        """构建商务报价/报价清单章节"""
        cost = analysis.get("cost_analysis", {})
        items = cost.get("items", [])
        eng = metadata.get("engineering", {})
        eq_list = eng.get("main_equipment_list", [])

        # 优先使用成本分析导出的 BOM
        if items:
            total_cost = cost.get("total_cost", 0)
            budget_status = cost.get("budget_status", "")
            self._add_body_paragraph(doc, f"报价总金额：¥{total_cost:,.2f} 元", indent=True)
            if budget_status:
                self._add_body_paragraph(doc, f"预算状态：{budget_status}", indent=True)
            self._add_body_paragraph(doc, "")

            headers = ["序号", "设备/物资名称", "规格参数", "数量", "单位", "参考单价(元)", "小计(元)"]
            rows = []
            for i, item in enumerate(items):
                rows.append([
                    str(i + 1),
                    item.get("name", ""),
                    str(item.get("spec_requirement", ""))[:50],
                    str(item.get("qty", "")),
                    item.get("unit", ""),
                    f"{item.get('ref_price', 0):,.2f}",
                    f"{item.get('subtotal', 0):,.2f}",
                ])
            rows.append(["", "合计", "", "", "", "", f"¥{total_cost:,.2f}"])
            self._create_table(doc, headers, rows)
            return True

        # 若无成本分析，但有工程 BOM 设备清单
        if eq_list:
            self._add_body_paragraph(doc, "标书提取的设备及采购需求清单如下：", indent=True)
            headers = ["序号", "物品名称", "规格参数", "数量", "单位"]
            rows = []
            for i, eq in enumerate(eq_list):
                name = eq.get("item_name") or eq.get("name", "")
                specs = eq.get("specifications") or eq.get("spec_requirement", "")
                if isinstance(specs, list):
                    specs = "; ".join(specs)
                rows.append([
                    str(i + 1),
                    name,
                    str(specs)[:50],
                    str(eq.get("quantity", eq.get("qty", ""))),
                    eq.get("unit", "")
                ])
            self._create_table(doc, headers, rows)
            return True

        return False

    def _build_risk_deviation(self, doc: Document, metadata: dict, analysis: dict) -> bool:
        """构建偏离表/风险声明章节"""
        risks = analysis.get("risks_analysis", [])
        if risks:
            headers = ["序号", "风险类型", "偏离/风险描述", "严重程度", "应对措施"]
            rows = []
            for i, risk in enumerate(risks):
                rows.append([
                    str(i + 1),
                    risk.get("risk_type", ""),
                    risk.get("description", ""),
                    risk.get("severity", ""),
                    "按招标文件要求响应",
                ])
            self._create_table(doc, headers, rows)
            return True

        self._add_body_paragraph(doc, "我方完全响应招标文件的全部要求，无任何商务及技术偏离。", indent=True)
        return True

    def _build_technical_section(self, doc: Document, metadata: dict, analysis: dict) -> bool:
        """构建技术方案章节"""
        eng = metadata.get("engineering", {})
        equipment_list = eng.get("main_equipment_list", [])
        special_conditions = eng.get("special_working_conditions", [])
        standards = eng.get("mandatory_standards", [])

        if not equipment_list and not special_conditions and not standards:
            return False

        if equipment_list:
            self._add_body_paragraph(doc, "一、主要设备及材料清单", indent=False)
            self._add_body_paragraph(doc, "")
            headers = ["序号", "设备名称", "规格参数", "数量", "单位"]
            rows = []
            for i, eq in enumerate(equipment_list):
                name = eq.get("item_name") or eq.get("name", "")
                specs = eq.get("specifications") or eq.get("spec_requirement", "")
                if isinstance(specs, list):
                    specs = "; ".join(specs)
                rows.append([
                    str(i + 1),
                    name,
                    str(specs)[:80],
                    str(eq.get("quantity", eq.get("qty", ""))),
                    eq.get("unit", ""),
                ])
            self._create_table(doc, headers, rows)

        if special_conditions:
            self._add_body_paragraph(doc, "")
            self._add_body_paragraph(doc, "二、特殊工况及应对措施", indent=False)
            for condition in special_conditions:
                text = condition.get("description") if isinstance(condition, dict) else str(condition)
                self._add_body_paragraph(doc, f"• {text}", indent=True)

        if standards:
            self._add_body_paragraph(doc, "")
            self._add_body_paragraph(doc, "三、执行技术标准", indent=False)
            for std in standards:
                self._add_body_paragraph(doc, f"• {std}", indent=True)

        return True

    def _build_service_section(self, doc: Document, metadata: dict, analysis: dict) -> bool:
        """构建售后服务承诺章节"""
        eval_data = metadata.get("evaluation", {})
        service_reqs = eval_data.get("hard_service_requirements", {})

        if not service_reqs:
            return False

        self._add_body_paragraph(doc, "我方承诺严格履行以下售后服务条款：", indent=True)
        self._add_body_paragraph(doc, "")

        headers = ["序号", "服务项目", "承诺内容"]
        rows = []
        for i, (key, value) in enumerate(service_reqs.items()):
            rows.append([str(i + 1), str(key), str(value)])
        self._create_table(doc, headers, rows)
        return True

    def _build_personnel_section(self, doc: Document, metadata: dict, analysis: dict) -> bool:
        """构建人员配备章节"""
        qual = metadata.get("qualification", {})
        personnel = qual.get("personnel_requirements", [])

        if not personnel:
            return False

        self._add_body_paragraph(doc, "拟派项目团队人员要求如下：", indent=True)
        self._add_body_paragraph(doc, "")

        headers = ["序号", "岗位/角色", "要求说明", "拟派人员", "资质证书"]
        rows = []
        for i, person in enumerate(personnel):
            if isinstance(person, dict):
                rows.append([
                    str(i + 1),
                    person.get("role", person.get("position", "核心人员")),
                    person.get("requirement", person.get("description", str(person))),
                    "[按招标文件填报]",
                    "[按招标文件填报]",
                ])
            else:
                rows.append([str(i + 1), "核心人员", str(person), "[按招标文件填报]", "[按招标文件填报]"])
        self._create_table(doc, headers, rows)
        return True

    def _build_performance_section(self, doc: Document, metadata: dict, analysis: dict) -> bool:
        """构建业绩清单章节"""
        qual = metadata.get("qualification", {})
        performance = qual.get("performance_requirements", [])

        if not performance:
            return False

        self._add_body_paragraph(doc, "招标文件要求的历史同类业绩条件：", indent=True)
        for perf in performance:
            text = perf.get("requirement") if isinstance(perf, dict) else str(perf)
            self._add_body_paragraph(doc, f"• {text}", indent=True)

        self._add_body_paragraph(doc, "")
        self._add_body_paragraph(doc, "我方类似项目业绩格式表：", indent=True)

        headers = ["序号", "项目名称", "合同金额", "完工时间", "业主单位"]
        rows = [
            ["1", "[填写历史项目名称]", "[填写金额]", "[填写时间]", "[填写单位]"],
            ["2", "[填写历史项目名称]", "[填写金额]", "[填写时间]", "[填写单位]"],
        ]
        self._create_table(doc, headers, rows)
        return True

    def _build_financial_section(self, doc: Document, metadata: dict, analysis: dict) -> bool:
        """构建财务相关章节"""
        fin = metadata.get("financial", {})
        if not fin:
            return False

        items_to_show = [
            ("投标保证金", fin.get("bid_bond")),
            ("履约保证金", fin.get("performance_bond")),
            ("质保金", fin.get("warranty_bond")),
            ("预付款比例", fin.get("advance_payment_ratio")),
        ]
        has_any = False
        for label, value in items_to_show:
            if value:
                if not has_any:
                    self._add_body_paragraph(doc, "招标文件财务控制要求：", indent=True)
                    has_any = True
                display_val = value if isinstance(value, str) else str(value)
                self._add_body_paragraph(doc, f"• {label}：{display_val}", indent=True)

        return has_any

    def _build_schedule_section(self, doc: Document, metadata: dict, analysis: dict) -> bool:
        """构建工期计划章节"""
        timeline = metadata.get("timeline", {})
        period_days = timeline.get("construction_period_days")
        period_desc = timeline.get("construction_period_description")

        if period_desc:
            self._add_body_paragraph(doc, f"工期要求：{period_desc}", indent=True)
            return True
        elif period_days:
            self._add_body_paragraph(doc, f"工期要求：{period_days} 个日历天", indent=True)
            return True

        return False

    def _build_safety_section(self, doc: Document, metadata: dict, analysis: dict) -> bool:
        """构建安全生产方案章节"""
        eng = metadata.get("engineering", {})
        safety = eng.get("safety_and_env_requirements", [])

        if not safety:
            return False

        self._add_body_paragraph(doc, "我方将严格遵守以下安全生产要求：", indent=True)
        self._add_body_paragraph(doc, "")
        for item in safety:
            text = item if isinstance(item, str) else str(item)
            self._add_body_paragraph(doc, f"• {text}", indent=True)
        return True

    # ----------------------------------------------------------
    # 5.4 mapping_hint 路由器
    # ----------------------------------------------------------

    BUILDER_MAP: dict[str, str] = {
        "bid_letter": "_build_bid_letter",
        "authorization": "_build_authorization",
        "qualification": "_build_qualification_section",
        "pricing": "_build_cost_section",
        "cost": "_build_cost_section",
        "technical": "_build_technical_section",
        "deviation": "_build_risk_deviation",
        "risk": "_build_risk_deviation",
        "service": "_build_service_section",
        "warranty": "_build_service_section",
        "personnel": "_build_personnel_section",
        "performance": "_build_performance_section",
        "financial": "_build_financial_section",
        "schedule": "_build_schedule_section",
        "safety": "_build_safety_section",
    }

    def _render_section(
        self,
        doc: Document,
        item: OutlineItem,
        metadata: dict,
        analysis: dict,
        chapter_results: Optional[dict] = None,
        level: int = 1,
    ):
        """
        渲染单个目录条目。
        优先通过 Executor 填空产出的 chapter_results 渲染，次之通过数据分析填充，无数据时展示 content_hint。
        """
        title_text = f"{item.number} {item.title}" if item.number else item.title
        self._add_heading(doc, title_text, level=level)

        hint = (item.mapping_hint or "").strip().lower()

        # 优先看 chapter_results 中是否有 Executor 填充的结果
        matched_res = None
        if chapter_results:
            for tid, cres in chapter_results.items():
                chint = (cres.get("mapping_hint") or "").strip().lower()
                ctitle = cres.get("chapter_title") or ""
                if (hint and chint == hint) or (item.title and ctitle and item.title in ctitle):
                    matched_res = cres
                    break

        has_rendered = False
        if matched_res:
            filled_content = matched_res.get("filled_content", "")
            table_rows = matched_res.get("table_rows", [])

            if filled_content and not filled_content.startswith("[此处按招标文件要求手动补充"):
                for line in filled_content.splitlines():
                    line_s = line.strip()
                    if line_s.startswith("#"):
                        self._add_heading(doc, line_s.lstrip("#").strip(), level=min(level + 1, 3))
                    elif line_s:
                        self._add_body_paragraph(doc, line_s, indent=True)
                    else:
                        self._add_body_paragraph(doc, "")
                has_rendered = True

            if table_rows and not has_rendered:
                if hint in ["qualification"] or "资格" in (item.title or ""):
                    headers = ["序号", "证书/资质名称", "级别", "有效期至", "持证公司"]
                    rows = []
                    for i, q in enumerate(table_rows):
                        rows.append([
                            str(i + 1),
                            q.get("name", ""),
                            q.get("level", ""),
                            q.get("expiry", ""),
                            q.get("company", "")
                        ])
                    self._create_table(doc, headers, rows)
                    has_rendered = True

        if not has_rendered:
            builder_name = self.BUILDER_MAP.get(hint)
            if builder_name:
                builder_fn = getattr(self, builder_name, None)
                if builder_fn:
                    try:
                        has_rendered = builder_fn(doc, metadata, analysis)
                    except Exception as e:
                        logger.error(f"构建章节 [{item.title}] 异常: {e}")

        if not has_rendered:
            if item.content_hint:
                self._add_body_paragraph(doc, f"【招标说明/格式要求】: {item.content_hint}", indent=True)
            else:
                self._add_placeholder(doc, item.title)

        # 递归渲染子条目
        for sub_item in item.sub_items:
            self._render_section(doc, sub_item, metadata, analysis, chapter_results=chapter_results, level=2)

    # ----------------------------------------------------------
    # 5.5 主入口
    # ----------------------------------------------------------

    def generate_bidding_draft(
        self,
        outline: BidDocOutline,
        metadata: dict,
        analysis: dict,
        chapter_results: Optional[dict] = None,
    ) -> bytes:
        """
        主入口：根据目录结构生成投标书 Word 文档。

        :param outline: 从格式章节提取的目录骨架 + 排版规范
        :param metadata: 5大元数据字典 (timeline, financial, engineering, qualification, evaluation)
        :param analysis: 分析结果字典 (qualifications_analysis, risks_analysis, cost_analysis, company_quals)
        :param chapter_results: Executor 填空节点返回的各章节填空与数据结果
        :return: .docx 文件的字节流
        """
        doc = Document()

        # 应用页面设置（不设置页眉页脚）
        section = doc.sections[0]
        section.top_margin = self.styles.margin_top
        section.bottom_margin = self.styles.margin_bottom
        section.left_margin = self.styles.margin_left
        section.right_margin = self.styles.margin_right

        # 封面标题
        timeline = metadata.get("timeline", {})
        project_name = timeline.get("project_name") or "[项目名称]"

        # 封面
        for _ in range(6):
            doc.add_paragraph("")  # 留白

        cover_para = doc.add_paragraph()
        cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_run = cover_para.add_run(project_name)
        self._set_run_font(cover_run, self.styles.heading_font, 22.0, bold=True)

        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run("投 标 文 件")
        self._set_run_font(subtitle_run, self.styles.heading_font, 26.0, bold=True)

        # 投标方信息
        for _ in range(4):
            doc.add_paragraph("")

        info_lines = [
            "投标人：[公司全称]",
            "日  期：    年    月    日",
        ]
        for line in info_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line)
            self._set_run_font(r, self.styles.body_font, self.styles.h2_size_pt)

        # 分页 → 进入正文
        doc.add_page_break()

        # 按目录逐章生成
        if outline.outline:
            logger.info(f"按照甲方目录生成投标书，共 {len(outline.outline)} 个顶级章节")
            for item in outline.outline:
                self._render_section(doc, item, metadata, analysis, chapter_results=chapter_results, level=1)
        else:
            # 降级：无目录时使用默认结构
            logger.warning("未提取到格式章节目录，使用默认投标书模板")
            self._generate_default_structure(doc, metadata, analysis, chapter_results=chapter_results)

        # 导出为字节流
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        logger.info(f"投标书草稿生成完成，文档大小: {file_stream.getbuffer().nbytes} 字节")
        return file_stream.getvalue()

    # ----------------------------------------------------------
    # 5.6 降级：默认模板结构（RAG 未命中格式章节时）
    # ----------------------------------------------------------

    def _generate_default_structure(self, doc: Document, metadata: dict, analysis: dict, chapter_results: Optional[dict] = None):
        """当无法提取甲方目录时，使用通用默认章节结构"""
        default_sections = [
            OutlineItem(number="一", title="投标函", mapping_hint="bid_letter"),
            OutlineItem(number="二", title="法定代表人授权书", mapping_hint="authorization"),
            OutlineItem(number="三", title="资格审查资料", mapping_hint="qualification"),
            OutlineItem(number="四", title="商务报价", mapping_hint="cost"),
            OutlineItem(number="五", title="技术方案", mapping_hint="technical"),
            OutlineItem(number="六", title="偏离表", mapping_hint="deviation"),
            OutlineItem(number="七", title="拟投入人员", mapping_hint="personnel"),
            OutlineItem(number="八", title="业绩清单", mapping_hint="performance"),
            OutlineItem(number="九", title="售后服务承诺", mapping_hint="service"),
        ]
        for item in default_sections:
            self._render_section(doc, item, metadata, analysis, chapter_results=chapter_results, level=1)
