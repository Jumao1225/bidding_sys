"""
招投标全量数据库直查工具集 (bid_db_tools.py)

功能：
提供供 LangChain / LangGraph Agent 调用的标准 @tool 查库工具。
涵盖系统全量数据库表结构 (company_profiles, company_qualifications, market_price_references, 
timeline_metadata, financial_metadata, qualification_metadata, engineering_metadata, 
evaluation_metadata, cost_estimates, risk_items, qualification_matches, documents, projects, users)，
实现零幻觉的绝对 SQL 直查，直接返回数据库存储的真实企业档案、资质证书、项目元数据、财务报价及合规要求。

遵循项目规范：
1. 全面使用中文注释与 Docstrings；
2. 全面使用 Type Hints 类型提示；
3. 使用 loguru 进行超详细调试日志记录；
4. 防御性编程与尽早返回 (Early Return)。
"""

import os
import re
import json
from contextvars import ContextVar
from typing import Dict, Any, List, Optional, Tuple, Set
from loguru import logger
from sqlalchemy.orm import Session
from langchain_core.tools import tool

from app.db.session import SessionLocal
from app.db.models.business import CompanyProfileModel, CompanyQualification, MarketPriceReference
from app.db.models.metadata import (
    QualificationMetadata,
    FinancialMetadata,
    TimelineMetadata,
    EngineeringMetadata,
    EvaluationMetadata,
)
from app.db.models.project import Project as ProjectModel, Document as DocumentModel
from app.db.models.ai_analysis import QualificationMatch, RiskItem, CostEstimate
from app.db.models.user import User as UserModel
from app.utils.rmb_formatter import number_to_chinese_rmb

# 运行时上下文变量：标书撰写时指定使用的企业档案 ID
# 由 BidFillerAgent 入口设置，Worker 子进程/线程继承
current_profile_id: ContextVar[Optional[str]] = ContextVar("current_profile_id", default=None)


def resolve_company_profile(
    db: Session,
    profile_id: Optional[str] = None,
) -> Optional[CompanyProfileModel]:
    """按指定主体、默认主体、最早创建主体的顺序解析企业档案。

    统一所有 Agent 读取企业档案的兜底规则，避免直接使用无序的
    ``query(...).first()`` 导致不同主体之间发生串用。
    """
    if profile_id:
        selected_profile = db.query(CompanyProfileModel).filter(
            CompanyProfileModel.id == profile_id
        ).first()
        if selected_profile:
            logger.info(
                "🛠️ [企业档案解析] 使用指定主体: id={}, name='{}', company='{}'",
                profile_id,
                selected_profile.profile_name,
                selected_profile.company_name,
            )
            return selected_profile
        logger.warning(
            "🛠️ [企业档案解析] 指定主体不存在，准备回退默认主体: profile_id={}",
            profile_id,
        )

    default_profile = db.query(CompanyProfileModel).filter(
        CompanyProfileModel.is_default == True
    ).order_by(
        CompanyProfileModel.created_at.asc(),
        CompanyProfileModel.id.asc(),
    ).first()
    if default_profile:
        logger.info(
            "🛠️ [企业档案解析] 使用默认主体: id={}, name='{}', company='{}'",
            default_profile.id,
            default_profile.profile_name,
            default_profile.company_name,
        )
        return default_profile

    first_profile = db.query(CompanyProfileModel).order_by(
        CompanyProfileModel.created_at.asc(),
        CompanyProfileModel.id.asc(),
    ).first()
    if first_profile:
        logger.warning(
            "🛠️ [企业档案解析] 没有默认主体，使用最早创建主体: id={}, name='{}', company='{}'",
            first_profile.id,
            first_profile.profile_name,
            first_profile.company_name,
        )
    else:
        logger.warning("🛠️ [企业档案解析] 数据库中没有任何企业档案")
    return first_profile


# ============================================================
# 全量数据库字段别名同义词映射字典 (Comprehensive Alias Mapping)
# ============================================================

ALIAS_MAP = {
    # 1. 企业基础工商档案 (company_profiles)
    "company_name": ["company_name", "公司名称", "投标人名称", "单位名称", "投标人全称", "投标人", "单位", "响应人名称", "投标单位", "企业名称"],
    "legal_representative": ["legal_representative", "法定代表人", "法人", "法人代表", "单位负责人", "法定代表人（签字）", "法定代表人或其委托代理人", "法定代表人姓名"],
    "authorized_delegate": ["authorized_delegate", "授权代表", "被授权人", "授权委托人", "签字代表", "委托代理人", "代表（签字）", "投标单位代表姓名", "代表姓名", "受托人", "代理人", "项目代表"],
    "credit_code": ["credit_code", "统一社会信用代码", "纳税人识别号", "税号", "营业执照注册号", "信用代码", "社会信用代码", "机构代码"],
    "registered_address": ["registered_address", "注册地址", "公司地址", "住所", "单位地址", "地址", "注册地", "经营地址", "详细地址", "通讯地址"],
    "contact_phone": ["contact_phone", "联系电话", "手机号", "电话", "联系方式", "固定电话", "联系人电话"],
    "email": ["email", "电子邮箱", "邮箱", "email", "e-mail", "联系邮箱", "企业邮箱"],
    "bank_name": ["bank_name", "开户银行", "基本户开户行", "结算银行", "开户行", "基本开户银行", "存款银行", "基本户银行"],
    "bank_account": ["bank_account", "银行账号", "基本户账号", "结算账号", "账号", "基本户银行账号", "银行卡号", "开户账号"],

    # 2. 项目与时间节点元数据 (timeline_metadata)
    "project_name": ["project_name", "项目名称", "招标项目名称", "采购项目名称", "工程名称", "项目全称", "招标工程名称"],
    "project_code": ["project_code", "project_id_code", "项目编号", "招标编号", "包件号", "标段编号", "采购编号", "项目代码", "招标代码"],
    "tender_segment": ["tender_segment", "标段名称", "包件名称", "标段", "包件", "分包名称"],
    "bid_deadline": ["bid_deadline", "投标截止时间", "开标时间", "递交截止时间", "截标时间", "开标日期"],
    "bid_validity_days": ["bid_validity_days", "投标有效期", "有效期", "投标有效期天数"],
    "construction_period": ["construction_period", "construction_period_description", "工期", "建设周期", "交货期", "服务期", "承诺工期", "交货期限", "完成时间", "服务期限"],

    # 3. 财务与价格元数据 (financial_metadata & cost_estimates)
    "bid_price_numeric": ["bid_price_numeric", "total_price_numeric", "投标总价", "投标金额", "投标小写", "报价小写", "总报价", "投标总金额", "最终报价"],
    "bid_price_chinese": ["bid_price_chinese", "total_price_chinese", "投标大写", "报价大写", "汉字大写", "大写金额", "投标总金额大写"],
    "bid_bond": ["bid_bond", "投标保证金", "保证金", "投标担保", "保证金金额"],
    "performance_bond": ["performance_bond", "履约保证金", "履约担保", "履约保证"],
    "warranty_bond": ["warranty_bond", "质保金", "质量保证金", "保修金", "质保保证金"],
    "payment_milestones": ["payment_milestones", "付款方式", "付款节点", "支付方式", "结算方式", "付款条款", "支付节点"],
    "tax_rate_requirement": ["tax_rate_requirement", "税率", "增值税率", "税率要求"],
    "contract_price_type": ["contract_price_type", "计价方式", "合同计价形式", "承包方式", "价格类型"],

    # 4. 资格要求与核心团队 (qualification_metadata)
    "mandatory_qualifications": ["mandatory_qualifications", "资质门槛", "企业资质要求", "必备资质"],
    "personnel_requirements": ["personnel_requirements", "核心人员", "项目经理", "技术负责人", "项目团队", "人员配置"],
    "performance_requirements": ["performance_requirements", "历史业绩", "类似业绩", "同类业绩", "业绩门槛"],
    "invalid_bid_clauses": ["invalid_bid_clauses", "废标条款", "否决投标条款", "无效投标条款", "一票否决"],

    # 5. 工程与技术规范 (engineering_metadata)
    "main_equipment_list": ["main_equipment_list", "设备清单", "主要设备", "材料清单"],
    "mandatory_standards": ["mandatory_standards", "技术标准", "强制性标准", "规范要求"],

    # 6. 评审标准与办法 (evaluation_metadata)
    "evaluation_method": ["evaluation_method", "评标方法", "评标办法", "评审标准", "评标标准"],

    # 7. BOM 成本与分项报价模型字段 (CostEstimate)
    "item_name": ["item_name", "标的物名称", "标的物", "设备名称", "货物名称", "品名", "产品名称", "材料名称", "名称"],
    "spec": ["spec", "规格型号", "规格", "型号", "技术参数", "参数要求", "技术规格"],
    "brand": ["brand", "品牌", "制造品牌", "商标"],
    "manufacturer": ["manufacturer", "生产厂家", "制造厂家", "生产企业", "制造商", "厂家", "产地"],
    "unit": ["unit", "单位", "计量单位"],
    "quantity": ["quantity", "数量", "工程量", "规模"],
    "unit_price": ["unit_price", "单价", "综合单价", "单价(元)", "单价（元）", "投标单价"],
    "calculated_total": ["calculated_total", "合价", "总价", "小计", "合价(元)", "合价（元）", "总价(元)", "总价（元）", "金额"],
    "remark": ["remark", "备注", "说明", "备注说明"],
}


def _match_alias_key(user_key: str) -> str:
    """按最长有效别名归一化字段，避免短别名抢占更具体的标签。"""
    key_clean = str(user_key or "").lower().strip()
    normalized_key = re.sub(r"[\s:：_＿（）()\[\]［］]", "", key_clean)
    if not normalized_key:
        return key_clean

    matched_key = ""
    matched_alias_length = 0
    for std_key, aliases in ALIAS_MAP.items():
        for alias in aliases:
            normalized_alias = re.sub(
                r"[\s:：_＿（）()\[\]［］]",
                "",
                str(alias).lower(),
            )
            if not normalized_alias:
                continue
            # 仅允许“完整别名包含于实际标签”，避免“委托人”这类短标签
            # 被扩展别名“委托代理人”反向命中，造成企业字段误填。
            if normalized_alias in normalized_key:
                # 具体别名优先于“单位”“代表”等短别名，保证复合标签归入正确字段。
                alias_length = len(normalized_alias)
                if alias_length > matched_alias_length:
                    matched_key = std_key
                    matched_alias_length = alias_length

    return matched_key or key_clean


# ============================================================
# LangChain / LangGraph Agent 可调用的工具集合 (@tool)
# ============================================================

@tool
def query_company_profile_tool(field_key: str) -> str:
    """
    [数据库直查工具] 查询本公司的企业基础工商档案 (company_profiles 表)。
    涵盖全量字段：公司名称、法定代表人、授权代表、统一社会信用代码、注册地址、联系电话、电子邮箱、开户银行、银行账号等。

    :param field_key: 欲查询的字段名或中文引导词 (如 'company_name', '统一社会信用代码', '开户银行')
    :return: 数据库中存储的真实原值字符串
    """
    logger.info(f"🛠️ [DB Tool] query_company_profile_tool 被调用, 字段请求: '{field_key}'")
    std_key = _match_alias_key(field_key)
    logger.debug(f"🛠️ [DB Tool] 别名归一化对齐结果: '{field_key}' -> '{std_key}'")

    db: Session = SessionLocal()
    try:
        # 按运行时主体上下文解析，禁止使用无序查询结果作为主体依据。
        profile = resolve_company_profile(db, current_profile_id.get())

        if profile:
            val = getattr(profile, std_key, None)
            if val and str(val).strip():
                logger.info(f"🛠️ [DB Tool] 成功从数据库 (company_profiles.{std_key}) 查询到真实数据: '{val}'")
                return str(val).strip()

        logger.warning(f"🛠️ [DB Tool] 字段 '{field_key}' 在企业档案数据库中尚未录入")
        return f"[待补充: {field_key}]"

    except Exception as e:
        logger.exception(f"🛠️ [DB Tool] query_company_profile_tool 执行异常: {str(e)}")
        return f"[查询异常: {str(e)}]"
    finally:
        db.close()


def resolve_qualification_image_path(file_url: Optional[str]) -> Tuple[Optional[str], bool]:
    """
    解析资质证书 file_url，转换为本地磁盘绝对路径并校验是否存在。

    :param file_url: 数据库中存放的 file_url (如 '/uploads/qualifications/xxxx.png' 或 'xxxx.png')
    :return: (local_abs_path, exists)
    """
    if not file_url or not str(file_url).strip():
        return None, False

    clean_url = str(file_url).strip()
    # bid_db_tools.py 位于 backend/app/agents/tools/ -> 向上 4 层到达 backend 目录
    backend_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
    qual_dir = os.path.join(backend_dir, "uploads", "qualifications")

    filename = os.path.basename(clean_url)
    local_path = os.path.join(qual_dir, filename)

    if os.path.exists(local_path):
        return os.path.abspath(local_path), True

    if os.path.exists(clean_url):
        return os.path.abspath(clean_url), True

    return os.path.abspath(local_path), False


@tool
def query_company_qualification_tool(cert_keyword: str = "", tenant_id: Optional[str] = None) -> str:
    """
    [数据库直查工具] 查询企业资质与证书数据库 (company_qualifications 表)。
    支持 Agent 自主检索：可输入关键字搜索，若关键字为空则自动列出所有有效资质证书及其磁盘图片路径。

    :param cert_keyword: 证书名称关键字（可选，如 '营业执照', '电力工程', 'ISO9001' 等）
    :param tenant_id: 租户 ID（可选）。有当前会话租户时强制使用会话租户，防止 Agent 参数越权。
    :return: 匹配的资质证书记录列表（包含资质名称、等级、到期日、图片物理路径 local_image_path）
    """
    logger.info(f"🛠️ [DB Tool] query_company_qualification_tool 被调用, 关键字: '{cert_keyword}'")
    db: Session = SessionLocal()
    try:
        kw = cert_keyword.strip() if cert_keyword else ""

        from app.core.context import current_tenant_id
        context_tenant = current_tenant_id.get()
        effective_tenant = context_tenant or tenant_id or "default-tenant"

        base_query = db.query(CompanyQualification).filter(
            CompanyQualification.tenant_id == effective_tenant
        )
        if kw:
            quals = base_query.filter(
                CompanyQualification.name.ilike(f"%{kw}%")
            ).order_by(CompanyQualification.created_at.desc()).all()
        else:
            quals = base_query.order_by(CompanyQualification.created_at.desc()).all()

        # 若精准匹配未查到，自动回退全表列表供 Agent 自主挑选
        if not quals and kw:
            logger.info(f"🛠️ [DB Tool] 关键字 '{kw}' 未直接匹配，回退查询租户全量资质供 Agent 自主匹配...")
            quals = base_query.order_by(CompanyQualification.created_at.desc()).all()

        if quals:
            res_list = []
            for q in quals:
                exp_str = q.expiry_date.strftime("%Y-%m-%d") if q.expiry_date else "长期有效"
                lvl_str = q.level or "通用"
                img_path, img_exists = resolve_qualification_image_path(q.file_url)
                status_str = f"【图片路径: {img_path}】(物理存在: {img_exists})" if img_path else "【未上传图片文件】"

                res_list.append(
                    f"• 证书名称: {q.name} | 等级/范围: {lvl_str} | 有效期: {exp_str} | {status_str}"
                )

            result_str = "\n".join(res_list)
            logger.info(f"🛠️ [DB Tool] 成功检索到 {len(quals)} 条资质证书记录 (tenant_id={effective_tenant})")
            return f"查询到以下资质证书记录：\n{result_str}"

        logger.warning(f"🛠️ [DB Tool] 数据库中未找到符合条件的资质证书记录")
        return f"[待手动补充资质证书: {cert_keyword or '资质文件'}]"

    except Exception as e:
        logger.exception(f"🛠️ [DB Tool] query_company_qualification_tool 执行异常: {str(e)}")
        return f"[查询资质异常: {str(e)}]"
    finally:
        db.close()


def insert_paragraph_after(p, doc=None):
    """在指定段落 p 下方插入一个新的空段落"""
    try:
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        from docx.text.paragraph import Paragraph
        new_p_elem = parse_xml(r'<w:p %s/>' % nsdecls('w'))
        p._element.addnext(new_p_elem)
        return Paragraph(new_p_elem, p._parent)
    except Exception:
        return doc.add_paragraph() if doc else p


import threading

_DOC_FILE_LOCKS: Dict[str, threading.Lock] = {}
_DOC_LOCK_MUTEX = threading.Lock()

def get_doc_file_lock(docx_path: str) -> threading.Lock:
    """获取指定 Word 文档物理路径的线程排队锁 (防止多 Worker 线程并发写盘冲突)"""
    norm_path = os.path.normpath(docx_path).lower()
    with _DOC_LOCK_MUTEX:
        if norm_path not in _DOC_FILE_LOCKS:
            _DOC_FILE_LOCKS[norm_path] = threading.Lock()
        return _DOC_FILE_LOCKS[norm_path]


def _safe_save_doc(doc, docx_path: str, max_retries: int = 10) -> bool:
    """
    安全保存 Word 文档，具备 Windows 文件锁 PermissionError 中转保存、线程排队锁与 10 次长退避重试保护机制
    """
    file_lock = get_doc_file_lock(docx_path)
    with file_lock:
        try:
            doc.save(docx_path)
            return True
        except PermissionError:
            logger.warning(f"⚠️ [File Lock] 保存文件 {os.path.basename(docx_path)} 被占用 (PermissionError)，主动触发孤儿进程清理与长退避重试...")
            import tempfile
            import shutil
            import time
            import random
            from app.services.office_cli_service import office_cli_service

            # 尝试强 kill 残留的 officecli / soffice 孤儿进程以释放文件独占锁
            try:
                office_cli_service.kill_lingering_processes()
            except Exception as kill_err:
                logger.debug(f"尝试清理孤儿进程提示: {kill_err}")

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
                temp_out = tf.name

            try:
                doc.save(temp_out)
                for attempt in range(1, max_retries + 1):
                    try:
                        shutil.copyfile(temp_out, docx_path)
                        logger.info(f"✅ 成功通过临时文件中转写入目标 Word: {os.path.basename(docx_path)} (第 {attempt} 次重试成功)")
                        return True
                    except PermissionError:
                        if attempt % 3 == 0:
                            try:
                                office_cli_service.kill_lingering_processes()
                            except Exception:
                                pass
                        sleep_time = 0.4 * attempt + random.uniform(0.1, 0.3)
                        time.sleep(sleep_time)
                logger.error(f"❌ 重试 {max_retries} 次仍无法覆盖目标文件 {docx_path} (文件正被 OfficeCLI/系统进程锁定)")
                return False
            finally:
                if os.path.exists(temp_out):
                    try:
                        os.remove(temp_out)
                    except Exception:
                        pass
        except Exception as e:
            logger.error(f"❌ 保存 Word 文档产生未预期异常: {e}")
            return False


def check_qual_matches_paragraph(item_name: str, text: str) -> bool:
    """精准判定资质名称 item_name 是否与条款段落文本 text 匹配（杜绝因公共词导致的误配过配）"""
    name = item_name.strip()
    if not name or not text:
        return False
    if "营业执照" in name or "执照" in name:
        return "营业执照" in text or "执照" in text or "法人" in text

    # 动态剥离行政级别与通用公文修饰后缀，提取核心业务特征词进行比对
    noise_patterns = [
        "中华人民共和国", "国家", "省", "市", "施工总承包", "专业承包", "总承包",
        "特级", "一级", "二级", "三级", "甲级", "乙级", "丙级",
        "资质证书", "证书", "资质", "执照", "许可证"
    ]
    core = name
    for pat in noise_patterns:
        core = core.replace(pat, "")
    core = core.strip()

    if len(core) >= 2 and core in text:
        return True
    return False


def _is_explicit_qualification_evidence_request(text: str) -> bool:
    """判断段落是否明确要求提供/提交资质证明材料，而非仅提到资质关键词。"""
    normalized = re.sub(r"\s+", "", text or "")
    if not normalized:
        return False

    # 这些短语本身就是明确的插图/附件意图。
    explicit_markers = [
        "请插入图片", "插入图片", "附资质图片", "附证书图片",
        "图片见附件", "证明材料见附件", "证明文件见附件",
    ]
    if any(marker in normalized for marker in explicit_markers):
        return True

    evidence_terms = r"(?:营业执照|执照|资质|资质证书|证书|许可证|体系认证|证明文件|证明材料|身份证明|扫描件|复印件|原件)"
    request_actions = r"(?:提供|提交|附上|附具|随附|上传|递交|报送|取得|持有|具有|应附|须附|需附|应提供|须提供|需提供|应提交|须提交|需提交|应上传|须上传|需上传)"

    # 只有“动作 + 证明材料”同时出现时，才允许纯文本条款触发兜底插图。
    if re.search(request_actions, normalized) and re.search(evidence_terms, normalized):
        return True

    # 兼容“营业执照等证明文件”“资质证书扫描件”等结构化材料描述。
    if re.search(rf"{evidence_terms}.{{0,10}}(?:证明文件|证明材料|扫描件|复印件|原件)", normalized):
        return True
    if re.search(rf"(?:证明文件|证明材料|扫描件|复印件|原件).{{0,10}}{evidence_terms}", normalized):
        return True

    return False


def auto_embed_qualification_images_in_docx(docx_path: str, tenant_id: Optional[str] = None) -> int:
    """
    自动扫描 Word 文档中的所有资质证书占位符与资质文件章节/表格条款，
    从数据库 company_qualifications 表中寻找匹配的证书记录，
    自动在占位符或纯文本资质条款下方原位嵌入 backend/uploads/qualifications 真实证书图片。

    :param docx_path: Word 文档 (.docx) 物理路径
    :param tenant_id: 可选租户 ID
    :return: 成功嵌入的资质图片数量
    """
    if not docx_path or not os.path.exists(docx_path):
        return 0

    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from sqlalchemy.orm import Session
        from app.db.session import SessionLocal
        from app.db.models.business import CompanyQualification

        from app.core.context import current_tenant_id
        context_tenant = current_tenant_id.get()
        effective_tenant = (
            tenant_id
            if tenant_id and tenant_id != "default-tenant"
            else (context_tenant or tenant_id or "default-tenant")
        )

        db: Session = SessionLocal()
        try:
            query = db.query(CompanyQualification)
            query = query.filter(CompanyQualification.tenant_id == effective_tenant)
            quals = query.all()

            valid_quals = []
            for q in quals:
                img_path = None
                file_url_val = getattr(q, 'file_url', None) or getattr(q, 'file_path', None)
                if file_url_val:
                    p_res, exists = resolve_qualification_image_path(file_url_val)
                    if exists:
                        img_path = p_res

                if img_path:
                    valid_quals.append({
                        "name": getattr(q, 'name', '') or "企业资质证书",
                        "code": getattr(q, 'code', '') or "",
                        "level": getattr(q, 'level', '') or "通用",
                        "image_path": img_path,
                        "used_paragraphs": set(),
                        "used": False
                    })
        finally:
            db.close()

        if not valid_quals:
            logger.info("数据库中未查询到物理磁盘存在的有效企业资质证书图片，跳过自动图嵌入")
            return 0

        doc = Document(docx_path)
        embedded_count = 0
        embedded_image_paths_in_doc = set()

        # 预扫描文档中已经存在的资质图注文本，预判并标记已插入的物理图片
        for p_item in doc.paragraphs:
            p_item_text = p_item.text.strip() if p_item.text else ""
            if p_item_text.startswith("图："):
                for item in valid_quals:
                    name_core = item["name"].replace("证书", "").strip()
                    if name_core and name_core in p_item_text:
                        embedded_image_paths_in_doc.add(item["image_path"])

        def paragraph_or_next_has_inserted_qual_image(p) -> bool:
            if not p:
                return False
            try:
                curr = p
                for _ in range(3):
                    next_elem = curr._element.getnext()
                    if next_elem is None or not next_elem.tag.endswith('p'):
                        break
                    from docx.text.paragraph import Paragraph
                    next_p = Paragraph(next_elem, p._parent)
                    next_text = (next_p.text or "").strip()
                    if next_text.startswith("图：") and any(k in next_text for k in ["资质", "证书", "执照", "许可证"]):
                        return True
                    if next_p._element.xpath('.//w:drawing | .//w:pict'):
                        return True
                    curr = next_p
            except Exception:
                pass
            return False

        def find_best_qual(target_text: str) -> Optional[Dict[str, Any]]:
            best_item = None
            max_score = 0
            for item in valid_quals:
                score = 0
                q_name = item["name"]
                if q_name in target_text:
                    score += 10
                elif check_qual_matches_paragraph(q_name, target_text):
                    score += 5
                if score > max_score:
                    max_score = score
                    best_item = item
            return best_item

        def insert_paragraph_after(paragraph, doc_obj):
            new_p = doc_obj.add_paragraph()
            paragraph._p.addnext(new_p._p)
            return new_p

        def replace_paragraph_with_image(p, qual_item: Dict[str, Any]):
            p.text = ""
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(qual_item["image_path"], width=Inches(5.5))
            cap_run = p.add_run(f"\n图：{qual_item['name']}（等级/范围: {qual_item['level']}）")
            cap_run.font.size = Pt(10.5)
            cap_run.font.bold = True

        # 1. 扫描正文段落（增加目录识别与过滤，严禁在目录列表中插入图片）
        in_toc = False
        toc_item_count = 0
        in_qual_section = False
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if not text:
                continue

            # 目录区域判定
            if any(k in text for k in ["目录", "目 录", "文件目录"]):
                in_toc = True
                continue
            if in_toc:
                # 当在目录后再次遇到正文第一章（如 一、封面 / 第一章 / 【一】）时，目录结束
                if bool(re.match(r'^\s*(?:[一1][、\.．\s]|第一[章节部分篇]|【[一1]】)', text)) and toc_item_count > 3:
                    in_toc = False
                else:
                    toc_item_count += 1
                    continue

            qualification_section_markers = ["资格证明", "资质文件", "资格审查", "资质证明", "企业资质"]

            # 离开资质章节判定：遇到后续大章标题且非资质相关时重置 in_qual_section
            # 普通的“1.”、“8.”资格条款不是新章节，不能因此退出资格证明章节。
            is_new_chapter_title = bool(re.match(r'^\s*(?:[一二三四五六七八九十百]+[、\.．\s]|第[一二三四五六七八九十\d]+[章节部分篇]|【[一二三四五六七八九十\d]+】)', text)) or (p.style and p.style.name and p.style.name.startswith("Heading"))
            if is_new_chapter_title and not any(h in text for h in qualification_section_markers):
                in_qual_section = False

            if any(h in text for h in qualification_section_markers):
                in_qual_section = True

            # 场景 A: 包含显式占位符 [待手动补充资质证书: xxx] 或 [待...]
            if "[待" in text and ("资质" in text or "证书" in text or "执照" in text or "证明" in text):
                matched_q = find_best_qual(text)
                if matched_q and matched_q["image_path"] not in embedded_image_paths_in_doc:
                    replace_paragraph_with_image(p, matched_q)
                    embedded_image_paths_in_doc.add(matched_q["image_path"])
                    embedded_count += 1
                    logger.info(f"   🖼️ 已原位替换段落占位符 '{text[:30]}' -> 资质图片 {matched_q['name']}")
                continue

            # 场景 B: 纯文本资质要求条款。
            # 仅有“资质/证书/执照”等关键词不代表需要插图，必须同时具备明确的材料提供语义。
            if in_qual_section and _is_explicit_qualification_evidence_request(text):
                # 若该条款下方已存在插入好的资质图片或图注，坚决跳过，防重复插入
                if paragraph_or_next_has_inserted_qual_image(p):
                    logger.info(f"   ⏩ [跳过已存图片] 条款 '{text[:30]}' 下方已存在资质图片或图注")
                    continue

                matched_quals_for_p = [item for item in valid_quals if check_qual_matches_paragraph(item["name"], text)]
                if not matched_quals_for_p:
                    continue

                # 按物理图片路径进行去重与合并组装（支持一证多资质共享图片）
                grouped_by_img = {}
                for item in matched_quals_for_p:
                    img_p_key = item["image_path"]
                    if img_p_key not in grouped_by_img:
                        grouped_by_img[img_p_key] = []
                    grouped_by_img[img_p_key].append(item)

                # 限制单个段落下方最多只挑选 1 个最匹配的图片组落盘，防止连续砸入多张独立大图
                available_img_groups = [
                    (img_key, items) for img_key, items in grouped_by_img.items()
                    if img_key not in embedded_image_paths_in_doc
                ]
                if not available_img_groups:
                    continue

                img_p_key, items = available_img_groups[0]
                last_p = p

                # 若同一条款匹配到了同一张证书上的多个资质，合并渲染一张图片与综合图注
                if len(items) > 1:
                    merged_names = " / ".join(list(dict.fromkeys([it["name"] for it in items])))
                    merged_levels = " / ".join(list(dict.fromkeys([it["level"] for it in items if it["level"] and it["level"] != "通用"]))) or "通用"
                    composite_item = {
                        "name": f"综合资质证书（涵盖: {merged_names}）",
                        "level": merged_levels,
                        "image_path": img_p_key,
                        "used_paragraphs": set()
                    }
                    img_p = insert_paragraph_after(last_p, doc)
                    replace_paragraph_with_image(img_p, composite_item)
                    embedded_image_paths_in_doc.add(img_p_key)
                    for it in items:
                        it["used_paragraphs"].add(p)
                        it["used"] = True
                    embedded_count += 1
                    logger.info(f"   🖼️ 已在条款 '{text[:30]}' 下方插入一证多资质证书图片: {composite_item['name']}")
                else:
                    q_item = items[0]
                    img_p = insert_paragraph_after(last_p, doc)
                    replace_paragraph_with_image(img_p, q_item)
                    embedded_image_paths_in_doc.add(img_p_key)
                    q_item["used_paragraphs"].add(p)
                    q_item["used"] = True
                    embedded_count += 1
                    logger.info(f"   🖼️ 已在条款 '{text[:30]}' 下方自动附带插入资质图片: {q_item['name']}")

            # 2. 扫描表格单元格
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            text = p.text.strip()
                            if "[待" in text and ("资质" in text or "证书" in text or "执照" in text or "证明" in text):
                                matched_q = find_best_qual(text)
                                if matched_q and matched_q["image_path"] not in embedded_image_paths_in_doc:
                                    replace_paragraph_with_image(p, matched_q)
                                    embedded_image_paths_in_doc.add(matched_q["image_path"])
                                    embedded_count += 1
                                    logger.info(f"   🖼️ 已原位替换表格单元格占位符 '{text[:30]}' -> 资质图片 {matched_q['name']}")

            if embedded_count > 0:
                _safe_save_doc(doc, docx_path)
                logger.info(f"✅ [Auto Image Embedder] 成功完成 {embedded_count} 张资质证明图片的落盘嵌入！文档: {docx_path}")

        return embedded_count
    except Exception as e:
        logger.exception(f"❌ [Auto Image Embedder] 自动嵌入资质图片产生异常: {e}")
        return 0




@tool
def query_project_metadata_tool(document_id: str, field_key: str) -> str:
    """
    [数据库直查工具] 全量直查本次招标项目的元数据信息。
    支持跨表覆盖：
    - timeline_metadata (项目名称, 招标编号, 标段, 递交截止时间, 投标有效期, 工期描述)
    - qualification_metadata (强制性资质门槛, 核心人员要求, 历史业绩门槛, 废标条款)
    - engineering_metadata (主要设备清单, 强制性技术标准, 现场限制说明)
    - evaluation_metadata (评标办法, 总分与权重分配)
    - documents (包含解析元数据 parsed_metadata)

    :param document_id: 关联的招标文件 Document ID
    :param field_key: 查询字段 (如 'project_name', 'project_code', 'construction_period', 'bid_deadline', 'personnel_requirements', 'invalid_bid_clauses')
    :return: 项目元数据真实字符串
    """
    logger.info(f"🛠️ [DB Tool] query_project_metadata_tool 被调用, doc_id: '{document_id}', 字段: '{field_key}'")
    db: Session = SessionLocal()
    try:
        key_lower = field_key.lower()

        # 1. 查时间节点与基本信息 (timeline_metadata)
        meta = db.query(TimelineMetadata).filter(TimelineMetadata.document_id == document_id).first()
        if meta:
            if any(k in key_lower for k in ["项目名称", "工程名称", "project_name"]):
                if meta.project_name:
                    return meta.project_name
            elif any(k in key_lower for k in ["编号", "代码", "project_code", "project_id"]):
                if meta.project_id_code:
                    return meta.project_id_code
            elif any(k in key_lower for k in ["工期", "建设周期", "交货期", "construction_period"]):
                if meta.construction_period_description:
                    return meta.construction_period_description
                elif meta.construction_period_days:
                    return f"{meta.construction_period_days} 日历天"
            elif any(k in key_lower for k in ["截止时间", "开标时间", "bid_deadline"]):
                if meta.bid_deadline:
                    return str(meta.bid_deadline)
            elif any(k in key_lower for k in ["有效期", "validity"]):
                if meta.bid_validity_days:
                    return f"{meta.bid_validity_days} 天"
            elif any(k in key_lower for k in ["标段", "包件", "tender_segment"]):
                if meta.tender_segment:
                    return meta.tender_segment

        # 2. 查资格与人员要求 (qualification_metadata)
        qual_meta = db.query(QualificationMetadata).filter(QualificationMetadata.document_id == document_id).first()
        if qual_meta:
            if any(k in key_lower for k in ["人员", "项目经理", "技术负责人", "personnel"]):
                if qual_meta.personnel_requirements:
                    return json.dumps(qual_meta.personnel_requirements, ensure_ascii=False)
            elif any(k in key_lower for k in ["废标", "否决", "invalid_bid"]):
                if qual_meta.invalid_bid_clauses:
                    return json.dumps(qual_meta.invalid_bid_clauses, ensure_ascii=False)
            elif any(k in key_lower for k in ["业绩", "performance"]):
                if qual_meta.performance_requirements:
                    return json.dumps(qual_meta.performance_requirements, ensure_ascii=False)

        # 3. 查工程与设备清单 (engineering_metadata)
        eng_meta = db.query(EngineeringMetadata).filter(EngineeringMetadata.document_id == document_id).first()
        if eng_meta:
            if any(k in key_lower for k in ["设备", "清单", "equipment"]):
                if eng_meta.main_equipment_list:
                    return json.dumps(eng_meta.main_equipment_list, ensure_ascii=False)
            elif any(k in key_lower for k in ["标准", "规范", "standard"]):
                if eng_meta.mandatory_standards:
                    return json.dumps(eng_meta.mandatory_standards, ensure_ascii=False)
            elif any(k in key_lower for k in ["工程", "engineering", "技术", "特殊工况", "现场"]):
                # 通用工程查询 → 返回完整 engineering_metadata
                result = {k: v for k, v in eng_meta.__dict__.items() if not k.startswith('_')}
                return json.dumps(result, ensure_ascii=False, default=str)

        # 4. 查评标方法与权重 (evaluation_metadata)
        eval_meta = db.query(EvaluationMetadata).filter(EvaluationMetadata.document_id == document_id).first()
        if eval_meta:
            if any(k in key_lower for k in ["评标", "评审", "evaluation", "评分", "权重"]):
                if eval_meta.evaluation_method:
                    return eval_meta.evaluation_method

        # 5. 查 Document 实体及 parsed_metadata
        doc = db.query(DocumentModel).filter(DocumentModel.id == document_id).first()
        if doc:
            if any(k in key_lower for k in ["项目名称", "工程名称", "project_name"]):
                if getattr(doc, "project_name", None):
                    return doc.project_name
                elif doc.filename:
                    clean_name = doc.filename.replace("Frontend Uploads (", "").replace(")", "").replace(".docx", "").replace(".pdf", "")
                    return clean_name.strip()
            elif any(k in key_lower for k in ["编号", "代码", "project_code", "project_id"]):
                if getattr(doc, "project_code", None):
                    return doc.project_code
            elif doc.parsed_metadata:
                val = doc.parsed_metadata.get(field_key) or doc.parsed_metadata.get("project_name")
                if val:
                    return str(val)

        logger.warning(f"🛠️ [DB Tool] 项目元数据中未查到 '{field_key}'")
        return f"[待补充: {field_key}]"

    except Exception as e:
        logger.exception(f"🛠️ [DB Tool] query_project_metadata_tool 执行异常: {str(e)}")
        return f"[查询元数据异常: {str(e)}]"
    finally:
        db.close()


def sort_cost_items_by_scope_and_hierarchy(cost_items: List[Any]) -> List[Any]:
    """
    【区域聚类优先、区域内保持前端顺序的排序算法 — 零行业与零设备硬编码】
    检测到多个区域/标段时，先按区域/标段聚类，再在每个区域内严格按照
    前端 BOM 的 sort_order 排序；没有多个区域时，完全保持前端顺序。
    对没有 sort_order 的历史数据，使用原始列表位置作为区域内顺序。
    根据清单项的通用区域/标段结构特征（如第X标段、第X包、第X期、第X区等）进行聚合，
    并严格保持同区域/标段内部原招标文件抽取时的自然相对次序（Original Relative Order），
    使同区域/同标段项目连续聚合，杜绝多标段交叉穿插，适用于任何招投标行业与品类。
    """
    if not cost_items or len(cost_items) <= 1:
        return cost_items

    import re

    def _extract_scope_key(item: Any) -> tuple:
        name = str(getattr(item, 'item_name', '') or '')
        sec = str(getattr(item, 'section_name', '') or '')
        combined = f"{sec} {name}"

        scope_num = 999
        scope_str = ""

        # 通用区域/标段结构模式匹配（支持中英文括号或词尾区划标识，如 1标段、包2、一期、二区等）
        scope_match = re.search(r'[\(（]([^\)）]*(?:区|期|标|包|段|厂|站)[^\)）]*)[\)）]', combined)
        if not scope_match:
            scope_match = re.search(r'([一二三四五六七八九十\d]+(?:区|期|标段|包|标|分厂|车间))', combined)

        if scope_match:
            scope_str = scope_match.group(1).strip()
            # 通用中文数字映射
            cn_num_map = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9, '十': 10}
            digit_m = re.search(r'(\d+)', scope_str)
            if digit_m:
                scope_num = int(digit_m.group(1))
            else:
                for cn_char, val in cn_num_map.items():
                    if cn_char in scope_str:
                        scope_num = val
                        break

        return (scope_num, scope_str)

    # cost_estimates.sort_order 是前端 CostTable 的原始列表索引。
    # 先按区域聚类，再用 sort_order 作为区域内的稳定顺序，避免聚类破坏前端顺序。
    has_explicit_order = any(getattr(item, "sort_order", None) is not None for item in cost_items)
    has_scope = any(_extract_scope_key(item)[1] for item in cost_items)

    def _frontend_order(item: Any, original_index: int) -> tuple:
        raw_order = getattr(item, "sort_order", None)
        if raw_order is None:
            return (1, original_index)
        try:
            return (0, int(raw_order))
        except (TypeError, ValueError):
            return (1, original_index)

    # 没有区域/标段时，完全按前端顺序；有区域时才执行聚类。
    indexed_items = list(enumerate(cost_items))
    if has_scope:
        indexed_items.sort(key=lambda pair: (
            _extract_scope_key(pair[1])[0],
            _extract_scope_key(pair[1])[1],
            _frontend_order(pair[1], pair[0]) if has_explicit_order else (0, pair[0]),
        ))
    elif has_explicit_order:
        indexed_items.sort(key=lambda pair: _frontend_order(pair[1], pair[0]))
    return [pair[1] for pair in indexed_items]


def build_dynamic_matrix_for_header(cost_items: List[Any], header_columns: Optional[List[str]] = None) -> List[List[str]]:
    """
    【纯 ORM 字段名直映射引擎 — 语义映射由 Worker LLM 负责】
    接收 Worker LLM 预推理好的 ORM 物理字段名列表 (header_columns)，
    直接按列名顺序从 CostEstimate 实体中提取数据生成二维矩阵。
    
    支持的特殊标记列：
      - "__INDEX__"      → 自动填充 1..N 递增序号
      - "__BRAND_SPEC__" → 将 brand 和 spec 字段合并为一列
    
    此函数不包含任何中文关键字映射逻辑，所有中文表头→ORM字段名的
    语义推理全部由上游 Worker LLM Agent 在调用前完成。
    """
    if not cost_items:
        return []

    import re

    # 0. 自动执行区域与分项智能聚类排序
    cost_items = sort_cost_items_by_scope_and_hierarchy(cost_items)

    # 1. 动态获取 ORM 模型物理字段列表（排除系统字段）
    from app.db.models.ai_analysis import CostEstimate
    model_columns = [col.key for col in CostEstimate.__table__.columns
                     if not col.primary_key and col.key not in (
                         'document_id', 'project_id', 'reference_price_id',
                         'created_at', 'updated_at', 'tenant_id'
                     )]

    # 2. 若未提供映射后的字段列表，反射模型所有业务字段
    if not header_columns:
        sample = cost_items[0]
        header_columns = [col for col in model_columns
                          if hasattr(sample, col) and getattr(sample, col) is not None]
        if not header_columns:
            header_columns = model_columns

    # 3. 按区域进行通用聚合分组
    from itertools import groupby
    def _get_scope_name(it: Any) -> str:
        name = str(getattr(it, 'item_name', '') or '')
        sec = str(getattr(it, 'section_name', '') or '')
        combined = f"{sec} {name}"
        scope_match = re.search(r'[\(（]([^\)）]*(?:区|期|标|包|段|厂|站)[^\)）]*)[\)）]', combined)
        if not scope_match:
            scope_match = re.search(r'([一二三四五六七八九十\d]+(?:区|期|标段|包|标|分厂|车间))', combined)
        return scope_match.group(1).strip() if scope_match else ""

    groups = []
    for k, g in groupby(cost_items, key=_get_scope_name):
        groups.append((k, list(g)))

    # 识别名称列与总价列索引
    name_col_idx = 1
    total_col_idx = -1
    for i, col in enumerate(header_columns):
        if col in ("item_name", "product_name"):
            name_col_idx = i
        if col in ("calculated_total", "total_price", "amount", "total"):
            total_col_idx = i
    if total_col_idx == -1:
        total_col_idx = len(header_columns) - 1

    has_multi_scopes = len(groups) > 1 and any(g[0] for g in groups)

    def _extract_cell_value(item: Any, col_name: str, row_idx: int, g_prefix: str = "") -> str:
        col_key = str(col_name).strip()
        col_lower = col_key.lower().replace(" ", "").replace("_", "").replace("（", "(").replace("）", ")")

        # 1. 序号列
        if col_key == "__INDEX__" or any(kw in col_lower for kw in ["序号", "no", "no.", "项号", "num"]):
            return f"{g_prefix}{row_idx}" if g_prefix else str(row_idx)

        # 2. 品牌规格型号合并列
        if col_key == "__BRAND_SPEC__" or ("品牌" in col_lower and ("规格" in col_lower or "型号" in col_lower)):
            brand = str(getattr(item, 'brand', '') or '').strip()
            spec = str(getattr(item, 'spec', '') or '').strip()
            spec_clean = re.sub(r'[\r\n\t]+', ' ', spec).strip()
            if brand and spec_clean:
                if spec_clean.startswith(brand):
                    return spec_clean
                return f"{brand} {spec_clean}".strip()
            return brand or spec_clean

        # 3. 生产厂家列（支持多级回溯解析：item.manufacturer -> item.brand）
        if col_key in ("manufacturer", "mfg", "producer", "factory") or any(kw in col_lower for kw in ["生产厂家", "制造厂家", "制造厂商", "生产厂商", "制造商", "厂家", "生产企业", "制造企业", "产地"]):
            mfg_val = str(getattr(item, 'manufacturer', '') or '').strip()
            if not mfg_val:
                mfg_val = str(getattr(item, 'brand', '') or '').strip()
            return mfg_val

        # 4. 品牌列
        if col_key in ("brand", "trademark") or any(kw in col_lower for kw in ["品牌", "商标"]):
            b_val = str(getattr(item, 'brand', '') or '').strip()
            if not b_val:
                b_val = str(getattr(item, 'manufacturer', '') or '').strip()
            return b_val

        # 5. 名称/标的物/货物列
        if col_key in ("item_name", "product_name", "name") or any(kw in col_lower for kw in ["标的物", "货物名称", "设备名称", "材料名称", "项目名称", "品名", "名称"]):
            val = getattr(item, 'item_name', '') or ''
            return re.sub(r'[\r\n\t]+', ' ', str(val)).strip()

        # 6. 规格型号列
        if col_key in ("spec", "specification") or any(kw in col_lower for kw in ["规格型号", "规格及型号", "规格", "型号", "技术参数", "技术要求"]):
            val = getattr(item, 'spec', '') or ''
            return re.sub(r'[\r\n\t]+', ' ', str(val)).strip()

        # 7. 计量单位列
        if col_key in ("unit", "measure_unit") or any(kw in col_lower for kw in ["单位", "计量单位"]):
            return str(getattr(item, 'unit', '') or '').strip()

        # 8. 数量/工程量列
        if col_key in ("quantity", "qty", "amount_count") or any(kw in col_lower for kw in ["数量", "工程量", "采购量"]):
            val = getattr(item, 'quantity', None)
            if val is not None:
                return f"{val:g}" if isinstance(val, (int, float)) else str(val).strip()
            return ""

        # 9. 单价列
        if col_key in ("unit_price", "price") or any(kw in col_lower for kw in ["单价", "综合单价", "投标单价", "参考单价"]):
            val = getattr(item, 'unit_price', None)
            if val is not None:
                if isinstance(val, (int, float)):
                    return f"{val:,.2f}"
                return str(val).strip()
            return ""

        # 10. 合价/总价列
        if col_key in ("calculated_total", "total_price", "amount", "total") or any(kw in col_lower for kw in ["合价", "总价", "分项总价", "金额", "小计"]):
            val = getattr(item, 'calculated_total', None)
            if val is not None:
                if isinstance(val, (int, float)):
                    return f"{val:,.2f}"
                return str(val).strip()
            return ""

        # 11. 备注/说明列
        if col_key in ("remark", "note", "desc") or any(kw in col_lower for kw in ["备注", "说明"]):
            val = getattr(item, 'remark', '') or ''
            return re.sub(r'[\r\n\t]+', ' ', str(val)).strip()

        # 12. 反射模型属性兜底
        if hasattr(item, col_key):
            val = getattr(item, col_key, None)
            if val is not None:
                return re.sub(r'[\r\n\t]+', ' ', str(val)).strip()

        return ""

    matrix = []

    # 4. 单一区域降级模式（无分部结构）
    if not has_multi_scopes:
        for idx, item in enumerate(cost_items, start=1):
            row = [_extract_cell_value(item, col_name, idx) for col_name in header_columns]
            matrix.append(row)
        return matrix

    # 5. 多区域模式：插入区域分部标题行、分部层级序号与小计行
    cn_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]

    for g_i, (scope_name, group_items) in enumerate(groups):
        section_cn = cn_nums[g_i] if g_i < len(cn_nums) else str(g_i + 1)
        scope_display = scope_name or f"分部工程 {section_cn}"

        # 5.1 区域分部标题行
        sec_header_row = ["" for _ in header_columns]
        if "__INDEX__" in header_columns:
            sec_header_row[header_columns.index("__INDEX__")] = f"{section_cn}、"
        if 0 <= name_col_idx < len(header_columns):
            sec_header_row[name_col_idx] = f"{scope_display}"
        matrix.append(sec_header_row)

        # 5.2 区域明细数据行
        group_total = 0.0
        for item_i, item in enumerate(group_items, start=1):
            item_total = float(getattr(item, 'calculated_total', 0) or 0)
            group_total += item_total
            row = [_extract_cell_value(item, col_name, item_i, g_prefix=f"{g_i + 1}.") for col_name in header_columns]
            matrix.append(row)

        # 5.3 区域分部小计行
        sec_subtotal_row = ["" for _ in header_columns]
        if 0 <= name_col_idx < len(header_columns):
            sec_subtotal_row[name_col_idx] = f"{scope_display} 小计"
        if 0 <= total_col_idx < len(header_columns):
            sec_subtotal_row[total_col_idx] = f"{group_total:,.2f}"
        matrix.append(sec_subtotal_row)

    return matrix


@tool
def query_financial_quotation_tool(document_id: str, field_key: str = "cost_estimates", header_columns_json: Optional[str] = None) -> str:
    """
    [数据库直查工具] 全量查询财务报价、BOM 清单、分项造价、保证金及付款条款数据。

    :param document_id: 招标文件 ID
    :param field_key: 查询类型 ('cost_estimates', 'bom_list', 'cost_estimates_json_matrix', 'total_price_numeric', 'total_price_chinese', 'bid_bond', 'performance_bond', 'payment_milestones')，默认值为 'cost_estimates'。
        【重要提示】：若你的目标是为 Word 表格生成数据并调用 officecli_fill_table_rows 写盘，必须直接指定 field_key='cost_estimates_json_matrix' 并传入 header_columns_json，严禁先查 cost_estimates 纯文本再查 matrix！
    :param header_columns_json: 【重要 - 获取 matrix 时强烈建议传入】JSON 字符串数组，包含你根据 Word 表格实际表头推理映射后的 ORM 字段名列表。
        可用的 ORM 字段名（按实际需要选用、排列）：
          - "__INDEX__"       → 自动生成 1..N 递增序号
          - "item_name"       → 名称/标的物/设备/货物/品名
          - "brand"           → 品牌
          - "manufacturer"    → 生产厂家/制造厂商/生产企业
          - "spec"            → 规格/型号/技术参数
          - "__BRAND_SPEC__"  → 品牌+规格合并为一列 (当表头为"品牌、规格、型号"时使用)
          - "unit"            → 单位/计量单位
          - "quantity"        → 数量/工程量
          - "unit_price"      → 单价/综合单价
          - "calculated_total"→ 总价/合价/小计/金额
          - "remark"          → 备注/说明
        示例：若表头为 [序号, 货物名称, 规格型号, 生产厂家, 单位, 数量, 单价, 合价, 备注]
              → header_columns_json = '["__INDEX__", "item_name", "spec", "manufacturer", "unit", "quantity", "unit_price", "calculated_total", "remark"]'
    :return: 分项报价明细列表、价格、保证金或付款条款字符串
    """
    logger.info(f"🛠️ [DB Tool] query_financial_quotation_tool 被调用, doc_id: '{document_id}', 字段: '{field_key}'")
    db: Session = SessionLocal()
    try:
        key_lower = str(field_key or "cost_estimates").lower()

        # 1. 查财务元数据中的保证金与付款节点
        fin_meta = db.query(FinancialMetadata).filter(FinancialMetadata.document_id == document_id).first()
        if fin_meta:
            if "bond" in key_lower or "保证金" in key_lower:
                if fin_meta.bid_bond:
                    return json.dumps(fin_meta.bid_bond, ensure_ascii=False)
            elif "payment" in key_lower or "付款" in key_lower or "支付" in key_lower:
                if fin_meta.payment_milestones:
                    return json.dumps(fin_meta.payment_milestones, ensure_ascii=False)

        # 2. 查询成本测算总价与大写转换
        cost_items = (
            db.query(CostEstimate)
            .filter(CostEstimate.document_id == document_id)
            .order_by(
                CostEstimate.sort_order.asc().nullslast(),
                CostEstimate.created_at.asc(),
            )
            .all()
        )
        # 有 sort_order 时由排序函数严格保持前端顺序；旧数据仍兼容区域聚类。
        cost_items = sort_cost_items_by_scope_and_hierarchy(cost_items)
        
        # 针对分项清单、BOM 表或分项单价/合价查询，返回每行精细列表
        if any(k in key_lower for k in ["cost", "bom", "item", "清单", "明细", "分项", "配置", "设备", "sub", "quote", "报价"]):
            if not cost_items:
                return "[待补充: 成本测算与 BOM 分项清单数据库尚未录入]"
            
            # 如果请求 JSON 二维矩阵格式（供 officecli_fill_table_rows 直接填充表格）
            if any(k in key_lower for k in ["matrix", "json", "grid", "table", "矩阵", "二维"]):
                hdr_cols = None
                if header_columns_json:
                    try:
                        hdr_cols = json.loads(header_columns_json)
                    except Exception:
                        pass
                
                rows_matrix = build_dynamic_matrix_for_header(cost_items, hdr_cols)
                logger.info(f"🛠️ [DB Tool] 成功生成 {len(rows_matrix)} 行 BOM 自适应表格 JSON 矩阵 (列数: {len(rows_matrix[0]) if rows_matrix else 0})")
                return json.dumps(rows_matrix, ensure_ascii=False)

            res_items = []
            for item in cost_items:
                brand_str = f" [品牌: {item.brand}]" if getattr(item, 'brand', None) else ""
                spec_str = f" [规格: {item.spec}]" if getattr(item, 'spec', None) else ""
                unit_str = str(item.unit) if item.unit else ""
                res_items.append(
                    f"- {item.item_name}{brand_str}{spec_str} | 数量: {item.quantity}{unit_str} | 参考单价: {item.unit_price}元 | 测算合计合价(总价): {item.calculated_total}元 | 备注: {getattr(item, 'remark', '')}"
                )
            logger.info(f"🛠️ [DB Tool] 成功查得并回传 {len(res_items)} 条 BOM 分项成本报价明细")
            return "\n".join(res_items) + f"\n\n【分项展开提示】共检索到 {len(res_items)} 项具体分项清单。在分项报价表中，必须在建设费等汇总大类下方将全部具体细项逐行完整展开（按 2.1, 2.2... 2.K 编号），严禁只填大类总额。"

        if not cost_items:
            return "[待补充: 财务总报价与分项测算数据尚未录入]"

        total_price = sum(item.calculated_total for item in cost_items)

        if any(k in key_lower for k in ["大写", "chinese"]):
            chinese_upper = number_to_chinese_rmb(total_price)
            logger.info(f"🛠️ [DB Tool] 成功计算投标总价汉字大写: {total_price} -> '{chinese_upper}'")
            return chinese_upper
        else:
            num_str = f"{total_price:,.2f}"
            logger.info(f"🛠️ [DB Tool] 成功查询投标总价阿拉伯数字: '{num_str}' 元")
            return num_str

    except Exception as e:
        logger.exception(f"🛠️ [DB Tool] query_financial_quotation_tool 执行异常: {str(e)}")
        return f"[查询财务报价异常: {str(e)}]"
    finally:
        db.close()


@tool
def query_market_price_reference_tool(item_name: str) -> str:
    """
    [数据库直查工具] 查询市场设备/材料参考指导价以及 BOM 历史报价数据库 (market_price_references 及 cost_estimates 表)。
    适用于查询特定设备（如某核心主机、某配套总成）、软硬件、材料、耗材的参考单价、品牌、规格型号以及系统测算的单价和分项合价。

    :param item_name: 品目/设备/材料名称关键字（支持单关键词或复合关键词，如 'XXX主设备'、'XXX配件 某推荐品牌' 或 '某设备名称 XXX技术规范'）
    :return: 包含品名、品牌、规格、参考单价、建议数量与总价等完整维度的指导意见
    """
    logger.info(f"🛠️ [DB Tool] query_market_price_reference_tool 被调用, 品目关键字: '{item_name}'")
    if not item_name:
        return "[错误: 品目关键字不能为空]"

    db: Session = SessionLocal()
    try:
        clean_kw = item_name.strip()
        # 1. 优先尝试完全精准模糊匹配
        items = db.query(MarketPriceReference).filter(
            MarketPriceReference.item_name.ilike(f"%{clean_kw}%")
        ).all()

        # 2. 如果无精准结果，则采取多关键字切分匹配（如 "XXX核心设备 某厂牌 XXX型号" -> ["XXX核心设备", "某厂牌"]）
        if not items and (" " in clean_kw or "/" in clean_kw or "—" in clean_kw or "-" in clean_kw):
            tokens = [t for t in clean_kw.replace("/", " ").replace("—", " ").replace("-", " ").split() if len(t) >= 2]
            for t in tokens:
                sub_res = db.query(MarketPriceReference).filter(
                    MarketPriceReference.item_name.ilike(f"%{t}%")
                ).all()
                for r in sub_res:
                    if r.id not in [x.id for x in items]:
                        items.append(r)

        res_list = []
        if items:
            for it in items:
                brand_str = f" [品牌: {it.brand}]" if it.brand else ""
                spec_str = f" [规格: {it.spec}]" if it.spec else ""
                res_list.append(f"【市场参考单价】{it.item_name}{brand_str}{spec_str}: 单价 {it.unit_price}元/{it.unit} (说明: {getattr(it, 'remark', '')})")

        # 3. 联合直查本系统 CostEstimate 成本估算底单（能够提供真实的数量与合计合价/总价，弥补缺失）
        first_word = clean_kw.split()[0] if " " in clean_kw else clean_kw
        cost_matches = db.query(CostEstimate).filter(
            CostEstimate.item_name.ilike(f"%{first_word}%")
        ).all()
        if cost_matches:
            for c in cost_matches:
                brand_str = f" [品牌: {c.brand}]" if getattr(c, 'brand', None) else ""
                res_list.append(
                    f"【本期项目BOM实测记录】{c.item_name}{brand_str} | 推荐数量: {c.quantity}{c.unit} | 测算单价: {c.unit_price}元 | 分项估算总价(合价): {c.calculated_total}元"
                )

        if res_list:
            return "\n".join(res_list)

        return f"[未查到与 '{item_name}' 相关的参考指导价与合价]"

    except Exception as e:
        logger.exception(f"🛠️ [DB Tool] query_market_price_reference_tool 执行异常: {str(e)}")
        return f"[查询市场指导价异常: {str(e)}]"
    finally:
        db.close()


@tool
def query_evaluation_method_tool(document_id: str, detail_type: str = "method") -> str:
    """
    [数据库直查工具] 独立直查本招投标项目的评标方法、评标办法与评审打分细则 (evaluation_metadata 表)。
    可单独查询：
    - evaluation_method: 评标办法名称
    - total_score: 评标总分
    - weight_distribution: 商务/技术/价格权重分值分配 JSON
    - score_tree: 详细打分细则树 JSON

    :param document_id: 关联的招标文件 Document ID
    :param detail_type: 查询类型 ('method', 'weight', 'score_tree', 'all')
    :return: 评标办法与打分细则描述文本
    """
    logger.info(f"🛠️ [DB Tool] query_evaluation_method_tool 独立工具被调用, doc_id: '{document_id}', detail_type: '{detail_type}'")
    db: Session = SessionLocal()
    try:
        eval_meta = db.query(EvaluationMetadata).filter(EvaluationMetadata.document_id == document_id).first()
        if eval_meta:
            if (detail_type == "weight" or "weight" in detail_type) and eval_meta.weight_distribution:
                return json.dumps(eval_meta.weight_distribution, ensure_ascii=False)
            elif (detail_type == "score_tree" or "tree" in detail_type) and eval_meta.score_tree:
                return json.dumps(eval_meta.score_tree, ensure_ascii=False)
            elif eval_meta.evaluation_method:
                score_str = f" (总分: {eval_meta.total_score}分)" if eval_meta.total_score else ""
                logger.info(f"🛠️ [DB Tool] 查得真实评标办法: '{eval_meta.evaluation_method}{score_str}'")
                return f"{eval_meta.evaluation_method}{score_str}"

        logger.warning("🛠️ [DB Tool] DB 中未录入 evaluation_metadata，返回待补充标记")
        return "[待补充: 评标办法未录入系统]"

    except Exception as e:
        logger.exception(f"🛠️ [DB Tool] query_evaluation_method_tool 执行异常: {str(e)}")
        return f"[查询评标办法异常: {str(e)}]"
    finally:
        db.close()


def get_all_bid_db_tools() -> List[Any]:
    """获取所有 6 大数据库直查工具实例列表"""
    return [
        query_company_profile_tool,
        query_company_qualification_tool,
        query_project_metadata_tool,
        query_financial_quotation_tool,
        query_market_price_reference_tool,
        query_evaluation_method_tool,
    ]
