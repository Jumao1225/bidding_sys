"""
BidFiller Worker 工厂模块 (bid_filler_workers.py)

Worker Agent 职责：读文档 → 查 DB → 产出结构化 FillProposal。
不直接写 Word —— 所有写盘由 Review Agent 审查后统一执行。
"""

import json as _json
from typing import Dict, Any, List, Optional
from loguru import logger
from langgraph.prebuilt import create_react_agent
from langchain_core.messages import HumanMessage, SystemMessage

from app.services.llm_service import llm_service
from app.agents.tools.bid_db_tools import get_all_bid_db_tools

import threading as _threading
import os
import time
import re

# 全局提案收集池（线程安全，key 为 document_id）
_PROPOSALS_LOCK = _threading.Lock()
_WORKER_PROPOSALS: Dict[str, List[Dict[str, Any]]] = {}


def get_worker_proposals(document_id: str) -> List[Dict[str, Any]]:
    """获取指定文档的所有 Worker 填写提案"""
    with _PROPOSALS_LOCK:
        return list(_WORKER_PROPOSALS.get(document_id, []))


def clear_worker_proposals(document_id: str) -> None:
    """清理指定文档的提案数据"""
    with _PROPOSALS_LOCK:
        _WORKER_PROPOSALS.pop(document_id, None)


def _read_target_paragraph_text(docx_path: str, target_path: str) -> str:
    """动态读取图片目标节点的当前文本，避免依赖固定段落编号。"""
    if not docx_path or not os.path.exists(docx_path) or not target_path:
        return ""

    try:
        from docx import Document

        document = Document(docx_path)
        para_id_match = re.search(r"@paraId=([A-Fa-f0-9]+)", target_path)
        if para_id_match:
            expected_id = para_id_match.group(1).upper()
            paragraphs = list(document.paragraphs)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.extend(cell.paragraphs)
            for paragraph in paragraphs:
                for key, value in paragraph._element.attrib.items():
                    if str(key).lower().endswith("paraid") and str(value).upper() == expected_id:
                        return str(paragraph.text or "").strip()

        body_path_match = re.search(r"/body/p\[(\d+)\]", target_path)
        if body_path_match:
            paragraph_index = int(body_path_match.group(1)) - 1
            if 0 <= paragraph_index < len(document.paragraphs):
                return str(document.paragraphs[paragraph_index].text or "").strip()
    except (OSError, ValueError, ImportError) as read_error:
        logger.warning(f"读取图片目标节点原文失败: {read_error}")

    return ""


def _proposal_merge_key(proposal: Dict[str, Any]) -> str:
    """为提案生成动态去重键，允许同一条款挂载多张不同图片。"""
    path = str(proposal.get("path", "")).strip()
    if str(proposal.get("type", "")).strip() == "image":
        image_value = str(
            proposal.get("proposed_text")
            if proposal.get("proposed_text") is not None
            else proposal.get("value", "")
        ).strip()
        return f"{path}::image::{image_value}"
    return path


def _filter_dom_scope(raw_structure: str, target_chapter: str, keyword: str, window_size: int = 3) -> str:
    """
    [视口切碎保持器 (Strict Scope Splicing)]
    当文档包含大量节点时，仅精准保留目标点及邻近 ±window_size 段的纯净视口，彻底消除跨章节信息过载与注意力稀释。
    硬上限约束：最多保留 25 行、总长度不超过 4000 字符，杜绝超大 DOM 导致 LLM API 超时。
    """
    if not raw_structure:
        return ""
    lines = [line.strip() for line in raw_structure.split("\n") if line.strip()]
    if len(lines) <= 30 and not keyword and len(raw_structure) <= 3500:
        return raw_structure

    search_term = keyword.strip() if keyword else target_chapter.strip()
    if not search_term:
        safe_lines = lines[:25]
        return "\n".join(safe_lines) + ("\n...(全文过长，默认仅截取前25行DOM)..." if len(lines) > 25 else "")

    matched_indices = [i for i, l in enumerate(lines) if search_term in l]
    if not matched_indices and len(search_term) > 3:
        short_term = search_term[:4]
        matched_indices = [i for i, l in enumerate(lines) if short_term in l]

    if not matched_indices:
        # 若为常规不涉及精准命中的表，回退保护不过载
        logger.info(f"   🔎 [DOM视口] 未直接命中关键词 '{search_term}'，回退截取前 3000 字符安全视口")
        return raw_structure[:3000] + ("\n...(部分跨章冗余节点已按规则保护折叠)..." if len(raw_structure) > 3000 else "")

    selected_indices = set()
    for idx in matched_indices:
        for w in range(max(0, idx - window_size), min(len(lines), idx + window_size + 1)):
            selected_indices.add(w)

    sorted_indices = sorted(selected_indices)
    filtered_lines = [lines[i] for i in sorted_indices]

    # 安全防爆 1：若匹配行数过多，仅保留前 25 行
    if len(filtered_lines) > 25:
        logger.warning(f"   ⚠️ [DOM视口保护] 匹配节点行数过多 ({len(filtered_lines)} 行)，自动修剪至前 25 行")
        filtered_lines = filtered_lines[:25] + ["...(其余冗余匹配行已自动截断以防 Token 溢出)..."]

    summary_hdr = (
        f"✂️ [切碎视口绝缘池 (Scope Spliced)]: 全文件 {len(lines)} 个 DOM 节点 -> 依据 '{search_term}' "
        f"聚合锁定前后 ±{window_size} 段 ({len(filtered_lines)} 个目标行)：\n"
    )
    result_text = summary_hdr + "\n".join(filtered_lines)

    # 安全防爆 2：硬字符上限 4000 字符
    if len(result_text) > 4000:
        logger.warning(f"   ⚠️ [DOM视口保护] 视口字符数达到 {len(result_text)} 字符，自动截断至 4000 字符")
        result_text = result_text[:4000] + "\n...(超出 4000 字符部分已安全截断)..."

    logger.info(f"   🔎 [DOM视口交付] 原始 {len(raw_structure)} 字符 -> 最终裁切交付 Agent: {len(result_text)} 字符 ({len(filtered_lines)} 行)")
    return result_text


def _build_worker_tools(
    docx_temp_path: str,
    chapter_title: str = "",
    collected_proposals: Optional[List[Dict[str, Any]]] = None,
    mapping_hint: str = "",
    category: str = "",
) -> List[Any]:
    """
    按填报范围的专属角色动态组装最精简的只读+直写工具集（Tool Pruning），
    剔除无关工具定义，降低每轮 Prompt 开销。
    """
    if collected_proposals is None:
        collected_proposals = []

    hint = (mapping_hint or "").lower().strip()
    title_lower = (chapter_title or "").lower().strip()
    cat = (category or "").lower().strip()

    is_pricing = (hint in ("pricing", "cost")) or any(k in title_lower for k in ["报价", "清单", "分项", "开标一览", "主要材料"])
    is_qualification = (hint == "qualification") or any(k in title_lower for k in ["资格", "资质", "执照", "证明文件", "安全生产", "承装"])
    is_deviation = (hint in ("deviation", "technical")) or any(k in title_lower for k in ["偏离", "响应", "技术偏离", "商务偏离", "条款偏离"])
    is_letter_or_form = (not is_pricing and not is_qualification and not is_deviation)

    from app.agents.tools.rag_tools import get_full_chapter_text, search_bidding_document
    from app.agents.tools.office_cli_agent_tools import (
        officecli_query_structure_tool,
        officecli_write_slot_value_tool,
        officecli_batch_fill_sentence_tool,
        officecli_fill_table_rows_tool,
        officecli_add_table_row_tool,
        officecli_insert_image_tool,
    )
    import asyncio
    import concurrent.futures
    from langchain_core.tools import tool

    def _sync_call_async(async_fn, *args, **kwargs):
        """线程安全的同步调用异步函数 Helper，彻底防范 RuntimeError: no running event loop"""
        try:
            loop = asyncio.get_running_loop()
        except RuntimeError:
            return asyncio.run(async_fn(*args, **kwargs))

        if loop.is_running():
            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                return executor.submit(lambda: asyncio.run(async_fn(*args, **kwargs))).result()
        else:
            return loop.run_until_complete(async_fn(*args, **kwargs))

    dom_cache: Dict[str, str] = {}

    @tool
    def officecli_query_structure(selector: str = "paragraph", keyword_filter: str = "", window: int = 3) -> str:
        """
        [填报范围结构一键提取工具] 一次性完整提取当前填报范围内的全部段落与表格 DOM 结构。
        调用一次即可获得全部待填节点与物理路径（如 /body/p[N] 或 /body/tbl[M]），无需且严禁按关键词多次重复调用！
        参数：
        - selector: 元素类型，可选 'paragraph'（段落）/ 'table'（表格）/ 'all'（全部段落与表格）
        - keyword_filter: 辅助过滤词（通常留空即可，默认自动提供当前章节全部 100% 完整结构）
        - window: 关联段数（默认 3）
        """
        cache_key = f"{selector}_{keyword_filter}_{window}"
        if cache_key in dom_cache:
            logger.debug(f"   ⚡ [章节视野缓存复用] 章节 [{chapter_title}] 命中已提取的 DOM 缓存 (selector='{selector}')")
            return dom_cache[cache_key]

        logger.info(f"   [Worker 视野] 查询结构 (selector='{selector}', kw='{keyword_filter}')")

        # 资格证明材料可能由多个相邻表单组成，不能以单一章节边界截断。
        # 直接读取完整 DOM，确保授权、承诺、证明及落款等动态节点都进入同一填报范围。
        if is_qualification:
            raw_text = _sync_call_async(
                officecli_query_structure_tool.coroutine,
                file_path=docx_temp_path,
                selector=selector,
            )
            raw_str = str(raw_text)
            result = (
                f"【系统通知：已完整提取资格证明材料相关文档 DOM 节点（当前请求：{chapter_title}）】\n"
                "⚠️ 核心指引：下方包含完整文档的段落、表格及物理路径。请依据运行时原文、字段标签和上下文，仅填报属于资格证明材料的节点；不得因附件标题、表单名称或跨章节位置跳过任何相关表单。\n\n"
                f"{raw_str}"
            )
            if selector in ("table", "all"):
                tbl_info = extract_docx_tables_summary(docx_temp_path, chapter_title)
                if tbl_info:
                    result = f"📊 【运行时检测到的表格结构】\n{tbl_info}\n\n{result}"
            dom_cache[cache_key] = result
            return result

        # 优先使用精准章节提取器（100% 完整无损提取当前章节专属 DOM 节点）
        from app.utils.table_utils import extract_chapter_dom_structure
        chapter_dom = extract_chapter_dom_structure(docx_temp_path, chapter_title, selector=selector)
        if chapter_dom and len(chapter_dom) > 50:
            logger.info(f"   🎯 [章节专属视野命中] 成功提取章节 [{chapter_title}] 100% 完整 DOM 结构 ({len(chapter_dom)} 字符)，零噪音且无信息丢失！")
            result = (
                f"【系统通知：已 100% 完整提取当前章节《{chapter_title}》的全部 DOM 节点（共 {len(chapter_dom)} 字符）】\n"
                f"⚠️ 核心指引：本章节的所有待填段落与表格节点已全量呈现在下方（零信息丢失，且无任何其他隐藏段落）。"
                f"请直接研读并利用下方的节点路径（如 /body/p[N] 或 /body/tbl[M]）进行数据填充与写盘，严禁针对本章节重复发起结构查询！\n\n"
                f"{chapter_dom}"
            )
            dom_cache[cache_key] = result
            # 兼容不同 keyword_filter 的重复命中，直接复用整章缓存
            dom_cache[f"{selector}___3"] = result
            return result

        # 降级：调用 OfficeCLI 并通过视口剪裁
        raw_text = _sync_call_async(officecli_query_structure_tool.coroutine, file_path=docx_temp_path, selector=selector)
        raw_str = str(raw_text)
        filtered = _filter_dom_scope(raw_str, chapter_title, keyword_filter, window)
        if selector in ("table", "all"):
            tbl_info = extract_docx_tables_summary(docx_temp_path, chapter_title)
            if tbl_info:
                filtered = f"📊 【当前具体表格的真实表头与列定义】\n{tbl_info}\n\n{filtered}"
        dom_cache[cache_key] = filtered
        return filtered

    @tool
    def officecli_write_slot_value(path: str, value: str) -> str:
        """
        [原位节点单槽位提案工具] 提议对 Word 指定节点 Path 进行 100% 格式继承的原位值替换。
        若有多个槽位待填，强烈建议优先使用 officecli_batch_write_slots 一次性批量提交！
        """
        p_path = str(path).strip()
        p_val = str(value).strip()
        logger.info(f"   [Worker 提案注册] 节点 {p_path} -> {p_val[:60]}")
        if p_path and p_val is not None:
            collected_proposals.append({
                "path": p_path,
                "proposed_text": p_val,
                "value": p_val,
                "type": "text",
                "status": "success"
            })
        return f"成功提交节点 {p_path} 的替换提案，已进入主控集中刷盘队列"

    @tool
    def officecli_batch_write_slots(slots_json_str: str) -> str:
        """
        [全章槽位一次性批量写盘提案工具 - 强烈推荐]
        在收集齐本章节所有待填字段后，一次性提交全部槽位替换提案！
        彻底避免逐个槽位多次调用工具带来的网络延迟与 Token 膨胀。
        参数 slots_json_str 格式：'[{"path": "/body/p[2]", "value": "纯数据值1"}, {"path": "/body/p[3]", "value": "纯数据值2"}]'
        """
        logger.info(f"   [Worker 批量槽位提案注册] 提交批量槽位替换提案: {str(slots_json_str)[:150]}...")
        count = 0
        if slots_json_str:
            try:
                parsed_list = _json.loads(slots_json_str) if isinstance(slots_json_str, str) else slots_json_str
                if isinstance(parsed_list, list):
                    for it in parsed_list:
                        if isinstance(it, dict) and "path" in it:
                            p_path = str(it["path"]).strip()
                            val = it.get("value") if it.get("value") is not None else (it.get("text") or it.get("proposed_text") or "")
                            p_val = str(val).strip()
                            if p_path and p_val is not None:
                                collected_proposals.append({
                                    "path": p_path,
                                    "proposed_text": p_val,
                                    "value": p_val,
                                    "type": "text",
                                    "status": "success"
                                })
                                count += 1
            except Exception as je:
                logger.warning(f"   解析 slots_json_str 异常: {je}")
        return f"成功批量提交 {count} 个槽位的替换提案，已全量进入主控集中刷盘队列"

    @tool
    def officecli_batch_fill_sentence(updates_json_str: str) -> str:
        """
        [长句/段落原子批处理提案工具] 在收集齐该章节长段落的所有字段后，一次性提交更新提案。
        参数 updates_json_str 格式：'[{"path": "/body/p[2]", "value": "字段标签：[抽象数据内容]"}, ...]'
        """
        logger.info(f"   [Worker 原子提案注册] 提交长句批处理提案: {str(updates_json_str)[:150]}...")
        if updates_json_str:
            try:
                parsed_list = _json.loads(updates_json_str) if isinstance(updates_json_str, str) else updates_json_str
                if isinstance(parsed_list, list):
                    for it in parsed_list:
                        if isinstance(it, dict) and "path" in it:
                            val = it.get("value") or it.get("text") or it.get("proposed_text") or ""
                            collected_proposals.append({
                                "path": str(it["path"]).strip(),
                                "proposed_text": str(val).strip(),
                                "value": str(val).strip(),
                                # 长句提案必须按完整段落处理，避免同一段中的字段发生错位。
                                "type": "sentence_batch",
                                "original_context": str(it.get("original_context", "")).strip(),
                                "status": "success"
                            })
            except Exception as je:
                logger.warning(f"   解析 updates_json_str 异常: {je}")
        return "成功提交段落批处理提案，已进入主控集中刷盘队列"

    @tool
    def officecli_fill_table_rows(table_path: str, rows_json_str: str, auto_index: bool = True) -> str:
        """
        [表格全量追加填充提案工具] 批量填充表格行，自动保留 row[1] 表头不变，并在第一列自动生成 1..N 递增序号。
        参数 rows_json_str 格式：'[["数据项1", "数据项2"], ["数据项3", "数据项4"]]'
        """
        t_path = str(table_path).strip()
        logger.info(f"   [Worker 表格提案注册] 向表格 {t_path} 提交批量填充行提案")
        val_str = rows_json_str if isinstance(rows_json_str, str) else _json.dumps(rows_json_str, ensure_ascii=False)
        if t_path and val_str:
            collected_proposals.append({
                "path": t_path,
                "proposed_text": val_str,
                "value": val_str,
                "type": "table_rows",
                "status": "success"
            })
        return f"成功提交表格 {t_path} 的数据行提案，已进入主控集中刷盘队列"

    @tool
    def officecli_insert_image(
        target_path: str,
        image_path: str,
        width_inches: float = 5.5,
        caption: str = "",
        anchor_text: str = "",
    ) -> str:
        """
        [资质证明与图片嵌入提案工具] 在 Word 指定节点 Path (如 '/body/p[12]' 或 '/body/tbl[1]/row[2]/cell[1]') 提议插入资质证明/证书图片。
        参数：
        - target_path: Word 中的物理 DOM 节点 Path
        - image_path: 资质证书图片的磁盘绝对路径 (可通过 query_company_qualification_tool 查库获取)
        - width_inches: 图片宽度 (默认 5.5 英寸)
        - caption: 图片说明图注（由资质库名称动态提供）
        - anchor_text: 图片对应的原文条款，由当前文档动态提取
        """
        tg_path = str(target_path).strip()
        img_path = str(image_path).strip()
        logger.info(f"   [Worker 图片提案注册] 节点 {tg_path} -> 提议嵌入图片 {img_path}")
        if tg_path and img_path:
            # 写入目标段落的当前原文，供主控写盘阶段做语义锚点校验。
            resolved_anchor = str(anchor_text or _read_target_paragraph_text(docx_temp_path, tg_path)).strip()
            collected_proposals.append({
                "path": tg_path,
                "proposed_text": img_path,
                "value": img_path,
                "type": "image",
                "caption": str(caption or "").strip(),
                "anchor_text": resolved_anchor,
                "original_context": resolved_anchor,
                "status": "success"
            })
        return f"已登记节点 {tg_path} 的资质图片提案，待主控写盘与回读验证"

    from app.agents.tools.style_extractor_tool import extract_text_by_style
    from app.agents.tools.bid_db_tools import (
        query_company_profile_tool,
        query_company_qualification_tool,
        query_project_metadata_tool,
        query_financial_quotation_tool,
        query_market_price_reference_tool,
        query_evaluation_method_tool,
    )

    # 按角色动态裁剪工具包（Tool Pruning），仅保留当前章节真正需要的 2~6 个核心工具
    if is_pricing:
        worker_tools = [
            officecli_query_structure,
            extract_text_by_style,
            query_financial_quotation_tool,
            officecli_batch_write_slots,
            officecli_write_slot_value,
            officecli_fill_table_rows,
            query_company_profile_tool,
            query_project_metadata_tool,
        ]
        logger.info(f"   🛠️ [Worker 工具包裁剪] 为报价章节 [{chapter_title}] 裁剪装配 8 个专用工具 (含样式识别 + pricing + company_info)")
    elif is_deviation:
        worker_tools = [
            officecli_query_structure,
            extract_text_by_style,
            get_full_chapter_text,
            search_bidding_document,
            officecli_fill_table_rows,
        ]
        logger.info(f"   🛠️ [Worker 工具包裁剪] 为偏离表章节 [{chapter_title}] 裁剪装配 5 个专用工具 (含样式识别 + RAG + fill_table_rows)")
    elif is_qualification:
        worker_tools = [
            officecli_query_structure,
            extract_text_by_style,
            query_company_qualification_tool,
            query_company_profile_tool,
            officecli_insert_image,
            officecli_batch_write_slots,
            officecli_batch_fill_sentence,
            officecli_write_slot_value,
        ]
        logger.info(f"   🛠️ [Worker 工具包裁剪] 为资质证明章节 [{chapter_title}] 裁剪装配 8 个专用工具 (含样式识别 + qualification + insert_image + sentence_batch)")
    elif is_letter_or_form:
        worker_tools = [
            officecli_query_structure,
            extract_text_by_style,
            officecli_batch_write_slots,
            officecli_write_slot_value,
            officecli_batch_fill_sentence,
            query_company_profile_tool,
            query_project_metadata_tool,
        ]
        logger.info(f"   🛠️ [Worker 工具包裁剪] 为公文表单章节 [{chapter_title}] 裁剪装配 7 个专用轻量工具 (含样式识别 + batch_write_slots + basic_info)")
    else:
        # 通用兜底
        worker_tools = [
            officecli_query_structure,
            officecli_batch_write_slots,
            officecli_write_slot_value,
            officecli_batch_fill_sentence,
            officecli_fill_table_rows,
            officecli_insert_image,
            get_full_chapter_text,
            search_bidding_document,
            extract_text_by_style,
            query_company_profile_tool,
            query_company_qualification_tool,
            query_project_metadata_tool,
            query_financial_quotation_tool,
            query_market_price_reference_tool,
            query_evaluation_method_tool,
        ]
        logger.info(f"   🛠️ [Worker 工具包全量兜底] 为通用章节 [{chapter_title}] 组装全量工具集: {len(worker_tools)} 个工具")

    return worker_tools



# ============================================================
def extract_docx_tables_summary(docx_path: str, chapter_title: str = "") -> str:
    """
    按章节精准扫描当前 Word 文档中属于当前具体章节的待填表格物理路径、行列规模与真实表头定义。
    仅将当前章节专属的具体表格表头呈现给该 Worker，彻底消除跨章节异表干扰。
    """
    if not docx_path or not os.path.exists(docx_path):
        return ""
    try:
        from docx import Document
        from app.utils.table_utils import (
            detect_table_header_rows,
            get_merged_header_texts,
            get_chapter_specific_table_indices,
            is_fixed_slot_summary_table,
        )
        doc = Document(docx_path)
        if not doc.tables:
            return ""

        target_tbl_indices = get_chapter_specific_table_indices(doc, chapter_title)
        if not target_tbl_indices:
            return ""

        tables_info = []
        for i, tbl_idx in enumerate(target_tbl_indices):
            if 0 <= tbl_idx < len(doc.tables):
                table = doc.tables[tbl_idx]
                if not table.rows:
                    continue
                hdr_count = detect_table_header_rows(table)
                headers = get_merged_header_texts(table, hdr_count)
                if not any(headers):
                    continue
                tbl_path = f"/body/tbl[{tbl_idx + 1}]"
                headers_str = " | ".join(headers)
                hdr_desc = f"包含 {hdr_count} 行复合表头, " if hdr_count > 1 else ""
                tbl_label = f"目标表格 ({i+1}/{len(target_tbl_indices)})" if len(target_tbl_indices) > 1 else "唯一目标表格"

                is_fixed = is_fixed_slot_summary_table(table, chapter_title)
                if is_fixed:
                    type_rule = "👉 【表格类型：固定格式表单（严禁增加行/严禁插行）】此表已有具体的项目行与大写总价行，只需调用 `officecli_batch_write_slots` 在对应空白单元格填入数值与大写，**绝对禁止调用 `officecli_fill_table_rows` 插入新行！绝对禁止将多条 BOM 设备/细项拆行插入！**"
                elif "商务条款" in (chapter_title or "") and any(k in (chapter_title or "") for k in ["响应", "偏离"]):
                    type_rule = "👉 【商务条款精准对照表】只提取与本表逐项对应的商务条款，不得扩展提取技术参数、项目需求或其它章节内容；必须先按实际表头动态映射字段，每行提交与实际表格列数一致的数据，禁止空行、短行或人为补充不存在的列。"
                else:
                    type_rule = "👉 【表格类型：动态多行清单展开表（空位大，需增加行全量展开）】此表需要展开具体设备材料明细/响应条款，必须根据数据库中全部明细数据，调用 `officecli_fill_table_rows(table_path, rows_json_str)` 全量覆写并展开所有数据行！"

                tables_info.append(f"- 【本章节专属{tbl_label}】：`{tbl_path}`（共 {len(headers)} 列, {hdr_desc}预置 {len(table.rows)} 行）：真实表头定义为 `[{headers_str}]`。\n  {type_rule}")

        if tables_info:
            return "\n".join(tables_info)
    except Exception as e:
        logger.warning(f"按章节提取表格表头概要异常: {e}")
    return ""


def build_worker_prompt(
    chapter_title: str,
    category: str,
    template_text: str,
    content_hint: str,
    document_id: str,
    docx_temp_path: str = "",
    mapping_hint: str = "",
    extra_instructions: str = "",
    repair_instructions: str = "",
    prefetched_metadata: Optional[Dict[str, Any]] = None,
) -> tuple:
    """构建章节 Worker Agent 的针对性专家 System Prompt 与 User Prompt（支持四类专家角色分治、真实表头注入与专项修复）。

    :param mapping_hint: 章节分类标签（如 pricing / qualification / deviation / bid_letter / authorization 等）
    :param extra_instructions: 用户自定义额外指令
    :param repair_instructions: Supervisor 下发的专项修复反馈指令
    :param prefetched_metadata: 预读取的企业档案与项目核心元数据（仅在公文函件类章节定向按需注入）
    """
    cat = (category or "needs_fill").lower().strip()
    hint = (mapping_hint or "").lower().strip()
    title_lower = (chapter_title or "").lower().strip()

    # 动态提取当前填报范围内的目标表格结构与真实表头定义。
    tables_summary = extract_docx_tables_summary(docx_temp_path, chapter_title) if docx_temp_path else ""

    # 1. 判定专家角色类型
    is_pricing = (hint in ("pricing", "cost")) or any(k in title_lower for k in ["报价", "清单", "分项", "开标一览", "主要材料"])
    is_qualification = (hint == "qualification") or any(k in title_lower for k in ["资格", "资质", "执照", "证明文件", "安全生产", "承装"])
    is_deviation = (hint in ("deviation", "technical")) or any(k in title_lower for k in ["偏离", "响应", "技术偏离", "商务偏离", "条款偏离"])
    is_business_deviation = "商务" in title_lower and (
        "偏离" in title_lower or "响应" in title_lower
    )
    is_letter_or_form = (not is_pricing and not is_qualification and not is_deviation)
    structure_scope_rule = (
        "当前资格证明材料集合（允许跨附件、跨表单读取相关节点）"
        if is_qualification
        else f"当前章节《{chapter_title}》"
    )

    # 2. 差异化专家工作流与职责
    if is_pricing:
        role_title = "造价工程师与分项报价专家"
        domain_workflow = f"""【造价工程师与分项报价专家工作流 — 场景分流与全量展开铁律】

1. **【表格类型智能识别与填报原则】**：
   - **固定格式表单（已有具体单元格填报要求）**：
     * **特征**：表格预置总行数固定（如仅有 1 行标的物汇总行 + 1 行大写金额/落款合并行），或各单元格已有明确的预置描述与填报槽位；
     * **填报规范**：**只需要原位填写空白单元格，绝对严禁增加行或插入新行！** 严禁将多条明细拆行插入到固定表单中；
     * **填报方式**：必须调用 `officecli_batch_write_slots` 对空白单元格进行【原位赋值】：
       - 项目汇总行：填入标的物名称、技术要求、阿拉伯数字总价金额（纯数字值，如金额数据）、备注；
       - 大写金额行：对应单元格直接填入人民币汉字大写金额；
   - **动态多行清单展开表（空位较大，需展开多条明细）**：
     * **特征**：表格为分项清单、设备明细、偏离对照、人员清单等明细表格，模板中通常留有较大空白占位行；
     * **填报规范**：**必须增加行并全量展开（Full Matrix Expand）！** 将数据库中的全部明细条目从第 1 项到第 N 项逐行完整展开列出；
     * **执行步骤**：按以下第 2、3 步执行 2D 矩阵全量直查与覆写。

2. **表头感知与 2D 数据矩阵一步直查（支持生产厂家与多列智能映射）**：
   - 若 User Prompt 中【文档中检测到的实际表格与真实表头定义】已包含目标表格路径及真实表头，**严禁再次调用 `officecli_query_structure` 重复扫描结构**；
   - 仔细研读真实表头各列名称，将其映射为 ORM 字段名称列表 JSON：
     * 可用 ORM 字段：`__INDEX__`, `item_name`, `spec`, `manufacturer`, `brand`, `__BRAND_SPEC__`, `unit`, `quantity`, `unit_price`, `calculated_total`, `remark`；
     * **【生产厂家映射要求】**：若表头包含“生产厂家”、“制造厂商”、“生产企业”、“制造商”等列，**必须明确映射为 `"manufacturer"`**，严禁映射为 remark 或留空！
     * 示例：若表头为 [序号, 货物名称, 规格型号, 生产厂家, 单位, 数量, 单价, 合价, 备注]
       → `header_columns_json = '["__INDEX__", "item_name", "spec", "manufacturer", "unit", "quantity", "unit_price", "calculated_total", "remark"]'`；
   - **必须且仅调用一次** `query_financial_quotation_tool(document_id='{document_id}', field_key='cost_estimates_json_matrix', header_columns_json='...')` 获取与表头完全对齐的 2D 数据矩阵；
   - **【严禁冗余查询】**：在分项清单表格填报任务中，**绝对禁止同时或重复调用 `field_key='cost_estimates'` 纯文本字段**，直接使用返回的 2D 矩阵！

3. **【分项清单全量展开与对齐规范】**：
   - **细项逐行展开**：若表格模板包含汇总大类及占位行，必须在汇总大类下方，将数据库中查得的全部具体标的物细项逐行完整展开并按顺序编号排列，严禁省略二级细项；
   - **各列严格分离对齐**：第 1 列【序号】填入层级编号（如 1、2 等），第 2 列【项目/标的物名称】填纯名称（严禁重复前缀或序号）；
   - **【单价】与【合价/总价】列严格分离**：
     * **设备材料等采购细项**：【单价】列填单价数值，【分项总价】列填数量 × 单价之合价；
     * **按项包干/工程安装/未细分单价项**：【单价】列填破折号 `"——"` 或留空，【分项总价】列填写该项的总金额，严禁将整项包干大额总金额错误复制到单价列；
     * **包含在总价内/不单独计价项**：【单价】列填 `"——"`（或 `0.00`），【分项总价】列填 `0.00`；
     * **各列严格独立对齐**：严格按表头列序逐列对应，确保【单价】与【分项总价】两列数据精准独立，绝不错列、串列；
   - **金额层级平衡**：所有二级细项合价之和必须精准等于所属大类总额，所有一级大类总额之和必须精准等于表尾【合计总价】（大写与小写一致）。
   - **整表 2D 矩阵一次性写盘**：确认矩阵数据完整后，直接调用 `officecli_fill_table_rows(table_path, rows_json_str)` 一次性提交写盘，原位覆盖并彻底清除模板原有的空白行和占位符！"""
    elif is_qualification:
        role_title = "资格审查与资质证明专家"
        domain_workflow = f"""【资格审查与资质证明专项工作流 — 文字与图片全量填报】
1. **一次扫描全部材料内容**：调用 `officecli_query_structure(selector='all')` 获取当前资格证明材料范围内的全部段落、表格和空白字段。凡是该材料范围内出现的声明、身份证明、授权委托、承诺书、落款和日期栏，均属于本 Worker 的填报对象，不得因附件标题、表单名称或段落位置而跳过。
2. **先建立动态槽位清单**：逐个记录运行时返回的 DOM 路径、原文、字段标签、占位符类型和所在段落上下文。严禁凭固定段落序号、固定路径或预设字段顺序写入。
3. **文字字段全量填报**：
   - 对公司、人员、地址、联系方式、项目、日期、期限、盖章和签字等字段，从企业库、人员库和项目库动态检索真实值；
   - 同一段落包含多个字段时，必须提交覆盖该段落的完整 `sentence_batch`，保留所有固定标签和原文，只替换空白槽位；
   - 字段缺少真实数据时，写入动态生成的待补充标记并标记为人工补充，不得使用其他字段代替，不得捏造数据。
4. **图片与文字协同填报**：
   - 对明确要求提供证书、执照、许可证、身份证明或复印件的条款，调用 `query_company_qualification_tool` 动态检索匹配图片；
   - 调用 `officecli_insert_image` 时必须同时提交当前条款的原文 `anchor_text`，图片只能插入该条款之后；
   - 图片提案只能使用运行时获取的 `paraId` 或完整 DOM 路径，严禁使用固定段落编号；
   - 严禁把图片物理路径写进正文，严禁把图片插入日期、落款或其他无关段落。
5. **提交后必须闭环**：文字和图片提案提交后，主控必须重新读取 Word，逐项检查资格材料范围内的空白字段、字段错位、图片图注与条款锚点；任何一项未通过都不能报告本任务完成。"""
    elif is_deviation:
        role_title = "商务合规与技术响应专家"
        if is_business_deviation:
            domain_workflow = f"""【商务条款精准响应专项工作流 — 只提取对应商务条款】
1. **限定检索范围**：只查找与当前表格逐项对应的商务条款，例如交货、付款、质保、售后、履约、合同、投标有效期、保证金、报价约束等；不得扩展提取技术参数、项目需求、设备明细或其它章节的内容。
2. **按表格逐项匹配**：先读取当前表格的真实表头和已有条款线索，再使用 `search_bidding_document` 或 `get_full_chapter_text` 精准检索对应商务条款。只保留能与当前表格对应的条款，不要把整章原文全部搬入表格，也不要为了凑行数生成无关条款。
3. **按实际表头动态装配**：逐列理解当前表格真实表头，将对应商务条款、承诺、响应状态、偏离说明等内容放入实际存在的列；不得假定固定列号、固定列名或固定列数。没有对应列时不要虚构数据；需要序号时可由工具自动生成。
4. **完整提交且不留空列**：每一行必须按实际表格列数提交完整数据，缺少的字段应根据该列表头补充合理的商务响应，不能提交短行或空行。
5. **一次性提交**：使用 `officecli_fill_table_rows(table_path, rows_json_str)` 一次性提交匹配后的二维矩阵；不得提交技术要求、实质性需求或与商务条款无关的额外记录。"""
        else:
            domain_workflow = f"""【偏离表与实质性条款响应专项工作流 — 允许全量生成与覆盖重写】
1. **扫描识别目标表格结构**：若未预置表格定义，使用 `officecli_query_structure(selector='table')` 一次性获取本章表格路径与列定义（严禁使用关键词重复查询）。
2. **原文件整章全量阅读与交叉检索**：
   - **必须调用 `get_full_chapter_text('{document_id}', chapter_name)` 检索原招标文件对应章节全量原文**（如需求、技术规格、合同条款或商务要求章节），获取全部条款与技术规格细节！
3. **【表格全量覆盖重写授权 (Full Table Matrix Overwrite)】**：
   - **全面重写与覆盖原表格数据行**：你有完全的权限使用 `officecli_fill_table_rows(table_path, rows_json_str)` 从序号 1 开始将所有梳理出的条款生成完整的二维矩阵，一次性覆写替换原表格的所有明细行！
   - **杜绝单点打补丁**：严禁受限于原模板历史残留的空行或部分行，必须从序号 1 到 N 全量生成整张完整规范的对照表！
4. **条款逐条独立拆分填报规范（一事一行，严禁合并挤压，覆盖全部表格行）**：
   - **一事一行独立呈现**：严禁将多个独立条款强行合并压缩到同一行！必须将原招标文件中的所有各项独立条款与参数指标约束（包括各项商务约束、技术参数、履约要求等）**逐条独立拆分为单独的数据行**，保证表格所有行均得到充分、清晰的逐条响应！
   - **各列严格对齐填报规范**：
     * **第 1 列 `tc[1]`【序号】**：填入连续纯序号数字 `1, 2, 3...`，严禁长文本混入；
     * **第 2 列 `tc[2]`【招标文件要求】**：完整原样呈现原招标文件对应条款条文，严禁使用 `...` 截断；
     * **第 3 列 `tc[3]`【服务承诺与技术响应/是否响应】**：若表头为“是否响应（填是或者否）”，统一填写“是”；若表头为“响应情况/技术承诺”，详细列明所投指标响应及法律承诺；
     * **第 4 列 `tc[4]`【有无偏离】**（若有）：统一填写“无偏离”或“无”；
     * **第 5 列 `tc[5]`【偏离内容及原因】**（若有）：统一填写“完全响应招标文件要求，无偏离。”。
5. **高效提报方式（首选 officecli_fill_table_rows 一次性全量写盘）**：
   - 必须优先直接调用 `officecli_fill_table_rows(table_path, rows_json_str)` 一次性提交整张表格的全部数据行（每行按表头列数装配 `["序号", "招标文件要求", "是/服务承诺", ...]`，底层引擎会自动保留表头并以全新数据行全量覆写原表格数据区）；
   - 严禁对表格单元格零敲碎打地调用 write_slot_value 局部打补丁！"""
    else:
        role_title = "公文函件与表单填报专家"
        domain_workflow = f"""【法定公文函件与表单填报专项工作流 — 原位切片注入】
1. **全形态槽位一键扫描识别（下划线 / 括号 / 纯空格留白）**：
   - 使用 `officecli_query_structure(selector='all')` 一次性获取本章节内的所有待填槽位（单次即可获得全章所有节点，严禁使用关键词重复查询）；
   - **必须覆盖全部空白留白形态**：
     * **符号形态**：下划线 `______`、括号 `( )` 或 `[ ]`；
     * **空格留白形态**：属性标签或冒号后的**连续空格、制表符留白**（如 `通讯地址：              `、`联系电话：          `）；
     * **日期留白形态**：年月日之间的留白空格（如 `    年    月    日`）；
     凡是属于待填信息的空白区域，一律属于合法填报目标！
2. **企业档案与项目信息装配**：
   - 优先直接使用下方【已定向提取的企业主档案与项目关键元数据】；若个别生僻字段未涵盖，可调用企业信息库（`query_company_profile_tool`）或项目元数据库（`query_project_metadata_tool`）补齐；
   - 严格根据招标文件上下文区分收件单位（如致代理机构或采购人）、组织单位与投标方主体，准确填入官方全称。
3. **全章槽位一次性批量写盘（严禁单个槽位多次重复调工具）**：
   - **必须收集齐本章节所有待填字段后，优先调用 `officecli_batch_write_slots(slots_json_str)` 一次性提交全部槽位替换提案！**
   - 绝对禁止针对每个槽位分别单独单次调用 `officecli_write_slot_value` 进行几十次工具往返！
   - 提交的数据必须是**纯数据值**（绝对不包含前缀标签），底层引擎会自动将冒号后的纯空格/下划线精准替换为该数据值并附带下划线。"""

    system_prompt = f"""你是标书【{role_title}】，负责直接对 Word 标书文档的【{chapter_title}】填报范围进行深度信息检索与原位写盘操作。

【格式样式识别工具 — 所有标书章节均可调用】
- 当需要确认招标原文中被特别标记的强制性要求时，调用 `extract_text_by_style`，优先传入当前 `document_id`，并设置 `chapter_keyword` 为当前章节标题。
- 支持 `italic_underline`（斜体且下划线）、`italic`（斜体）、`underline`（下划线）、`bold`（加粗）和 `bold_red`（红色加粗）。
- 对“实质性要求响应对照表”、偏离表及其他要求清单，必须先核对原文样式标记，再生成响应内容；不得仅凭普通文本推断哪些条款属于重点要求。

【最高铁律 — 原文零改动与高质量填报法则】
1. **模板原文 100% 盲守**：绝对严禁删除、篡改、润色、删减或遗漏任何模板原文（包括前缀标签如“项目名称：”、“招标编号：”、“致：”、标点符号及授权声明等全部固定文本）！
2. **仅精准替换占位符**：只针对模板中的下划线 `______`、括号 `( )`、`[待填]` 槽位填充检索到的真实数据，非占位符的原文必须 100% 原封不动完整保留！
3. **表头感知与逐列精准对齐填报（按表头给出完整提案）**：
   - **认真研读表头定义**：必须根据 User Prompt 中【文档中检测到的实际表格与真实表头定义】，明确当前表格的列数与每一列的中文表头名称；
   - **严格按列装配真实数据**：调用数据库或原文检索工具获取数据后，**必须严格按表头定义的列序逐列对齐装配数据**（每一列分别对应表头名称，严禁错列、漏列、跨列挤压）；
   - **完整提交结构化提案**：必须通过调用 `officecli_fill_table_rows(table_path, rows_json_str)` 工具提交二维数据矩阵，或输出标准 JSON 提案列表提交所有单元格，确保表头下的所有行与列全部完整填满，严禁留下空白单元格！
4. **零容忍任何省略号与伪装标记（严禁 `...` / `……` / `…` / `（完整技术要求）`）**：
   在生成任何条款、响应或表格内容时，**绝对禁止在句子开头、句中连接处或末尾使用任何形式的省略号（包括 `…`、`……`、`...`、`..`）**！
   - 严禁使用省略号截断或连接多个指标；
   - 多个技术指标或分项承诺必须使用中文逗号 `，` 或分号 `；` 完整书写连接，严禁添加『（完整技术要求）』等任何摘要假标签！每一条响应必须是一字一句、语法完整、表述严谨的完整中文法律承诺闭合语句！
5. **【绝对禁止携带原文前缀与标签】**：
   若目标段落/槽位原文本包含字段属性名标签（例如 `"XXX名称：______"` 或 `"XXX编号：______"`），提交的替换数据 `value` **绝对禁止包含"XXX名称："等前缀标签，仅允许填入纯粹的数据值！**
   - 正确写法 (纯数据)：`value = "XXX内容值"`
6. **查无结果立刻止步**：若 DB 工具返回 "未找到..."、"尚未录入" 或空记录，**严禁换用类似关键词重复循环调库**！应当立即将该槽位标记为 "[待补充: <字段名>]"，并直接完成该句/表单写盘。绝对严禁捏造假数据！
7. **【严禁重复盲目查询与批量提交铁律】**：
   - 当前 Worker 负责{structure_scope_rule}，调用 `officecli_query_structure(selector='...')` 一次即可获得该填报范围 100% 完整的全部段落与表格 DOM 结构；
   - **绝对禁止使用不同关键词多次重复调用 `officecli_query_structure`**；
   - 若 User Prompt 中已预先注入了【文档中检测到的实际表格与真实表头定义】（包含 `/body/tbl[N]` 路径及真实表头列名），**严禁再次调用 `officecli_query_structure` 重复扫描结构**；
   - **多槽位必须批量提交**：必须优先使用 `officecli_batch_write_slots` 或 `officecli_fill_table_rows` 一次性提交整章所有槽位/表格行，严禁对每个槽位单独单次调用工具！获取数据后迅速闭环！
8. **【严禁在填报结果中添加任何说明性元数据或保留注释】**：
   - **绝对禁止**在提议内容、扩写结果或总结表格中附加任何类似 `（原文无槽位，零改动保留）`、`（固定原文，零改动保留）`、`（模板固定原文）`、`（零改动）`、`（原样保留）` 等说明性注释！
   - 对于无需修改的固定原文段落，保持原样即可，**严禁在总结表格中伪造扩写结果或将说明性括号写入提案**。

{domain_workflow}

【输出总结格式要求 — 必须包含 Markdown 表格】
在完成所有读写工具调用后，请给出一份操作总结，**必须在总结末尾输出如下格式的 Markdown 明细表格**：
| 序号 | DOM 节点路径 | 替换前模板原文 | 实际填入/扩写结果 | 提议类型 | 写盘状态 |
- 第 3 列 (替换前模板原文)：填入替换前未修饰的原始模板文本（如 `"XXX属性：______"` 或表格单元格原文）；
- 第 4 列 (实际填入/扩写结果)：【纯数据填充】如果提议类型是 `text`，仅允许填写纯数据值；如果是 `image`，填入图片绝对路径；如果是 `sentence_batch`，填入覆盖重写后的完整新段落。**绝对禁止附加任何说明性括号（如“（原文无槽位，零改动保留）”）**；无修改的固定段落无需写入提案表格。严禁使用 `**` 加粗标记；
- 第 5 列 (提议类型)：必须严格填写以下之一："text"、"image"、"sentence_batch"。"""

    if extra_instructions:
        system_prompt += f"""

【用户单章节专属重新生成与微调指令 — 最高优先级】
当前任务为单章节重新生成与覆写，必须与主流程初次生成完全一致：
1. **源头全量检索**：根据具体的表格表头与章节要求，主动调用查原文工具（`get_full_chapter_text`、`search_bidding_document`）或企业数据库工具，从源头完整检索所有条款、技术参数与业务数据；
2. **整表全量覆写**：针对表格章节，必须从序号 1 开始将所检索出的全部条文装配为完整的二维数据矩阵，调用 `officecli_fill_table_rows(table_path, rows_json_str)` 全量覆写原表格，严禁受历史半成品残留影响做局部打补丁！
3. **用户微调要求**：
{extra_instructions}"""

    if repair_instructions:
        system_prompt += f"""

【专项修复紧急指令 — Supervisor 质量审核反馈】
Supervisor 在上一轮审核中发现以下问题。你必须先重新查询当前节点和章节上下文，再判断问题根因：
1. 不得直接相信上一轮提案中的 expected value，也不得把 actual value 当作正确值；
2. 必须重新核对招标原文、企业/项目数据源以及当前 Word 节点的行列语义；
3. 如果目标路径正确但内容被覆盖，提交针对该路径的修复提案；如果目标路径错误，重新定位真实节点后再提交；
4. 如果数据源无法证明应填具体值，保留 `[待补充: 字段]`，不得编造；
5. 写入后必须再次查询目标节点确认实际落位，不能只依据工具返回的“提交成功”判断。

【终审问题明细】
{repair_instructions}"""

    tables_part = f"\n\n【文档中检测到的实际表格与真实表头定义】\n{tables_summary}" if tables_summary else ""

    # 针对公文函件类与开标汇总类章节精准定向按需注入基础企业档案与项目元数据（避免无关章节产生 Token 浪费）
    prefetched_context_part = ""
    if (is_letter_or_form or is_pricing) and prefetched_metadata:
        meta_items = []
        if prefetched_metadata.get("company_name"):
            meta_items.append(f"- 投标人全称: {prefetched_metadata['company_name']}")
        if prefetched_metadata.get("credit_code"):
            meta_items.append(f"- 统一社会信用代码: {prefetched_metadata['credit_code']}")
        if prefetched_metadata.get("legal_person"):
            meta_items.append(f"- 法定代表人: {prefetched_metadata['legal_person']}")
        if prefetched_metadata.get("address"):
            meta_items.append(f"- 注册/通讯地址: {prefetched_metadata['address']}")
        if prefetched_metadata.get("postal_code"):
            meta_items.append(f"- 邮政编码: {prefetched_metadata['postal_code']}")
        if prefetched_metadata.get("phone"):
            meta_items.append(f"- 联系电话: {prefetched_metadata['phone']}")
        if prefetched_metadata.get("fax"):
            meta_items.append(f"- 传真号码: {prefetched_metadata['fax']}")
        if prefetched_metadata.get("email"):
            meta_items.append(f"- 电子邮箱: {prefetched_metadata['email']}")
        if prefetched_metadata.get("project_name"):
            meta_items.append(f"- 投标项目名称: {prefetched_metadata['project_name']}")
        if prefetched_metadata.get("project_code"):
            meta_items.append(f"- 招标/项目编号: {prefetched_metadata['project_code']}")
        if prefetched_metadata.get("total_price_str"):
            w_str = f" (大写: {prefetched_metadata['total_price_words']})" if prefetched_metadata.get("total_price_words") else ""
            meta_items.append(f"- 投标总报价: {prefetched_metadata['total_price_str']}{w_str}")
        if prefetched_metadata.get("delivery_period"):
            meta_items.append(f"- 承诺工期/交货期: {prefetched_metadata['delivery_period']}")
        if prefetched_metadata.get("quality_standard"):
            meta_items.append(f"- 质量标准承诺: {prefetched_metadata['quality_standard']}")

        if meta_items:
            prefetched_context_part = (
                "\n\n【已定向提取的企业主档案与项目关键元数据 — 优先直接使用】\n"
                + "\n".join(meta_items)
                + "\n👉 核心指引：本章节待填的基础数据已在上方完整提供！"
                + "请直接优先收集上方数据，调用 officecli_batch_write_slots 一次性批量提交所有槽位替换提案，无需重复调用数据库查询工具！"
            )

    user_prompt = f"""【撰写任务】
- 文档 ID: {document_id}
- 章节标题: {chapter_title}
- 任务类别: {category}
- 映射标签: {mapping_hint or '通用'}{tables_part}{prefetched_context_part}

【甲方原文模板】
{template_text or '（按招标要求智能撰写）'}

【填写说明】
{content_hint or '（无特殊说明）'}

请根据专家专项工作流开启工具调取与写盘，完成【{chapter_title}】章节的智能撰写。"""

    return system_prompt, user_prompt



# ============================================================
# Worker 执行器
# ============================================================

def run_chapter_worker(
    chapter_title: str,
    chapter_number: str,
    mapping_hint: str,
    category: str,
    document_id: str,
    docx_temp_path: str,
    template_text: str = "",
    content_hint: str = "",
    extra_instructions: str = "",
    repair_instructions: str = "",
    prefetched_metadata: Optional[Dict[str, Any]] = None,
    tenant_id: Optional[str] = None,
) -> Dict[str, Any]:
    """
    为单个章节创建独立 ReAct Agent 并直接执行读写 Word 盘块操作。

    :param extra_instructions: 用户自定义额外指令
    :param repair_instructions: Supervisor 质量审核反馈的专项修复指令
    :param prefetched_metadata: 预读取的企业档案与项目元数据（定向按需注入）
    :param tenant_id: 当前任务所属租户 ID，必须显式传入以支持并发线程安全读取模型配置
    :return: {chapter_title, mapping_hint, status, summary, error}
    """
    cat = (category or "needs_fill").lower().strip()
    logger.info(f"[Worker Direct-Fill] 启动撰写 Agent → [{chapter_title}] (类别: {cat})")
    if repair_instructions:
        logger.warning(f"[Worker 专项修复模式] 接收到 Supervisor 反馈指令: {repair_instructions[:100]}...")

    if cat in ("needs_writing", "skip"):
        logger.info(f"⏩ [Worker] [{chapter_title}] 属于 {cat}，跳过撰写")
        return {
            "chapter_title": chapter_title, "mapping_hint": mapping_hint,
            "category": cat, "status": "skipped",
            "proposals": [], "summary": f"跳过 ({cat})",
        }

    worker_llm = llm_service.get_llm(temperature=0.3, json_mode=False, tenant_id=tenant_id)
    if worker_llm is None:
        return {"chapter_title": chapter_title, "mapping_hint": mapping_hint,
                "category": cat, "status": "failed", "proposals": [],
                "error": "LLM not initialized"}

    try:
        chapter_collected_proposals: List[Dict[str, Any]] = []
        worker_tools = _build_worker_tools(
            docx_temp_path=docx_temp_path,
            chapter_title=chapter_title,
            collected_proposals=chapter_collected_proposals,
            mapping_hint=mapping_hint,
            category=cat
        )
        system_prompt, user_prompt = build_worker_prompt(
            chapter_title=chapter_title, category=cat,
            template_text=template_text, content_hint=content_hint,
            document_id=document_id,
            docx_temp_path=docx_temp_path,
            mapping_hint=mapping_hint,
            extra_instructions=extra_instructions,
            repair_instructions=repair_instructions,
            prefetched_metadata=prefetched_metadata,
        )


        # [优化点1：零度确定性控制] 常规表单与表格清单填写必须无限强行死扣于 `temperature=0.0`；长文本限制于0.2
        target_temp = 0.0 if cat in ("needs_fill", "needs_data", "skip") else 0.2
        worker_llm = llm_service.get_llm(
            temperature=target_temp,
            json_mode=False,
            tenant_id=tenant_id,
        )
        if not worker_llm:
            return {
                "chapter_title": chapter_title,
                "mapping_hint": mapping_hint,
                "category": cat,
                "status": "failed",
                "proposals": [],
                "error": "当前租户未配置可用的大模型",
            }
        logger.info(f"Worker [{chapter_title}] ({cat}) → 分配模型温度 (temperature={target_temp})")

        # 详细打印大模型初始输入（System Prompt 与 User Prompt 概况）
        logger.info(
            f"🚀 [LLM Prompt 准备发送] [{chapter_title}] | System Prompt: {len(system_prompt)} 字符 | "
            f"User Prompt: {len(user_prompt)} 字符 | 工具集数量: {len(worker_tools)}"
        )
        logger.info(f"   📋 [User Prompt 完整内容]:\n{user_prompt[:100]}")
        if extra_instructions:
            logger.info(f"   🎯 [注入的微调提示词]: '{extra_instructions}'")

        agent = create_react_agent(worker_llm, worker_tools)
        import time
        t_start = time.time()

        # 自动重试机制（针对网络波动与大模型 API 连接限流进行容错退避）
        max_retries = 3
        result = None
        for attempt in range(1, max_retries + 1):
            try:
                result = agent.invoke(
                    {"messages": [SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)]},
                    config={"recursion_limit": 50}
                )
                break
            except Exception as err:
                err_str = str(err).lower()
                is_conn_err = any(k in err_str for k in ["connection", "timeout", "reset", "disconnected", "http", "rate", "500", "502", "503", "504"])
                logger.error(f"[Worker 异常调用报错] [{chapter_title}] 第 {attempt} 次失败: {err}")
                if is_conn_err and attempt < max_retries:
                    backoff = attempt * 2.0
                    logger.warning(f"[Worker 网络重试] [{chapter_title}] 等待 {backoff:.1f}s 后自动重试...")
                    time.sleep(backoff)
                else:
                    raise err

        t_end = time.time()
        final_msg = result["messages"][-1].content
        
        # 提取全量 ReAct 中间思考步骤 (Intermediate Thought Steps)
        thought_steps = []
        step_idx = 1
        for msg in result.get("messages", []):
            msg_type = type(msg).__name__
            if msg_type == "AIMessage":
                content = getattr(msg, "content", "")
                tool_calls = getattr(msg, "tool_calls", [])
                if content or tool_calls:
                    thought_steps.append({
                        "step": step_idx,
                        "type": "thought",
                        "thought": content,
                        "tool_calls": tool_calls
                    })
                    step_idx += 1
            elif msg_type == "ToolMessage":
                thought_steps.append({
                    "step": step_idx - 1,
                    "type": "tool_result",
                    "name": getattr(msg, "name", "tool"),
                    "output": str(getattr(msg, "content", ""))[:1500]
                })

        tool_calls_count = sum(1 for m in result["messages"] if hasattr(m, 'tool_calls') and m.tool_calls)

        # 1. 解析文本输出中的提案
        text_proposals = _parse_proposals(final_msg)

        # 2. 双路融合：Tool-First 原则（工具调用捕获的结构化提案为最高权威，绝对禁止被文本概括覆写）
        proposals_dict = {}

        # 先合入文本提取的辅助提案
        for p in text_proposals:
            proposal_key = _proposal_merge_key(p)
            if proposal_key:
                proposals_dict[proposal_key] = p

        # 随后合入工具调用捕获的权威提案（具有最高优先级）
        for p in chapter_collected_proposals:
            proposal_key = _proposal_merge_key(p)
            if proposal_key:
                proposals_dict[proposal_key] = p
                # 不再因为存在整表矩阵就删除同表的单元格提案。
                #
                # Worker 可能只提交“剩余数据行”的局部矩阵，同时文本回复中仍包含
                # 其它行的合法单元格提案。旧逻辑会把这些提案全部删掉，随后 DOM
                # 写入器又会按局部矩阵清空后续行，最终形成“界面内容完整、Word 空格”的
                # 数据丢失链路。整表提案仍然保留优先级，但单元格提案交由最终 DOM
                # 写入阶段按路径覆盖，避免合法数据被提前丢弃。

        proposals = list(proposals_dict.values())
        logger.info(f"   [Worker 提案汇聚] [{chapter_title}] 工具捕获 {len(chapter_collected_proposals)} 条 + 文本解析 {len(text_proposals)} 条 -> 融合去重后共 {len(proposals)} 条有效写盘提案")
        n = len(proposals)

        # 2. 提取 Token 消耗与审计事件记录
        p_tok, c_tok = 0, 0
        for m in result.get("messages", []):
            if hasattr(m, "response_metadata") and isinstance(m.response_metadata, dict):
                usage = m.response_metadata.get("token_usage") or m.response_metadata.get("usage") or {}
                p_tok += usage.get("prompt_tokens", 0)
                c_tok += usage.get("completion_tokens", 0)

        from app.services.audit_service import audit_service
        audit_service.log_event(
            action_type="llm_call_worker",
            node_name=f"BidFillerWorker-{chapter_title}",
            inputs={"chapter_title": chapter_title, "category": cat, "document_id": document_id},
            outputs={
                "proposals_count": n,
                "proposals": proposals,
                "summary": final_msg,
                "thought_steps": thought_steps
            },
            prompt_tokens=p_tok,
            completion_tokens=c_tok,
            execution_time_ms=int((t_end - t_start) * 1000),
            status="success"
        )

        # 记录 Worker 完整诊断上下文（供导出日志排查）
        _record_worker_context(
            doc_id=document_id,
            chapter_title=chapter_title,
            category=cat,
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            final_msg=final_msg,
            tool_calls=tool_calls_count,
            proposals=proposals
        )

        # 存入全局提案池（供 Review Agent 读取）
        if proposals:
            with _PROPOSALS_LOCK:
                if document_id not in _WORKER_PROPOSALS:
                    _WORKER_PROPOSALS[document_id] = []
                for p in proposals:
                    p["chapter_title"] = chapter_title
                    p["mapping_hint"] = mapping_hint
                _WORKER_PROPOSALS[document_id].extend(proposals)

        logger.info(
            f"[Worker Agent 完成] [{chapter_title}] | 耗时: {int((t_end - t_start) * 1000)}ms | "
            f"工具调用: {tool_calls_count} 次 | 提案: {n} 个 | "
            f"Prompt Tokens: {p_tok:,} | Completion Tokens: {c_tok:,} | Total: {p_tok + c_tok:,}"
        )
        return {
            "chapter_title": chapter_title, "mapping_hint": mapping_hint,
            "category": cat, "status": "success",
            "tool_calls": tool_calls_count,
            "proposals_count": n,
            "proposals": proposals,
            "summary": final_msg,
        }

    except Exception as e:
        logger.error(f"[Worker] [{chapter_title}] 失败: {e}")
        try:
            from app.services.audit_service import audit_service
            audit_service.log_event(
                action_type="llm_call_worker",
                node_name=f"BidFillerWorker-{chapter_title}",
                inputs={"chapter_title": chapter_title, "category": cat, "document_id": document_id},
                outputs={
                    "summary": f"Worker 章节填报发生异常: {str(e)[:200]}",
                    "error": str(e)[:500]
                },
                status="failed"
            )
        except Exception as log_err:
            logger.warning(f"写入 Worker 异常审计日志失败: {log_err}")

        return {
            "chapter_title": chapter_title, "mapping_hint": mapping_hint,
            "category": cat, "status": "failed", "proposals": [],
            "error": str(e)[:500],
        }


def _normalize_proposal_item(item: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    """归一化单个提案对象，统一对齐 proposed_text 与 value 键名，并过滤虚假概括行与零改动无操作提案"""
    if not isinstance(item, dict):
        return None
    import re
    from app.agents.review_engine import clean_zero_change_annotations, is_zero_change_or_no_op_proposal

    path = str(item.get("path", "")).strip().replace("`", "")
    if not path:
        return None
    # 过滤包含 ~ 或 .. 的概括性假路径 (如 /body/tbl[2]/tr[2]~tr[17] 或 /body/tbl[2]/tr[2..17])
    if "~" in path or ".." in path:
        return None

    # 提取有效文本内容：优先 proposed_text，次选 value，再选 text
    val = item.get("proposed_text")
    if val is None:
        val = item.get("value")
    if val is None:
        val = item.get("text")
    if val is None:
        val = ""

    val_str = str(val).strip().replace("`", "").replace("**", "")
    prop_type = str(item.get("type", "")).strip()

    # 剥离说明性元数据注释
    if prop_type != "image":
        val_str = clean_zero_change_annotations(val_str)

    # 过滤零改动无操作提案
    orig_ctx = str(item.get("original_context", "")).strip()
    if is_zero_change_or_no_op_proposal(val_str, orig_ctx, prop_type):
        return None

    # 若内容为空且非特殊类型，丢弃
    if not val_str and prop_type != "image":
        return None

    res = dict(item)
    res["path"] = path
    res["proposed_text"] = val_str
    res["value"] = val_str
    if "status" not in res:
        res["status"] = "success"
    return res


def _repair_json_unescaped_quotes(json_str: str) -> str:
    """自动对 JSON 字符串值中未经转义的半角双引号进行容错替换"""
    import re
    def fix_field_val(m):
        prefix = m.group(1)   # `"reasoning": "`
        content = m.group(2)  # `原文"招标编号..."`
        suffix = m.group(3)   # `"`
        fixed_content = content.replace('"', '”')
        return f'{prefix}{fixed_content}{suffix}'
    pattern = r'("(?:path|original_context|source_data|source_tool|proposed_text|value|text|reasoning|chapter_title|mapping_hint)"\s*:\s*")([\s\S]*?)("\s*[,\}])'
    return re.sub(pattern, fix_field_val, json_str)


def _parse_proposals(raw_text: str) -> List[Dict[str, Any]]:
    """从 Worker 的最终回复中提取 JSON 提案列表（包含多重智能容错、双向键名归一化与融合提取机制）"""
    import re
    if not raw_text:
        return []
    cleaned = raw_text.strip()
    raw_json_candidates = []

    # 策略1: 提取 ```json ... ``` 代码块
    if "```" in cleaned:
        for match in re.finditer(r'```(?:json)?\s*\n?([\s\S]*?)\n?```', cleaned):
            cand_str = match.group(1).strip()
            if cand_str:
                raw_json_candidates.append(cand_str)

    # 策略2: 直接解析全文
    raw_json_candidates.append(cleaned)

    # 策略3: 从混合文本中提取 JSON 数组（贪婪匹配）
    for pattern in [r'\[\s*\{[\s\S]*\}\s*\]', r'\[\s*\{[\s\S]*?\}\s*\]']:
        m = re.search(pattern, cleaned)
        if m:
            raw_json_candidates.append(m.group(0))

    parsed_json_items: List[Dict[str, Any]] = []
    for cand in raw_json_candidates:
        if not cand:
            continue
        # 尝试标准解析
        try:
            data = _json.loads(cand)
            if isinstance(data, list):
                parsed_json_items = data
                break
            if isinstance(data, dict) and "proposals" in data and isinstance(data["proposals"], list):
                parsed_json_items = data["proposals"]
                break
        except Exception:
            pass

        # 尝试修复尾部逗号与未转义双引号后解析
        try:
            fixed = re.sub(r',\s*\]', ']', cand)
            fixed = re.sub(r',\s*\}', '}', fixed)
            data = _json.loads(fixed)
            if isinstance(data, list):
                parsed_json_items = data
                break
        except Exception:
            pass

        try:
            repaired = _repair_json_unescaped_quotes(cand)
            repaired = re.sub(r',\s*\]', ']', repaired)
            repaired = re.sub(r',\s*\}', '}', repaired)
            data = _json.loads(repaired)
            if isinstance(data, list):
                parsed_json_items = data
                break
        except Exception:
            pass

    # 策略4: 对象级逐项恢复机制
    if not parsed_json_items:
        obj_matches = re.finditer(r'\{\s*"path"\s*:[\s\S]*?\}', raw_text)
        for m in obj_matches:
            candidate = m.group(0)
            try:
                item = _json.loads(candidate)
                if isinstance(item, dict) and "path" in item:
                    parsed_json_items.append(item)
                    continue
            except Exception:
                pass
            try:
                repaired_cand = _repair_json_unescaped_quotes(candidate)
                item = _json.loads(repaired_cand)
                if isinstance(item, dict) and "path" in item:
                    parsed_json_items.append(item)
            except Exception:
                pass

    # 策略5: 从 Markdown 表格行提取提案明细（补充提取与融合）
    markdown_table_items: List[Dict[str, Any]] = []
    tbl_row_counters: Dict[str, int] = {}

    if "|" in raw_text:
        for line in raw_text.split("\n"):
            l_str = line.strip()
            if not l_str.startswith("|") or "---" in l_str or "序号" in l_str or "DOM 节点" in l_str or "替换前" in l_str or "检索结果" in l_str or "数据项" in l_str:
                continue
            # 跳过仅作为保护声明的表头保护行
            if any(k in l_str for k in ["表头原样保留", "已保护", "未改动", "未修改", "原样保留"]):
                continue
            cells = [c.replace("`", "").replace("**", "").strip() for c in l_str.split("|") if c.strip()]
            if len(cells) >= 3:
                path_col_idx = -1
                for idx, cell in enumerate(cells):
                    if re.search(r'(/[a-zA-Z0-9\[\]\@\=\:\-\_\.\~]+)+', cell):
                        path_col_idx = idx
                        break

                if path_col_idx == -1:
                    continue

                path_val = cells[path_col_idx]
                path_m = re.search(r'(/[a-zA-Z0-9\[\]\@\=\:\-\_\.\~]+)+', path_val)
                if path_m:
                    path_val = path_m.group(0)

                # 智能自愈：若大模型在 Markdown 表格中使用了范围概括路径 (如 /body/tbl[6]/tr[2]~tr[19] 或 /body/tbl[6]/tr[2..19])
                # 自动按表格出现顺序展开为连续递增的物理行路径 tr[2], tr[3], tr[4]... 彻底杜绝多行数据被误删丢弃！
                m_range = re.search(r'(/body/tbl\[\d+\])(?:/tr\[\d+\])?(?:~|\.\.)tr\[(\d+)\]', path_val)
                m_base_tbl = re.search(r'(/body/tbl\[\d+\])', path_val)
                if m_range or ("~" in path_val or ".." in path_val):
                    tbl_base = m_range.group(1) if m_range else (m_base_tbl.group(1) if m_base_tbl else "/body/tbl[1]")
                    curr_r = tbl_row_counters.get(tbl_base, 2)
                    path_val = f"{tbl_base}/tr[{curr_r}]"
                    tbl_row_counters[tbl_base] = curr_r + 1
                elif m_base_tbl and "/tr[" in path_val and "/tc[" not in path_val:
                    # 若连续多行使用相同的单行路径 (如都是 /body/tbl[6]/tr[2])，自动递增行号
                    tbl_base = m_base_tbl.group(1)
                    curr_r = tbl_row_counters.get(tbl_base, 2)
                    path_val = f"{tbl_base}/tr[{curr_r}]"
                    tbl_row_counters[tbl_base] = curr_r + 1
                elif m_base_tbl:
                    tbl_base = m_base_tbl.group(1)
                    m_tr_num = re.search(r'/tr\[(\d+)\]', path_val)
                    if m_tr_num:
                        tbl_row_counters[tbl_base] = max(tbl_row_counters.get(tbl_base, 2), int(m_tr_num.group(1)) + 1)

                # 自动自愈尾部缺失右括号
                path_val = re.sub(r'/tc\[(\d+)$', r'/tc[\1]', path_val)
                path_val = re.sub(r'/cell\[(\d+)$', r'/cell[\1]', path_val)
                path_val = re.sub(r'/tr\[(\d+)$', r'/tr[\1]', path_val)

                orig_val = cells[path_col_idx + 1] if path_col_idx + 1 < len(cells) else ""
                prop_val = cells[path_col_idx + 2] if path_col_idx + 2 < len(cells) else ""

                prop_type = "text"
                is_cell_path = bool(re.search(r'/(?:tr|row|tc|cell)\[\d+\]', path_val))
                if path_col_idx + 3 < len(cells):
                    t_str = cells[path_col_idx + 3].lower()
                    if "image" in t_str or "图片" in t_str:
                        prop_type = "image"
                    elif "sentence_batch" in t_str or "覆盖" in t_str:
                        prop_type = "sentence_batch"
                    elif ("table_rows" in t_str or "插行" in t_str) and not is_cell_path:
                        prop_type = "table_rows"

                prop_val = prop_val.replace("**", "").replace("`", "").strip()

                # 通用规则 1：过滤大模型在总结时输出的运算符伪拼接表达式 (如 'A + B + C' 或 '"A" + "B"')
                if re.search(r'(?:\S\s*\+\s*\S|\"\s*\+\s*\")', prop_val):
                    logger.warning(f"   [通用提案过滤] 拦截到伪拼接表达式文本，拒绝提取: path={path_val}, val={prop_val[:50]}")
                    continue

                # 通用规则 2：剥离零改动等说明性注释，并过滤纯元数据状态词或无操作行
                from app.agents.review_engine import clean_zero_change_annotations, is_zero_change_or_no_op_proposal
                if prop_type != "image":
                    prop_val = clean_zero_change_annotations(prop_val)

                if is_zero_change_or_no_op_proposal(prop_val, orig_val, prop_type):
                    continue

                if prop_val in ("—", "-", "--", "同上", "略", "无变更", "原样保留") or re.match(r'^(?:已|未)[^\s]{1,6}(?:队列|保护|修改|变更|填报)$', prop_val):
                    continue

                # 通用规则 3：表格路径与类型严格归一化
                # 如果是整表路径 (/body/tbl[N])，必须能够成功反序列化为合法的 2D 矩阵 (list of lists)，否则判定为无效表格提案
                if bool(re.search(r'^/body/tbl\[\d+\]$', path_val)):
                    try:
                        parsed_m = _json.loads(prop_val)
                        if isinstance(parsed_m, list) and parsed_m and isinstance(parsed_m[0], list):
                            prop_type = "table_rows"
                        else:
                            continue
                    except Exception:
                        continue
                elif bool(re.search(r'/tbl\[\d+\]/(?:tr|row)\[\d+\]$', path_val)):
                    # 若非 sentence_batch 且不含多列分隔符，过滤行级模糊单文本路径防止错列
                    if prop_type not in ("sentence_batch", "table_rows") and not any(s in prop_val for s in ["|", "｜", "\t"]):
                        logger.warning(f"   [通用提案过滤] 行级路径未指定具体单元格 (tc/cell)，跳过单文本提炼: {path_val}")
                        continue

                markdown_table_items.append({
                    "path": path_val,
                    "original_context": orig_val,
                    "proposed_text": prop_val,
                    "value": prop_val,
                    "type": prop_type,
                    "status": "success"
                })

    # 6. 双向归一化与融合去重（优先保留更精确的提议类型与上下文）
    normalized_dict: Dict[str, Dict[str, Any]] = {}

    for it in parsed_json_items:
        norm = _normalize_proposal_item(it)
        if norm:
            p_key = norm["path"]
            if p_key not in normalized_dict:
                normalized_dict[p_key] = norm

    # 融合 Markdown 表格提取项（补充缺失行，并提升 JSON 中可能被泛化的提议类型）
    for it in markdown_table_items:
        norm = _normalize_proposal_item(it)
        if norm:
            p_key = norm["path"]
            if p_key not in normalized_dict:
                normalized_dict[p_key] = norm
            else:
                existing = normalized_dict[p_key]
                # 若 JSON 项为通用 "text" 且 Markdown 表格标记了更具体的 sentence_batch / table_rows / image，提升之
                if existing.get("type") in ("text", "", None) and norm.get("type") in ("sentence_batch", "table_rows", "image"):
                    existing["type"] = norm.get("type")
                if not existing.get("original_context") and norm.get("original_context"):
                    existing["original_context"] = norm.get("original_context")

    normalized_list = list(normalized_dict.values())

    if normalized_list:
        logger.info(f"   [Worker 提案提炼成功] 归一化并融合产出 {len(normalized_list)} 条合法提案明细")
        return normalized_list

    logger.warning(f"   [Worker] 无法解析提案 JSON ({len(raw_text)} 字符):\n{raw_text}...")
    return []


# ============================================================
# Worker 上下文日志导出管理
# ============================================================
_WORKER_CONTEXT_LOGS: Dict[str, List[Dict[str, Any]]] = {}
_CONTEXT_LOCK = _threading.Lock()

def _record_worker_context(
    doc_id: str,
    chapter_title: str,
    category: str,
    system_prompt: str,
    user_prompt: str,
    final_msg: str,
    tool_calls: int,
    proposals: List[Dict[str, Any]]
):
    with _CONTEXT_LOCK:
        if doc_id not in _WORKER_CONTEXT_LOGS:
            _WORKER_CONTEXT_LOGS[doc_id] = []
        _WORKER_CONTEXT_LOGS[doc_id].append({
            "chapter_title": chapter_title,
            "category": category,
            "system_prompt": system_prompt,
            "user_prompt": user_prompt,
            "final_msg": final_msg,
            "tool_calls": tool_calls,
            "proposals_count": len(proposals),
            "proposals": proposals,
        })

def export_worker_context_log(doc_id: str) -> str:
    """将指定文档的所有 Worker 子 Agent 运行时上下文导出为独立的 Markdown 审计日志文件"""
    with _CONTEXT_LOCK:
        logs = _WORKER_CONTEXT_LOGS.get(doc_id, [])

    output_dir = os.path.join(os.getcwd(), "outputs", "human_fill_results")
    os.makedirs(output_dir, exist_ok=True)
    report_file = os.path.join(output_dir, f"worker_agent_context_log_{doc_id[:8]}.md")

    md_lines = [
        f"# Worker 子 Agent 运行时完整上下文诊断报告",
        f"- **文档 ID**: `{doc_id}`",
        f"- **已完成 Worker 数**: `{len(logs)}`",
        f"- **生成时间**: `{time.strftime('%Y-%m-%d %H:%M:%S')}`",
        "",
        "---",
        ""
    ]

    for idx, ctx in enumerate(logs, 1):
        md_lines.append(f"## [{idx}] Worker 章节: {ctx['chapter_title']} (类别: {ctx['category']})")
        md_lines.append(f"- **工具调用次数**: `{ctx['tool_calls']}`")
        md_lines.append(f"- **产出提案数**: `{ctx['proposals_count']}`")
        md_lines.append("")
        md_lines.append("### 1. System Prompt (系统提示词)")
        md_lines.append("```text")
        md_lines.append(ctx['system_prompt'])
        md_lines.append("```")
        md_lines.append("")
        md_lines.append("### 2. User Prompt (用户任务与模板输入)")
        md_lines.append("```text")
        md_lines.append(ctx['user_prompt'])
        md_lines.append("```")
        md_lines.append("")
        md_lines.append("### 3. ReAct LLM 终端回复原文")
        md_lines.append("```text")
        md_lines.append(ctx['final_msg'])
        md_lines.append("```")
        md_lines.append("")
        md_lines.append("### 4. 最终提炼的 Proposals 提案清单")
        md_lines.append("```json")
        md_lines.append(_json.dumps(ctx['proposals'], ensure_ascii=False, indent=2))
        md_lines.append("```")
        md_lines.append("")
        md_lines.append("---")
        md_lines.append("")

    content = "\n".join(md_lines)
    with open(report_file, "w", encoding="utf-8") as f:
        f.write(content)

    logger.info(f"   📄 [Worker 上下文诊断日志已导出]: {report_file}")
    return report_file
