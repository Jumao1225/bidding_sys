"""
Review Engine — 标书填报深度质检引擎

包含 5 条规则化审查管线（纯 Python，零 LLM 消耗）+ 1 个 LLM 自动修正管线。
所有管线输出统一的 ReviewFinding 数据结构，供 review_node 消费。

管线清单：
  R1: check_unfilled_slots      — 必填项遗漏检测（扫描 Word 文档中残留的占位符）
  R2: check_data_consistency    — 数据交叉一致性校验（同字段跨章节填写值是否一致）
  R3: check_financial_accuracy  — 财务数据准确性校验（总价 vs 分项求和、大写 vs 小写）
  R4: check_format_compliance   — 格式合规性校验（日期格式、金额格式、公司名完整性）
  R5: check_scoring_coverage    — 评分项覆盖率检查（与 score_tree 做结构化详细比对）
  R6: auto_fix_proposals        — LLM 自动修正管线（对问题 proposals 进行修正）
"""

import os
import re
import json
from typing import Dict, Any, List, Optional, Tuple
from loguru import logger
from docx import Document

from app.utils.rmb_formatter import number_to_chinese_rmb


# ============================================================
# ReviewFinding 数据结构（纯 dict，避免循环依赖）
# ============================================================

def make_finding(
    rule_id: str,
    severity: str,
    path: str,
    description: str,
    current_value: str = "",
    expected_value: str = "",
    auto_fixable: bool = False,
    fix_proposal: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """
    构造统一的审查发现字典。

    :param rule_id: 管线 ID，如 "R1-UNFILLED"
    :param severity: 严重级别 "error" | "warning" | "info"
    :param path: 问题所在的 Word XPath
    :param description: 人类可读的问题描述
    :param current_value: 当前填写值
    :param expected_value: 期望值（如有）
    :param auto_fixable: 是否可自动修正
    :param fix_proposal: 修正后的 proposal（如有）
    :return: ReviewFinding 字典
    """
    return {
        "rule_id": rule_id,
        "severity": severity,
        "path": path,
        "description": description,
        "current_value": current_value,
        "expected_value": expected_value,
        "auto_fixable": auto_fixable,
        "fix_proposal": fix_proposal,
    }


# ============================================================
# R1: 必填项遗漏检测
# ============================================================

# 匹配连续 3 个及以上下划线、[待补充...]、[建议人工...]、[查询...] 等占位符
_UNFILLED_PATTERNS = [
    re.compile(r'_{3,}'),                        # 连续下划线占位符
    re.compile(r'\[待补充[^\]]*\]'),              # [待补充] / [待补充: xxx]
    re.compile(r'\[建议人工[^\]]*\]'),            # [建议人工...]
    re.compile(r'\[查询[^\]]*\]'),                # [查询异常...]
    re.compile(r'\[错误[^\]]*\]'),                # [错误...]
    re.compile(r'\[待手动补充[^\]]*\]'),          # [待手动补充: xxx]
]


def check_unfilled_slots(
    docx_path: str,
    proposals: Optional[List[Dict[str, Any]]] = None,
) -> List[Dict[str, Any]]:
    """
    R1: 必填项遗漏检测 — 扫描 Word 文档中残留的占位符与空白单元格。

    扫描范围：所有正文段落 + 所有表格单元格。
    检测目标：连续下划线 ___、[待补充]、[建议人工]、空白单元格。

    :param docx_path: Word 文档路径
    :param proposals: Worker 提案列表（可选，用于标记哪些字段被跳过）
    :return: ReviewFinding 列表
    """
    logger.info("🔍 [R1] 启动必填项遗漏检测...")
    findings: List[Dict[str, Any]] = []

    try:
        doc = Document(docx_path)
    except Exception as exc:
        logger.error(f"🔍 [R1] 无法打开文档 {docx_path}: {exc}")
        return [make_finding(
            rule_id="R1-UNFILLED", severity="error", path="N/A",
            description=f"无法打开文档进行遗漏检测: {exc}",
        )]

    # 扫描正文段落
    for idx, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        if not text:
            continue
        xpath = f"/body/p[{idx}]"
        for pattern in _UNFILLED_PATTERNS:
            if pattern.search(text):
                findings.append(make_finding(
                    rule_id="R1-UNFILLED", severity="warning", path=xpath,
                    description=f"正文段落残留占位符: '{text[:80]}'",
                    current_value=text[:100],
                    auto_fixable=True,
                ))
                break  # 每个段落只报一次

    # 扫描表格单元格 (按 XML <w:tc> 物理节点遍历，防止合并单元格产生错位)
    for tbl_idx, table in enumerate(doc.tables, start=1):
        for row_idx, row in enumerate(table.rows, start=1):
            tc_elements = [c for c in row._tr.iterchildren() if c.tag.endswith('tc')]
            for col_idx, tc_elem in enumerate(tc_elements, start=1):
                cell_text = "".join(tc_elem.itertext()).strip()
                xpath = f"/body/tbl[{tbl_idx}]/tr[{row_idx}]/tc[{col_idx}]"

                # 检查空白单元格（跳过表头行，即第一行）
                if row_idx > 1 and not cell_text:
                    findings.append(make_finding(
                        rule_id="R1-EMPTY-CELL", severity="info", path=xpath,
                        description=f"表格第{tbl_idx}个表 第{row_idx}行 第{col_idx}列为空白单元格",
                        auto_fixable=True,
                    ))
                    continue

                # 检查占位符残留
                for pattern in _UNFILLED_PATTERNS:
                    if pattern.search(cell_text):
                        findings.append(make_finding(
                            rule_id="R1-UNFILLED", severity="warning", path=xpath,
                            description=f"表格单元格残留占位符: '{cell_text[:60]}'",
                            current_value=cell_text[:100],
                            auto_fixable=True,
                        ))
                        break

    logger.info(f"🔍 [R1] 必填项遗漏检测完成，发现 {len(findings)} 个问题")
    return findings


# ============================================================
# R2: 数据交叉一致性校验
# ============================================================

# 需要跨章节保持一致的关键字段（source_tool + 语义关键词）
_CONSISTENCY_FIELDS = {
    "company_name": ["公司名称", "投标人名称", "投标人", "单位名称"],
    "legal_representative": ["法定代表人", "法人代表", "法人"],
    "project_name": ["项目名称", "工程名称"],
    "project_code": ["项目编号", "招标编号"],
    "total_price": ["投标总价", "总报价", "投标金额"],
    "credit_code": ["统一社会信用代码", "信用代码"],
}


def _select_best_clean_value(field_key: str, values: List[str]) -> str:
    """智能选出最合规、最纯净的字段值（自动剔除超长段落污染与问答语气词）"""
    if not values:
        return ""

    clean_candidates = []
    for val in values:
        c = re.sub(r'^[^:：]*[:：]\s*', '', val).strip()
        # 剔除超长文本与带有段落句式动词的污染值
        if len(c) > 35 or any(bad in c for bad in ["根据", "授权", "招标文件", "处理", "事宜", "特此"]):
            continue
        # 项目名称、公司名称、法人姓名拒绝纯数字/纯金额
        if field_key in ("project_name", "company_name", "legal_representative"):
            if re.match(r'^\d+(\.\d+)?$', c):
                continue
        clean_candidates.append(c)

    if clean_candidates:
        # 优先选择最短且符合常理的实体名称
        return min(clean_candidates, key=len)

    # 若全被污染，尝试正则捕获真实实体名（如 XXX有限公司）
    for val in values:
        m = re.search(r'([\u4e00-\u9fa5]{2,30}(?:有限责任公司|股份有限公司|有限公司|集团|公司))', val)
        if m:
            return m.group(1)

    return re.sub(r'^[^:：]*[:：]\s*', '', values[0]).strip()


def check_data_consistency(
    proposals: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """
    R2: 数据交叉一致性校验 — 同一字段在不同章节的填写值是否一致。

    从所有 proposals 中按语义分组，比对同一语义字段在不同 path 的 proposed_text。

    :param proposals: Worker 提案列表
    :return: ReviewFinding 列表
    """
    logger.info("🔍 [R2] 启动数据交叉一致性校验...")
    findings: List[Dict[str, Any]] = []

    if not proposals:
        return findings

    # 按语义字段分组收集提案
    field_groups: Dict[str, List[Dict[str, Any]]] = {}
    for p in proposals:
        text = str(p.get("proposed_text", "")).strip()
        orig = str(p.get("original_context", "")).strip()
        source = str(p.get("source_tool", "")).strip()

        # 跳过无效提案
        if not text or text.startswith("["):
            continue

        # 尝试匹配到语义字段
        combined = f"{orig} {source} {text}".lower()
        for field_key, keywords in _CONSISTENCY_FIELDS.items():
            if any(kw in combined for kw in keywords):
                if field_key not in field_groups:
                    field_groups[field_key] = []
                field_groups[field_key].append(p)
                break

    # 检查每个字段组内的一致性
    for field_key, group in field_groups.items():
        if len(group) < 2:
            continue

        # 提取所有去除前缀标签后的纯值
        values_map: Dict[str, List[str]] = {}
        for p in group:
            raw_text = str(p.get("proposed_text", "")).strip()
            # 去除标签前缀（如 "投标人名称（盖章）：" ），保留纯值
            clean = re.sub(r'^[^:：]*[:：]\s*', '', raw_text).strip()
            if not clean:
                clean = raw_text
            path = str(p.get("path", ""))
            if clean not in values_map:
                values_map[clean] = []
            values_map[clean].append(path)

        # 如果只有一种值，说明一致
        if len(values_map) <= 1:
            continue

        # 发现不一致
        all_values = list(values_map.keys())
        all_paths = [path for paths in values_map.values() for path in paths]
        # 选出最佳合法纯净基准值，拒绝段落大长文基准
        best_correct_value = _select_best_clean_value(field_key, all_values)

        findings.append(make_finding(
            rule_id="R2-INCONSISTENT", severity="error",
            path="; ".join(all_paths[:5]),
            description=f"字段 [{field_key}] 在不同章节填写值不一致: {all_values}",
            current_value=json.dumps(all_values, ensure_ascii=False),
            expected_value=best_correct_value,
            auto_fixable=True,
            fix_proposal={"field_key": field_key, "correct_value": best_correct_value},
        ))

    logger.info(f"🔍 [R2] 数据交叉一致性校验完成，发现 {len(findings)} 个问题")
    return findings


# ============================================================
# R3: 财务数据准确性校验
# ============================================================

def check_financial_accuracy(
    proposals: List[Dict[str, Any]],
    document_id: str,
) -> List[Dict[str, Any]]:
    """
    R3: 财务数据准确性校验 — 验证总价与分项之和的一致性、大写金额的正确性。

    校验维度：
    1. 投标总价（小写）是否等于所有分项报价之和
    2. 投标总价（大写）是否与小写金额一致
    3. 大写金额格式是否符合人民币大写规范

    :param proposals: Worker 提案列表
    :param document_id: 文档 ID
    :return: ReviewFinding 列表
    """
    logger.info("🔍 [R3] 启动财务数据准确性校验...")
    findings: List[Dict[str, Any]] = []

    # 从 DB 查询实际的分项报价数据
    try:
        from app.db.session import SessionLocal
        from app.db.models.ai_analysis import CostEstimate

        db = SessionLocal()
        try:
            cost_items = db.query(CostEstimate).filter(
                CostEstimate.document_id == document_id
            ).all()

            if not cost_items:
                logger.info("🔍 [R3] 无分项报价数据，跳过财务校验")
                return findings

            # 计算分项报价之和
            db_total = sum(item.calculated_total for item in cost_items)
        finally:
            db.close()
    except Exception as exc:
        logger.warning(f"🔍 [R3] 查询分项报价异常: {exc}")
        return findings

    # 从 proposals 中找到总价相关的提案
    total_price_proposal = None
    chinese_price_proposal = None

    for p in proposals:
        text = str(p.get("proposed_text", "")).strip()
        orig = str(p.get("original_context", "")).strip()
        source = str(p.get("source_tool", "")).strip()
        combined = f"{orig} {source}".lower()

        # 匹配小写总价
        if any(kw in combined for kw in ["投标总价", "总报价", "投标金额", "total_price_numeric"]):
            # 尝试提取数字值
            num_match = re.search(r'[\d,]+\.?\d*', text.replace(",", ""))
            if num_match:
                total_price_proposal = {
                    "value": float(num_match.group().replace(",", "")),
                    "path": p.get("path", ""),
                    "text": text,
                }

        # 匹配大写总价
        if any(kw in combined for kw in ["大写", "chinese", "汉字大写"]):
            chinese_price_proposal = {
                "value": text,
                "path": p.get("path", ""),
            }

    # 校验 1: 总价与分项之和
    if total_price_proposal:
        proposal_total = total_price_proposal["value"]
        # 允许 0.01 元的浮点误差
        if abs(proposal_total - db_total) > 0.01:
            findings.append(make_finding(
                rule_id="R3-SUM-MISMATCH", severity="error",
                path=total_price_proposal["path"],
                description=f"投标总价 ({proposal_total:.2f}) 与分项报价之和 ({db_total:.2f}) 不一致，差额: {abs(proposal_total - db_total):.2f} 元",
                current_value=f"{proposal_total:.2f}",
                expected_value=f"{db_total:.2f}",
                auto_fixable=True,
                fix_proposal={
                    "path": total_price_proposal["path"],
                    "proposed_text": f"{db_total:.2f}",
                    "original_context": total_price_proposal["text"],
                    "source_tool": "review_engine_fix",
                    "source_data": f"分项报价之和: {db_total:.2f}",
                    "reasoning": f"R3 自动修正: 投标总价修正为分项报价之和 {db_total:.2f}",
                },
            ))

    # 校验 2: 大写金额与小写金额一致性
    if chinese_price_proposal and total_price_proposal:
        expected_chinese = number_to_chinese_rmb(total_price_proposal["value"])
        actual_chinese = chinese_price_proposal["value"]
        # 清理对比（去除空格、标点差异）
        clean_expected = re.sub(r'\s+', '', expected_chinese)
        clean_actual = re.sub(r'\s+', '', actual_chinese)
        if clean_expected != clean_actual:
            findings.append(make_finding(
                rule_id="R3-CHINESE-MISMATCH", severity="error",
                path=chinese_price_proposal["path"],
                description=f"大写金额 '{actual_chinese}' 与小写金额 '{total_price_proposal['value']:.2f}' 换算不一致",
                current_value=actual_chinese,
                expected_value=expected_chinese,
                auto_fixable=True,
                fix_proposal={
                    "path": chinese_price_proposal["path"],
                    "proposed_text": expected_chinese,
                    "original_context": actual_chinese,
                    "source_tool": "review_engine_fix",
                    "source_data": f"小写金额: {total_price_proposal['value']:.2f}",
                    "reasoning": f"R3 自动修正: 大写金额修正为 {expected_chinese}",
                },
            ))

    logger.info(f"🔍 [R3] 财务数据准确性校验完成，发现 {len(findings)} 个问题")
    return findings


# ============================================================
# R4: 格式合规性校验
# ============================================================

# 日期格式正则：YYYY年MM月DD日
_DATE_PATTERN = re.compile(r'\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日')
# 非法日期格式（如 YYYY-MM-DD、YYYY/MM/DD）
_BAD_DATE_PATTERNS = [
    re.compile(r'\d{4}[-/]\d{1,2}[-/]\d{1,2}'),    # YYYY-MM-DD 或 YYYY/MM/DD
    re.compile(r'\d{1,2}/\d{1,2}/\d{4}'),            # MM/DD/YYYY
]

# 公司名称应包含的法律实体后缀
_COMPANY_SUFFIXES = ["有限公司", "有限责任公司", "股份有限公司", "集团", "合伙企业"]


def check_format_compliance(
    proposals: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """
    R4: 格式合规性校验 — 校验日期格式、金额格式、公司名完整性。

    校验维度：
    1. 日期应为 YYYY年MM月DD日 格式，不应出现 YYYY-MM-DD
    2. 公司名称应包含法律实体后缀（有限公司 等）
    3. 金额数值不应有异常格式

    :param proposals: Worker 提案列表
    :return: ReviewFinding 列表
    """
    logger.info("🔍 [R4] 启动格式合规性校验...")
    findings: List[Dict[str, Any]] = []

    for p in proposals:
        text = str(p.get("proposed_text", "")).strip()
        orig = str(p.get("original_context", "")).strip()
        path = str(p.get("path", ""))

        if not text or text.startswith("["):
            continue

        combined_ctx = f"{orig}".lower()

        # 校验 1: 日期格式
        if any(kw in combined_ctx for kw in ["日期", "时间", "年", "月", "日", "截止", "deadline"]):
            for bad_pat in _BAD_DATE_PATTERNS:
                m = bad_pat.search(text)
                if m:
                    # 尝试转换为标准中文日期格式
                    raw_date = m.group()
                    parts = re.split(r'[-/]', raw_date)
                    if len(parts) == 3:
                        fixed = f"{parts[0]}年{int(parts[1])}月{int(parts[2])}日"
                        findings.append(make_finding(
                            rule_id="R4-DATE-FORMAT", severity="warning",
                            path=path,
                            description=f"日期格式不规范: '{raw_date}'，标书应使用 'YYYY年MM月DD日' 格式",
                            current_value=text,
                            expected_value=text.replace(raw_date, fixed),
                            auto_fixable=True,
                            fix_proposal={
                                "path": path,
                                "proposed_text": text.replace(raw_date, fixed),
                                "original_context": orig,
                                "source_tool": "review_engine_fix",
                                "source_data": text,
                                "reasoning": f"R4 自动修正: 日期格式规范化 '{raw_date}' -> '{fixed}'",
                            },
                        ))

        # 校验 2: 公司名完整性
        if any(kw in combined_ctx for kw in ["公司名", "投标人名", "单位名", "company_name"]):
            # 检查填写值中是否包含公司法律实体后缀
            if text and not any(suffix in text for suffix in _COMPANY_SUFFIXES):
                # 排除非公司名字段（如 "XXX姓名" 法定代表人）
                if len(text) > 4 and not any(kw in combined_ctx for kw in ["法定代表人", "法人", "代理人"]):
                    findings.append(make_finding(
                        rule_id="R4-COMPANY-NAME", severity="warning",
                        path=path,
                        description=f"公司名称可能不完整，缺少法律实体后缀: '{text}'",
                        current_value=text,
                        auto_fixable=False,
                    ))

    logger.info(f"🔍 [R4] 格式合规性校验完成，发现 {len(findings)} 个问题")
    return findings


# ============================================================
# R5: 评分项覆盖率检查（与 score_tree 结构化详细比对）
# ============================================================

def _extract_score_tree_keywords(score_tree: Any) -> List[Dict[str, Any]]:
    """
    从评标办法 score_tree 中递归提取所有得分项的关键词与分值。

    :param score_tree: 评标办法打分细则树（JSON 对象或列表）
    :return: [{name, keywords, score, path}] 扁平化列表
    """
    items: List[Dict[str, Any]] = []

    if isinstance(score_tree, dict):
        # 典型结构: {"name": "技术方案", "score": 30, "children": [...], "criteria": "..."}
        name = score_tree.get("name", "") or score_tree.get("item", "") or ""
        score = score_tree.get("score", 0) or score_tree.get("max_score", 0) or 0
        criteria = score_tree.get("criteria", "") or score_tree.get("description", "") or ""
        path = score_tree.get("path", "")

        # 提取关键词：从名称 + 评分标准中提取中文词组
        keywords = set()
        for text in [name, criteria]:
            if text:
                # 提取连续中文字符词组（2字以上）
                words = re.findall(r'[\u4e00-\u9fa5]{2,}', text)
                keywords.update(words)

        if name and keywords:
            items.append({
                "name": name,
                "keywords": list(keywords),
                "score": float(score) if score else 0.0,
                "criteria": criteria,
                "path": path,
            })

        # 递归处理子节点
        children = score_tree.get("children", []) or score_tree.get("items", []) or []
        if isinstance(children, list):
            for child in children:
                items.extend(_extract_score_tree_keywords(child))

    elif isinstance(score_tree, list):
        for item in score_tree:
            items.extend(_extract_score_tree_keywords(item))

    return items


def check_scoring_coverage(
    proposals: List[Dict[str, Any]],
    document_id: str,
) -> List[Dict[str, Any]]:
    """
    R5: 评分项覆盖率检查 — 将填写内容与评标办法 score_tree 做结构化详细比对。

    详细比对逻辑：
    1. 从 DB 查询 evaluation_metadata.score_tree 完整评分细则
    2. 递归展开 score_tree 为扁平评分项列表
    3. 对每个评分项，检查 proposals 中是否有对应内容覆盖
    4. 输出各评分项的覆盖状态与总覆盖率

    :param proposals: Worker 提案列表
    :param document_id: 文档 ID
    :return: ReviewFinding 列表
    """
    logger.info("🔍 [R5] 启动评分项覆盖率检查（详细比对 score_tree）...")
    findings: List[Dict[str, Any]] = []

    # 从 DB 查询评标办法数据
    try:
        from app.db.session import SessionLocal
        from app.db.models.metadata import EvaluationMetadata

        db = SessionLocal()
        try:
            eval_meta = db.query(EvaluationMetadata).filter(
                EvaluationMetadata.document_id == document_id
            ).first()

            if not eval_meta:
                logger.info("🔍 [R5] 无评标办法数据，跳过评分项覆盖率检查")
                return findings

            score_tree = eval_meta.score_tree
            weight_distribution = eval_meta.weight_distribution
            total_score = eval_meta.total_score or 100.0
        finally:
            db.close()
    except Exception as exc:
        logger.warning(f"🔍 [R5] 查询评标办法异常: {exc}")
        return findings

    if not score_tree:
        logger.info("🔍 [R5] score_tree 为空，跳过评分项覆盖率检查")
        return findings

    # 展开 score_tree 为扁平评分项列表
    score_items = _extract_score_tree_keywords(score_tree)
    if not score_items:
        logger.info("🔍 [R5] 从 score_tree 中未提取到评分项，跳过")
        return findings

    logger.info(f"🔍 [R5] 从 score_tree 提取 {len(score_items)} 个评分项，开始逐项比对...")

    # 构建 proposals 的全文索引（用于快速关键词匹配）
    all_proposal_text = " ".join(
        str(p.get("proposed_text", "")) + " " + str(p.get("original_context", ""))
        for p in proposals
        if str(p.get("proposed_text", "")).strip() and not str(p.get("proposed_text", "")).startswith("[")
    )

    # 逐项比对覆盖率
    covered_count = 0
    covered_score = 0.0
    uncovered_items: List[Dict[str, Any]] = []

    for item in score_items:
        name = item["name"]
        keywords = item["keywords"]
        score = item["score"]

        # 关键词命中率计算：至少 30% 的关键词出现在 proposals 中则视为已覆盖
        hit_count = sum(1 for kw in keywords if kw in all_proposal_text)
        hit_rate = hit_count / len(keywords) if keywords else 0.0

        if hit_rate >= 0.3:
            covered_count += 1
            covered_score += score
        else:
            uncovered_items.append({
                "name": name,
                "score": score,
                "keywords": keywords[:5],  # 只取前 5 个关键词
                "hit_rate": round(hit_rate, 2),
                "criteria": item.get("criteria", "")[:100],
            })

    # 计算总覆盖率
    total_items = len(score_items)
    coverage_rate = covered_count / total_items if total_items > 0 else 0.0
    score_coverage_rate = covered_score / total_score if total_score > 0 else 0.0

    # 为每个未覆盖的评分项生成 finding
    for item in uncovered_items:
        severity = "error" if item["score"] >= 10 else "warning" if item["score"] >= 5 else "info"
        findings.append(make_finding(
            rule_id="R5-UNCOVERED", severity=severity,
            path="N/A",
            description=(
                f"评分项 [{item['name']}] ({item['score']}分) 未被充分覆盖，"
                f"关键词命中率: {item['hit_rate']:.0%}，"
                f"评分标准: {item['criteria'][:80]}"
            ),
            current_value=f"命中率: {item['hit_rate']:.0%}",
            expected_value=f"≥ 30% 关键词覆盖",
            auto_fixable=False,
        ))

    # 输出整体覆盖率摘要
    findings.append(make_finding(
        rule_id="R5-COVERAGE-SUMMARY", severity="info",
        path="N/A",
        description=(
            f"评分项覆盖率统计: {covered_count}/{total_items} 项已覆盖 ({coverage_rate:.0%})，"
            f"分值覆盖: {covered_score:.1f}/{total_score:.1f} ({score_coverage_rate:.0%})"
        ),
        current_value=f"{coverage_rate:.0%}",
        expected_value="100%",
    ))

    logger.info(
        f"🔍 [R5] 评分项覆盖率检查完成: {covered_count}/{total_items} 项已覆盖 "
        f"({coverage_rate:.0%})，分值覆盖 {covered_score:.1f}/{total_score:.1f}"
    )
    return findings


# ============================================================
# R7: 表格单元格污染检测（剥离 "序号X行，答：" 等污染前缀）
# ============================================================

# 匹配表格单元格污染前缀的正则表达式
_CELL_POLLUTION_PATTERNS = [
    re.compile(r'^\s*序号\s*\d+\s*[行条项]?\s*[:：,，\s]*[答：:\s]*', re.IGNORECASE),
    re.compile(r'^\s*第\s*\d+\s*[行条项]\s*[:：,，\s]*[答：:\s]*', re.IGNORECASE),
    re.compile(r'^\s*答\s*[:：,，]\s*'),
    re.compile(r'^[^\(\（\n]*[（\(]\s*表头\s*tc[:：]?.*?[）\)]\s*[:：]?\s*', re.IGNORECASE),
    re.compile(r'^\s*表头\s*tc[:：]?\s*', re.IGNORECASE),
]

# 匹配各种说明性元数据与零改动注释括号
_ZERO_CHANGE_ANNOTATION_PATTERNS = [
    re.compile(r'[\(（]\s*(?:原文|模板|固定|正文|章节|标题|大标题|引言|注|甲方)?\s*(?:无槽位|固定原文|零改动|原样保留|无需改动|无需修改|无需写盘|盲守|100%|已写|保留原文|待线下|线下盖章|线下签字|线下签署|线下办理)[^()（）]*[\)）]', re.IGNORECASE),
    re.compile(r'[\(（]\s*(?:原文|模板|正文|章节|标题)?\s*(?:100%|盲守|零改动|原样保留|保留原文)[^()（）]*[\)）]', re.IGNORECASE),
    re.compile(r'[\(（]\s*正文已写[，,\s]*零改动\s*[\)）]', re.IGNORECASE),
    re.compile(r'[\(（]\s*原文无槽位[，,\s]*零改动保留\s*[\)）]', re.IGNORECASE),
    re.compile(r'[\(（]\s*固定原文[，,\s]*零改动保留\s*[\)）]', re.IGNORECASE),
    re.compile(r'[\(（]\s*(?:原文)?无槽位[，,\s]*(?:零改动|原样)?保留\s*[\)）]', re.IGNORECASE),
]


def clean_zero_change_annotations(text: str) -> str:
    """
    程序化剥离文本中残留的各类零改动、无槽位、固定原文等说明性注释括号。
    例如：
    - "致太湖咨询：（原文无槽位，零改动保留）" -> "致太湖咨询："
    - "据此函，签字人宣布如下：（固定原文，零改动保留）" -> "据此函，签字人宣布如下："
    """
    if not text:
        return text
    cleaned = text
    for pat in _ZERO_CHANGE_ANNOTATION_PATTERNS:
        cleaned = pat.sub('', cleaned)
    return cleaned.strip()


def is_zero_change_or_no_op_proposal(proposed_text: str, original_context: str = "", prop_type: str = "") -> bool:
    """
    判定提案是否属于“零改动/无需修改”的无操作提案 (No-Op Proposal)。
    如果提案文本剥离注释后为空、为元数据标记词、或与原文模板完全一致，则返回 True，指示无需写盘改写。
    """
    if prop_type in ("image", "table_rows"):
        return False
    if not proposed_text:
        return True

    cleaned_p = clean_zero_change_annotations(proposed_text).strip()
    # 纯元数据标记检测
    if not cleaned_p or any(cleaned_p == k for k in ["—", "-", "--", "同上", "略", "无变更", "原样保留", "无需写盘", "零改动", "固定原文", "无槽位"]):
        return True
    if any(k in proposed_text for k in ["（原文无槽位", "(原文无槽位", "（固定原文", "(固定原文", "（无需写盘", "（零改动", "(零改动", "（原样保留", "(原样保留"]) and not cleaned_p:
        return True

    if original_context:
        cleaned_orig = clean_zero_change_annotations(original_context).strip()
        # 骨架去标点空白比较
        p_core = re.sub(r'[\s:：_＿\[\]［］\(\)（）\.\,，。；;、“”"\'`]', '', cleaned_p)
        o_core = re.sub(r'[\s:：_＿\[\]［］\(\)（）\.\,，。；;、“”"\'`]', '', cleaned_orig)
        if p_core and o_core and p_core == o_core:
            return True
        if cleaned_p == cleaned_orig:
            return True

    return False


def clean_all_ellipsis(text: str) -> str:
    """
    全方位清除与智能自愈文本中的所有省略号（包括首尾截断、句中偷懒连接符及伪装标记）。
    将句中的 `…` / `……` / `...` 自愈为中文逗号，确保语义连贯完整。
    """
    if not text:
        return text
    had_trailing_ellipsis = bool(re.search(r'(?:[…\.]{2,}|…+|\.{2,}|。{2,})\s*$', text) or text.endswith("..."))
    # 1. 移除伪装标记如 （完整技术要求）
    text = re.sub(r'[\(（]完整技术要求[\)）]', '', text)
    # 2. 清理首部省略号与孤立点号
    text = re.sub(r'^[…\.]+', '', text)
    # 清理尾部省略号（连续点号、连续省略号、连续句号），保留单个合法句末标点
    text = re.sub(r'(?:[…\.]{2,}|…+|\.{2,}|。{2,})\s*$', '', text)
    # 3. 将句中的省略号（…、……、...、..）智能转换为中文标点
    # 若省略号紧挨着标点，直接规整
    text = re.sub(r'[，,；;、]\s*[…\.]+', '，', text)
    text = re.sub(r'[…\.]+\s*[，,；;、]', '，', text)
    # 若位于文字之间，转换为逗号
    text = re.sub(r'([^\s，,；;、])\s*[…\.]{2,}\s*([^\s，,；;、])', r'\1，\2', text)
    text = re.sub(r'([^\s，,；;、])\s*…\s*([^\s，,；;、])', r'\1，\2', text)
    # 去除多余连续逗号
    text = re.sub(r'，{2,}', '，', text)
    text = text.strip()
    if had_trailing_ellipsis and text and not text.endswith(("。", "！", "？", "；", ")", "）")):
        text += "。"
    if text and re.match(r'^\s*\.{1,6}\s*$', text):
        return ""
    return text


def clean_cell_text_value(text: str, ctx: str = "") -> str:
    """
    程序化清理表格单元格文本中的污染前缀、零改动注释与多余指示词。
    """
    if not text or text.startswith("["):
        return text

    cleaned = text.strip()
    # 递归剥离匹配到的所有污染前缀
    for pat in _CELL_POLLUTION_PATTERNS:
        while pat.search(cleaned):
            cleaned = pat.sub('', cleaned).strip()

    # 剥离零改动等说明性元数据注释
    cleaned = clean_zero_change_annotations(cleaned)

    # 特殊规范化：对【有无偏离】单元格文本进行标准判定与提取
    combined = f"{text} {ctx}".lower()
    if "偏离" in combined or "有无" in combined:
        if "无" in cleaned or "完全响应" in cleaned or "符合" in cleaned:
            if len(cleaned) <= 15:
                return "无"
        elif "有" in cleaned or "偏离" in cleaned:
            if len(cleaned) <= 15:
                return "有"

    # 全方位智能清洗句中与尾部省略号
    cleaned = clean_all_ellipsis(cleaned)

    return cleaned


def check_table_cell_cleanliness(
    proposals: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """
    R7: 表格单元格污染检测 — 自动发现并修正单元格中残留的 "序号 1 行，答：" 等前缀污染。

    :param proposals: Worker 提案列表
    :return: ReviewFinding 列表
    """
    logger.info("🔍 [R7] 启动表格单元格污染检测管线...")
    findings: List[Dict[str, Any]] = []

    for p in proposals:
        text = str(p.get("proposed_text", "")).strip()
        orig = str(p.get("original_context", "")).strip()
        path = str(p.get("path", ""))

        if not text or text.startswith("["):
            continue

        # 判断是否为表格路径/上下文
        is_table = bool(re.search(r'/(?:tbl|tc|tr)\[', path)) or "表格" in orig or "行" in orig

        if is_table or any(pat.search(text) for pat in _CELL_POLLUTION_PATTERNS):
            cleaned = clean_cell_text_value(text, orig)
            if cleaned != text:
                findings.append(make_finding(
                    rule_id="R7-CELL-POLLUTION", severity="warning",
                    path=path,
                    description=f"表格单元格文本包含污染前缀: '{text[:50]}' -> 纯化为 '{cleaned[:50]}'",
                    current_value=text,
                    expected_value=cleaned,
                    auto_fixable=True,
                    fix_proposal={
                        "path": path,
                        "proposed_text": cleaned,
                        "original_context": orig,
                        "source_tool": "review_engine_r7_clean",
                        "source_data": text,
                        "reasoning": f"R7 自动纯化: 剥离污染前缀 '{text}' -> '{cleaned}'",
                    },
                ))

    logger.info(f"🔍 [R7] 表格单元格污染检测完成，发现并标记 {len(findings)} 个污染字段")
    return findings


# ============================================================
# R8: 短字段长文本错填检测（拦截公司名/人名被整段话覆写污染）
# ============================================================

_SHORT_FIELD_KEYWORDS = ["名称", "公司", "投标人", "单位", "法定代表人", "法人", "代理人", "项目编号", "招标编号", "电话", "邮箱", "邮编", "传真"]
_SENTENCE_POLLUTION_KEYWORDS = ["根据", "授权", "招标文件", "处理", "事宜", "特此", "代表我方"]


def check_short_field_text_overflow(
    proposals: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """
    R8: 短字段长文本错填检测 — 发现并自动纠正将整个段落或授权大段话填入 "投标单位名称:" / "投标人名称:" 的严重错填。

    :param proposals: Worker 提案列表
    :return: ReviewFinding 列表
    """
    logger.info("🔍 [R8] 启动短字段长文本错填检测管线...")
    findings: List[Dict[str, Any]] = []

    for p in proposals:
        text = str(p.get("proposed_text", "")).strip()
        orig = str(p.get("original_context", "")).strip()
        path = str(p.get("path", ""))

        if not text or text.startswith("["):
            continue

        combined_ctx = f"{orig}".lower()
        # 剥离可能包含的前缀标签
        clean_text = re.sub(r'^[^:：]*[:：]\s*', '', text).strip()

        # 检查是否为结构化短字段（投标人名称、单位名称、法定代表人等）
        is_short_field = any(kw in combined_ctx for kw in _SHORT_FIELD_KEYWORDS)

        if is_short_field:
            # 判断是否包含整段长文或句式关联词
            is_overflow = len(clean_text) > 35 or any(bad in clean_text for bad in _SENTENCE_POLLUTION_KEYWORDS)

            if is_overflow:
                # 尝试从长文中提取纯净的公司名/实体名
                pure_extracted = None
                m_company = re.search(r'([\u4e00-\u9fa5]{2,30}(?:有限责任公司|股份有限公司|有限公司|集团|公司))', clean_text)
                if m_company:
                    pure_extracted = m_company.group(1)

                prefix_match = re.match(r'^([^:：]*[:：]\s*)', text)
                prefix = prefix_match.group(1) if prefix_match else ""

                expected_fixed = (prefix + pure_extracted) if (prefix and pure_extracted) else (pure_extracted or clean_text[:30])

                findings.append(make_finding(
                    rule_id="R8-TEXT-OVERFLOW", severity="error",
                    path=path,
                    description=f"结构化短字段错填长文段落: '{text[:60]}...' -> 建议纯化为 '{expected_fixed}'",
                    current_value=text,
                    expected_value=expected_fixed,
                    auto_fixable=True,
                    fix_proposal={
                        "path": path,
                        "proposed_text": expected_fixed,
                        "original_context": orig,
                        "source_tool": "review_engine_r8_fix",
                "source_data": clean_text,
                        "reasoning": f"R8 自动纠偏: 将结构化短字段从错填长文中提炼为 '{expected_fixed}'",
                    },
                ))

    logger.info(f"🔍 [R8] 短字段长文本错填检测完成，发现并标记 {len(findings)} 个严重错填字段")
    return findings


# ============================================================
# R10: 填写前后 DOM 文本 Diff 比对熔断器（比对写前 vs 写后）
# ============================================================

def check_and_rollback_single_node(
    p_node: Any,
    before_text: str,
    path: str,
) -> bool:
    """
    单节点实时 Diff 对比与回滚校验器：
    在修改单独节点后立即触发，计算该节点的模板文字保留率，若低于阈值立刻回滚还原！

    :return: True (校验通过), False (模板破坏已回滚还原)
    """
    if not before_text or len(before_text.strip()) < 20:
        return True

    after_text = p_node.text or ""
    clean_before = re.sub(r'(_{2,}|\[[^\]]+\]|［[^］]+］|【[^】]+】)', '', before_text).strip()

    if len(clean_before) <= 20:
        return True

    words = [w for w in re.findall(r'[\u4e00-\u9fa5]{2,}', clean_before) if len(w) >= 2]
    if not words:
        return True

    retained_count = sum(1 for w in words if w in after_text)
    retention_rate = retained_count / len(words)

    if retention_rate >= 0.6:
        logger.info(f"   🔍 [R10 节点检测] 路径 {path} | 模板保留率: {retention_rate:.0%} (✅ 合规保留)")
        return True

    logger.warning(
        f"   🚨 [R10 节点熔断] 路径 {path} 模板保留率异常低下 ({retention_rate:.0%})！"
        f"\n      写前: '{before_text[:50]}...'\n      写后: '{after_text[:50]}...'"
        f"\n      已触发单节点实时回滚还原！"
    )
    p_node.text = before_text
    return False


def compare_text_before_and_after_filling(
    docx_path: str,
    before_snapshots: Dict[str, str],
) -> List[Dict[str, Any]]:
    """
    R10: 填写前后 DOM 文本 Diff 比对熔断器 — 比对 Word 填报前的原始快照与填报后的实际文本。
    计算模板文字保留率 (Template Retention Ratio)，自动拦截并回滚丢失原文的段落。

    :param docx_path: 填报后的 Word 文档路径
    :param before_snapshots: 写盘前的 DOM 节点快照 {path: original_text}
    :return: ReviewFinding 列表 (包含需要自动回滚恢复的项)
    """
    logger.info(f"🔍 [R10] 启动填写前后 DOM 文本 Diff 对比与模板保留率检测 (共受影响 {len(before_snapshots)} 处节点)...")
    findings: List[Dict[str, Any]] = []

    if not docx_path or not os.path.exists(docx_path) or not before_snapshots:
        return findings

    try:
        doc = Document(docx_path)
    except Exception as e:
        logger.error(f"🔍 [R10] 无法读取文档进行前后比对: {e}")
        return findings

    from app.agents.bid_filler_agent import _find_paragraph_by_path

    for path, before_text in before_snapshots.items():
        if not before_text or len(before_text.strip()) < 20:
            continue

        p_node = _find_paragraph_by_path(doc, path)
        if not p_node:
            continue

        passed = check_and_rollback_single_node(p_node, before_text, path)
        if not passed:
            doc.save(docx_path)
            findings.append(make_finding(
                rule_id="R10-TEXT-DIFF-LOSS", severity="error", path=path,
                description=f"填写前后比对发现模板保留率低，已自动回滚还原原文！",
                current_value=p_node.text,
                expected_value=before_text,
                auto_fixable=True,
            ))

    logger.info(f"🔍 [R10] 填写前后对比完成: 检查 {len(before_snapshots)} 处，拦截并自动回滚 {len(findings)} 处破坏模板原文的节点")
    return findings

def check_paragraph_text_destruction(
    proposals: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """
    R9: 模板原文破坏熔断检测 — 自动拦截试图用短词覆盖多于 25 字的大段模板声明的提案。
    彻底防止丢字、丢话、丢失标书原文。
    """
    logger.info("🔍 [R9] 启动模板原文破坏熔断检测...")
    findings: List[Dict[str, Any]] = []

    for idx, p in enumerate(proposals):
        path = str(p.get("path", "")).strip()
        text = str(p.get("proposed_text", "")).strip()
        orig = str(p.get("original_context", "")).strip()

        # 仅对正文段落执行检测（跳过表格单元格）
        if "/tbl[" in path or "/tc[" in path:
            continue

        # 如果原文是一个多于 25 字的模板长句
        if len(orig) > 25:
            # 检查 proposed_text 是否严重缺失原模版中的长词句 (长度不足原文一半且原文无占位符)
            if len(text) < len(orig) * 0.5 and not re.search(r'_[_]+|\[[^\]]+\]', orig):
                if any(kw in orig for kw in ["根据", "授权", "代表我方", "同意", "责任", "义务"]):
                    findings.append(make_finding(
                        rule_id="R9-TEMPLATE-DESTRUCTION", severity="error", path=path,
                        description=f"提案企图擦除 {len(orig)} 字的模板声明原文！自动熔断保护原文！",
                        current_value=text,
                        expected_value=orig,
                        auto_fixable=True,
                        fix_proposal={
                            "path": path,
                            "proposed_text": orig,  # 强制恢复完整原文
                            "reasoning": "[R9熔断修复] 拦截模板破坏，强行保留大段声明原文",
                        }
                    ))

    logger.info(f"🔍 [R9] 模板原文破坏熔断检测完成: 拦截 {len(findings)} 处模板破坏提案")
    return findings


# ============================================================
# R6: LLM 自动修正管线
# ============================================================

def auto_fix_proposals(
    proposals: List[Dict[str, Any]],
    findings: List[Dict[str, Any]],
    document_id: str,
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    """
    R6: LLM 自动修正管线 — 根据审查发现，修正问题 Proposals。

    修正策略：
    1. 对于有 fix_proposal 的 finding（R2/R3/R4 等规则引擎已提供修正方案）→ 直接应用
    2. 对于无 fix_proposal 但 auto_fixable=True 的 finding → 调用 LLM + DB 工具重新查询
    3. 对于 auto_fixable=False 的 finding → 跳过该字段不写入，标记「待人工补充」

    :param proposals: Worker 提案列表
    :param findings: 审查发现列表
    :param document_id: 文档 ID
    :return: (修正后的 proposals 列表, 被跳过的 proposals 路径列表)
    """
    logger.info(f"🔧 [R6] 启动自动修正管线，审查发现 {len(findings)} 个问题...")

    # 筛选出需要处理的 findings
    fixable_findings = [f for f in findings if f.get("auto_fixable") and f.get("severity") in ("error", "warning")]
    unfixable_findings = [f for f in findings if not f.get("auto_fixable") and f.get("severity") in ("error",)]

    if not fixable_findings and not unfixable_findings:
        logger.info("🔧 [R6] 无需修正的问题，跳过自动修正")
        return proposals, []

    # 构建路径到 proposal 的索引
    proposal_by_path: Dict[str, int] = {}
    for idx, p in enumerate(proposals):
        path = str(p.get("path", "")).strip()
        if path:
            proposal_by_path[path] = idx

    fixed_count = 0
    skipped_paths: List[Dict[str, Any]] = []
    modified_proposals = [dict(p) for p in proposals]  # 深拷贝

    # 步骤 1: 应用规则引擎已提供的 fix_proposal（R2/R3/R4 产出的确定性修正）
    for finding in fixable_findings:
        fix = finding.get("fix_proposal")
        if not fix:
            continue

        fix_path = str(fix.get("path", finding.get("path", ""))).strip()
        if not fix_path or fix_path == "N/A":
            continue

        fix_text = str(fix.get("proposed_text", "")).strip()
        if not fix_text:
            continue

        # 在 proposals 中找到对应路径并替换
        if fix_path in proposal_by_path:
            idx = proposal_by_path[fix_path]
            old_text = modified_proposals[idx].get("proposed_text", "")
            modified_proposals[idx]["proposed_text"] = fix_text
            modified_proposals[idx]["reasoning"] = (
                f"{finding.get('rule_id', 'R6')}: {finding.get('description', '自动修正')}"
            )
            logger.info(
                f"   🔧 [R6] 自动修正 {fix_path}: "
                f"'{str(old_text)[:40]}' → '{fix_text[:40]}'"
            )
            fixed_count += 1
        else:
            # 路径不在现有 proposals 中，作为新 proposal 追加
            modified_proposals.append({
                "path": fix_path,
                "proposed_text": fix_text,
                "original_context": fix.get("original_context", ""),
                "source_tool": "review_engine_fix",
                "source_data": fix.get("source_data", ""),
                "reasoning": f"{finding.get('rule_id', 'R6')}: {finding.get('description', '自动修正')}",
            })
            logger.info(f"   🔧 [R6] 新增修正 proposal: {fix_path} → '{fix_text[:40]}'")
            fixed_count += 1

    # 步骤 2: 对于 R2 一致性问题，统一修正所有相关路径
    for finding in fixable_findings:
        if finding.get("rule_id") != "R2-INCONSISTENT":
            continue
        fix = finding.get("fix_proposal", {})
        correct_value = fix.get("correct_value", "")
        field_key = fix.get("field_key", "")
        if not correct_value or not field_key:
            continue

        # 找到所有属于该字段的 proposals，统一修正
        for idx, p in enumerate(modified_proposals):
            orig = str(p.get("original_context", "")).lower()
            source = str(p.get("source_tool", "")).lower()
            combined = f"{orig} {source}"
            keywords = _CONSISTENCY_FIELDS.get(field_key, [])
            if any(kw in combined for kw in keywords):
                old_text = str(p.get("proposed_text", ""))
                # 保留原文前缀标签
                prefix_match = re.match(r'^([^:：]*[:：]\s*)', old_text)
                if prefix_match:
                    modified_proposals[idx]["proposed_text"] = prefix_match.group(1) + correct_value
                else:
                    modified_proposals[idx]["proposed_text"] = correct_value

    # 步骤 3: 标记不可修正的 findings 对应的 proposals 为跳过（用户选择 B 策略）
    for finding in unfixable_findings:
        finding_path = finding.get("path", "")
        if finding_path and finding_path != "N/A" and finding_path in proposal_by_path:
            idx = proposal_by_path[finding_path]
            skipped_proposals_item = {
                "path": finding_path,
                "reason": finding.get("description", "审查未通过"),
                "rule_id": finding.get("rule_id", "UNKNOWN"),
            }
            skipped_paths.append(skipped_proposals_item)
            # 将 proposed_text 标记为待人工补充
            modified_proposals[idx]["proposed_text"] = "[待人工补充]"
            modified_proposals[idx]["reasoning"] = (
                f"[R6 跳过] {finding.get('description', '审查未通过，需人工介入')}"
            )
            logger.info(f"   ⏭️ [R6] 跳过不可修正字段: {finding_path} → 标记为 [待人工补充]")

    logger.info(f"🔧 [R6] 自动修正完成: {fixed_count} 个已修正, {len(skipped_paths)} 个跳过待人工补充")
    return modified_proposals, skipped_paths


# ============================================================
# R11: 报价表格单价与分项总价列分离与防错列检测
# ============================================================

def check_pricing_table_column_alignment(
    proposals: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """
    R11: 报价表格单价与分项总价列分离与防错列质检管线：
    检测 2D 报价表格矩阵提案中是否存在将包干/工程总价重复填入单价列，或单价与分项总价错列的情况。
    自动将包干/汇总项的单价纠偏为破折号 "——"，保留分项总价。

    :param proposals: Worker 提案列表
    :return: ReviewFinding 列表
    """
    logger.info("🔍 [R11] 启动报价表格单价与分项总价列分离质检管线...")
    findings: List[Dict[str, Any]] = []

    pkg_kws = ["费", "工程", "系统", "加固", "防水", "敷设", "安装", "调试", "服务", "培训", "大类", "购置", "总承包", "支架", "桥架", "辅材", "电缆", "柜"]

    for p in proposals:
        p_path = str(p.get("path", "")).strip()
        raw_val = p.get("proposed_text") if p.get("proposed_text") is not None else p.get("value", "")
        p_val = str(raw_val).strip() if raw_val is not None else ""

        # 检查是否为 2D 矩阵提案
        matrix = None
        if isinstance(raw_val, list) and raw_val and isinstance(raw_val[0], list):
            matrix = raw_val
        elif p_val.startswith("[") and p_val.endswith("]"):
            try:
                parsed = json.loads(p_val)
                if isinstance(parsed, list) and parsed and isinstance(parsed[0], list):
                    matrix = parsed
            except Exception:
                pass

        if not matrix:
            continue

        # 检查是否为 5 列的报价数据矩阵
        has_duplicate_pricing = False
        new_matrix = []
        for row in matrix:
            if not isinstance(row, list):
                new_matrix.append(row)
                continue
            r_copy = list(row)
            if len(r_copy) >= 4:
                name_val = str(r_copy[1]).strip()
                unit_val = str(r_copy[2]).strip()
                total_val = str(r_copy[3]).strip()
                # 检查单价与总价相同且非空且非破折号且非 0.00
                if unit_val and total_val and unit_val == total_val and unit_val not in ("—", "——", "/", "-", "0", "0.00"):
                    if any(kw in name_val for kw in pkg_kws):
                        r_copy[2] = "——"
                        has_duplicate_pricing = True
            new_matrix.append(r_copy)

        if has_duplicate_pricing:
            new_json_str = json.dumps(new_matrix, ensure_ascii=False)
            findings.append(make_finding(
                rule_id="R11-PRICING-COL-DUPLICATE",
                severity="warning",
                path=p_path,
                description=f"报价表格存在包干/工程项将总价重复填入单价列，已自动纯化单价列为破折号",
                current_value=p_val[:100],
                expected_value=new_json_str[:100],
                auto_fixable=True,
                fix_proposal={
                    "path": p_path,
                    "proposed_text": new_json_str,
                    "value": new_json_str,
                    "type": p.get("type", "table_rows"),
                    "source_tool": "review_engine_r11_fix",
                    "reasoning": "R11 自动纠偏: 修复报价表单价列与分项总价列重复填报，规范包干项单价为破折号",
                }
            ))

    logger.info(f"🔍 [R11] 报价表列分离质检完成，发现并标记 {len(findings)} 处表格错列")
    return findings


# ============================================================
# 总入口：执行全部审查管线
# ============================================================

def run_all_review_pipelines(
    proposals: List[Dict[str, Any]],
    document_id: str,
    docx_path: Optional[str] = None,
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]], List[Dict[str, Any]]]:
    """
    执行全部规则审查管线，并调用 R6 自动修正管线。

    :param proposals: Worker 提案列表
    :param document_id: 文档 ID
    :param docx_path: Word 文档路径（R1 遗漏检测需要，可选）
    :return: (修正后的 proposals, 所有 findings, 跳过的 proposals)
    """
    logger.info("=" * 60)
    logger.info("🔬 [Review Engine] 启动全量深度质检审查管线...")
    logger.info("=" * 60)

    all_findings: List[Dict[str, Any]] = []

    # R2: 数据交叉一致性校验
    r2_findings = check_data_consistency(proposals)
    all_findings.extend(r2_findings)

    # R3: 财务数据准确性校验
    r3_findings = check_financial_accuracy(proposals, document_id)
    all_findings.extend(r3_findings)

    # R4: 格式合规性校验
    r4_findings = check_format_compliance(proposals)
    all_findings.extend(r4_findings)

    # R5: 评分项覆盖率检查
    r5_findings = check_scoring_coverage(proposals, document_id)
    all_findings.extend(r5_findings)

    # R7: 表格单元格污染检测
    r7_findings = check_table_cell_cleanliness(proposals)
    all_findings.extend(r7_findings)

    # R8: 短字段长文本错填检测
    r8_findings = check_short_field_text_overflow(proposals)
    all_findings.extend(r8_findings)

    # R11: 报价表格单价与分项总价列分离与防错列检测
    r11_findings = check_pricing_table_column_alignment(proposals)
    all_findings.extend(r11_findings)

    # 统计各级别数量
    errors = sum(1 for f in all_findings if f["severity"] == "error")
    warnings = sum(1 for f in all_findings if f["severity"] == "warning")
    infos = sum(1 for f in all_findings if f["severity"] == "info")
    logger.info(
        f"🔬 [Review Engine] 规则审查完成: "
        f"{errors} errors, {warnings} warnings, {infos} infos (共 {len(all_findings)} 条)"
    )

    # R6: 自动修正管线
    modified_proposals, skipped = auto_fix_proposals(
        proposals, all_findings, document_id
    )

    # R1: 此处不执行遗漏检测（留待写盘后二次扫描）
    # R1 findings 会在 review_node 写盘后单独追加

    logger.info("=" * 60)
    logger.info("🔬 [Review Engine] 全量质检完毕")
    logger.info("=" * 60)

    return modified_proposals, all_findings, skipped
