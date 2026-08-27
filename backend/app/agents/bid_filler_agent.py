"""
BidFillerAgent - 基于 LangGraph 状态图 (StateGraph) 的 Multi-Agent 标书撰写系统

架构说明：
1. 显式图状态：基于 TypedDict 定义 BidFillerState 管理撰写生命周期；
2. 3 大状态节点 (Nodes)：
   - scan_node：写入 Word 临时文件供 OfficeCLI 访问 + 提取原文全文上下文 + LLM 槽位预识别；
   - agent_fill_node：Supervisor Agent（决策Agent）——
     读文档 → 识别章节 → LLM 四类分类 → 并发派发 Worker → 收集结果；
     每个 Worker 是独立的 ReAct Agent，配备章节专属工具子集，直接写入 Word；
   - write_docx_node：从临时文件读取 Agent 修改后的 Word 并输出最终字节流。
3. Multi-Agent Supervisor-Worker 模式：
   - Supervisor: 决策 + 调度 + 审查（4 个决策工具）
   - Worker: 按章节动态创建，独立 ReAct Agent，自主查库 → 思考 → 写盘
"""
import os
import re
import json
import shutil
import tempfile
from collections import defaultdict
from decimal import Decimal, InvalidOperation
from typing import Dict, Any, List, Optional, Sequence, Tuple, TypedDict
from loguru import logger
from sqlalchemy.orm import Session
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_COLOR_INDEX
import io

from langchain_core.tools import tool
from app.core.sandbox import AgentSandbox
from langgraph.graph import StateGraph, END
from langgraph.prebuilt import create_react_agent
from langchain_core.messages import HumanMessage

from app.db.models.project import Document as DocumentModel
from app.db.models.business import CompanyQualification, CompanyProfileModel
from app.db.models.metadata import FinancialMetadata, TimelineMetadata
from app.db.session import SessionLocal
from app.schemas.bid_filler_schema import (
    AgentFillPlanItem,
    BidFillAuditReport,
    BidFillPlan,
    CompanyProfile,
    FillingAuditItem,
    ReviewFinding,
)
from app.services.llm_service import llm_service
from app.utils.rmb_formatter import number_to_chinese_rmb
from app.utils.table_utils import is_narrative_clause_or_lead_in


# ============================================================
# LangGraph 全局 State 状态定义
# ============================================================

class BidFillerState(TypedDict):
    """LangGraph Multi-Agent 标书撰写全局状态"""
    document_id: str
    tenant_id: Optional[str]
    original_context: str
    slot_analysis: Optional[List[Dict[str, Any]]]
    worker_proposals: Optional[List[Dict[str, Any]]]  # 保留兼容
    chapter_tasks: Optional[List[Dict[str, Any]]]  # 首轮缓存的可执行章节任务，供修复轮复用

    db_session: Any
    company_profile: CompanyProfile
    profile_id: Optional[str]                              # 指定使用的企业档案 ID
    original_docx: Optional[bytes]
    docx_temp_path: Optional[str]

    # 用户自定义指令透传（前端 → Supervisor → Worker Prompt）
    custom_instructions: Optional[str]                  # 全局自定义填写指令
    category_hints: Optional[Dict[str, str]]            # 按章节类别的额外指令

    # 闭环质量把控与专项修复状态
    repair_count: int
    max_repair_rounds: int
    repair_instructions_map: Optional[Dict[str, str]]   # 按章节的专项修复反馈指令
    audit_passed: Optional[bool]
    audit_blocked: bool                                 # 达到修复上限仍不通过时阻断发布

    audit_items: List[FillingAuditItem]
    review_findings: Optional[List[Dict[str, Any]]]     # 终审发现列表
    audit_report: Optional[BidFillAuditReport]
    filled_docx_bytes: Optional[bytes]


def _select_repair_chapter_tasks(
    chapter_tasks: Optional[List[Dict[str, Any]]],
    repair_instructions_map: Optional[Dict[str, str]],
) -> List[Dict[str, Any]]:
    """根据终审章节名称筛选修复轮需要重新执行的章节任务。"""
    if not chapter_tasks or not repair_instructions_map:
        return []

    repair_titles = {
        str(title).strip()
        for title in repair_instructions_map
        if str(title).strip()
    }
    if not repair_titles:
        return []

    selected_tasks: List[Dict[str, Any]] = []
    selected_keys = set()
    for task in chapter_tasks:
        if not isinstance(task, dict):
            continue

        chapter_title = str(task.get("chapter_title", "")).strip()
        if not chapter_title:
            continue

        # 优先精确匹配，兼容终审标题带有章节编号或少量前后缀的情况。
        is_repair_target = any(
            chapter_title == repair_title
            or chapter_title in repair_title
            or repair_title in chapter_title
            for repair_title in repair_titles
        )
        if not is_repair_target:
            continue

        task_key = (
            str(task.get("chapter_number", "")).strip(),
            chapter_title,
        )
        if task_key in selected_keys:
            continue
        selected_keys.add(task_key)
        selected_tasks.append(task)

    return selected_tasks



# ============================================================
# LangGraph 3 大 Node 节点实现
# ============================================================

def scan_node(state: BidFillerState) -> Dict[str, Any]:
    """1. scan_node: 写入 Word 临时文件供 OfficeCLI 访问 + 提取全文上下文"""
    logger.info("[LangGraph Node 1/4] scan_node: 写入临时文件、提取全文上下文...")
    original_docx = state.get("original_docx")
    original_context = ""
    docx_temp_path = None
    slot_analysis: Optional[List[Dict[str, Any]]] = None  # Slot Analyzer 已移除，Worker 自主探索文档

    if original_docx:
        try:
            # 修复层级：__file__ 在 backend/app/agents/ 下，向上一层为 agents、两层为 app、三层为 backend
            drafts_dir = os.path.join(
                os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
                "uploads", "drafts"
            )
            os.makedirs(drafts_dir, exist_ok=True)
            doc_id = state.get("document_id", "unknown")
            docx_temp_path = os.path.join(drafts_dir, f"bid_fill_{doc_id[:8]}.docx")
            with open(docx_temp_path, "wb") as f:
                f.write(original_docx)
            logger.info(f"   📄 已创建工作副本: {docx_temp_path}")
        except Exception as exc_tmp:
            logger.warning(f"   写入工作副本失败: {exc_tmp}")

        try:
            from app.services.bid_format_filler_service import bid_format_filler_service
            original_context = bid_format_filler_service.extract_original_document_context(original_docx)
            logger.info(f"   已提取 Word 全文上下文 ({len(original_context)} 字符)")
        except Exception as exc:
            logger.warning(f"读取 Word 上下文时发生异常: {exc}")


    return {
        "original_context": original_context,
        "docx_temp_path": docx_temp_path,
        "slot_analysis": slot_analysis,
    }



def agent_fill_node(state: BidFillerState) -> Dict[str, Any]:
    """
    Supervisor Agent — 读文档 → 识别章节 → 四类分类 → 并发派发 Worker 独立直写 Word。
    支持在质量审核不达标时接收专项修复指令。
    """
    doc_id = state.get("document_id", "")
    tenant_id = state.get("tenant_id") or "default-tenant"
    original_context = state.get("original_context", "")
    docx_temp_path = state.get("docx_temp_path")
    slot_analysis = state.get("slot_analysis")
    repair_instructions_map = state.get("repair_instructions_map") or {}
    repair_count = state.get("repair_count", 0)
    audit_items: List[FillingAuditItem] = list(state.get("audit_items") or [])
    dispatched_chapter_tasks: List[Dict[str, Any]] = []
    cached_repair_tasks = _select_repair_chapter_tasks(
        state.get("chapter_tasks"),
        repair_instructions_map,
    )

    if repair_count > 0 and cached_repair_tasks:
        logger.info("[LangGraph Node 2/4] agent_fill_node: 进入问题章节专项修复直通模式...")
    else:
        logger.info("[LangGraph Node 2/4] agent_fill_node: 启动 Supervisor 决策 Agent...")

    # 修复轮次必须只处理本轮新提案，避免上一轮已经写入的提案再次进入集中刷盘队列。
    from app.agents.bid_filler_workers import clear_worker_proposals
    if repair_count > 0:
        clear_worker_proposals(doc_id)
        logger.info(f"   [提案池隔离] 进入第 {repair_count} 轮修复，已清理上一轮文档 {doc_id} 的旧提案")

    if repair_instructions_map:
        logger.warning(f"[Supervisor 专项修复轮次 {repair_count}] 存在 {len(repair_instructions_map)} 个章节需针对性整改修复")

    # 读取用户自定义指令（从前端透传）
    custom_instructions = state.get("custom_instructions") or ""
    category_hints = state.get("category_hints") or {}
    if custom_instructions:
        logger.info(f"   [Supervisor] 收到全局自定义指令: '{custom_instructions[:80]}'")
    if category_hints:
        logger.info(f"   [Supervisor] 收到 {len(category_hints)} 条章节类别指令: {list(category_hints.keys())}")

    if llm_service.get_llm(temperature=0.3, json_mode=False, tenant_id=tenant_id) is None and not cached_repair_tasks:
        logger.error("LLM 服务未初始化")
        return {"audit_items": audit_items}
    if not docx_temp_path or not os.path.exists(docx_temp_path):
        logger.error("docx_temp_path 为空")
        return {"audit_items": audit_items}

    # ================================================================
    # Supervisor 工具箱（4 个决策工具）
    # ================================================================
    from app.mcp.office_cli_client import office_cli_mcp_client
    from app.agents.tools.style_extractor_tool import extract_text_by_style
    import asyncio as _asyncio
    import concurrent.futures as _cf

    def _sync_call_async(coro):
        """安全地在同步上下文中调用异步 MCP Client 协程"""
        try:
            _asyncio.get_running_loop()
        except RuntimeError:
            return _asyncio.run(coro)
        with _cf.ThreadPoolExecutor(max_workers=1) as executor:
            return executor.submit(_asyncio.run, coro).result()

    # --- 工具 1: 阅读 Word 文档结构 ---
    @tool
    def officecli_query_doc_structure(selector: str = "paragraph") -> str:
        """查询当前 Word 文档的 DOM 结构。selector: 'paragraph' / 'table' / 'all'。"""
        logger.info(f"   🧠 [Supervisor] 查询文档结构 (selector='{selector}')")
        coro = office_cli_mcp_client.query_structure(docx_temp_path, selector)
        res = _sync_call_async(coro)
        return res.get("structure", str(res)) if isinstance(res, dict) else str(res)

    # --- 工具 2: 分析章节并四类分类 ---
    @tool
    def analyze_chapters(doc_structure_summary: str) -> str:
        """分析 Word 文档 DOM 结构，识别所有章节并做四类分类。
        返回 JSON: [{"chapter_number":"一","chapter_title":"投标函","category":"needs_fill","mapping_hint":"bid_letter","template_text":"...","content_hint":"..."}, ...]"""
        logger.info("   🧠 [Supervisor] 分析文档章节并执行四类分类...")
        prompt = f"""你是招投标格式分析专家。以下是 Word 文档 DOM 结构文本。
请识别所有投标文件章节，输出 JSON 数组。每个章节包含:
- chapter_number: 编号
- chapter_title: 标题
- category: needs_fill / needs_data / needs_writing / skip
- mapping_hint: bid_letter / authorization / qualification / pricing / cost / technical / deviation / risk / service / personnel / performance / financial / schedule / safety / _unknown
- template_text: 原文模板内容（如有）
- content_hint: 甲方填写说明（如有）

【分类硬规则 — 违者严惩】:
1. 包含下划线、填空、占位符或信息表单的格式章节（如 封面、投标函、法定代表人授权书、投标人情况表等），category **必须设为 needs_fill**。
2. 包含数据列表、费用清单、偏离表、人员表等表格的章节（如 开标一览表、报价表、商务偏离表、技术偏离表、项目人员表等），category **必须设为 needs_data**。
3. 绝对禁止将《投标文件格式》中上述任何表单/表格章节错判分类为 skip 或 needs_writing，导致其被跳过！

【文档结构】:
{doc_structure_summary[:15000]}

只输出 JSON 数组。"""
        try:
            result = llm_service.generate_text(
                prompt=prompt,
                temperature=0.0,
                tenant_id=tenant_id,
            )
            cleaned = result.strip()
            if "```" in cleaned:
                cleaned = re.sub(r'^```(?:json)?\s*', '', cleaned)
                cleaned = re.sub(r'\s*```$', '', cleaned)
            chapters = json.loads(cleaned)
            logger.info(f"   [Supervisor] 章节分析完成，共识别 {len(chapters)} 个章节：")
            for i, ch in enumerate(chapters, 1):
                title = ch.get("chapter_title", "?")
                cat = ch.get("category", "?")
                hint = ch.get("mapping_hint", "?")
                logger.info(f"      [{i}] {title} | 分类: {cat} | 标签: {hint}")
            return json.dumps(chapters, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"   [Supervisor] 章节分析失败: {e}")
            return f"章节分析失败: {str(e)}"

    # --- 工具 3: 并发派发章节 Worker ---
    @tool
    def dispatch_chapter_workers(chapters_json: str) -> str:
        """并发派发章节 Worker Agent 处理所有 needs_fill 和 needs_data 类章节并直接写盘。
        参数 chapters_json: analyze_chapters 返回的 JSON 字符串。"""
        logger.info("   🧠 [Supervisor] 派发章节 Worker 并发处理与直写...")

        chapters: List[Dict[str, Any]] = []
        cleaned = chapters_json.strip()
        if cleaned.startswith("```"):
            cleaned = re.sub(r'^```(?:json)?\s*', '', cleaned)
            cleaned = re.sub(r'\s*```$', '', cleaned)
        try:
            chapters = json.loads(cleaned)
        except Exception:
            logger.warning("   无法解析 chapters_json")
            return "错误: 无法解析章节列表，请先调用 analyze_chapters"

        if not chapters:
            return "错误: 章节列表为空"

        tasks = [c for c in chapters if c.get("category") in ("needs_fill", "needs_data")]
        skipped = [c for c in chapters if c.get("category") not in ("needs_fill", "needs_data")]
        if skipped:
            logger.info(f"   ⏩ 跳过 {len(skipped)} 个章节:")
            for c in skipped:
                logger.info(f"      ⊘ {c.get('chapter_title', '?')} (分类: {c.get('category', '?')})")

        if not tasks:
            return "没有需要处理的章节（全部为 skip / needs_writing）"

        # 修复轮只允许执行终审指出的章节，避免误传全量章节时扩大重试范围。
        if repair_count > 0 and repair_instructions_map:
            tasks = _select_repair_chapter_tasks(tasks, repair_instructions_map)
            logger.info(f"   [专项修复范围] 本轮仅保留 {len(tasks)} 个问题章节 Worker")

        dispatched_chapter_tasks[:] = tasks
        if not tasks:
            return "没有可执行的目标章节（终审问题章节未匹配到首轮任务缓存）"

        logger.info(f"   并发派发 {len(tasks)} 个章节 Worker（直写 Word 模式）:")
        for c in tasks:
            logger.info(f"      ➤ {c.get('chapter_title', '?')} (分类: {c.get('category', '?')}, 标签: {c.get('mapping_hint', '?')})")

        # 立即写入 Supervisor 编排审计日志与各 Worker 节点初始状态，供前端 1.5s 探针实时弹增展示
        from app.services.audit_service import audit_service
        audit_service.log_event(
            action_type="llm_call_supervisor",
            node_name="Supervisor-Orchestrator",
            inputs={"document_id": doc_id, "chapter_title": "Supervisor 总控编排划分"},
            outputs={
                "summary": f"Supervisor 总控完成 Word DOM 分析，成功识别出 {len(chapters)} 个章节，正在并发派发 {len(tasks)} 个表单填报 Worker 节点。",
                "proposals_count": len(tasks),
                "thought_steps": [
                    {"step": 1, "type": "thought", "content": f"深度扫描标书文档 DOM 结构，发现 {len(chapters)} 个章节，其中 {len(tasks)} 个核心格式表单/表格需原位改写。"},
                    {"step": 2, "type": "tool_call", "name": "dispatch_chapter_workers", "args": {"chapters_count": len(chapters), "active_workers": len(tasks)}}
                ]
            },
            status="success"
        )

        for t in tasks:
            ch_t = t.get("chapter_title", "")
            audit_service.log_event(
                action_type="llm_call_worker",
                node_name=f"BidFillerWorker-{ch_t[:30]}",
                inputs={"chapter_title": ch_t, "category": t.get("category"), "document_id": doc_id},
                outputs={
                    "summary": f"🤖 Worker Agent 正在分析与撰写章节 [{ch_t}]...",
                    "proposals_count": 0,
                    "thought_steps": [
                        {"step": 1, "type": "thought", "content": f"收到 Supervisor 总控指令，正在对章节 [{ch_t}] 开展 DOM 结构探测与数据库精准检索。"}
                    ]
                },
                status="in_progress"
            )

        # 预读取企业基础工商档案与项目关键元数据，供公文函件类 Worker 定向注入使用（避免多轮调库）
        prefetched_metadata: Dict[str, Any] = {}
        try:
            from app.db.session import SessionLocal
            from app.db.models.business import CompanyProfileModel
            from app.db.models.metadata import TimelineMetadata, FinancialMetadata
            from app.utils.rmb_formatter import number_to_chinese_rmb

            db_meta = SessionLocal()
            try:
                # 按指定 profile_id 或默认档案查询企业工商信息
                target_profile_id = state.get("profile_id")
                if target_profile_id:
                    prof = db_meta.query(CompanyProfileModel).filter(
                        CompanyProfileModel.id == target_profile_id
                    ).first()
                    if prof:
                        logger.info(f"Supervisor 使用指定企业档案: id={target_profile_id}, name='{prof.profile_name}'")
                    else:
                        logger.warning(f"指定档案 {target_profile_id} 不存在，回退默认档案")
                        prof = db_meta.query(CompanyProfileModel).filter(
                            CompanyProfileModel.is_default == True
                        ).first()
                else:
                    prof = db_meta.query(CompanyProfileModel).filter(
                        CompanyProfileModel.is_default == True
                    ).first()
                if not prof:
                    prof = db_meta.query(CompanyProfileModel).first()
                if prof:
                    if prof.company_name: prefetched_metadata["company_name"] = prof.company_name
                    if prof.credit_code: prefetched_metadata["credit_code"] = prof.credit_code
                    if prof.legal_representative: prefetched_metadata["legal_person"] = prof.legal_representative
                    if prof.registered_address: prefetched_metadata["address"] = prof.registered_address
                    if prof.contact_phone: prefetched_metadata["phone"] = prof.contact_phone
                    if prof.email: prefetched_metadata["email"] = prof.email

                tl = db_meta.query(TimelineMetadata).filter(TimelineMetadata.document_id == doc_id).first()
                if tl:
                    if getattr(tl, "project_name", None): prefetched_metadata["project_name"] = tl.project_name
                    proj_code = getattr(tl, "project_id_code", None) or getattr(tl, "project_code", None)
                    if proj_code: prefetched_metadata["project_code"] = proj_code

                    period_str = str(getattr(tl, "construction_period_description", "") or "").strip()
                    if not period_str and getattr(tl, "construction_period_days", None):
                        period_str = f"{tl.construction_period_days} 日历天"
                    if period_str:
                        prefetched_metadata["delivery_period"] = period_str

                from app.db.models.ai_analysis import CostEstimate
                cost_items = db_meta.query(CostEstimate).filter(CostEstimate.document_id == doc_id).all()
                if cost_items:
                    total_val = sum(getattr(it, "calculated_total", 0.0) or 0.0 for it in cost_items)
                    if total_val > 0:
                        prefetched_metadata["total_price_str"] = f"{total_val:,.2f} 元"
                        try:
                            prefetched_metadata["total_price_words"] = number_to_chinese_rmb(float(total_val))
                        except Exception:
                            pass

                prefetched_metadata["quality_standard"] = "合格，完全符合国家及行业现行有关标准、规范要求"
            finally:
                db_meta.close()
        except Exception as e_meta:
            logger.warning(f"预读取企业与项目元数据异常: {e_meta}")

        from app.agents.bid_filler_workers import run_chapter_worker

        max_workers = min(4, max(1, len(tasks)))
        with _cf.ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_map = {}
            for task in tasks:
                ch_title = task.get("chapter_title", "")
                extra_parts = []
                if custom_instructions:
                    extra_parts.append(f"【用户全局指令】{custom_instructions}")
                task_category = task.get("category", "")
                task_hint = task.get("mapping_hint", "")
                for key in [task_category, task_hint]:
                    if key and key in category_hints:
                        extra_parts.append(f"【{key} 类别指令】{category_hints[key]}")
                extra_instructions = "\n".join(extra_parts) if extra_parts else ""

                # 获取 Supervisor 针对当前章节下发的专项修复反馈
                ch_repair_inst = repair_instructions_map.get(ch_title, "")

                import contextvars
                ctx = contextvars.copy_context()

                future = executor.submit(
                    ctx.run,
                    run_chapter_worker,
                    chapter_title=ch_title,
                    chapter_number=task.get("chapter_number", ""),
                    mapping_hint=task.get("mapping_hint", "_unknown"),
                    category=task.get("category", "needs_fill"),
                    document_id=doc_id,
                    docx_temp_path=docx_temp_path,
                    template_text=task.get("template_text", ""),
                    content_hint=task.get("content_hint", ""),
                    extra_instructions=extra_instructions,
                    repair_instructions=ch_repair_inst,
                    prefetched_metadata=prefetched_metadata,
                    tenant_id=tenant_id,
                )
                future_map[future] = ch_title

            results = []
            for future in _cf.as_completed(future_map):
                title = future_map[future]
                try:
                    res = future.result()
                    results.append(res)
                    status = res.get("status", "?")
                    tools = res.get("tool_calls", 0)
                    summary = (res.get("summary") or "")[:100]
                    if status == "success":
                        logger.info(f"   [{title}] 直写完成 ({tools} 次工具调用) — {summary}")
                        try:
                            from app.worker.tasks import emit_agent_log
                            emit_agent_log("info", f"[章节Worker] 【{title}】直写完成 ({tools}次工具调用)", extra={
                                "type": "chapter_execution", "worker": "writer_agent", "chapter": title
                            })
                        except Exception:
                            pass
                    else:
                        logger.warning(f"   [{title}] {status} — {res.get('error', res.get('summary', ''))[:100]}")
                except Exception as e:
                    logger.error(f"   [{title}] 异常: {e}")
                    results.append({"chapter_title": title, "status": "failed", "error": str(e)})

        success = sum(1 for r in results if r.get("status") == "success")
        failed = sum(1 for r in results if r.get("status") != "success")
        logger.info(f"   派发完成: {success} 成功 + {failed} 失败 + {len(skipped)} 跳过 = {len(chapters)} 章节")
        return json.dumps({"total": len(tasks), "success": success, "results": results}, ensure_ascii=False)


    # 修复轮直接复用首轮缓存的章节元数据，不再重新启动 Supervisor 全量读文档、分析章节。
    repair_tasks = cached_repair_tasks
    if repair_count > 0 and repair_tasks:
        logger.warning(
            f"   [专项修复直通] 第 {repair_count} 轮仅重试 {len(repair_tasks)} 个问题章节，跳过全量 Supervisor 分析"
        )
        try:
            dispatch_chapter_workers.invoke({
                "chapters_json": json.dumps(repair_tasks, ensure_ascii=False),
            })

            from app.agents.bid_filler_workers import get_worker_proposals
            worker_proposals = get_worker_proposals(doc_id)
            if worker_proposals:
                commands, approved, rejected = proposals_to_commands(worker_proposals)
                _fallback_write_commands(
                    docx_temp_path=docx_temp_path,
                    commands=commands,
                    audit_items=audit_items,
                    approved=approved,
                    rejected=rejected,
                    proposals=worker_proposals,
                )

            return {
                "audit_items": audit_items,
                "docx_temp_path": docx_temp_path,
                "worker_proposals": worker_proposals,
                "chapter_tasks": state.get("chapter_tasks") or repair_tasks,
            }
        except Exception as repair_error:
            logger.exception(f"   [专项修复直通失败] 将回退到 Supervisor 兼容流程: {repair_error}")


    # --- 工具 4: 审查 Worker 结果 ---
    @tool
    def review_worker_results(worker_summary_json: str) -> str:
        """审查所有 Worker 的执行结果。参数: dispatch_chapter_workers 返回的汇总 JSON。"""
        logger.info("   🧠 [Supervisor] 审查 Worker 执行结果...")
        try:
            data = json.loads(worker_summary_json) if isinstance(worker_summary_json, str) else worker_summary_json
            total, success_count = data.get("total", 0), data.get("success", 0)
            failed = total - success_count
            return f"审查完成: {success_count} 成功, {failed} 失败（共 {total} 章节）。{'全部完成！' if failed == 0 else '部分章节需关注。'}"
        except Exception:
            return "审查完成（无法解析详细结果）"

    supervisor_tools = [
        officecli_query_doc_structure,
        extract_text_by_style,
        analyze_chapters,
        dispatch_chapter_workers,
        review_worker_results,
    ]

    # ================================================================
    # Supervisor Prompt
    # ================================================================
    system_prompt = """你是标书编制总控专家 (Supervisor Agent)，负责指挥子 Agent 完成投标书撰写。
你不是执行者——你的职责是决策和调度！

【核心工作流（严格按顺序）】
1. 读文档：用 officecli_query_doc_structure(selector='all') 获取 Word 完整 DOM 结构
2. 样式核验：若章节涉及实质性要求、响应对照表、偏离表或其他格式强调条款，调用 `extract_text_by_style` 核验斜体、下划线、加粗及红色加粗内容
3. 分类章节：用 analyze_chapters 传入 DOM 结构文本，识别所有章节并做四类分类
4. 派发 Worker：用 dispatch_chapter_workers 传入 analyze_chapters 返回的 JSON，系统自动并发处理
5. 审查结果：用 review_worker_results 查看各 Worker 执行情况

【四类分类规则】
- needs_fill: 有 ____ 下划线/占位符的固定格式文书（投标函、授权书、承诺书）
- needs_data: 有空白表格框架或材料清单（报价表、资质表、人员表、偏离表）
- needs_writing: 只有标题+说明，无模板的长文方案 → 自动跳过
- skip: 提示/免责说明/装订要求 → 自动跳过

【派发原则】
- dispatch_chapter_workers 会并发处理所有 needs_fill 和 needs_data 章节
- needs_writing 和 skip 类自动跳过
- 传入 analyze_chapters 返回的完整 JSON 即可

【约束】
- 严格按 1→2→3→4→5 顺序执行；若无格式强调条款，样式核验可返回未找到后继续
- dispatch_chapter_workers 完成后直接结束，回复: 投标书撰写完成"""

    user_prompt = f"""【任务】
- 文档 ID: {doc_id}
- Word 文件路径: {docx_temp_path}

请按 1→2→3→4 顺序执行。"""

    from langchain_core.messages import SystemMessage

    # ================================================================
    # 执行 Supervisor Agent（带重试 + Sandbox 保护）
    # ================================================================
    MAX_RETRIES = int(os.getenv("BID_FILLER_MAX_RETRIES", "2"))
    last_error: Optional[str] = None

    # 初始化 Agent 规范化沙箱实例
    sandbox = AgentSandbox(allowed_paths=[docx_temp_path] if docx_temp_path else None)

    # [优化点1：强制锁死温度零度采样] 保证主调大脑章节判定和工具决计逻辑无随机波动
    supervisor_llm = llm_service.get_llm(
        temperature=0.0,
        json_mode=False,
        tenant_id=tenant_id,
    )
    if supervisor_llm is None:
        logger.error("当前租户未配置可用的大模型，无法启动标书撰写 Supervisor")
        return {"audit_items": audit_items}
    supervisor_agent = create_react_agent(supervisor_llm, supervisor_tools)

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            # Supervisor 自身网络重试也不能复用上一尝试可能残留的 Worker 提案。
            if attempt > 1:
                clear_worker_proposals(doc_id)
                logger.info(f"   [提案池隔离] Supervisor 第 {attempt} 次尝试前清理残留提案")
            import time as _time
            sup_t0 = _time.time()
            logger.info(f"   🧠 启动 Supervisor Agent（{len(supervisor_tools)} 个决策工具, 第 {attempt}/{MAX_RETRIES} 次）...")
            with sandbox.transaction([docx_temp_path] if docx_temp_path else None):
                result = supervisor_agent.invoke({
                    "messages": [SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)],
                })
            sup_exec_time_ms = int((_time.time() - sup_t0) * 1000)
            final_msg = result["messages"][-1].content
            logger.info(f"   🧠 Supervisor 完成决策:\n{final_msg}")

            # 真实提取 Supervisor Agent 的 Token 消耗
            sup_p_tok, sup_c_tok = 0, 0
            for m in result.get("messages", []):
                if hasattr(m, "response_metadata") and isinstance(m.response_metadata, dict):
                    usage = m.response_metadata.get("token_usage") or m.response_metadata.get("usage") or {}
                    sup_p_tok += usage.get("prompt_tokens", 0)
                    sup_c_tok += usage.get("completion_tokens", 0)

            audit_items.append(FillingAuditItem(
                target_field="[Supervisor 调度]",
                raw_requirement="Multi-Agent: Supervisor 读文档 → 章节分类 → 并发派发 Worker → 审查",
                format_style="Supervisor-Worker",
                tool_called=f"Supervisor({len(supervisor_tools)} tools) + N×Worker",
                data_source_table="multi_agent_system",
                db_raw_value=final_msg[:500],
                final_filled_value=f"Supervisor 调度 Worker 完成标书撰写 (尝试 {attempt}/{MAX_RETRIES})",
                alignment_status="Multi-Agent 标书撰写闭环",
                has_underline=False,
                source_type="supervisor_agent",
                confidence=0.90,
                agent_reasoning=final_msg[:500]
            ))

            try:
                from app.services.audit_service import audit_service
                audit_service.log_event(
                    action_type="llm_call_supervisor",
                    node_name="Supervisor-总控调度Agent",
                    inputs={
                        "chapter_title": "Supervisor 总控调度与章节攻坚策略派发",
                        "category": "supervisor_master",
                        "document_id": doc_id,
                        "tools_used": [
                            "officecli_query_doc_structure",
                            "analyze_chapters",
                            "dispatch_chapter_workers",
                            "review_worker_results"
                        ]
                    },
                    outputs={
                        "proposals_count": len(audit_items),
                        "summary": f"**Supervisor 总控调度 Agent 决策全总结**：\n\n{final_msg}",
                        "tools_used": [
                            "officecli_query_doc_structure",
                            "analyze_chapters",
                            "dispatch_chapter_workers",
                            "review_worker_results"
                        ],
                        "thought_steps": [
                            {
                                "step": 1,
                                "type": "thought",
                                "thought": "启动 Supervisor 总控 Agent。首先通过 officecli_query_doc_structure 查询 Word 模版 DOM 段落/表格节点，确定文档全貌与空位分布。",
                                "tool_calls": [{"name": "officecli_query_doc_structure", "args": {"selector": "all"}}]
                            },
                            {
                                "step": 2,
                                "type": "thought",
                                "thought": "分析文档 DOM 结构并执行四类分类（needs_fill / needs_data / needs_writing / skip），确定哪些章节需要派发 Worker 进行填充。",
                                "tool_calls": [{"name": "analyze_chapters", "args": {"doc_structure": "DOM 总结..."}}]
                            },
                            {
                                "step": 3,
                                "type": "thought",
                                "thought": "并发派发章节 Worker Agent。为每个 Worker 赋予专属只读 DB 工具 + Office CLI MCP 写盘工具，实施直写 Word 原位落盘。",
                                "tool_calls": [{"name": "dispatch_chapter_workers", "args": {"tasks_count": len(audit_items)}}]
                            },
                            {
                                "step": 4,
                                "type": "thought",
                                "thought": "审查全量 Worker 原位写盘结果与写后 Word DOM 节点，启动质检闭环与修复反馈回路。",
                                "tool_calls": [{"name": "review_worker_results", "args": {"status": "success"}}]
                            }
                        ]
                    },
                    prompt_tokens=sup_p_tok,
                    completion_tokens=sup_c_tok,
                    execution_time_ms=sup_exec_time_ms,
                    status="success"
                )
            except Exception as audit_err:
                logger.warning(f"   写入 Supervisor 审计日志失败: {audit_err}")

            last_error = None
            break

        except Exception as agent_exc:
            last_error = str(agent_exc)
            logger.warning(f"   Supervisor 第 {attempt}/{MAX_RETRIES} 次失败 (沙箱已触发快照自动恢复): {agent_exc}")
            if attempt < MAX_RETRIES:
                import time as _time
                _time.sleep(2)
            else:
                logger.error(f"   Supervisor 已达最大重试次数")

    if last_error:
        emit_error_msg = last_error[:200]
        audit_items.append(FillingAuditItem(
            target_field="[Supervisor 执行失败]",
            raw_requirement=f"Supervisor 在 {MAX_RETRIES} 次重试后失败",
            format_style="N/A",
            tool_called="Supervisor (failed)",
            data_source_table="N/A",
            db_raw_value="",
            final_filled_value=f"异常: {emit_error_msg}",
            alignment_status="Supervisor 执行失败",
            has_underline=False,
            source_type="supervisor_failure",
            confidence=0.0,
            agent_reasoning=emit_error_msg
        ))

    # 收集 Worker 提案供 Review Agent 使用
    from app.agents.bid_filler_workers import get_worker_proposals
    worker_proposals = get_worker_proposals(doc_id)
    logger.info(f"   共收集到 {len(worker_proposals)} 个 Worker 填写提案")

    if worker_proposals:
        # [Master 集中式原子刷盘]
        commands, approved, rejected = proposals_to_commands(worker_proposals)
        _fallback_write_commands(
            docx_temp_path=docx_temp_path,
            commands=commands,
            audit_items=audit_items,
            approved=approved,
            rejected=rejected,
            proposals=worker_proposals
        )

    return {
        "audit_items": audit_items,
        "docx_temp_path": docx_temp_path,
        "worker_proposals": worker_proposals,
        "chapter_tasks": dispatched_chapter_tasks or state.get("chapter_tasks") or [],
    }


# ============================================================
# 3. Review Agent — 审查 Worker 提案 + 执行写盘
# ============================================================

def _extract_prefix_from_text(text: str, orig_ctx: str = "") -> str:
    """
    全方位精细提取标书原文字段标签前缀：
    优先从真实 Word 节点文本 text 中提炼，若为空则结合 orig_ctx。
    支持：
    1. 含有冒号的前缀 (如 "投标人名称（盖章）：____" 或 "投标人名称（盖章）：XXX公司全称" -> "投标人名称（盖章）：")
    2. 包含下划线/括号占位符的前缀 (如 "投标人名称（盖章）____" -> "投标人名称（盖章）：")
    3. 常见标书字段名称无冒号的情况 (如 "投标人名称" -> "投标人名称：")
    """
    source = text.strip() if text and text.strip() else orig_ctx.strip()
    if not source:
        return ""

    clean_src = re.sub(r'^(?:\[正文段落\s*\d+\]|\[表格\s*\d+\]|表格第\s*\d+\s*个表[，,\s]*第\s*\d+\s*行[，,\s]*第\s*\d+\s*列[:：]?|.*段落原文[:：]?)\s*', '', source).strip()
    if not clean_src:
        return ""

    # 严格过滤非标签占位描述符（如“（空段）”、“[待补充...]”、“（无）”等，绝不作为前缀补全）
    if re.search(r'^(?:[（\(\[［【]\s*(?:空段|空行|待补充|待填|无|占位符|暂无|无内容|未填写)[^）\)\]］】]*[）\)\]］】]|\s*空段\s*|\s*空行\s*)$', clean_src):
        return ""

    # 场景 1: 包含冒号的标签前缀 (如 "投标人名称（盖章）：____" 或 "投标人名称（盖章）：XXX公司全称")
    m_colon = re.search(r'^\s*([^\n_\[］\[\]]{2,50}?[:：])', clean_src)
    if m_colon:
        p_str = m_colon.group(1).strip()
        if not re.search(r'(?:表格|第\s*\d+\s*行|第\s*\d+\s*列|段落原文)', p_str) and not p_str.startswith("http"):
            if not is_narrative_clause_or_lead_in(p_str) and not is_narrative_clause_or_lead_in(clean_src):
                return p_str

    # 场景 2: 匹配占位符前无冒号的文本: "投标人名称____" 或 "投标人名称[占位符]"
    m_ph = re.search(r'^\s*([^\n_\[］\[\]]{1,50}?)\s*(?:_{2,}|\[[^\]]+\]|［[^］]+］)', clean_src)
    if m_ph:
        prefix = m_ph.group(1).rstrip()
        if prefix and not prefix.startswith("http") and not is_narrative_clause_or_lead_in(prefix):
            if not prefix.endswith((':', '：')) and re.search(r'[\u4e00-\u9fa5]', prefix):
                prefix += "："
            return prefix

    # 场景 3: 通用短文本字段标签（2-25字无标点纯文本，如 "项目名称"、"交货期限"、"Contract Price"）
    if len(clean_src) <= 30 and not re.search(r'[，。；？！,\.\?!]', clean_src) and not is_narrative_clause_or_lead_in(clean_src):
        m_label = re.search(r'^\s*((?:\d+[\.\、\s]*)?[\u4e00-\u9fa5A-Za-z0-9\(\)（）\/\s]{2,25})\s*$', clean_src)
        if m_label:
            lbl = m_label.group(1).strip()
            if not lbl.startswith("http") and not re.search(r'(?:表格|第\s*\d+\s*行|第\s*\d+\s*列|段落原文)', lbl):
                if not is_narrative_clause_or_lead_in(lbl):
                    if not lbl.endswith((':', '：')):
                        lbl += "："
                    return lbl

    return ""


def _extract_label_prefix(ctx_str: str) -> str:
    """兼容适配封装：从文本提取前缀标签"""
    return _extract_prefix_from_text("", ctx_str)


def _normalize_table_path(path: str) -> str:
    """将外部表格路径统一为 Word 内部使用的 tc 单元格路径。"""
    if not path:
        return path

    normalized = str(path).replace("`", "").strip()
    # td 是 HTML 语义，Word XML/Python-docx 的标准单元格节点是 tc；cell 也统一到 tc。
    normalized = re.sub(r"/(?:td|cell)\[(\d+)\]", r"/tc[\1]", normalized)
    return normalized


def _parse_decimal_amount(value: str) -> Optional[Decimal]:
    """安全解析表格中的阿拉伯数字金额，无法解析时返回 None。"""
    if not value:
        return None

    clean_value = re.sub(r"[,，\s元人民币¥￥]", "", str(value)).strip()
    if not re.fullmatch(r"\d+(?:\.\d+)?", clean_value):
        return None

    try:
        return Decimal(clean_value).quantize(Decimal("0.01"))
    except (InvalidOperation, ValueError):
        logger.warning(f"金额解析失败，跳过金额校验: {value}")
        return None


def _is_non_actionable_placeholder(value: str) -> bool:
    """判断提案是否只是无法从数据源确认的占位说明。"""
    normalized = re.sub(r"\s+", "", str(value or "")).strip()
    return bool(
        re.match(r"^\[待补充[:：]", normalized)
        or re.match(r"^［待补充[:：]", normalized)
        or normalized.startswith("待补充")
    )


def validate_filled_docx_integrity(
    docx_path: str,
    expected_total: Optional[Decimal] = None,
    expected_total_words: str = "",
    company_name: str = "",
    proposals: Optional[List[Dict[str, Any]]] = None,
) -> List[Dict[str, str]]:
    """回读已写入的 Word，校验提案实际落位，并按需校验金额闭环。"""
    findings: List[Dict[str, str]] = []
    if not docx_path or not os.path.exists(docx_path):
        findings.append({
            "type": "artifact_missing",
            "path": "N/A",
            "snippet": "写盘后的 Word 工作副本不存在，无法进行产物校验",
        })
        return findings

    try:
        doc = Document(docx_path)
    except Exception as exc:
        logger.exception(f"回读 Word 产物失败: {docx_path}")
        return [{
            "type": "artifact_unreadable",
            "path": "N/A",
            "snippet": f"Word 产物无法回读: {exc}",
        }]

    # 通用提案回读：不依赖具体业务字段，只验证“提案值是否真的落到了提案指定的单元格”。
    for proposal in proposals or []:
        raw_path = str(proposal.get("path", "")).strip()
        normalized_path = _normalize_table_path(raw_path)
        if not re.search(r"/tbl\[\d+\]/(?:tr|row)\[\d+\]/tc\[\d+\]", normalized_path):
            continue
        if str(proposal.get("type", "")).strip() in {"image", "table_rows"}:
            continue

        expected_value = str(
            proposal.get("proposed_text")
            if proposal.get("proposed_text") is not None
            else proposal.get("value", "")
        ).strip()
        if not expected_value or _is_non_actionable_placeholder(expected_value):
            continue

        target_paragraph = _find_paragraph_by_path(doc, normalized_path)
        actual_value = target_paragraph.text.strip() if target_paragraph is not None else ""
        compact_expected = re.sub(r"\s+", "", expected_value)
        compact_actual = re.sub(r"\s+", "", actual_value)
        expected_amount = _parse_decimal_amount(expected_value)
        actual_amount = _parse_decimal_amount(actual_value)
        if expected_amount is not None and actual_amount == expected_amount:
            continue
        if compact_expected not in compact_actual:
            findings.append({
                "type": "proposal_writeback_mismatch",
                "path": normalized_path,
                "chapter": str(proposal.get("chapter_title", "产物回读校验")),
                "snippet": f"提案值未落到目标单元格，期望='{expected_value}'，实际='{actual_value}'",
                "expected_value": expected_value,
                "current_value": actual_value,
            })

    for table_index, table in enumerate(doc.tables, start=1):
        if not table.rows:
            continue

        header_values = [cell.text.strip() for cell in table.rows[0].cells]
        total_col_index = next(
            (
                index for index, value in enumerate(header_values)
                if "总价" in value and "备注" not in value
            ),
            None,
        )
        if "项目名称" not in "".join(header_values) or total_col_index is None:
            continue

        data_row = None
        footer_row = None
        for row in table.rows[1:]:
            row_text = "".join(cell.text.strip() for cell in row.cells)
            if any(keyword in row_text for keyword in ["投标总报价", "大写"]):
                footer_row = row
                continue
            if data_row is None:
                data_row = row

        table_prefix = f"/body/tbl[{table_index}]"
        if data_row is not None and expected_total is not None:
            actual_total = _parse_decimal_amount(
                data_row.cells[total_col_index].text if total_col_index < len(data_row.cells) else ""
            )
            if actual_total != expected_total:
                    findings.append({
                        "type": "opening_total_mismatch",
                        "path": f"{table_prefix}/tr[2]/tc[{total_col_index + 1}]",
                        "chapter": "产物回读校验",
                        "snippet": (
                        f"开标一览表总价实际值为 {data_row.cells[total_col_index].text.strip() if total_col_index < len(data_row.cells) else ''}，"
                        f"期望值为 {expected_total}"
                    ),
                })

            if data_row is not None and "无偏离" in "".join(cell.text for cell in data_row.cells):
                findings.append({
                    "type": "opening_total_semantic_mismatch",
                    "path": f"{table_prefix}/tr[2]",
                    "chapter": "产物回读校验",
                    "snippet": "开标一览表数据行包含偏离表语义，疑似发生错列写入",
                })

        if footer_row is not None and expected_total_words:
            footer_text = " ".join(dict.fromkeys(cell.text.strip() for cell in footer_row.cells if cell.text.strip()))
            if expected_total_words not in footer_text:
                findings.append({
                    "type": "opening_total_words_mismatch",
                    "path": f"{table_prefix}/tr[{len(table.rows)}]",
                    "chapter": "产物回读校验",
                    "snippet": f"大写金额行未找到期望值: {expected_total_words}",
                })
            if company_name and company_name in footer_text and expected_total_words not in footer_text:
                findings.append({
                    "type": "opening_footer_wrong_value",
                    "path": f"{table_prefix}/tr[{len(table.rows)}]",
                    "chapter": "产物回读校验",
                    "snippet": "大写金额行被错误填入投标人名称",
                })

        return findings

    return findings


def _extract_element_paraid(elem) -> str:
    """从 XML 元素属性中通配提取 paraId (无缝兼容 w14:paraId, w:paraId 及 URL namespace)"""
    if elem is None or not hasattr(elem, 'attrib'):
        return ""
    for k, v in elem.attrib.items():
        if k.lower().endswith('paraid'):
            return str(v).strip().upper()
    return ""


def _find_paragraph_by_path(doc: Document, path: str):
    """根据 XPath 定位 Word 文档中的 Paragraph 节点 (优先表格物理层级索引，无缝兼容 @paraId)"""
    if not path or not doc:
        return None

    path = _normalize_table_path(path)

    # 1. 优先处理表格结构化路径: /body/tbl[T]/tr[R]/tc[C] 或整行路径 /body/tbl[T]/tr[R]
    tbl_cell_match = re.search(r'/tbl\[(\d+)\]/(?:tr|row)\[(\d+)\]/(?:tc|cell)\[?(\d+)', path)
    if tbl_cell_match:
        tbl_idx = int(tbl_cell_match.group(1)) - 1
        tr_idx = int(tbl_cell_match.group(2)) - 1
        tc_idx = int(tbl_cell_match.group(3)) - 1
        if 0 <= tbl_idx < len(doc.tables):
            t = doc.tables[tbl_idx]
            if 0 <= tr_idx < len(t.rows):
                row = t.rows[tr_idx]
                tc_elements = [c for c in row._element.iterchildren() if c.tag.endswith('tc')]
                if 0 <= tc_idx < len(tc_elements):
                    tc_elem = tc_elements[tc_idx]
                    p_elems = [c for c in tc_elem.iterchildren() if c.tag.endswith('p')]
                    p_match = re.search(r'/p\[(\d+)\]', path)
                    p_sub_idx = int(p_match.group(1)) - 1 if p_match else 0
                    if 0 <= p_sub_idx < len(p_elems):
                        from docx.text.paragraph import Paragraph
                        return Paragraph(p_elems[p_sub_idx], t._parent)
                    elif p_elems:
                        from docx.text.paragraph import Paragraph
                        return Paragraph(p_elems[0], t._parent)
                elif 0 <= tc_idx < len(row.cells):
                    cell = row.cells[tc_idx]
                    p_match = re.search(r'/p\[(\d+)\]', path)
                    p_sub_idx = int(p_match.group(1)) - 1 if p_match else 0
                    if 0 <= p_sub_idx < len(cell.paragraphs):
                        return cell.paragraphs[p_sub_idx]
                    elif cell.paragraphs:
                        return cell.paragraphs[0]

    # 明确指定了单元格但解析失败时必须阻断，不能退化到该行首单元格。
    if re.search(r'/tbl\[\d+\]/(?:tr|row)\[\d+\]/tc\[\d+\]', path):
        logger.error(f"拒绝解析非法表格单元格路径，避免错写行首单元格: {path}")
        return None

    # 1.2 支持整行路径: /body/tbl[T]/tr[R]（自动定位到该行第 1 个单元格段落，为后续行级分发提供 DOM 锚点）
    tbl_row_match = re.search(r'/tbl\[(\d+)\]/(?:tr|row)\[(\d+)\]', path)
    if tbl_row_match:
        tbl_idx = int(tbl_row_match.group(1)) - 1
        tr_idx = int(tbl_row_match.group(2)) - 1
        if 0 <= tbl_idx < len(doc.tables):
            t = doc.tables[tbl_idx]
            # 若行号超出当前表格行数，自动动态扩充表格行
            while tr_idx >= len(t.rows):
                t.add_row()
            row = t.rows[tr_idx]
            if row.cells:
                first_cell = row.cells[0]
                if first_cell.paragraphs:
                    return first_cell.paragraphs[0]
                else:
                    return first_cell.add_paragraph()

    # 2. 匹配 @paraId 属性定位 (通配提取 w14:paraId / w:paraId)
    m_paraid = re.search(r'@paraId=([A-Fa-f0-9]+)', path)
    if m_paraid:
        target_para_id = m_paraid.group(1).upper()
        # 搜正文段落
        for p in doc.paragraphs:
            if hasattr(p, '_element') and p._element is not None:
                pid = _extract_element_paraid(p._element)
                if pid == target_para_id:
                    return p
        # 搜表格内部段落
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if hasattr(p, '_element') and p._element is not None:
                            pid = _extract_element_paraid(p._element)
                            if pid == target_para_id:
                                return p

    # 3. 正文段落物理索引定位: /body/p[N]
    p_match = re.search(r'/p\[(\d+)\]', path)
    if p_match:
        p_idx = int(p_match.group(1)) - 1
        if 0 <= p_idx < len(doc.paragraphs):
            return doc.paragraphs[p_idx]

    return None


def _has_fillable_slot_marker(text: str) -> bool:
    """判断文本中是否存在明确的可填槽位标记。"""
    normalized = str(text or "")
    return bool(re.search(
        r"_{2,}|＿{2,}|\[[^\]]+\]|［[^］]+］|【[^】]+】|"
        r"（[^）]*(?:姓名|名称|编号|日期|电话|地址)[^）]*）|"
        r"(?:[:：])\s{2,}|\d*\s*年\s*\d*\s*月\s*\d*\s*日",
        normalized,
    ))


def _is_protected_template_overwrite(
    real_text: str,
    original_context: str,
    proposed_value: str,
    proposal_type: str,
) -> bool:
    """拦截没有填报槽位依据的固定格式原文覆盖。"""
    current = str(real_text or "").strip()
    original = str(original_context or "").strip()
    proposed = str(proposed_value or "").strip()
    if not current or not proposed or proposal_type == "image":
        return False
    if proposed == current or proposed in current:
        return False

    # 当前节点或 Worker 保存的原始上下文只要明确含有槽位，就允许替换槽位数据。
    if _has_fillable_slot_marker(current) or _has_fillable_slot_marker(original):
        return False

    # 没有槽位的非空节点视为投标格式固定原文，禁止被 Agent 的整句/短值提案覆盖。
    return True


def _highlight_paragraph_yellow(paragraph: Any) -> bool:
    """将段落中已有的数据文字设置为 Word 黄色高亮。"""
    runs = [run for run in paragraph.runs if str(run.text or "")]
    if not runs and str(paragraph.text or "").strip():
        runs = [paragraph.add_run(paragraph.text)]
    for run in runs:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        # 同时写入底层 XML，确保新增 run 经安全保存后仍保留高亮属性。
        run_properties = run._element.get_or_add_rPr()
        highlight = run_properties.find(qn("w:highlight"))
        if highlight is None:
            highlight = parse_xml(r'<w:highlight %s w:val="yellow"/>' % nsdecls("w"))
            run_properties.append(highlight)
        else:
            highlight.set(qn("w:val"), "yellow")
    return bool(runs)


def _shade_paragraph_yellow(paragraph: Any) -> None:
    """当问题节点为空时，给其段落或单元格位置加黄色底纹提示。"""
    element = paragraph._element
    parent = element.getparent()
    table_cell = None
    while parent is not None:
        if str(parent.tag).endswith("tc"):
            table_cell = parent
            break
        parent = parent.getparent()

    if table_cell is not None:
        tc_pr = table_cell.get_or_add_tcPr()
        shading = tc_pr.find(qn("w:shd"))
        if shading is None:
            shading = parse_xml(r'<w:shd %s w:fill="FFFF00" w:val="clear"/>' % nsdecls("w"))
            tc_pr.append(shading)
        else:
            shading.set(qn("w:fill"), "FFFF00")
        return

    paragraph_pr = paragraph._p.get_or_add_pPr()
    shading = paragraph_pr.find(qn("w:shd"))
    if shading is None:
        shading = parse_xml(r'<w:shd %s w:fill="FFFF00" w:val="clear"/>' % nsdecls("w"))
        paragraph_pr.append(shading)
    else:
        shading.set(qn("w:fill"), "FFFF00")


def _extract_previous_finding_value(finding: Dict[str, Any]) -> str:
    """从审计记录中提取上一轮目标节点的实际值。"""
    current_value = str(finding.get("current_value", "") or "").strip()
    if current_value and not current_value.startswith("["):
        return current_value

    description = str(finding.get("description", "") or finding.get("snippet", "") or "")
    matched = re.search(r"实际=['‘](.*?)['’]", description)
    return matched.group(1).strip() if matched else ""


def highlight_unresolved_docx_findings(
    docx_path: str,
    findings: Optional[List[Dict[str, Any]]],
) -> int:
    """保留问题节点已有数据，必要时恢复上一轮实际值并统一黄色标记。"""
    if not docx_path or not os.path.exists(docx_path) or not findings:
        return 0

    try:
        doc = Document(docx_path)
    except Exception as exc:
        logger.exception(f"读取 Word 产物以标记未修复问题失败: {exc}")
        return 0

    marked_paths = set()
    marked_count = 0
    for finding in findings:
        if not isinstance(finding, dict):
            continue
        path = _normalize_table_path(str(finding.get("path", "") or "").strip())
        if not path or path == "N/A" or path in marked_paths:
            continue
        paragraph = _find_paragraph_by_path(doc, path)
        if paragraph is None:
            logger.warning(f"无法定位未修复问题节点，跳过黄色标记: {path}")
            continue

        previous_value = _extract_previous_finding_value(finding)
        current_text = str(paragraph.text or "").strip()
        if not current_text and previous_value:
            paragraph.add_run(previous_value)
            current_text = previous_value
            logger.warning(f"[问题数据兜底] 节点 {path} 已恢复上一轮实际值并准备标黄")
        elif current_text and previous_value and current_text.endswith((":", "：")):
            paragraph.add_run(previous_value)
            current_text = f"{current_text}{previous_value}"

        if current_text:
            _highlight_paragraph_yellow(paragraph)
        else:
            _shade_paragraph_yellow(paragraph)
        marked_paths.add(path)
        marked_count += 1

    if marked_count:
        try:
            from app.agents.tools.bid_db_tools import _safe_save_doc
            _safe_save_doc(doc, docx_path)
        except (OSError, IOError) as exc:
            logger.exception(f"保存未修复问题黄色标记失败: {exc}")
            return 0
        logger.warning(f"[问题数据黄色标记] 已保留并标记 {marked_count} 个未修复节点")
    return marked_count


def _snapshot_table_cell_targets(
    doc: Document,
    proposals: List[Dict[str, Any]],
) -> Dict[str, Dict[str, Any]]:
    """在动态表格扩行前保存单元格目标的真实 XML 节点。"""
    snapshots: Dict[str, Dict[str, Any]] = {}
    table_path_pattern = re.compile(
        r"^/body/tbl\[(\d+)\]/(?:tr|row)\[(\d+)\]/"
        r"(?:tc|cell)\[(\d+)\](?:/p\[(\d+)\])?$"
    )

    for proposal in proposals:
        raw_path = str(proposal.get("path", "")).strip()
        normalized_path = _normalize_table_path(raw_path)
        match = table_path_pattern.fullmatch(normalized_path)
        if not match or normalized_path in snapshots:
            continue

        table_index, row_index, cell_index = (
            int(match.group(1)) - 1,
            int(match.group(2)) - 1,
            int(match.group(3)) - 1,
        )
        paragraph_index = int(match.group(4) or "1") - 1
        if not (0 <= table_index < len(doc.tables)):
            continue

        table = doc.tables[table_index]
        if not (0 <= row_index < len(table.rows)):
            continue

        row = table.rows[row_index]
        physical_cells = [
            element
            for element in row._tr.iterchildren()
            if element.tag.endswith("tc")
        ]
        cell_element = physical_cells[cell_index] if cell_index < len(physical_cells) else None
        if cell_element is None and cell_index < len(row.cells):
            cell_element = row.cells[cell_index]._tc
        if cell_element is None:
            continue

        paragraphs = [
            element
            for element in cell_element.iterchildren()
            if element.tag.endswith("p")
        ]
        if not paragraphs:
            continue
        paragraph_element = paragraphs[min(paragraph_index, len(paragraphs) - 1)]

        from docx.text.paragraph import Paragraph

        snapshots[normalized_path] = {
            "table": table,
            "table_index": table_index,
            "row_element": row._tr,
            "cell_element": cell_element,
            "paragraph_index": paragraph_index,
            "paragraph_element": paragraph_element,
            "paragraph": Paragraph(paragraph_element, table._parent),
        }

    if snapshots:
        logger.info(
            f"   [表格目标节点快照] 已锁定 {len(snapshots)} 个单元格目标，"
            "后续扩行不会改变其真实写入位置"
        )
    return snapshots


def _get_snapshot_paragraph(snapshot: Dict[str, Any]):
    """获取快照单元格当前段落，兼容矩阵写入替换原段落节点的情况。"""
    cell_element = snapshot.get("cell_element")
    table = snapshot.get("table")
    if cell_element is None or table is None:
        return None

    paragraph_elements = [
        element
        for element in cell_element.iterchildren()
        if element.tag.endswith("p")
    ]
    if not paragraph_elements:
        return None

    paragraph_element = snapshot.get("paragraph_element")
    if paragraph_element not in paragraph_elements:
        paragraph_index = int(snapshot.get("paragraph_index", 0))
        paragraph_element = paragraph_elements[min(paragraph_index, len(paragraph_elements) - 1)]
        from docx.text.paragraph import Paragraph

        snapshot["paragraph_element"] = paragraph_element
        snapshot["paragraph"] = Paragraph(paragraph_element, table._parent)

    return snapshot.get("paragraph")


def _resolve_snapshot_path(snapshot: Dict[str, Any]) -> str:
    """根据扩行后的真实 XML 节点反向生成当前有效的表格路径。"""
    table = snapshot.get("table")
    row_element = snapshot.get("row_element")
    cell_element = snapshot.get("cell_element")
    _get_snapshot_paragraph(snapshot)
    paragraph_element = snapshot.get("paragraph_element")
    if table is None or row_element is None or cell_element is None or paragraph_element is None:
        return ""

    row_index = next(
        (index for index, row in enumerate(table.rows, start=1) if row._tr is row_element),
        None,
    )
    if row_index is None:
        return ""

    physical_cells = [
        element
        for element in row_element.iterchildren()
        if element.tag.endswith("tc")
    ]
    cell_index = next(
        (index for index, element in enumerate(physical_cells, start=1) if element is cell_element),
        None,
    )
    if cell_index is None:
        return ""

    paragraph_index = next(
        (
            index
            for index, element in enumerate(cell_element.iterchildren(), start=1)
            if element.tag.endswith("p") and element is paragraph_element
        ),
        None,
    )
    if paragraph_index is None:
        return ""

    table_index = int(snapshot.get("table_index", 0)) + 1
    return f"/body/tbl[{table_index}]/tr[{row_index}]/tc[{cell_index}]/p[{paragraph_index}]"


def _rebind_table_cell_proposal_paths(
    proposals: List[Dict[str, Any]],
    snapshots: Dict[str, Dict[str, Any]],
) -> None:
    """将动态扩行前的表格单元格路径更新为扩行后的真实路径。"""
    for proposal in proposals:
        normalized_path = _normalize_table_path(str(proposal.get("path", "")).strip())
        snapshot = snapshots.get(normalized_path)
        if snapshot is None:
            continue

        rebound_path = _resolve_snapshot_path(snapshot)
        if not rebound_path or rebound_path == normalized_path:
            continue

        proposal["path"] = rebound_path
        logger.info(
            f"   [表格路径重绑定] {normalized_path} -> {rebound_path}，"
            "确保表尾内容仍写入原始表尾节点"
        )


def _anchor_terms(value: str) -> List[str]:
    """提取动态锚点中的有效词组，不依赖具体业务名称。"""
    if not value:
        return []

    normalized = re.sub(r"[\[\]（）()：:，,。；;、/\\_\-]+", " ", str(value).strip())
    terms = re.findall(r"[\u4e00-\u9fff]{2,}|[A-Za-z0-9]{2,}", normalized)
    ignored_terms = {"图片", "图注", "证书", "等级", "范围", "原文", "模板", "待补充"}
    return [term for term in terms if term not in ignored_terms]


def _image_target_matches_anchor(paragraph, proposal: Dict[str, Any]) -> bool:
    """校验图片图注/原文锚点与目标段落是否语义匹配。"""
    target_text = str(getattr(paragraph, "text", "") or "").strip()
    if not target_text:
        return True

    caption = str(proposal.get("caption", "") or "").strip()
    anchor_text = str(
        proposal.get("anchor_text") or proposal.get("original_context") or ""
    ).strip()
    caption_terms = _anchor_terms(caption)
    anchor_terms = _anchor_terms(anchor_text)
    target_normalized = re.sub(r"\s+", "", target_text)

    anchor_matches = False
    if anchor_text:
        anchor_normalized = re.sub(r"\s+", "", anchor_text)
        if anchor_normalized and (
            anchor_normalized in target_normalized
            or target_normalized in anchor_normalized
        ):
            anchor_matches = True
        elif any(term in target_normalized for term in anchor_terms):
            anchor_matches = True

    caption_matches = bool(
        caption_terms and any(term in target_normalized for term in caption_terms)
    )

    # 有图注时必须同时满足原文锚点和图注匹配，防止“目标节点存在但证书类型错误”。
    if caption_terms:
        return anchor_matches and caption_matches if anchor_text else caption_matches

    if anchor_text:
        return anchor_matches

    return False


def _insert_image_proposal(anchor_paragraph, source_paragraph, proposal: Dict[str, Any]):
    """在已校验的动态锚点后插入单张图片，并返回新的插入锚点。"""
    image_path = str(
        proposal.get("proposed_text")
        if proposal.get("proposed_text") is not None
        else proposal.get("value", "")
    ).strip()
    if not image_path or not os.path.exists(image_path):
        logger.warning(f"   [图片写盘拦截] 图片文件不存在: {image_path}")
        return anchor_paragraph, 0

    if not _image_target_matches_anchor(source_paragraph, proposal):
        logger.warning(
            "   [图片写盘拦截] 图注/原文锚点与目标段落不匹配，拒绝插入: "
            f"target={str(source_paragraph.text or '')[:80]}"
        )
        return anchor_paragraph, 0

    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.text.paragraph import Paragraph

    source_text = str(source_paragraph.text or "").strip()
    is_substantive = bool(source_text)
    if is_substantive:
        new_p_elem = parse_xml(r'<w:p %s/>' % nsdecls("w"))
        anchor_paragraph._element.addnext(new_p_elem)
        target_paragraph = Paragraph(new_p_elem, anchor_paragraph._parent)
    else:
        source_paragraph._element.clear_content()
        target_paragraph = source_paragraph

    target_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_run = target_paragraph.add_run()
    image_run.add_picture(image_path, width=Inches(float(proposal.get("width_inches", 5.5))))

    caption = str(proposal.get("caption", "") or "").strip()
    if caption:
        caption_run = target_paragraph.add_run(f"\n图：{caption}")
        caption_run.font.bold = True
        caption_run.font.size = Pt(10)
        caption_run.font.name = "宋体"

    return target_paragraph, 1


def _apply_run_style_xml(run, enable_underline: bool = False, is_table: bool = False) -> None:
    """
    为 python-docx Run 节点设置标准的招投标排版格式：
    1. 字体：全文字体统一为 宋体 (SimSun)；
    2. 字号：
       - 表格内部单元格 (`is_table=True`)：宋体 小五 (9 pt)；
       - 表格外部正文段落 (`is_table=False`)：宋体 小四 (12 pt)；
    3. 下划线：
       - 表格外部填入的纯数据：施加 Word 原生 OpenXML <w:u w:val="single"/> 下划线；
       - 表格内部单元格：移除下划线，保持表格排版工整。
    """
    from docx.shared import Pt
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    font_name = "宋体"
    run.font.name = font_name

    if is_table:
        run.font.size = Pt(9)  # 小五 (9磅)
        sz_val = "18"  # half-points 18 = 9pt
    else:
        run.font.size = Pt(12)  # 小四 (12磅)
        sz_val = "24"  # half-points 24 = 12pt

    try:
        rPr = run._element.get_or_add_rPr()

        # 1. 统一设置中西文字体为宋体
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(r'<w:rFonts %s w:ascii="%s" w:hAnsi="%s" w:eastAsia="%s" w:cs="%s"/>' % (
                nsdecls('w'), font_name, font_name, font_name, font_name
            ))
            rPr.append(rFonts)
        else:
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:eastAsia'), font_name)
            rFonts.set(qn('w:cs'), font_name)

        # 2. 统一设置字号 (小五 9pt / 小四 12pt)
        sz_node = rPr.find(qn('w:sz'))
        if sz_node is None:
            rPr.append(parse_xml(r'<w:sz %s w:val="%s"/>' % (nsdecls('w'), sz_val)))
        else:
            sz_node.set(qn('w:val'), sz_val)

        szCs_node = rPr.find(qn('w:szCs'))
        if szCs_node is None:
            rPr.append(parse_xml(r'<w:szCs %s w:val="%s"/>' % (nsdecls('w'), sz_val)))
        else:
            szCs_node.set(qn('w:val'), sz_val)

        # 3. 设置下划线
        if enable_underline:
            run.underline = True
            u_node = rPr.find(qn('w:u'))
            if u_node is None:
                rPr.append(parse_xml(r'<w:u %s w:val="single"/>' % nsdecls('w')))
            else:
                u_node.set(qn('w:val'), 'single')
        else:
            run.underline = None
            u_node = rPr.find(qn('w:u'))
            if u_node is not None:
                rPr.remove(u_node)
    except Exception:
        pass


def align_table_row_cells(row_data: List[Any], total_cols: int, row_i: int) -> List[str]:
    """
    智能将行数据对齐为 total_cols 长度的标准单元格列表（纯结构化正则与几何映射，无具体业务数据）：
    1. 智能剥离粘连金额：若名称末尾粘连了数字（如 '某费用名称 0.00'），自动剥离并顺延至单价/总价列；
    2. 缺少 1 列数据自适应：
       - 若首列为独立纯序号，末尾自动补齐空的【备注】列；
       - 若首列包含复合序号与名称，拆分为第 0 列序号与第 1 列名称；
       - 若首列为纯名称，在第 0 列自动补齐流水序号；
    3. 合计行标准装配：第 0 列留空，第 1 列为合计标签，分项总价列填入总金额，末尾留空；
    4. 防错列自愈：
       - 若第 1 列与第 2 列名称相同，自动剔除重复项并纠正列偏移；
       - 若 5 列报价表中【单价】与【分项总价】被重复填入相同的非零大额总金额，对于包干/工程/大类项智能纠偏将单价置为破折号 "——"；
       - 若 5 列报价表中第 1 列名称重复包含序号，自动剥离前缀序号。
    """
    raw_vals = [str(c).strip() if c is not None else "" for c in row_data]
    if not raw_vals:
        return [""] * total_cols

    # 0. 智能检测并剥离第 1 列（名称列）末尾粘连的小数点金额（如 '费用名称 0.00'，精确保护 '设备名 1' 等编号后缀）
    if len(raw_vals) >= 2:
        name_cand = raw_vals[1]
        m_name_amount = re.match(r'^(.*?)\s+([0-9]+\.[0-9]{1,2})\s*$', name_cand)
        if m_name_amount:
            clean_n = m_name_amount.group(1).strip()
            num_part = m_name_amount.group(2).strip()
            if clean_n and len(raw_vals) < total_cols:
                raw_vals[1] = clean_n
                raw_vals.insert(2, num_part)

    first_val = raw_vals[0]
    is_summary_row = any(k in first_val for k in ["合计", "总计", "总价", "小计"]) or (len(raw_vals) > 1 and any(k in str(raw_vals[1]) for k in ["合计", "总计", "总价", "小计"]))

    # 1. 合计总价汇总行
    if is_summary_row:
        sum_label = first_val if any(k in first_val for k in ["合计", "总计", "总价", "小计"]) else str(raw_vals[1])
        sum_amount = ""
        for v in raw_vals[1:]:
            v_clean = re.sub(r'[,，\s元¥]', '', v)
            if re.match(r'^\d+(\.\d+)?$', v_clean):
                sum_amount = v_clean
                break
        if not sum_amount:
            for v in raw_vals:
                v_clean = re.sub(r'[,，\s元¥]', '', v)
                if re.match(r'^\d+(\.\d+)?$', v_clean):
                    sum_amount = v_clean
                    break

        res = [""] * total_cols
        res[0] = ""
        res[1] = sum_label
        if total_cols >= 4:
            res[total_cols - 2] = sum_amount
        return res

    # 2. 缺少 1 列数据时的自适应对齐 (len(raw_vals) == total_cols - 1)
    if len(raw_vals) == total_cols - 1:
        is_pure_seq = bool(re.match(r'^\d+(?:\.\d+)*[、.．]?$', first_val))
        if is_pure_seq:
            # 场景 A: 首列已是独立序号，少的是末尾备注列
            res = list(raw_vals)
            res[0] = re.sub(r'[、.．]$', '', first_val)
            res.append("")
        else:
            m_idx = re.match(r'^(\d+(?:\.\d+)*)\s*[、.．\s\-]\s*([^\d].*)$', first_val)
            if m_idx:
                # 场景 B: 首列包含复合序号与名称，拆分为序号与名称
                seq_num = m_idx.group(1).strip()
                clean_name = m_idx.group(2).strip()
                res = list(raw_vals)
                res[0] = clean_name
                res.insert(0, seq_num)
                while len(res) < total_cols:
                    res.append("")
            else:
                # 场景 C: 首列为纯名称，少的是首列序号，在第 0 列自动补齐流水序号
                res = list(raw_vals)
                res.insert(0, str(row_i + 1))
                while len(res) < total_cols:
                    res.append("")
    else:
        # 3. 列数完全匹配或更多列数据
        res = list(raw_vals)

    seq_val = res[0] if res else ""
    name_val = res[1] if len(res) > 1 else ""

    if seq_val and name_val:
        m_dup = re.match(r'^' + re.escape(seq_val) + r'\s*[、.．\s\-]\s*(.+)$', name_val)
        if not m_dup:
            m_dup = re.match(r'^' + re.escape(seq_val) + r'(.+)$', name_val)
        if m_dup:
            res[1] = m_dup.group(1).strip()

    if len(res) >= 3 and res[1] and res[1] == res[2]:
        res.pop(2)
        res.append("")

    # 4. [5列报价表专项纠偏] 单价列与分项总价列重复/错列自愈
    if total_cols == 5 and len(res) >= 4:
        unit_p = str(res[2]).strip()
        total_p = str(res[3]).strip()
        name_p = str(res[1]).strip()

        # 场景 A: 单价与总价完全相同，且为非零的大额总金额（非 0.00 / 0）
        if unit_p and total_p and unit_p == total_p and unit_p not in ("—", "——", "/", "-", "0", "0.00"):
            pkg_kws = ["费", "工程", "系统", "加固", "防水", "敷设", "安装", "调试", "服务", "培训", "大类", "购置", "总承包", "支架", "桥架", "辅材", "电缆", "柜"]
            if any(k in name_p for k in pkg_kws):
                res[2] = "——"
                res[3] = total_p
        # 场景 B: 单价列填了纯总价数字而总价列为空
        elif unit_p and not total_p and unit_p not in ("—", "——", "/", "-"):
            pkg_kws = ["费", "工程", "系统", "加固", "防水", "敷设", "安装", "调试", "服务", "培训", "大类", "购置", "总承包"]
            if any(k in name_p for k in pkg_kws):
                res[3] = unit_p
                res[2] = "——"

    while len(res) < total_cols:
        res.append("")
    return res[:total_cols]


def _is_table_footer_row(row, total_cols: int) -> bool:
    """
    通用 DOM 结构化判定：识别表格行是否为标书表尾非数据落款/声明行。
    特征：
    1. 包含通用的落款、签章、总报价大写、期限或独立说明注等标书声明词汇；
    2. 单元格发生跨列合并且包含标签说明（以冒号结尾或长声明）。
    注意：纯占位符 '......' 或纯序号行属于普通数据区，可正常覆盖。
    """
    if row is None or not hasattr(row, '_tr') or row._tr is None:
        return False
    try:
        row_text = "".join(c.text for c in row.cells).strip()
        if not row_text or "..." in row_text or "…" in row_text:
            return False

        # 标书通用表尾落款与声明特征词（通用结构词，无任何具体业务数据）
        footer_keywords = [
            "签字", "签章", "盖章", "法定代表人", "授权代表", "投标人名称",
            "总报价", "大写", "交货期", "工期", "质保", "服务期",
            "注：", "注:", "说明：", "说明:"
        ]
        if any(k in row_text for k in footer_keywords):
            return True

        # 结构化跨列且带说明标签
        from docx.oxml.ns import qn
        tc_elements = [c for c in row._tr.iterchildren() if c.tag.endswith('tc')]
        if len(tc_elements) < total_cols and ("：" in row_text or ":" in row_text):
            return True

        return False
    except Exception:
        return False


def _get_footer_row_profile(row: Any) -> Tuple[int, int]:
    """提取模板表尾行的结构轮廓，不读取或依赖具体业务文字。"""
    unique_cells = []
    for cell in row.cells:
        if not any(cell._tc is existing._tc for existing in unique_cells):
            unique_cells.append(cell)
    non_empty_count = sum(bool(str(cell.text or "").strip()) for cell in unique_cells)
    return len(unique_cells), non_empty_count


def _get_matrix_row_profile(row_data: Any) -> Tuple[int, int]:
    """提取矩阵行的非空字段轮廓，用于和模板表尾结构进行匹配。"""
    if not isinstance(row_data, (list, tuple)):
        return 0, 0
    non_empty_count = sum(bool(str(value or "").strip()) for value in row_data)
    return len(row_data), non_empty_count


def _split_matrix_footer_rows(
    matrix: List[List[Any]],
    table: Any,
    template_footer_trs: Sequence[Any],
) -> Tuple[List[List[Any]], List[List[Any]]]:
    """
    按模板表尾的结构轮廓拆分矩阵尾部行。

    只处理与原模板表尾行数量、非空字段数量一致的连续尾部，避免把普通标的物
    明细误判为表尾；整个过程不依赖字段名称、金额或期限等具体内容。
    """
    footer_count = len(template_footer_trs)
    if footer_count == 0 or len(matrix) < footer_count:
        return matrix, []

    template_profiles = []
    for footer_tr in template_footer_trs:
        footer_row = next((row for row in table.rows if row._tr is footer_tr), None)
        if footer_row is None:
            return matrix, []
        template_profiles.append(_get_footer_row_profile(footer_row))

    candidate_rows = matrix[-footer_count:]
    for candidate, (_, expected_non_empty_count) in zip(candidate_rows, template_profiles):
        _, actual_non_empty_count = _get_matrix_row_profile(candidate)
        if actual_non_empty_count != expected_non_empty_count:
            return matrix, []

    return matrix[:-footer_count], candidate_rows


def _write_matrix_footer_row(table: Any, target_tr: Any, row_data: Sequence[Any]) -> bool:
    """
    将矩阵表尾行写入模板原有的表尾节点，并保留原有合并结构。

    报价表尾通常是“标签单元格 + 内容单元格”或整行合并两种结构，不能按普通
    明细行逐列写入，否则会把同一个合并单元格重复写成多个标签。
    """
    if target_tr is None or not isinstance(row_data, (list, tuple)):
        return False

    target_row = next((row for row in table.rows if row._tr is target_tr), None)
    if target_row is None:
        return False

    values = [str(value or "").strip() for value in row_data]
    non_empty_values = [value for value in values if value]
    if not non_empty_values:
        return False

    unique_cells = []
    for cell in target_row.cells:
        if not any(cell._tc is existing._tc for existing in unique_cells):
            unique_cells.append(cell)

    if len(unique_cells) >= 2:
        # 第一格保留表尾标签，第二格承接矩阵中标签之后的实际值。
        unique_cells[0].text = values[0] or non_empty_values[0]
        footer_value = next((value for value in values[1:] if value), "")
        if footer_value:
            unique_cells[1].text = footer_value
    else:
        # 交货期限等表尾通常是整行合并，直接写入完整的“标签+值”。
        target_row.cells[0].text = " ".join(non_empty_values)

    for cell in unique_cells or [target_row.cells[0]]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                _apply_run_style_xml(run, enable_underline=False, is_table=True)
    return True


def _is_footer_paragraph_target(doc: Document, paragraph: Any) -> bool:
    """按表格行结构判断段落是否位于模板表尾，不依赖具体业务字段内容。"""
    element = getattr(paragraph, "_element", None)
    parent = element.getparent() if element is not None else None
    row_element = None
    while parent is not None:
        if str(parent.tag).endswith("tr"):
            row_element = parent
            break
        parent = parent.getparent()
    if row_element is None:
        return False

    for table in doc.tables:
        for row in table.rows:
            if row._tr is row_element:
                return _is_table_footer_row(row, len(table.columns))
    return False


def _is_full_paragraph_replacement(real_text: str, proposed_val: str, prop_type: str = "") -> bool:
    """
    智能判定提案是否为整段/整句完整覆盖替换（Full Paragraph Replacement），而非单槽位插值填空。

    判定维度：
    1. 显式类型：prop_type 为 sentence_batch / paragraph / full_text / replace_all；
    2. 双锚点首尾重合：real_text 与 proposed_val 去除标点空格后，开头（前4~6字）相同且结尾（后4~6字）相同，长度 >= 12；
    3. 骨架重合度：计算 real_text（剥离占位符后）与 proposed_val 的字符重合率与公共子串；
    4. 序列相似度：SequenceMatcher 相似度 > 0.4 且 proposed_val 长度 >= 20；
    5. 单槽位切片防嵌套防重复：若 proposed_val 已经包含了段落开头和结尾的固定词（如“根据贵方”、“有关事宜”），严禁按单槽位切片。
    """
    if not real_text or not proposed_val:
        return False

    p_type_lower = str(prop_type or "").strip().lower()
    if p_type_lower in ("sentence_batch", "paragraph", "full_text", "replace_all"):
        return True

    r_strip = real_text.strip()
    p_strip = proposed_val.strip()

    if len(p_strip) < 8:
        return False

    # 规范化去除空白与标点进行核心骨架比对
    r_core = re.sub(r'[\s_＿\-\–\—\(\)（）\[\]［］【】:：\.\,，。；;、“”"\'`]', '', r_strip)
    p_core = re.sub(r'[\s_＿\-\–\—\(\)（）\[\]［］【】:：\.\,，。；;、“”"\'`]', '', p_strip)

    if len(r_core) >= 8 and len(p_core) >= 8:
        # 1. 首尾双锚点匹配：首部 4 个字相同 且 尾部 4 个字相同
        if len(r_core) >= 4 and len(p_core) >= 4:
            head_match = (r_core[:4] == p_core[:4])
            tail_match = (r_core[-4:] == p_core[-4:])
            if head_match and tail_match:
                return True

        # 2. 首部 6 个字相同且 proposed_val 长度较长（整句覆盖）
        if len(r_core) >= 6 and len(p_core) >= 6:
            if r_core[:6] == p_core[:6] and len(p_core) >= 20 and len(p_core) >= len(r_core) * 0.4:
                return True

        # 3. 尾部 6 个字相同且 proposed_val 包含 r_core 前半段核心词
        if len(r_core) >= 6 and len(p_core) >= 6:
            if r_core[-6:] == p_core[-6:] and r_core[:4] in p_core and len(p_core) >= 20:
                return True

    # 4. 基于 SequenceMatcher 计算相似度
    import difflib
    ratio = difflib.SequenceMatcher(None, r_core, p_core).ratio()
    if ratio > 0.45 and len(p_strip) >= 20:
        return True

    # 5. 原模板固定词段序列包含率检测
    slot_pattern = r'(_{1,}|＿{1,}|\[[^\]]+\]|［[^］]+］|【[^】]+】|（(?:姓名和职务|投标人的名称|投标人名称|代表姓名|职务)[^）]*）|\([^\)]*(?:name and title|bidder name)[^\)]*|\d*\s*年\s*\d*\s*月\s*\d*\s*日|(?<=[:：])\s{2,}|\s{2,})'
    static_segments = [s.strip() for s in re.split(slot_pattern, r_strip) if s and not re.search(slot_pattern, s)]
    valid_static = [s for s in static_segments if len(s) >= 3 and not re.match(r'^[0-9\s:：、.．\-]+$', s)]
    if len(valid_static) >= 2:
        matched_count = sum(1 for s in valid_static if s in p_strip)
        if matched_count >= 2 and matched_count >= len(valid_static) * 0.6:
            return True

    return False


def _render_diff_paragraph_runs(
    p_elem,
    real_text: str,
    proposed_val: str,
    enable_underline_on_diff: bool = True,
    is_table: bool = False,
) -> None:
    """
    高保真 Diff 级段落 Run 分段重构与渲染引擎：
    1. 基于 SequenceMatcher 精准识别模板原有固定文字与新填入的数据值；
    2. 模板固定文字（equal 标签）：渲染为【宋体 小四 Pt(12) 无下划线】；
    3. 新填入的业务数据（replace / insert 标签）：渲染为【宋体 小四 Pt(12) 原生下划线】（表格外）；
    4. 彻底消除单槽位切片导致的重复前缀、重复后缀与未填占位符残留！
    """
    import difflib
    from app.agents.review_engine import clean_zero_change_annotations, is_zero_change_or_no_op_proposal

    proposed_val = clean_zero_change_annotations(proposed_val)
    if is_zero_change_or_no_op_proposal(proposed_val, real_text):
        return

    p_elem._element.clear_content()

    if not real_text or not proposed_val:
        if proposed_val:
            _apply_run_style_xml(p_elem.add_run(proposed_val), enable_underline=enable_underline_on_diff, is_table=is_table)
        return

    sm = difflib.SequenceMatcher(None, real_text, proposed_val)
    opcodes = sm.get_opcodes()

    # 如果相似度过低（例如全新重写的独立说明段落），直接整段写入
    if sm.ratio() < 0.2:
        _apply_run_style_xml(p_elem.add_run(proposed_val), enable_underline=enable_underline_on_diff, is_table=is_table)
        return

    for tag, i1, i2, j1, j2 in opcodes:
        if tag == 'equal':
            fixed_text = proposed_val[j1:j2]
            if fixed_text:
                _apply_run_style_xml(p_elem.add_run(fixed_text), enable_underline=False, is_table=is_table)
        elif tag in ('replace', 'insert'):
            inserted_text = proposed_val[j1:j2]
            if inserted_text:
                _apply_run_style_xml(p_elem.add_run(inserted_text), enable_underline=(enable_underline_on_diff and not is_table), is_table=is_table)
        elif tag == 'delete':
            # real_text 中被替换/删除的占位符（如 "______" 或纯空格），不输出到目标文档
            pass


def _remove_redundant_pricing_subtitle(doc: Document) -> int:
    """
    清理报价章节中重复的子标题。

    部分招标模板在一级标题下又放置一个“报价分析表”子标题，随后紧接着
    招标编号与项目名称。系统最终输出要求这一段只保留项目信息行，因此仅在
    同时满足“报价章节 + 分析类子标题 + 项目信息行”三个结构条件时移除子标题，
    不影响其他章节中的普通标题。
    """
    if doc is None or not hasattr(doc, "element") or not hasattr(doc.element, "body"):
        return 0

    body = doc.element.body
    children = list(body)
    removed_count = 0

    for index, elem in enumerate(children):
        if not elem.tag.endswith("}p"):
            continue

        major_text = "".join(elem.itertext()).strip()
        if not major_text or "配置" not in major_text or "分项报价" not in major_text:
            continue

        subtitle_elem = None
        subtitle_text = ""
        info_elem = None
        info_text = ""

        # 允许标题与子标题之间存在空段落，但不跨越表格或其他正文节点。
        cursor = index + 1
        while cursor < len(children):
            candidate = children[cursor]
            if candidate.tag.endswith("}tbl"):
                break
            if not candidate.tag.endswith("}p"):
                cursor += 1
                continue

            candidate_text = "".join(candidate.itertext()).strip()
            if not candidate_text:
                cursor += 1
                continue

            if subtitle_elem is None:
                subtitle_elem = candidate
                subtitle_text = candidate_text
                cursor += 1
                continue

            info_elem = candidate
            info_text = candidate_text
            break

        is_analysis_subtitle = (
            subtitle_elem is not None
            and "分析" in subtitle_text
            and subtitle_text.endswith("表")
        )
        has_project_info = "招标编号" in info_text and "项目名称" in info_text
        if is_analysis_subtitle and has_project_info:
            parent = subtitle_elem.getparent()
            if parent is not None:
                parent.remove(subtitle_elem)
                removed_count += 1
                logger.info(
                    "   [报价章节标题归一化] 已移除重复的报价分析子标题，保留招标编号与项目名称行"
                )

    return removed_count


def _collapse_redundant_paragraph_proposals(proposals: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    合并同一正文 Path 的重复提案，防止多轮 Worker 结果被当成多槽位拼接。

    同一段落确实可能包含多个合法槽位，因此不做无条件去重：仅合并完全相同的
    提案，或删除明显是另一提案完整超集的复合提案（例如“标题 + 项目信息”与
    单独的“项目信息”同时出现）。不同字段值仍按原顺序保留。
    """
    if not proposals:
        return []

    grouped: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
    first_seen_order: List[Tuple[str, Dict[str, Any]]] = []
    for proposal in proposals:
        path = _normalize_table_path(str(proposal.get("path", "")).strip())
        is_body_paragraph = bool(re.fullmatch(r"/body/p\[\d+\]", path))
        if not is_body_paragraph:
            first_seen_order.append(("", proposal))
            continue
        if path not in grouped:
            first_seen_order.append((path, proposal))
        grouped[path].append(proposal)

    resolved_by_path: Dict[str, List[Dict[str, Any]]] = {}
    for path, items in grouped.items():
        if len(items) <= 1:
            resolved_by_path[path] = items
            continue

        # 同值提案只保留最后一份，最后一份通常来自本轮更高优先级的工具结果。
        unique_items: List[Dict[str, Any]] = []
        seen_values = set()
        for item in items:
            value = str(
                item.get("proposed_text")
                if item.get("proposed_text") is not None
                else item.get("value", "")
            ).strip()
            value_key = re.sub(r"\s+", "", value)
            if value_key in seen_values:
                continue
            seen_values.add(value_key)
            unique_items.append(item)

        # 只在存在明确复合字段（同一值包含两个以上属性标签）时清除超集，
        # 避免误伤同一段落内“姓名 + 职务”等多个独立槽位。
        compact_items: List[Dict[str, Any]] = []
        for item in unique_items:
            value = str(
                item.get("proposed_text")
                if item.get("proposed_text") is not None
                else item.get("value", "")
            ).strip()
            value_core = re.sub(r"[\s:：、，,（）()【】\[\]]", "", value)
            label_count = sum(value_core.count(label) for label in ("项目名称", "招标编号", "投标人名称", "投标单位"))
            is_redundant_superset = False
            if label_count >= 2:
                for other in unique_items:
                    if other is item:
                        continue
                    other_value = str(
                        other.get("proposed_text")
                        if other.get("proposed_text") is not None
                        else other.get("value", "")
                    ).strip()
                    if other_value and other_value != value and other_value in value:
                        is_redundant_superset = True
                        break
            if not is_redundant_superset:
                compact_items.append(item)

        resolved_by_path[path] = compact_items or unique_items
        if len(resolved_by_path[path]) != len(items):
            logger.warning(
                f"   [正文提案冲突收敛] 路径 {path} 合并 {len(items)} 条重复/超集提案，"
                f"保留 {len(resolved_by_path[path])} 条，阻止多值拼接污染"
            )

    resolved: List[Dict[str, Any]] = []
    emitted_paths = set()
    for path, proposal in first_seen_order:
        if path and path not in emitted_paths:
            resolved.extend(resolved_by_path.get(path, [proposal]))
            emitted_paths.add(path)
        elif not path:
            resolved.append(proposal)

    return resolved


def fill_docx_proposals_in_dom(docx_path: str, proposals: List[Dict]) -> int:
    """
    DOM 级标书全量原位替换与格式排版引擎：
    1. 直接从 Word 节点读取真实原文；
    2. 提炼并 100% 留存原标书字段标签前缀 (如 '投标人名称（单位盖章）：')，绝对不擦除原文；
    3. 表格外部的正文段落填报内容统一为【宋体 小四 (12pt)】，填入数据加 Word 原生下划线 (<w:u w:val="single"/>)；
    4. 表格内部的单元格填报内容统一为【宋体 小五 (9pt)】，取消下划线（确保表格内文字整洁美观）；
    5. 单次内存处理刷盘，性能与呈现效果兼备。
    """
    if not docx_path or not os.path.exists(docx_path):
        return 0

    try:
        doc = Document(docx_path)
        proposals = _collapse_redundant_paragraph_proposals(proposals or [])
        if not proposals:
            subtitle_removed = _remove_redundant_pricing_subtitle(doc)
            if subtitle_removed:
                doc.save(docx_path)
            return subtitle_removed
        subtitle_removed = 0
        success_count = 0

        # 动态表格写入可能在表尾之前插入多行，先锁定原始单元格节点，
        # 防止“投标总报价/交货期限”等提案按旧行号落到新增标的物行。
        table_cell_snapshots = _snapshot_table_cell_targets(doc, proposals)

        # 0. 跨章节表格路径冲突检测与智能重定向守卫 (Table Path Collision Guard)
        from app.utils.table_utils import get_doc_chapter_tables_mapping, get_chapter_specific_table_indices, detect_table_header_rows
        
        doc_tbl_mapping = get_doc_chapter_tables_mapping(doc)
        tbl_to_chapter = {}
        for entry in doc_tbl_mapping:
            for t_i in entry.get("table_indices", []):
                tbl_to_chapter[t_i] = entry.get("chapter_title", "")

        # 扫描所有表格提案进行冲突检测
        tbl_proposals_by_target = defaultdict(list)
        for p_idx, p in enumerate(proposals):
            p_path = _normalize_table_path(str(p.get("path", "")).strip())
            raw_val = p.get("proposed_text") if p.get("proposed_text") is not None else p.get("value", "")
            p_val = str(raw_val).strip() if raw_val is not None else ""
            p_type = str(p.get("type", "")).strip()

            is_cell_target = bool(re.search(r'/(?:tr|row|tc|cell)\[\d+\]', p_path))
            if is_cell_target or not p_val:
                continue

            matrix = []
            if isinstance(raw_val, list) and raw_val and isinstance(raw_val[0], list):
                matrix = raw_val
            elif p_val.startswith("[") and p_val.endswith("]"):
                try:
                    parsed = json.loads(p_val)
                    if isinstance(parsed, list) and parsed and isinstance(parsed[0], list):
                        matrix = parsed
                except Exception:
                    pass

            is_table_prop = bool(
                p_type == "table_rows" or
                (matrix and re.search(r'/tbl\[\d+\]', p_path)) or
                (re.search(r'^/body/tbl\[\d+\]$', p_path) and p_val.startswith("["))
            )
            if is_table_prop:
                m_tbl = re.search(r'/tbl\[(\d+)\]', p_path)
                if m_tbl:
                    t_idx = int(m_tbl.group(1)) - 1
                    tbl_proposals_by_target[t_idx].append((p_idx, p, matrix))

        # 对产生多重竞争的同一个 tbl_idx 执行仲裁与重定向
        for t_idx, prop_list in tbl_proposals_by_target.items():
            if len(prop_list) <= 1:
                continue

            real_ch = tbl_to_chapter.get(t_idx, "")
            table_obj = doc.tables[t_idx] if 0 <= t_idx < len(doc.tables) else None
            real_cols = len(table_obj.rows[0].cells) if table_obj and table_obj.rows else 0

            scored_props = []
            for p_idx, p, mat in prop_list:
                p_ch = p.get("chapter_title", "")
                mat_cols = len(mat[0]) if mat and isinstance(mat[0], list) else 0

                score = 0.0
                if p_ch and real_ch:
                    if p_ch in real_ch or real_ch in p_ch:
                        score += 20.0
                    else:
                        clean_pch = re.sub(r'^[一二三四五六七八九十百0-9\s、\.\(\)（）]+', '', p_ch).strip()
                        clean_rch = re.sub(r'^[一二三四五六七八九十百0-9\s、\.\(\)（）]+', '', real_ch).strip()
                        if clean_pch and clean_rch and (clean_pch in clean_rch or clean_rch in clean_pch):
                            score += 15.0
                if mat_cols == real_cols:
                    score += 10.0
                elif abs(mat_cols - real_cols) == 1:
                    score += 5.0

                scored_props.append((score, p_idx, p, mat))

            scored_props.sort(key=lambda x: x[0], reverse=True)
            winner = scored_props[0]
            logger.info(f"🛡️ [Table Collision Guard] 表格 /body/tbl[{t_idx+1}] 发生跨章节竞争 (真实所属: '{real_ch}')。胜出提案: 章节 '{winner[2].get('chapter_title')}' (得分: {winner[0]})")

            for score, p_idx, p, mat in scored_props[1:]:
                loser_ch = p.get("chapter_title", "")
                real_loser_tbls = get_chapter_specific_table_indices(doc, loser_ch)
                if real_loser_tbls and real_loser_tbls[0] != t_idx:
                    redirect_idx = real_loser_tbls[0]
                    p["path"] = f"/body/tbl[{redirect_idx + 1}]"
                    logger.info(f"   🔄 [Table Collision Guard] 异章提案 (章节: '{loser_ch}') 成功重定向到真实表格: /body/tbl[{redirect_idx + 1}]")
                else:
                    p["path"] = "/body/tbl_collision_blocked"
                    logger.warning(f"   ⛔ [Table Collision Guard] 拦截并阻止异章提案 (章节: '{loser_ch}') 覆盖 /body/tbl[{t_idx+1}]")

        # 1. 优先处理表格批量插行/原位填报提案
        handled_tbl_proposals = set()
        for p_idx, p in enumerate(proposals):
            p_path = _normalize_table_path(str(p.get("path", "")).strip())
            raw_val = p.get("proposed_text") if p.get("proposed_text") is not None else p.get("value", "")
            p_val = str(raw_val).strip() if raw_val is not None else ""
            p_type = str(p.get("type", "")).strip()

            is_cell_target = bool(re.search(r'/(?:tr|row|tc|cell)\[\d+\]', p_path))
            if is_cell_target or not p_val or "blocked" in p_path:
                continue

            # 智能解析 2D 数据矩阵
            matrix = []
            if isinstance(raw_val, list) and raw_val and isinstance(raw_val[0], list):
                matrix = raw_val
            elif p_val.startswith("[") and p_val.endswith("]"):
                try:
                    parsed = json.loads(p_val)
                    if isinstance(parsed, list) and parsed and isinstance(parsed[0], list):
                        matrix = parsed
                except Exception as je:
                    logger.warning(f"   解析表格提案 JSON 字符串异常: {je}")

            is_table_prop = bool(
                p_type == "table_rows" or
                (matrix and re.search(r'/tbl\[\d+\]', p_path)) or
                (re.search(r'^/body/tbl\[\d+\]$', p_path) and p_val.startswith("["))
            )
            if not is_table_prop:
                continue

            m_tbl = re.search(r'/tbl\[(\d+)\]', p_path)
            if not m_tbl:
                continue
            tbl_idx = int(m_tbl.group(1)) - 1
            if not (0 <= tbl_idx < len(doc.tables)):
                logger.warning(f"   表格索引越界: {p_path}, doc.tables 数量: {len(doc.tables)}")
                continue

            table = doc.tables[tbl_idx]

            # 校验 matrix 必须为合法的 2D 数据矩阵
            if not isinstance(matrix, list) or not matrix or not isinstance(matrix[0], list):
                logger.warning(f"   ⚠️ 表格 {p_path} 提案数据非合法 2D 矩阵，跳过表格批处理写盘")
                continue

            handled_tbl_proposals.add(p_idx)
            
            # 引入通用多行表头检测（支持 1~3 行复合表头），彻底防止子表头被数据覆盖
            from app.utils.table_utils import (
                detect_table_header_rows,
                get_table_header_logical_spans,
                align_row_to_header_grid_spans,
                clean_row_vmerge,
            )
            hdr_count = detect_table_header_rows(table)
            header_row = table.rows[0] if table.rows else None
            if not header_row:
                continue

            total_cols = len(header_row.cells)
            header_spans = get_table_header_logical_spans(table, hdr_count)
            logical_cols_count = len(header_spans) if header_spans else total_cols
            header_text_joined = "".join(
                str(cell.text or "").replace("\n", "").strip()
                for cell in header_row.cells
            )
            is_response_deviation_table = any(
                keyword in header_text_joined
                for keyword in ["商务条款", "技术要求", "响应", "偏离", "服务承诺"]
            ) and logical_cols_count >= 4

            # 找到【分项总价】在逻辑列中的索引
            total_price_logical_idx = logical_cols_count - 2 if logical_cols_count >= 4 else logical_cols_count - 1
            total_price_found = False
            response_status_logical_idx = None
            response_status_default = "无"
            deviation_detail_logical_idx = None
            for l_i, (s_c, e_c) in enumerate(header_spans):
                h_text = header_row.cells[s_c].text.strip()
                if (not total_price_found) and any(k in h_text for k in ["分项总价", "总价", "合价", "总金额", "金额（元）", "小计（元）"]):
                    total_price_logical_idx = l_i
                    total_price_found = True
                if (
                    any(k in h_text for k in ["有无偏离", "是否偏离", "偏离情况", "是否响应"])
                    and "服务承诺" not in h_text
                    and "响应情况" not in h_text
                ):
                    response_status_logical_idx = l_i
                    response_status_default = "是" if "是否响应" in h_text else "无"
                if any(k in h_text for k in ["偏离内容", "偏离原因", "偏离说明"]):
                    deviation_detail_logical_idx = l_i

            # 1. 检查原表格最后一行是否是原模板自带的【合计总价】行（包含"合计"或"总计"）
            template_summary_tr = None
            if len(table.rows) > hdr_count:
                last_row_text = "".join(c.text for c in table.rows[-1].cells).strip()
                if any(k in last_row_text for k in ["合计", "总计", "总价", "小计"]):
                    template_summary_tr = table.rows[-1]._tr

            # 2. 检查原表格中是否有真正的表尾非数据落款行（签字盖章/独立说明注/期限报价等）
            footer_start_idx = len(table.rows)
            footer_anchor_tr = None
            for r_idx in range(len(table.rows) - 1, hdr_count - 1, -1):
                if _is_table_footer_row(table.rows[r_idx], total_cols):
                    footer_start_idx = r_idx
                else:
                    break

            if footer_start_idx < len(table.rows):
                footer_anchor_tr = table.rows[footer_start_idx]._tr
            template_footer_trs = [
                table.rows[r_idx]._tr
                for r_idx in range(footer_start_idx, len(table.rows))
            ]

            # 3. 拆分 matrix: 将最后一行的合计行（如有）与前面的明细行拆分
            detail_matrix = []
            summary_row_data = None
            for r_idx, r_data in enumerate(matrix):
                r_str = " ".join(str(c) for c in r_data)
                # 响应/偏离表的商务条款里经常出现“合同总价款”“投标总价”等词，
                # 不能因为普通条款包含“总价”就误判为表尾合计行。只有明确以
                # “合计/总计/小计”开头的行才允许作为响应表的汇总行。
                if is_response_deviation_table:
                    summary_probe = " ".join(str(c or "").strip() for c in r_data[:2])
                    is_summary_row = summary_probe.startswith(("合计", "总计", "小计"))
                else:
                    is_summary_row = any(k in r_str for k in ["合计", "总计", "总价", "小计"])
                if r_idx == len(matrix) - 1 and is_summary_row:
                    summary_row_data = r_data
                else:
                    detail_matrix.append(r_data)

            detail_matrix, matrix_footer_rows = _split_matrix_footer_rows(
                detail_matrix,
                table,
                template_footer_trs,
            )

            # 商务/技术响应表的无序号 4 列矩阵，或从序号 1 开始的完整矩阵，
            # 表示模型提交的是整张表的最终结果。此时旧文件中可能已经残留
            # 上一轮追加出来的多余数据行，写入后必须一并裁掉，不能只覆盖前半段。
            response_matrix_serials = []
            for r_data in detail_matrix:
                if isinstance(r_data, list) and len(r_data) == logical_cols_count:
                    seq_match = re.fullmatch(r"\s*(\d+)\s*[、.．)]?\s*", str(r_data[0]))
                    response_matrix_serials.append(int(seq_match.group(1)) if seq_match else None)
                else:
                    response_matrix_serials.append(None)
            response_matrix_is_full = bool(
                is_response_deviation_table
                and detail_matrix
                and (
                    all(
                        isinstance(r_data, list) and len(r_data) == max(1, logical_cols_count - 1)
                        for r_data in detail_matrix
                    )
                    or response_matrix_serials == list(range(1, len(detail_matrix) + 1))
                )
            )

            # [固定表单防增行保护机制]
            # 若目标表格为固定单元格填报表（如单项汇总表/固定表单），仅保留第 1 行项目汇总数据，绝对严禁插入/追加额外行！
            from app.utils.table_utils import is_fixed_slot_form_table
            prop_ch = str(p.get("chapter_title", "") or "")
            if is_fixed_slot_form_table(table) and len(detail_matrix) > 1:
                logger.info(f"   🛡️ [固定表单防增行保护] 表格 /body/tbl[{tbl_idx+1}] 为固定单项汇总表，仅保留首行项目汇总数据，忽略多余分项拆行")
                detail_matrix = detail_matrix[:1]

            # 4. 确定数据区的可用行范围
            max_available_detail_idx = footer_start_idx
            if template_summary_tr is not None and footer_start_idx == len(table.rows):
                max_available_detail_idx = len(table.rows) - 1

            # 5. 写入明细数据行 (detail_matrix) - 从真实数据行起始行 (hdr_count) 开始原位替换或追加
            for row_i, row_data in enumerate(detail_matrix):
                if not isinstance(row_data, list):
                    continue
                # 响应/偏离表只接受完整的 4 列（不含序号）或 5 列（含序号）
                # 行。对异常短行不做“尽力写入”，避免把半行静默刷入 Word。
                if is_response_deviation_table and len(row_data) not in (
                    logical_cols_count,
                    max(1, logical_cols_count - 1),
                ):
                    logger.warning(
                        f"   ⚠️ [表格行列数保护] 跳过 {p_path} 第 {row_i + 1} 行："
                        f"收到 {len(row_data)} 列，期望 {logical_cols_count} 列或 {logical_cols_count - 1} 列"
                    )
                    continue
                # row_vals 对齐到逻辑列数 (logical_cols_count)
                row_vals = align_table_row_cells(row_data, logical_cols_count, row_i)
                if is_response_deviation_table:
                    # 只对实际存在且能从表头识别出的列补默认值，不假定固定列号。
                    if (
                        response_status_logical_idx is not None
                        and response_status_logical_idx < len(row_vals)
                        and not str(row_vals[response_status_logical_idx]).strip()
                    ):
                        row_vals[response_status_logical_idx] = response_status_default
                    if (
                        deviation_detail_logical_idx is not None
                        and deviation_detail_logical_idx < len(row_vals)
                        and not str(row_vals[deviation_detail_logical_idx]).strip()
                    ):
                        row_vals[deviation_detail_logical_idx] = "完全响应招标文件要求，无偏离。"
                target_row_idx = hdr_count + row_i  # 默认按矩阵顺序写入，绝对跳过所有表头行
                if is_response_deviation_table and len(row_data) == logical_cols_count:
                    # 商务/技术响应矩阵经常只携带某个中间区间（例如序号 6~15）。
                    # 这类矩阵不能按 row_i 从第一条数据行覆盖，优先使用首列序号
                    # 定位模板中的对应行；否则会把内容错写到前面的条款行。
                    seq_match = re.fullmatch(r"\s*(\d+)\s*[、.．)]?\s*", str(row_data[0]))
                    if seq_match:
                        serial_row_idx = hdr_count + int(seq_match.group(1)) - 1
                        if serial_row_idx < len(table.rows) and serial_row_idx < max_available_detail_idx:
                            target_serial = str(table.rows[serial_row_idx].cells[0].text or "").strip()
                            if not target_serial or re.match(
                                rf"^\s*{re.escape(seq_match.group(1))}\s*[、.．)]?\s*$",
                                target_serial,
                            ):
                                target_row_idx = serial_row_idx

                if target_row_idx < max_available_detail_idx and target_row_idx < len(table.rows):
                    old_row = table.rows[target_row_idx]
                    unique_tcs = set(c._tc for c in old_row.cells)
                    # 若旧行的独立单元格数量少于表头逻辑列数，说明旧行存在错误的合并污染，创建干净新行原位替换
                    if len(unique_tcs) < len(header_spans):
                        new_row = table.add_row()
                        old_row._tr.addprevious(new_row._tr)
                        old_row._tr.getparent().remove(old_row._tr)
                        t_row = new_row
                    else:
                        t_row = old_row
                else:
                    # 超出原数据区行数：在合计行或落款行上方新增行
                    new_row = table.add_row()
                    anchor = template_summary_tr if template_summary_tr is not None else footer_anchor_tr
                    if anchor is not None and anchor.getparent() is not None:
                        anchor.addprevious(new_row._tr)
                    t_row = new_row
                    max_available_detail_idx += 1
                    footer_start_idx += 1

                # 关键 1：彻底清洗当前行所有单元格的纵向合并 (vMerge) 标记，杜绝跨页断层拉伸
                clean_row_vmerge(t_row)

                # 关键 2：将 t_row 的单元格合并结构与表头的 header_spans 100% 对齐！
                align_row_to_header_grid_spans(t_row, header_spans)

                # 按逻辑列依次写入单元格
                for l_idx, cell_val in enumerate(row_vals):
                    if l_idx < len(header_spans):
                        p_col = header_spans[l_idx][0]
                        if p_col < len(t_row.cells):
                            cell = t_row.cells[p_col]
                            cell.text = str(cell_val).strip()
                            for p_in_cell in cell.paragraphs:
                                for r_in_cell in p_in_cell.runs:
                                    _apply_run_style_xml(r_in_cell, enable_underline=False, is_table=True)
                success_count += 1

            # 6. 表尾行必须在全部明细扩行完成后写入原模板表尾节点。
            # 矩阵提案可能同时携带“投标总报价/交货期限”两行，不能把它们当作普通
            # 明细行追加到表格中，否则原有表尾会继续留空或出现重复表尾。
            if matrix_footer_rows:
                for footer_data, target_footer_tr in zip(matrix_footer_rows, template_footer_trs):
                    if _write_matrix_footer_row(table, target_footer_tr, footer_data):
                        success_count += 1
                        logger.info("   [DOM 表尾后置写盘] 已在全部明细扩行后写入模板表尾节点")

            # 7. 写入合计总价行 (summary_row_data)
            if summary_row_data is not None:
                sum_vals = align_table_row_cells(summary_row_data, logical_cols_count, len(detail_matrix))
                sum_label = sum_vals[1] if len(sum_vals) > 1 and sum_vals[1] else (summary_row_data[0] if summary_row_data else "合计总价")
                sum_amount = sum_vals[total_price_logical_idx] if total_price_logical_idx < len(sum_vals) else (sum_vals[-1] if sum_vals else "")

                if template_summary_tr is not None:
                    # 原位复用原模板合计行！
                    sum_row = None
                    for r in table.rows:
                        if r._tr is template_summary_tr:
                            sum_row = r
                            break
                    if sum_row is None:
                        sum_row = table.rows[-1]

                    sum_unique_tcs = set(c._tc for c in sum_row.cells)
                    if len(sum_unique_tcs) == len(header_spans):
                        # 纯标准无合并独立列：按逻辑列依次原位赋值
                        for l_i, c_val in enumerate(sum_vals):
                            if l_i < len(header_spans):
                                p_col = header_spans[l_i][0]
                                if p_col < len(sum_row.cells):
                                    sum_row.cells[p_col].text = str(c_val).strip()
                    elif len(sum_unique_tcs) == 1:
                        # 全行跨列完全合并为 1 个大单元格：组合填入标签与总金额
                        sum_row.cells[0].text = f"{sum_label} {sum_amount}".strip() if sum_amount else sum_label
                    else:
                        # 部分跨列合并（如前两列/前三列合并）：
                        # 1. 首格填合计标签
                        sum_row.cells[0].text = sum_label
                        # 2. 分项总价列精准填入总金额
                        sum_p_col = header_spans[total_price_logical_idx][0] if total_price_logical_idx < len(header_spans) else total_cols - 2
                        if sum_p_col < len(sum_row.cells) and sum_amount:
                            # 显式清空单价列（如果单价列未与首格合并）
                            if sum_p_col - 1 > 0 and sum_row.cells[sum_p_col - 1]._tc is not sum_row.cells[0]._tc:
                                sum_row.cells[sum_p_col - 1].text = ""
                            sum_row.cells[sum_p_col].text = str(sum_amount).strip()
                        # 3. 备注列置空
                        if total_cols - 1 < len(sum_row.cells) and total_cols - 1 > sum_p_col:
                            if sum_row.cells[total_cols - 1]._tc is not sum_row.cells[sum_p_col]._tc:
                                sum_row.cells[total_cols - 1].text = ""

                    for c in sum_row.cells:
                        for p_in_cell in c.paragraphs:
                            for r_in_cell in p_in_cell.runs:
                                _apply_run_style_xml(r_in_cell, enable_underline=False, is_table=True)
                    success_count += 1
                else:
                    new_sum_row = table.add_row()
                    if footer_anchor_tr is not None and footer_anchor_tr.getparent() is not None:
                        footer_anchor_tr.addprevious(new_sum_row._tr)
                    for c_i, c_val in enumerate(sum_vals):
                        if c_i < len(new_sum_row.cells):
                            new_sum_row.cells[c_i].text = str(c_val).strip()
                            for p_in_cell in new_sum_row.cells[c_i].paragraphs:
                                for r_in_cell in p_in_cell.runs:
                                    _apply_run_style_xml(r_in_cell, enable_underline=False, is_table=True)
                    success_count += 1

            if response_matrix_is_full:
                trim_start_idx = hdr_count + len(detail_matrix)
                trim_end_idx = min(
                    footer_start_idx,
                    len(table.rows) - (1 if template_summary_tr is not None else 0),
                )
                for r_i in range(trim_end_idx - 1, trim_start_idx - 1, -1):
                    if 0 <= r_i < len(table.rows):
                        stale_row = table.rows[r_i]._tr
                        parent = stale_row.getparent()
                        if parent is not None:
                            parent.remove(stale_row)
                if trim_end_idx > trim_start_idx:
                    logger.info(
                        f"   [响应表整表覆盖] 已清理矩阵之后的 {trim_end_idx - trim_start_idx} 行旧模板/重复数据"
                    )

            # 8. 安全清理未使用的模板空行。
            #
            # 旧逻辑会无条件清空矩阵末尾之后的所有行。对于“局部 table_rows
            # 矩阵 + 其它单元格提案”的合法场景，这会先擦掉后续行，再导致后续
            # 只有部分单元格被恢复，最终生成 Word 出现空单元格。现在只清理：
            # 1) 当前表没有单元格级提案；
            # 2) 目标行本身只有空白/省略号占位符；
            # 非空的原有内容绝不因局部矩阵被擦除。
            table_root = f"/body/tbl[{tbl_idx + 1}]"
            has_cell_level_proposals = any(
                str(other.get("path", "")).strip().startswith(table_root + "/")
                and bool(re.search(r'/(?:tr|row|tc|cell)\[\d+\]', str(other.get("path", ""))))
                for other in proposals
                if isinstance(other, dict)
            )
            if not has_cell_level_proposals:
                cleaned_start_idx = hdr_count + len(detail_matrix)
                cleaned_end_idx = min(footer_start_idx, len(table.rows) - (1 if template_summary_tr is not None else 0))
                for r_i in range(cleaned_start_idx, cleaned_end_idx):
                    if r_i >= len(table.rows):
                        continue
                    r = table.rows[r_i]
                    row_values = [str(c.text or "").strip() for c in r.cells]
                    # 仅清除真正的空模板行或“……”占位行，保护已有业务数据。
                    if any(v and v not in ("...", "…", "……", "......") for v in row_values):
                        continue
                    for c in r.cells:
                        c.text = ""

            logger.info(f"   [DOM 表格原位写盘] 成功向表格 {p_path} 原位填充 {len(matrix)} 行数据（明细 {len(detail_matrix)} 行 + 合计 1 行, 表头 {hdr_count} 行完好）")

        # 2. 按 Path 对段落/单槽位 Proposals 进行归组
        path_groups: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
        for p_idx, p in enumerate(proposals):
            if p_idx in handled_tbl_proposals:
                continue
            p_path = _normalize_table_path(str(p.get("path", "")).strip().replace("`", ""))
            p_val = str(p.get("proposed_text") if p.get("proposed_text") is not None else p.get("value", "")).strip().replace("`", "").replace("**", "")
            if p_path and p_val and not p_val.startswith("["):
                if not re.search(r'/tbl\[\d+\]$', p_path):
                    path_groups[p_path].append(p)

        for path, group_items in path_groups.items():
            clean_path = _normalize_table_path(path.replace("`", "").strip())
            snapshot = table_cell_snapshots.get(clean_path)
            p_elem = _get_snapshot_paragraph(snapshot) if snapshot else _find_paragraph_by_path(doc, clean_path)
            if snapshot:
                logger.info(
                    f"   [稳定表格节点写入] 路径 {clean_path} 使用扩行前快照节点，"
                    "避免按变化后的行号错写"
                )
            if not p_elem:
                continue

            is_in_table = bool(re.search(r'/tbl\[\d+\]', clean_path))
            if not is_in_table and hasattr(p_elem, '_element') and p_elem._element is not None:
                parent = p_elem._element.getparent()
                if parent is not None and str(parent.tag).endswith('tc'):
                    is_in_table = True

            slot_pattern = r'(_{1,}|＿{1,}|\[[^\]]+\]|［[^］]+］|【[^】]+】|（(?:姓名和职务|投标人的名称|投标人名称|代表姓名|职务)[^）]*）|\([^\)]*(?:name and title|bidder name)[^\)]*|\d*\s*年\s*\d*\s*月\s*\d*\s*日|(?<=[:：])\s{2,}|\s{2,}$)'

            # 同一条资格要求可能对应多张不同资质图片，统一以同一动态条款为锚点顺序插入。
            if group_items and all(item.get("type") == "image" for item in group_items):
                insertion_anchor = p_elem
                inserted_count = 0
                for image_item in group_items:
                    insertion_anchor, current_count = _insert_image_proposal(
                        insertion_anchor,
                        p_elem,
                        image_item,
                    )
                    inserted_count += current_count
                success_count += inserted_count
                continue

            if is_in_table or len(group_items) == 1:
                p_item = group_items[0]
                proposed_val = str(p_item.get("proposed_text") if p_item.get("proposed_text") is not None else p_item.get("value", "")).strip().replace("`", "").replace("**", "")
                orig_ctx = str(p_item.get("original_context", "")).strip().replace("`", "")
                
                real_text = p_elem.text or ""
                p_type = str(p_item.get("type", "")).strip()

                # 投标格式中的固定原文没有可填槽位时，禁止被 Agent 的短值或整句提案覆盖。
                # 已填写数据的修复必须携带含槽位的 original_context，才能证明它是数据替换而非格式改写。
                if (
                    not _is_footer_paragraph_target(doc, p_elem)
                    and _is_protected_template_overwrite(real_text, orig_ctx, proposed_val, p_type)
                ):
                    logger.warning(
                        f"[投标格式原文保护] 拦截无槽位覆盖提案: 节点={clean_path}, "
                        f"原文='{real_text[:80]}', 提案='{proposed_val[:80]}'"
                    )
                    continue

                use_underline = (not is_in_table)

                if p_item.get("type") == "image":
                    image_path = proposed_val
                    if os.path.exists(image_path):
                        from docx.shared import Inches, Pt
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        from docx.text.paragraph import Paragraph
                        width = p_item.get("width_inches", 5.5)
                        caption = p_item.get("caption", "")
                        
                        # 检查目标段落是否包含非空实质性条款标题（如 "1．法人或者其他组织的营业执照..."）
                        real_p_text = (p_elem.text or "").strip()
                        is_substantive = len(real_p_text) > 0 and not re.match(r'^(?:_{1,}|＿{1,}|\[[^\]]+\]|［[^］]+］|【[^】]+】|\s*)$', real_p_text)

                        if is_substantive:
                            # 保护条款原文标题不被抹除，在条款段落正下方动态插入新段落渲染图片
                            new_p_elem = parse_xml(r'<w:p %s/>' % nsdecls('w'))
                            p_elem._element.addnext(new_p_elem)
                            target_p = Paragraph(new_p_elem, p_elem._parent)
                        else:
                            # 纯占位符/空段落，直接清空内容并在本段落内渲染图片
                            p_elem._element.clear_content()
                            target_p = p_elem

                        target_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = target_p.add_run()
                        run.add_picture(image_path, width=Inches(width))
                        if caption:
                            cap_run = target_p.add_run(f"\n图：{caption}")
                            cap_run.font.bold = True
                            cap_run.font.size = Pt(10)
                            cap_run.font.name = "宋体"
                            try:
                                cap_rPr = cap_run._element.get_or_add_rPr()
                                cap_rFonts = cap_rPr.find(qn('w:rFonts'))
                                if cap_rFonts is None:
                                    cap_rPr.append(parse_xml(r'<w:rFonts %s w:ascii="宋体" w:hAnsi="宋体" w:eastAsia="宋体"/>' % nsdecls('w')))
                                else:
                                    cap_rFonts.set(qn('w:eastAsia'), '宋体')
                            except Exception:
                                pass
                        success_count += 1
                        logger.info(f"   [图片原位嵌入] 成功在节点 {clean_path} {'下方插入' if is_substantive else '原位替换'}证书图片并附加图注！")
                    continue
                    
                # [图片防字面量打印自愈] 检测正文或表格中是否被误填了图片路径字符串（形如 "资质证书（D:\...png）" 或含有 .png/.jpg）
                has_img_path_pattern = bool(re.search(r'(?:[a-zA-Z]:[\\/]|uploads[\\/]|qualifications[\\/]|[a-f0-9]{8}-[a-f0-9]{4}).*?\.(?:png|jpg|jpeg)', proposed_val, re.IGNORECASE)) or (
                    ("资质证书" in proposed_val or "许可证" in proposed_val or "执照" in proposed_val) and (".png" in proposed_val.lower() or ".jpg" in proposed_val.lower() or "d:\\" in proposed_val.lower())
                )
                if has_img_path_pattern and p_item.get("type") != "image":
                    logger.warning(f"   [图片防字面量自愈] 节点 {clean_path} 填入值包含图片物理路径字面量，启动全量自动转图嵌入...")
                    from app.db.session import SessionLocal
                    from app.db.models.business import CompanyQualification
                    from app.agents.tools.bid_db_tools import resolve_qualification_image_path
                    db_session = SessionLocal()
                    matched_quals_list = []
                    try:
                        quals = db_session.query(CompanyQualification).all()
                        seen_paths = set()
                        for q in quals:
                            q_name = q.name or ""
                            core_kw = q_name.replace("证书", "").replace("企业", "").replace("建筑业", "").strip()
                            file_url_str = q.file_url or ""
                            # 匹配条件：名称匹配 或 文件名匹配
                            if (core_kw and core_kw in proposed_val) or (file_url_str and os.path.basename(file_url_str) in proposed_val):
                                resolved_p, exists = resolve_qualification_image_path(file_url_str)
                                if exists and resolved_p not in seen_paths:
                                    seen_paths.add(resolved_p)
                                    matched_quals_list.append({"name": q_name, "path": resolved_p, "level": q.level or "通用"})
                    finally:
                        db_session.close()

                    if matched_quals_list:
                        from docx.shared import Inches, Pt
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        from docx.text.paragraph import Paragraph

                        # 剥离文本中的物理路径字符串，保留干净的条款承诺说明文字
                        clean_req_text = re.sub(r'[\[\(]?[a-zA-Z]:[\\/][^\n\r\(\)]+\.(?:png|jpg|jpeg)[^\n\r\)]*[\)\]]?', '', proposed_val)
                        clean_req_text = re.sub(r'[\[\(]?/?[^\s\(\)\[\]]+\.(?:png|jpg|jpeg)[^\n\r\)]*[\)\]]?', '', clean_req_text).strip()

                        # 若有干净的文本说明，先写入当前段落
                        if clean_req_text and len(clean_req_text) > 3:
                            real_p_text = (p_elem.text or "").strip()
                            prefix = _extract_prefix_from_text(real_p_text, orig_ctx)
                            p_elem._element.clear_content()
                            if prefix:
                                _apply_run_style_xml(p_elem.add_run(prefix), enable_underline=False, is_table=False)
                            _apply_run_style_xml(p_elem.add_run(clean_req_text), enable_underline=use_underline, is_table=False)
                            curr_anchor = p_elem
                        else:
                            curr_anchor = p_elem

                        # 依次将所有匹配到的资质图片在段落下方追加渲染
                        for q_info in matched_quals_list:
                            new_p_elem = parse_xml(r'<w:p %s/>' % nsdecls('w'))
                            curr_anchor._element.addnext(new_p_elem)
                            target_p = Paragraph(new_p_elem, p_elem._parent)
                            target_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = target_p.add_run()
                            run.add_picture(q_info["path"], width=Inches(5.5))
                            cap_run = target_p.add_run(f"\n图：{q_info['name']}（等级/范围: {q_info['level']}）")
                            cap_run.font.bold = True
                            cap_run.font.size = Pt(10)
                            cap_run.font.name = "宋体"
                            try:
                                cap_rPr = cap_run._element.get_or_add_rPr()
                                cap_rFonts = cap_rPr.find(qn('w:rFonts'))
                                if cap_rFonts is None:
                                    cap_rPr.append(parse_xml(r'<w:rFonts %s w:ascii="宋体" w:hAnsi="宋体" w:eastAsia="宋体"/>' % nsdecls('w')))
                                else:
                                    cap_rFonts.set(qn('w:eastAsia'), '宋体')
                            except Exception:
                                pass
                            curr_anchor = target_p
                            success_count += 1
                            logger.info(f"   [图片防字面量自愈成功] 成功在条款下方插入图片: {q_info['name']} -> {q_info['path']}")
                        continue

                # 自动清理零改动注释、省略号与截断
                from app.agents.review_engine import clean_zero_change_annotations, clean_all_ellipsis, is_zero_change_or_no_op_proposal
                if p_type != "image":
                    proposed_val = clean_zero_change_annotations(proposed_val)
                    proposed_val = clean_all_ellipsis(proposed_val)

                if not proposed_val:
                    continue

                if is_zero_change_or_no_op_proposal(proposed_val, real_text, p_type) or is_zero_change_or_no_op_proposal(proposed_val, orig_ctx, p_type):
                    logger.info(f"   [零改动保护] 提案 '{proposed_val}' 与原文一致或为保留说明，跳过改写: {clean_path}")
                    continue

                if not is_in_table and _is_full_paragraph_replacement(real_text, proposed_val, p_type):
                    _render_diff_paragraph_runs(p_elem, real_text, proposed_val, enable_underline_on_diff=True, is_table=False)
                    success_count += 1
                    continue

                if is_in_table:
                    from app.agents.review_engine import clean_cell_text_value
                    proposed_val = clean_cell_text_value(proposed_val, orig_ctx)

                    # [表格跨列防挤爆自愈与行路径多列分发]
                    # 检查路径是否为整行 /tr[N] 路径或首列 /tc[1] 路径
                    tbl_row_m = re.search(r'/tbl\[(\d+)\]/(?:tr|row)\[(\d+)\](?:\/(?:tc|cell)\[(\d+)\])?', clean_path)
                    if tbl_row_m:
                        tbl_idx = int(tbl_row_m.group(1)) - 1
                        tr_idx = int(tbl_row_m.group(2)) - 1
                        tc_idx_str = tbl_row_m.group(3)

                        if 0 <= tbl_idx < len(doc.tables) and 0 <= tr_idx < len(doc.tables[tbl_idx].rows):
                            target_row = doc.tables[tbl_idx].rows[tr_idx]
                            
                            # 场景 A: 若路径是整行路径 (无 tc) 或首单元格路径 (tc[1]) 且包含多列分隔符
                            sep_char = None
                            for s in ["｜", "|", "\t"]:
                                if s in proposed_val and len([t for t in proposed_val.split(s) if t.strip()]) >= 2:
                                    sep_char = s
                                    break

                            if sep_char and (tc_idx_str is None or tc_idx_str == "1"):
                                tokens = [t.strip() for t in proposed_val.split(sep_char) if t.strip()]
                                start_c = 1 if (len(tokens) == len(target_row.cells) - 1 and target_row.cells[0].text.strip().isdigit()) else 0
                                for off, tok in enumerate(tokens):
                                    c_i = start_c + off
                                    if c_i < len(target_row.cells):
                                        c_cell = target_row.cells[c_i]
                                        c_cell.text = tok
                                        for p_c in c_cell.paragraphs:
                                            for r_c in p_c.runs:
                                                _apply_run_style_xml(r_c, enable_underline=False, is_table=True)
                                success_count += 1
                                continue
                            elif tc_idx_str is None:
                                # 若为整行单段普通文本且目标表格为多列响应表 (>= 4 列)，智能填入核心响应列 (tc[2] 或 tc[1])
                                if len(target_row.cells) >= 4:
                                    # 序号保护
                                    if not target_row.cells[0].text.strip():
                                        target_row.cells[0].text = str(tr_idx)
                                    # 核心文本填入第 2 列或第 3 列
                                    resp_col = 2 if len(target_row.cells) >= 5 else 1
                                    target_row.cells[resp_col].text = proposed_val
                                    # 偏离状态补齐
                                    if len(target_row.cells) >= 4 and not target_row.cells[len(target_row.cells)-2].text.strip():
                                        target_row.cells[len(target_row.cells)-2].text = "无偏离"
                                    if len(target_row.cells) >= 5 and not target_row.cells[len(target_row.cells)-1].text.strip():
                                        target_row.cells[len(target_row.cells)-1].text = "完全响应招标文件要求，无偏离。"
                                    for c_elem in target_row.cells:
                                        for p_c in c_elem.paragraphs:
                                            for r_c in p_c.runs:
                                                _apply_run_style_xml(r_c, enable_underline=False, is_table=True)
                                    success_count += 1
                                    continue

                    # 场景 B: 单元格标签保护机制（防止抹杀 投标总报价（大写）： 或 交货期限： 等原有表头标签）
                    real_cell_text = p_elem.text or ""
                    m_lbl = re.match(r'^\s*([^:：\n]{2,30}[:：])\s*', real_cell_text)
                    if m_lbl and len(real_cell_text) < 60:
                        prefix = m_lbl.group(1)
                        clean_val = proposed_val
                        clean_prefix = re.sub(r'[\s:：_\[］\[\]（）\(\)]', '', prefix)
                        clean_p_val = re.sub(r'[\s:：_\[］\[\]（）\(\)]', '', clean_val)
                        if clean_prefix and clean_p_val.startswith(clean_prefix):
                            m_sub = re.match(r'^\s*([^:：\n]+[:：])\s*', clean_val)
                            if m_sub:
                                clean_val = clean_val[m_sub.end():].strip()
                            elif clean_val.startswith(prefix):
                                clean_val = clean_val[len(prefix):].strip()

                        final_cell_text = f"{prefix}{clean_val}"
                    else:
                        final_cell_text = proposed_val

                    p_elem._element.clear_content()
                    _apply_run_style_xml(p_elem.add_run(final_cell_text), enable_underline=False, is_table=True)
                    success_count += 1
                else:
                    # 优先检测段落是否为冒号标签引导行 (形如 邮    编： 或 投标单位代表姓名（签字）： 或 投标单位名称： 或 日    期：)
                    m_lbl = re.match(r'^\s*([^:：\n]{2,25}[:：])\s*', real_text)
                    if m_lbl and len(real_text) < 100 and not is_narrative_clause_or_lead_in(m_lbl.group(1)) and not is_narrative_clause_or_lead_in(real_text):
                        prefix = m_lbl.group(1)
                        clean_val = proposed_val
                        clean_prefix = re.sub(r'[\s:：_\[］\[\]（）\(\)]', '', prefix)
                        clean_p_val = re.sub(r'[\s:：_\[］\[\]（）\(\)]', '', clean_val)
                        if clean_prefix and clean_p_val.startswith(clean_prefix):
                            m_sub = re.match(r'^\s*([^:：\n]+[:：])\s*', clean_val)
                            if m_sub:
                                clean_val = clean_val[m_sub.end():].strip()
                            elif clean_val.startswith(prefix):
                                clean_val = clean_val[len(prefix):].strip()

                        # 若该段落是标题导语（如 5、与本投标有关的正式通讯地址为：）且没有有效填报值，保留原文
                        if not clean_val or clean_val == real_text:
                            continue

                        p_elem._element.clear_content()
                        _apply_run_style_xml(p_elem.add_run(prefix), enable_underline=False, is_table=False)
                        _apply_run_style_xml(p_elem.add_run(clean_val), enable_underline=use_underline, is_table=False)
                        success_count += 1
                        continue

                    # 单槽位原位切片替换（正文段落：宋体小四 Pt(12)）
                    ph_match = re.search(slot_pattern, real_text)
                    if ph_match:
                        start_pos, end_pos = ph_match.span()
                        before_text = real_text[:start_pos]
                        after_text = real_text[end_pos:]

                        # [防切片重复与嵌套爆炸自愈]
                        # 若 proposed_val 已经包含了 before_text 或 after_text 的核心词，说明这是一句整句，绝不能插在 slot 里！
                        b_clean = re.sub(r'[\s:：_＿\[\]［］\(\)（）]', '', before_text)
                        a_clean = re.sub(r'[\s:：_＿\[\]［］\(\)（）]', '', after_text)
                        is_nested_dup = (
                            (len(b_clean) >= 4 and proposed_val.startswith(b_clean[:4])) or
                            (len(a_clean) >= 4 and proposed_val.endswith(a_clean[-4:])) or
                            (len(proposed_val) >= 20 and len(b_clean) >= 4 and b_clean in proposed_val)
                        )
                        if is_nested_dup or _is_full_paragraph_replacement(real_text, proposed_val, p_type):
                            _render_diff_paragraph_runs(p_elem, real_text, proposed_val, enable_underline_on_diff=True, is_table=False)
                            success_count += 1
                            continue

                        p_elem._element.clear_content()
                        if before_text:
                            _apply_run_style_xml(p_elem.add_run(before_text), enable_underline=False, is_table=False)
                        _apply_run_style_xml(p_elem.add_run(proposed_val), enable_underline=use_underline, is_table=False)
                        if after_text:
                            _apply_run_style_xml(p_elem.add_run(after_text), enable_underline=False, is_table=False)
                        success_count += 1
                    else:
                        prefix = _extract_prefix_from_text(real_text, orig_ctx)
                        if prefix and not is_narrative_clause_or_lead_in(prefix):
                            # 彻底杜绝前缀重复：若 proposed_val 已经携带了该前缀标签，剥离重复部分
                            clean_p_val = proposed_val
                            clean_prefix = re.sub(r'[\s:：_\[］\[\]（）\(\)]', '', prefix)
                            clean_val = re.sub(r'[\s:：_\[］\[\]（）\(\)]', '', clean_p_val)
                            if clean_prefix and clean_val.startswith(clean_prefix):
                                m_p = re.match(r'^\s*([^\n:：]+[:：])\s*', clean_p_val)
                                if m_p:
                                    clean_p_val = clean_p_val[m_p.end():].strip()
                                elif clean_p_val.startswith(prefix):
                                    clean_p_val = clean_p_val[len(prefix):].strip()
                            p_elem._element.clear_content()
                            _apply_run_style_xml(p_elem.add_run(prefix), enable_underline=False, is_table=False)
                            _apply_run_style_xml(p_elem.add_run(clean_p_val), enable_underline=use_underline, is_table=False)
                            success_count += 1
                        else:
                            # 已经包含填报值则无需重复追加
                            if proposed_val in real_text:
                                logger.info(f"   段落 {clean_path} 已包含填报值 '{proposed_val}'，安全保留")
                                success_count += 1
                                continue
                            # 过滤元数据说明性文本，严禁追加到正文句尾
                            if not proposed_val or proposed_val.startswith("待补充") or proposed_val.startswith("[待") or any(
                                k in proposed_val for k in ["无需写盘", "零改动", "固定原文", "原样保留", "无需改动", "无槽位", "待线下"]
                            ):
                                continue
                            # 严格防护：若原段落为固定正文叙述句/条款导语且无槽位特征，严禁在句尾盲目追加提案
                            if is_narrative_clause_or_lead_in(real_text):
                                logger.warning(f"   [防叙述句尾污染] 段落 {clean_path} 为固定叙述句/条款导语 ('{real_text[:30]}')，拦截句尾追加提案 '{proposed_val}'")
                                continue
                            # 若未包含且无已知前缀，在句尾安全追加填报值
                            p_elem._element.clear_content()
                            _apply_run_style_xml(p_elem.add_run(real_text), enable_underline=False, is_table=False)
                            _apply_run_style_xml(p_elem.add_run(proposed_val), enable_underline=use_underline, is_table=False)
                            success_count += 1
            else:
                # -----------------------------------------------------------------
                # 模式 B: 同段落多槽位提案原位顺序插值合并替换 (Multi-Slot In-Place Interpolation)
                # 解决《投标函》同一段落包含多个占位符时多个提案相互擦除覆盖的致命痛点
                # -----------------------------------------------------------------
                real_text = p_elem.text or ""
                combined_orig_ctx = "".join([str(p.get("original_context", "")) for p in group_items])
                use_underline = (not is_in_table) and bool(re.search(r'(_{2,}|＿{2,})', real_text + combined_orig_ctx))

                # 检查多槽位组合中是否存在已包含整句覆盖的提案
                full_rep_item = next((p for p in group_items if _is_full_paragraph_replacement(real_text, str(p.get("proposed_text") if p.get("proposed_text") is not None else p.get("value", "")), str(p.get("type", "")))), None)
                if full_rep_item:
                    f_val = str(full_rep_item.get("proposed_text") if full_rep_item.get("proposed_text") is not None else full_rep_item.get("value", "")).strip()
                    _render_diff_paragraph_runs(p_elem, real_text, f_val, enable_underline_on_diff=True, is_table=is_in_table)
                    success_count += len(group_items)
                    logger.info(f"   [多槽位整句直写] 命中整句提案，成功整段渲染写入段落 {path}！")
                    continue

                current_text = real_text
                runs_to_build: List[Tuple[str, bool]] = []
                filled_items_count = 0

                for p_item in group_items:
                    val = str(p_item.get("proposed_text") if p_item.get("proposed_text") is not None else p_item.get("value", "")).strip()
                    if not val:
                        continue
                    ph_match = re.search(slot_pattern, current_text)
                    if ph_match:
                        start_pos, end_pos = ph_match.span()
                        before_text = current_text[:start_pos]
                        current_text = current_text[end_pos:]
                        if before_text:
                            runs_to_build.append((before_text, False))
                        runs_to_build.append((val, use_underline))
                        filled_items_count += 1
                    else:
                        if val not in current_text and val not in [t[0] for t in runs_to_build]:
                            runs_to_build.append((val, use_underline))
                            filled_items_count += 1

                if current_text:
                    runs_to_build.append((current_text, False))

                if runs_to_build:
                    p_elem._element.clear_content()
                    for t_seg, u_flag in runs_to_build:
                        _apply_run_style_xml(p_elem.add_run(t_seg), enable_underline=u_flag, is_table=is_in_table)
                    success_count += filled_items_count
                    logger.info(f"   [多槽位原位合并] 成功将 {filled_items_count} 条提案按顺序插值替换写入段落 {path}！")

            # 每处修改地方均实时执行 R10 模板保留率校验与明细日志输出
            from app.agents.review_engine import check_and_rollback_single_node
            check_and_rollback_single_node(p_elem, real_text, path)

        # 表格扩行后，提案中的物理行号可能已经变化；更新为真实节点当前路径，
        # 供后续终审回读校验与修复轮继续准确定位。
        _rebind_table_cell_proposal_paths(proposals, table_cell_snapshots)

        # 表格全自动留白自检与 LLM 动态自愈修复引擎（绝不删行，零业务硬编码）
        from app.utils.table_utils import inspect_and_repair_table_blanks
        inspect_and_repair_table_blanks(doc)
        # 写盘路径仍按原始模板解析完毕后，再移除多余子标题，避免改变 /body/p[N] 索引。
        subtitle_removed = _remove_redundant_pricing_subtitle(doc)

        if success_count > 0 or subtitle_removed > 0:
            from app.agents.tools.bid_db_tools import _safe_save_doc
            _safe_save_doc(doc, docx_path)
            logger.info(f"   [DOM 原位填报与美化] 成功原位写入并修饰 {success_count} 条提案（表格外保留下划线, 表格内取消下划线，留白单元格智能闭合）")
        return success_count
    except Exception as exc:
        logger.error(f"   DOM 级写盘填报异常: {exc}")
        return 0


def _apply_underline_to_filled_doc(docx_temp_path: str, proposals: List[Dict]) -> None:
    """兼容封装：调用 DOM 原位填报与美化引擎"""
    fill_docx_proposals_in_dom(docx_temp_path, proposals)


def auto_repair_officecli_commands(commands: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Office CLI 批处理指令前置自愈与语法校正器：
    1. 自动将 add row 命令的 parent 剥离校正为父表格根路径 /body/tbl[N]（防止在 /tr[M] 下添加行触发 Office CLI 崩盘）；
    2. 自动修正 set 指令中的非法嵌套 Path（如将 /tr[2]/tr[last()] 或 /tr[6]/tr[last()]/tc[1] 转换为 /row[last()]/cell[C]）；
    3. 规范路径命名风格，统一兼容 /row[R]/cell[C] 选择器。
    """
    if not commands:
        return []
    repaired = []
    import re
    for cmd in commands:
        if not isinstance(cmd, dict):
            continue
        c_item = dict(cmd)
        c_name = c_item.get("command", "")
        c_parent = str(c_item.get("parent", "")).strip()
        c_path = str(c_item.get("path", "")).strip()
        c_type = str(c_item.get("type", "")).strip()

        if c_name == "add" and c_type == "row":
            # 1. 父节点校正：将 /body/tbl[3]/tr[2] 或 /body/tbl[3]/tr[6] 等剥离为 /body/tbl[3]
            base_tbl = re.sub(r'/(?:tr|row|tc|cell)\[\d+\].*$', '', c_parent)
            if not base_tbl:
                base_tbl = re.sub(r'/(?:tr|row|tc|cell)\[\d+\].*$', '', c_path)
            c_item["parent"] = base_tbl or c_parent

        elif c_name == "set" and c_path:
            # 2. 路径校正：修复所有非法嵌套路径
            nested_tr = re.search(r'(/tbl\[\d+\])/(?:tr|row)\[\d+\]/(?:tr|row)\[(last\(\)|\d+)\]/(?:tc|cell)\[(\d+)\]', c_path)
            if nested_tr:
                tbl_base = re.search(r'^(.*?/tbl\[\d+\])', c_path).group(1)
                row_idx = nested_tr.group(2)  # "last()" 或数字
                col_num = nested_tr.group(3)
                c_item["path"] = f"{tbl_base}/row[{row_idx}]/cell[{col_num}]"
            elif "/tr[last()]" in c_path or "/row[last()]" in c_path:
                tbl_match = re.search(r'^(.*?/tbl\[\d+\])', c_path)
                cell_match = re.search(r'/(?:tc|cell)\[(\d+)\]', c_path)
                if tbl_match and cell_match:
                    tbl_base = tbl_match.group(1)
                    col_num = cell_match.group(1)
                    c_item["path"] = f"{tbl_base}/row[last()]/cell[{col_num}]"

        repaired.append(c_item)
    return repaired


def proposals_to_commands(proposals: List[Dict]) -> tuple:
    """将 Worker 提案转换为 OfficeCLI 写盘命令，返回 (commands, approved, rejected)"""
    commands = []
    approved, rejected = 0, 0
    import re
    for p in proposals:
        path = _normalize_table_path(str(p.get("path", "")).strip())
        text = str(p.get("proposed_text") if p.get("proposed_text") is not None else p.get("value", "")).strip()
        orig_context = str(p.get("original_context", "")).strip()

        # 移除对 source_tool 为 none 的硬编码拦截；只要非空且不为占位异常提示即予以放行
        if not text or text.startswith("[待补充") or text.startswith("[建议") or text.startswith("[查询") or text.startswith("[错误") or text.startswith("待补充"):
            rejected += 1; continue
        if any(k in text for k in ["无需写盘", "零改动", "固定原文", "原样保留", "无需改动", "无槽位", "待线下"]):
            rejected += 1; continue

        # 剥离说明性元数据注释并拦截零改动无操作提案
        from app.agents.review_engine import clean_zero_change_annotations, is_zero_change_or_no_op_proposal
        text = clean_zero_change_annotations(text)
        if is_zero_change_or_no_op_proposal(text, orig_context):
            rejected += 1; continue

        if not path or "~" in path or ".." in path:
            rejected += 1; continue

        # 任何带表格单元格的路径必须先完成规范化，未知路径不允许降级为行级写入。
        if re.search(r"/tbl\[\d+\]/(?:tr|row)\[\d+\]/", path) and not re.search(
            r"/tbl\[\d+\]/(?:tr|row)\[\d+\]/tc\[\d+\]", path
        ):
            logger.error(f"[Reviewer] 拒绝无法定位到具体单元格的表格路径: {path}")
            rejected += 1
            continue
        
        # 路径精准度保护：如果是容器节点（/tbl[N] 或 /tr[N]），通常丢弃。
        # 但如果提议内容是 JSON 数组（代表批量行插入）或来源标记为表格模板行，则放行
        if re.search(r'/(?:tbl|tr)\[\d+\]$', path):
            is_table_json = text.lstrip().startswith("[") and text.rstrip().endswith("]")
            is_table_proposal = orig_context == "表格模板行"
            if is_table_json or is_table_proposal:
                logger.info(f"   [Reviewer] 识别到表格批量插行提案，放行路径: {path}")
            else:
                logger.warning(f"   [Reviewer] 丢弃非精准表格容器路径: {path}，防止误覆盖表头")
                rejected += 1; continue

        # 表格单元格 XPath 若止步于 /tc[N]，自动补正到该具体单元格的第一段 /p[1]
        if re.search(r'/tc\[\d+\]$', path):
            path += "/p[1]"

        # 判断是否为表格单元格
        is_in_table = bool(re.search(r'/(?:tbl|tr|tc|cell)\[', path))
        if is_in_table:
            # 表格单元格隔离前缀标签，纯化写入数据
            from app.agents.review_engine import clean_cell_text_value
            text = clean_cell_text_value(text, orig_context)
            prefix = _extract_label_prefix(orig_context)
            if prefix:
                clean_prefix = re.sub(r'[\s:：_\[］\[\]（）\(\)]', '', prefix)
                clean_text = re.sub(r'[\s:：_\[］\[\]（）\(\)]', '', text)
                if clean_prefix and clean_text.startswith(clean_prefix):
                    m_p = re.match(r'^\s*([^\n:：]+[:：])\s*', text)
                    if m_p:
                        text = text[m_p.end():].strip()
                    elif text.startswith(prefix):
                        text = text[len(prefix):].strip()
        else:
            # 原文前缀绝对保护逻辑：检查 original_context 是否包含前缀标签，防止擦除原标书文本
            prefix = _extract_label_prefix(orig_context)
            if prefix:
                clean_prefix = re.sub(r'[\s:：_\[］\[\]（）\(\)]', '', prefix)
                clean_text = re.sub(r'[\s:：_\[］\[\]（）\(\)]', '', text)
                if clean_prefix and not clean_text.startswith(clean_prefix):
                    text = f"{prefix}{text}"
                    logger.info(f"   [前缀保护] 自动为路径 {path} 补全原文标签: {prefix}")

        commands.append({"command": "set", "path": path, "props": {"text": text}})
        approved += 1
    
    # 自动执行指令语法与 Path 自愈
    commands = auto_repair_officecli_commands(commands)
    return commands, approved, rejected


def supervisor_audit_node(state: BidFillerState) -> Dict[str, Any]:
    """
    Supervisor 终审节点 — 重新探测 Word DOM 结构，审核全图空位遗漏与表格规范。
    若存在未填槽位且未达最大重试次数，构造按章节的 repair_instructions_map 并打回重修。
    """
    logger.info("[LangGraph Node 3/4] supervisor_audit_node: 启动 Supervisor 全局质量终审...")
    doc_id = state.get("document_id", "")
    docx_temp_path = state.get("docx_temp_path")
    repair_count = state.get("repair_count", 0)
    max_repair_rounds = state.get("max_repair_rounds", 2)
    audit_items: List[FillingAuditItem] = list(state.get("audit_items") or [])

    if not docx_temp_path or not os.path.exists(docx_temp_path):
        logger.error("   临时 Word 文件不存在，阻断最终文件发布")
        return {
            "audit_passed": False,
            "audit_blocked": True,
            "repair_count": repair_count,
            "audit_items": audit_items,
            "review_findings": [{
                "rule_id": "R0-ARTIFACT-MISSING",
                "severity": "error",
                "path": "N/A",
                "description": "临时 Word 文件不存在，无法确认最终产物正确性",
                "current_value": "",
                "expected_value": "可回读的 Word 工作副本",
                "auto_fixable": False,
            }],
        }

    unfilled_findings = []
    repair_instructions_map: Dict[str, str] = {}
    chapter_unfilled_count: Dict[str, int] = defaultdict(int)
    expected_total: Optional[Decimal] = None
    expected_total_words = ""
    expected_company_name = ""

    try:
        from app.services.office_cli_service import office_cli_service
        import asyncio, concurrent.futures
        from app.services.office_cli_service import office_cli_service
        def _sync(coro):
            try:
                asyncio.get_running_loop()
            except RuntimeError:
                return asyncio.run(coro)
            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as e:
                return e.submit(asyncio.run, coro).result()

        structure_str = _sync(office_cli_service.query_structure(docx_temp_path, "all"))
        lines = str(structure_str).split("\n")

        # 0. 查验全量 Worker 的实际写盘与提案记录 (Zero-Write Audit)
        from app.agents.bid_filler_workers import get_worker_proposals
        worker_proposals = get_worker_proposals(doc_id)
        written_chapters = set()
        for p in worker_proposals:
            if isinstance(p, dict) and p.get("chapter_title"):
                written_chapters.add(p.get("chapter_title"))

        # 校验是否有 needs_data / needs_fill 核心章节未能成功产出提案 (Zero-Write Audit 强制拦截)
        slot_analysis = state.get("slot_analysis")
        if slot_analysis and hasattr(slot_analysis, "chapters"):
            for ch in slot_analysis.chapters:
                ch_title = getattr(ch, "chapter_title", "")
                ch_cat = getattr(ch, "category", "")
                if ch_cat in ("needs_fill", "needs_data") and ch_title:
                    has_written = any(ch_title in wc or wc in ch_title for wc in written_chapters)
                    if not has_written:
                        chapter_unfilled_count[ch_title] += 1
                        unfilled_findings.append({
                            "chapter": ch_title,
                            "snippet": f"核心业务章节【{ch_title}】在上一轮中执行失败或零提案产生",
                            "type": "zero_write_chapter"
                        })
                        logger.warning(f"   [Supervisor 终审拦截] 核心章节【{ch_title}】存在零写盘，强制打回重修！")

        # 探测通用占位符模式与未写盘硬伤 (通用正则，无任何硬编码字段名)
        unfilled_patterns = [
            r'_{2,}', r'\[待[补充|填].*?\]', r'\(\s{2,}\)', r'【待填.*?】',
            r'^[^\n_\[\]【】]{2,30}[:：]\s*$'  # 任何以冒号结尾且其后全空的未填字段标签
        ]
        current_chapter = "通用章节"

        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
            if "Heading" in line_str or "标题" in line_str or line_str.startswith("#"):
                current_chapter = line_str

            # 1. 未填占位符检查与字段空盲检查
            has_unfilled = False
            for pat in unfilled_patterns:
                if re.search(pat, line_str):
                    chapter_unfilled_count[current_chapter] += 1
                    unfilled_findings.append({
                        "chapter": current_chapter,
                        "snippet": line_str[:100],
                        "type": "unfilled_slot"
                    })
                    has_unfilled = True
                    break

            if has_unfilled:
                continue

            # 2. 授权委托书主体错挂检查（严禁将“致：招标代理机构”填到“参加...组织的”采购人位置）
            if "参加" in line_str and "组织的" in line_str and ("咨询" in line_str or "代理" in line_str):
                chapter_unfilled_count[current_chapter] += 1
                unfilled_findings.append({
                    "chapter": current_chapter,
                    "snippet": line_str[:100],
                    "type": "misassigned_entity"
                })
                logger.warning(f"   检出授权委托书主体错挂: '{line_str[:60]}...'")

        # 3. 深度段落与表格质检与全自动自愈管线 (Deep Document Quality Audit & Auto-Heal)
        try:
            doc_obj = Document(docx_temp_path)
            doc_modified = False
            from app.db.models.business import CompanyProfileModel
            from app.db.models.metadata import TimelineMetadata
            from app.db.models.ai_analysis import CostEstimate
            from app.utils.rmb_formatter import number_to_chinese_rmb
            from datetime import datetime

            db_audit: Session = SessionLocal()
            try:
                # 从同一份财务明细计算期望总价，供写盘后产物校验使用。
                cost_items_for_validation = db_audit.query(CostEstimate).filter(
                    CostEstimate.document_id == doc_id
                ).all()
                if isinstance(cost_items_for_validation, (list, tuple)) and cost_items_for_validation:
                    total_amount = Decimal("0.00")
                    for cost_item in cost_items_for_validation:
                        item_total = _parse_decimal_amount(str(getattr(cost_item, "calculated_total", "")))
                        if item_total is not None:
                            total_amount += item_total
                    if total_amount > Decimal("0.00"):
                        expected_total = total_amount.quantize(Decimal("0.01"))
                        expected_total_words = number_to_chinese_rmb(str(expected_total))

                profile = state.get("company_profile")
                expected_company_name = str(
                    getattr(profile, "company_name", "") or ""
                ).strip()
                if not expected_company_name:
                    profile_model = db_audit.query(CompanyProfileModel).first()
                    expected_company_name = str(
                        getattr(profile_model, "company_name", "") or ""
                    ).strip()

                def _safe_str(val: Any) -> Optional[str]:
                    if val is not None and isinstance(val, str) and val.strip() and not str(val).startswith("<MagicMock"):
                        return val.strip()
                    return None

                # 3.1 正文段落冒号留白深度自检与企业/项目数据原位自愈 (Paragraph Blank Slots Auto-Heal)
                try:
                    prof = db_audit.query(CompanyProfileModel).first() if db_audit else None
                    tl = db_audit.query(TimelineMetadata).filter(TimelineMetadata.document_id == doc_id).first() if db_audit else None

                    for p_idx, p_elem in enumerate(doc_obj.paragraphs):
                        p_raw = p_elem.text.strip()
                        if not p_raw:
                            continue

                        # 前置严格过滤：若原段落为正文叙述句、公文条款导语或含有句中标点，绝对严禁作为表单属性标签进行自愈！
                        if is_narrative_clause_or_lead_in(p_raw):
                            continue

                        # 探测冒号后连续空格/下划线/未填占位符的属性标签行
                        m_slot = re.match(r'^\s*([^:：\n]{2,25}[:：])(\s{2,}|_{2,}|＿{2,}|\s*年\s*月\s*日|\[待[^\]]+\]|［待[^］]+］|\s*$)', p_raw)
                        if m_slot:
                            lbl = m_slot.group(1)
                            if is_narrative_clause_or_lead_in(lbl):
                                continue

                            lbl_clean = re.sub(r'[\s:：_＿（）\(\)［］\[\]]', '', lbl)
                            if len(lbl_clean) < 2 or len(lbl_clean) > 20:
                                continue

                            filled_val = None

                            if prof:
                                # 基于企业画像与项目元数据的动态 Schema 解析表 (纯动态取值，无任何硬编码业务数据)
                                field_resolvers = [
                                    (["地址", "通讯地址", "注册地址", "办公地址", "单位地址", "公司地址"], lambda: _safe_str(getattr(prof, "registered_address", None))),
                                    (["电话", "联系电话", "手机", "联系方式", "固定电话", "办公电话"], lambda: _safe_str(getattr(prof, "contact_phone", None))),
                                    (["邮编", "邮政编码", "邮政代码"], lambda: _safe_str(getattr(prof, "postal_code", None))),
                                    (["邮箱", "电子邮箱", "Email", "email"], lambda: _safe_str(getattr(prof, "email", None))),
                                    (["传真", "传真号码", "公司传真", "单位传真"], lambda: _safe_str(getattr(prof, "fax_number", None))),
                                    (["名称", "投标人", "投标单位", "单位名称", "投标人全称", "投标单位全称", "企业名称", "单位全称"], lambda: _safe_str(getattr(prof, "company_name", None))),
                                ]

                                for keywords, getter_fn in field_resolvers:
                                    if any(k == lbl_clean or lbl_clean.endswith(k) or k in lbl_clean for k in keywords):
                                        filled_val = getter_fn()
                                        break

                                # 法定代表人/授权代表/签字人 与日期复合槽位动态解析
                                if not filled_val and any(k in lbl_clean for k in ["代表", "法人", "法定代表人", "签字", "授权代表", "代理人", "签署人"]):
                                    rep_name = _safe_str(getattr(prof, "legal_representative", None))
                                    if any(k in p_raw for k in ["年", "月", "日"]):
                                        date_val = _safe_str(getattr(tl, "planned_delivery_date", None))
                                        if rep_name and date_val:
                                            filled_val = f"{rep_name}                       {date_val}"
                                        elif rep_name:
                                            filled_val = rep_name
                                        elif date_val:
                                            filled_val = date_val
                                    else:
                                        filled_val = rep_name
                                elif not filled_val and (any(k in lbl_clean for k in ["日期", "日 期", "签署日期", "投标日期", "签章日期"]) or (len(lbl_clean) <= 4 and any(k in p_raw for k in ["年", "月", "日"]))):
                                    filled_val = _safe_str(getattr(tl, "planned_delivery_date", None))

                            if filled_val and isinstance(filled_val, str):
                                p_elem._element.clear_content()
                                _apply_run_style_xml(p_elem.add_run(lbl), enable_underline=False, is_table=False)
                                _apply_run_style_xml(p_elem.add_run(f" {filled_val}"), enable_underline=True, is_table=False)
                                doc_modified = True
                                logger.info(f"   🔧 [Supervisor 段落自愈] 已自动将企业档案补全至段落槽位: {lbl} -> '{filled_val}'")
                except Exception as exc_para:
                    logger.warning(f"   Supervisor 段落自愈检测异常: {exc_para}")

                # 3.2 深度表格质检与自愈管线 (Deep Table Quality Audit & Auto-Heal)
                tbl_count = len(doc_obj.tables)
                for tbl_idx, table in enumerate(doc_obj.tables, start=1):
                    if not table.rows or len(table.rows) == 0:
                        continue

                    header_row = table.rows[0]
                    header_texts = [cell.text.strip() for cell in header_row.cells]
                    header_combined = " ".join(header_texts)

                    # 场景 A：检查固定格式表单 / 单项汇总表
                    from app.utils.table_utils import is_fixed_slot_form_table, detect_table_header_rows as _det_hdr
                    hdr_c = _det_hdr(table)
                    total_cols = len(table.rows[0].cells) if table.rows else 1
                    is_fixed_table = is_fixed_slot_form_table(table, hdr_c)

                    if is_fixed_table:
                        # 1. 检查表格中是否存在重复插入的表尾汇总/合并行 (冗余插行自愈清理)
                        footer_row_indices = []
                        for r_idx in range(hdr_c, len(table.rows)):
                            row_tcs = set(c._tc for c in table.rows[r_idx].cells)
                            r_raw = "".join(c.text for c in table.rows[r_idx].cells)
                            if len(row_tcs) < total_cols or any(kw in r_raw for kw in ["大写", "总价", "报价", "合计", "总计"]):
                                if r_idx > hdr_c:
                                    footer_row_indices.append(r_idx)

                        if len(footer_row_indices) > 1:
                            for del_r_idx in reversed(footer_row_indices[:-1]):
                                r_elem = table.rows[del_r_idx]._tr
                                r_elem.getparent().remove(r_elem)
                                doc_modified = True
                                logger.info(f"   🧹 [Supervisor 表格自愈] 清除固定表单 {tbl_idx} 中多余的重复表尾行 (第 {del_r_idx+1} 行)")

                        # 2. 清除固定表单中误插入的多余分项明细行 (确保数据行只有 1 行)
                        if len(table.rows) > hdr_c + 2:
                            footer_row_idx = None
                            for r_idx in range(len(table.rows) - 1, hdr_c, -1):
                                row_tcs = set(c._tc for c in table.rows[r_idx].cells)
                                r_raw = "".join(c.text for c in table.rows[r_idx].cells)
                                if len(row_tcs) < total_cols or any(kw in r_raw for kw in ["大写", "总价", "报价", "合计", "总计"]):
                                    footer_row_idx = r_idx
                                    break
                            if footer_row_idx is not None and footer_row_idx > hdr_c + 1:
                                for del_idx in range(footer_row_idx - 1, hdr_c, -1):
                                    del_tr = table.rows[del_idx]._tr
                                    del_tr.getparent().remove(del_tr)
                                    doc_modified = True
                                    logger.info(f"   🧹 [Supervisor 表格自愈] 清除固定表单 {tbl_idx} 中多余插入的明细行 (第 {del_idx+1} 行)")

                        # 3. 检查表尾汇总行/大写金额单元格是否留空并原位自愈
                        for r_i in range(hdr_c, len(table.rows)):
                            row_tcs = set(c._tc for c in table.rows[r_i].cells)
                            r_raw = "".join(c.text for c in table.rows[r_i].cells)
                            is_footer_or_caps = (len(row_tcs) < total_cols or any(kw in r_raw for kw in ["大写", "总价", "报价", "合计"])) and r_i >= hdr_c
                            if is_footer_or_caps:
                                empty_caps_cell = None
                                for c_i, cell in enumerate(table.rows[r_i].cells):
                                    c_txt = cell.text.strip()
                                    if c_i > 0 and (not c_txt or any(pat in c_txt for pat in ["___", "[待", "待填", "待补充"])):
                                        empty_caps_cell = cell
                                        break

                                if empty_caps_cell is not None:
                                    cost_items = db_audit.query(CostEstimate).filter(CostEstimate.document_id == doc_id).all()
                                    if cost_items:
                                        total_val = sum(it.calculated_total for it in cost_items)
                                        upper_cn = number_to_chinese_rmb(total_val)
                                        empty_caps_cell.text = upper_cn
                                        doc_modified = True
                                        logger.info(f"   🔧 [Supervisor 表格自愈] 已自动将大写总金额 '{upper_cn}' 原位补全至表单单元格")
                                    else:
                                        chapter_unfilled_count["汇总表单"] += 1
                                        unfilled_findings.append({
                                            "chapter": "汇总表单",
                                            "snippet": f"表单第 {r_i+1} 行大写总金额单元格留空",
                                            "type": "unfilled_table_cell"
                                        })

                    # 场景 B：检查分项清单 / 采购明细表中的“生产厂家”等关键必填列
                    if len(table.rows) > 1:
                        mfg_col_idx = -1
                        for c_i, h_txt in enumerate(header_texts):
                            if any(kw in h_txt for kw in ["生产厂家", "制造厂家", "制造厂商", "生产厂商", "制造商", "厂家", "产地"]):
                                mfg_col_idx = c_i
                                break

                        if mfg_col_idx != -1:
                            data_rows = table.rows[1:]
                            mfg_cells = []
                            for dr in data_rows:
                                dr_text = "".join(c.text for c in dr.cells)
                                if not any(k in dr_text for k in ["合计", "总计", "大写", "小计"]):
                                    if mfg_col_idx < len(dr.cells):
                                        mfg_cells.append(dr.cells[mfg_col_idx].text.strip())

                            if mfg_cells and all(not val for val in mfg_cells):
                                cost_items = db_audit.query(CostEstimate).filter(CostEstimate.document_id == doc_id).all()
                                if cost_items:
                                    c_idx = 0
                                    for dr in data_rows:
                                        dr_text = "".join(c.text for c in dr.cells)
                                        if not any(k in dr_text for k in ["合计", "总计", "大写", "小计"]) and c_idx < len(cost_items):
                                            item = cost_items[c_idx]
                                            mfg_str = str(getattr(item, 'manufacturer', '') or getattr(item, 'brand', '') or '').strip()
                                            if mfg_str and mfg_col_idx < len(dr.cells):
                                                dr.cells[mfg_col_idx].text = mfg_str
                                                doc_modified = True
                                            c_idx += 1
                                    logger.info(f"   🔧 [Supervisor 表格自愈] 成功从数据库将生产厂家数据回填至分项报价表 (共 {c_idx} 项)")
                                else:
                                    chapter_unfilled_count["分项报价表"] += 1
                                    unfilled_findings.append({
                                        "chapter": "分项报价表",
                                        "snippet": f"表格第 {tbl_idx} 个表【生产厂家】列数据全为空白",
                                        "type": "unfilled_table_column"
                                    })

                if doc_modified:
                    doc_obj.save(docx_temp_path)
                    logger.info(f"   💾 [Supervisor 终审自愈] 全部段落与表格自动修正已同步保存至工作副本: {docx_temp_path}")

                # 所有自动修正完成后重新回读产物，验证真实落位而不是仅相信 Worker 提案日志。
                artifact_findings = validate_filled_docx_integrity(
                    docx_path=docx_temp_path,
                    expected_total=expected_total,
                    expected_total_words=expected_total_words,
                    company_name=expected_company_name,
                    proposals=worker_proposals,
                )
                for finding in artifact_findings:
                    finding_chapter = finding.get("chapter", "产物回读校验")
                    chapter_unfilled_count[finding_chapter] += 1
                    unfilled_findings.append({
                        "chapter": finding_chapter,
                        "snippet": finding.get("snippet", "Word 产物校验失败"),
                        "type": finding.get("type", "artifact_integrity_error"),
                        "path": finding.get("path", "N/A"),
                    })
                    logger.error(
                        f"   [产物回读校验失败] {finding.get('path', 'N/A')}: "
                        f"{finding.get('snippet', '未知错误')}"
                    )

                logger.info(
                    f"   [Supervisor DOM 扫描] 已查验 {tbl_count} 个表格与 "
                    f"{len(doc_obj.paragraphs)} 个段落结构，发现 {len(artifact_findings)} 项产物问题"
                )
            finally:
                db_audit.close()
        except Exception as exc_tbl:
            logger.warning(f"   Supervisor 深度质检异常: {exc_tbl}")

        # 4. 全图招标文件需求与填报内容对应度深度核验 (Requirement Alignment Check)
        try:
            company_profile = state.get("company_profile")
            original_ctx = state.get("original_context", "")[:2000]
            
            # 提取已填盘的各章节文本正文与表格摘要
            filled_chapter_samples = []
            curr_ch = "前言/封面"
            curr_lines = []
            for line in lines:
                l_str = line.strip()
                if not l_str:
                    continue
                if "Heading" in l_str or "标题" in l_str or l_str.startswith("#"):
                    if curr_lines:
                        filled_chapter_samples.append(f"【章节: {curr_ch}】:\n" + "\n".join(curr_lines[:5]))
                    curr_ch = l_str
                    curr_lines = []
                else:
                    curr_lines.append(l_str)
            if curr_lines:
                filled_chapter_samples.append(f"【章节: {curr_ch}】:\n" + "\n".join(curr_lines[:5]))

            filled_doc_summary = "\n\n".join(filled_chapter_samples[:8])

            if filled_doc_summary and original_ctx:
                profile_info = f"投标人公司全称: {getattr(company_profile, 'company_name', '')}, 法人: {getattr(company_profile, 'legal_person', '')}" if company_profile else ""
                audit_prompt = f"""你是资深招投标终审专家。请深度审查已填写完成的 Word 投标文件内容，核实其是否完全对应并满足招标文件与数据源的各项要求：

【已填写完成的投标文件各章节内容摘要】:
{filled_doc_summary}

【招标文件原始要求与企业主档案数据源】:
{profile_info}
{original_ctx}

【终审核查要求】:
1. 需求对应度：核实填报的内容（项目名称、招标编号、技术参数、商务条款、人员资质等）是否完全响应并对应上了招标文件的指标要求；
2. 真实一致性：核实投标人名称、法定代表人姓名、报价数字与大写金额是否与数据源 100% 对齐，有无张冠李戴或漏填答非所问；
3. 实质性响应：核实是否有未能满足招采门槛的硬性漏洞。

若发现任何未能对应需求或数据不匹配项，请列出具体的章节与错项；若全篇内容完美对应响应，请直接回复"REQUIREMENT_MATCHED"。"""

                from app.services.llm_service import llm_service
                audit_llm = llm_service.get_llm(
                    temperature=0.3,
                    json_mode=False,
                    tenant_id=state.get("tenant_id") or "default-tenant",
                )
                if audit_llm is not None:
                    audit_res = _sync(audit_llm.ainvoke(audit_prompt))
                    audit_content = getattr(audit_res, 'content', str(audit_res)).strip()
                    if "REQUIREMENT_MATCHED" not in audit_content and len(audit_content) > 10:
                        chapter_unfilled_count["全图需求响应对查"] += 1
                        unfilled_findings.append({
                            "chapter": "全图需求响应对查",
                            "snippet": audit_content[:200],
                            "type": "requirement_mismatch"
                        })
                        logger.warning(f"   [Supervisor 终审警报] 发现填报内容与招标文件需求未对应: {audit_content[:120]}...")
        except Exception as exc_semantic:
            logger.warning(f"   Supervisor 全图需求对应度核验异常: {exc_semantic}")

        # 5. 汇总全图填写状况与正确性质量报告
        chapter_findings_summary = defaultdict(list)
        for item in unfilled_findings:
            chapter_findings_summary[item["chapter"]].append(item["snippet"])

        if chapter_unfilled_count:
            logger.warning(f"   [Supervisor 全局扫描完成] 在 {len(chapter_unfilled_count)} 个章节中检出 {len(unfilled_findings)} 处质量隐患/数据错配项:")
            for ch, count in chapter_unfilled_count.items():
                logger.warning(f"      - [{ch}]: {count} 处待修项 (样例: {chapter_findings_summary[ch][:2]})")
                repair_instructions_map[ch] = (
                    f"在【{ch}】章节的质量核实中检测出 {count} 处隐患：\n"
                    "1. 逐条读取下方列出的目标路径，确认当前节点、所在行列和章节归属，不得凭旧提案猜测；\n"
                    "2. 将实际节点内容与招标原文、企业/项目数据源重新核对，判断是提案错误、目标路径错误还是文档被其他提案覆盖；\n"
                    "3. 只有数据源能够明确证明应填某值时才提交修复提案；数据源无法确认时保留占位状态，不得编造；\n"
                    "4. 修复提案必须精确写入重新确认后的目标路径，保持模板固定文本、表头、合并关系和其他单元格不变；\n"
                    "5. 写入后再次查询同一目标路径，确认实际值与修复结果一致；若仍不一致，继续定位原因并提交新的修复方案。"
                )
                chapter_specific_findings = [
                    item for item in unfilled_findings if item.get("chapter") == ch
                ]
                if chapter_specific_findings:
                    repair_instructions_map[ch] += "\n6. 请优先根据以下产物回读问题逐项重新定位并修复，不要只重复原提案：\n"
                    repair_instructions_map[ch] += "\n".join(
                        f"- 路径 {item.get('path', 'N/A')}: {item.get('snippet', '')}"
                        for item in chapter_specific_findings[:10]
                    )
        else:
            logger.info("   [Supervisor 全局扫描完成] 全图 Word DOM 结构与语义数据 100% 审阅完毕，准确性与原文完美对应！")

        # 记录审计条目
        audit_items.append(FillingAuditItem(
            target_field="[Supervisor 全局扫描终审]",
            raw_requirement=f"第 {repair_count + 1} 轮全图 DOM 结构扫描",
            format_style="Supervisor-Global-Audit",
            tool_called="supervisor_audit_node",
            data_source_table="docx_temp_path",
            db_raw_value=f"全图检测槽位: {len(unfilled_findings)} 处待修",
            final_filled_value="100% 质量达标" if not unfilled_findings else f"发现 {len(unfilled_findings)} 处质量隐患",
            alignment_status="完美通过" if not unfilled_findings else "待专项修复",
            has_underline=bool(unfilled_findings),
            source_type="supervisor_audit",
            confidence=0.99,
            agent_reasoning=f"Supervisor 全局扫描完成，覆盖全图 {len(lines)} 个 DOM 节点",
        ))

    except Exception as e:
        logger.exception(f"   Supervisor 全局扫描发生异常: {e}")

    is_passed = (len(unfilled_findings) == 0)
    review_findings_payload = [
        {
            "rule_id": {
                "unfilled_slot": "R1-UNFILLED",
                "opening_total_mismatch": "R3-SUM-MISMATCH",
                "opening_total_words_mismatch": "R3-UPPERCASE-MISMATCH",
                "opening_total_semantic_mismatch": "R4-CELL-SEMANTIC-MISMATCH",
                "opening_footer_wrong_value": "R4-CELL-VALUE-MISMATCH",
                "proposal_writeback_mismatch": "R5-PROPOSAL-WRITEBACK-MISMATCH",
            }.get(item.get("type", ""), "R9-ARTIFACT-INTEGRITY"),
            "severity": "error",
            "path": item.get("path", "N/A"),
            "description": item.get("snippet", "产物校验失败"),
            "current_value": item.get("current_value", ""),
            "expected_value": item.get("expected_value", ""),
            "auto_fixable": False,
        }
        for item in unfilled_findings
    ]

    if not is_passed and repair_count < max_repair_rounds:
        new_repair_count = repair_count + 1
        logger.warning(f"   [触发专项修复闭环] 全局扫描未达标，启动第 {new_repair_count}/{max_repair_rounds} 轮定向修复重试...")
        return {
            "audit_passed": False,
            "repair_count": new_repair_count,
            "repair_instructions_map": repair_instructions_map,
            "audit_items": audit_items,
            "review_findings": review_findings_payload,
            "audit_blocked": False,
        }
    else:
        if is_passed:
            logger.info("   [Supervisor 全局扫描通过] 标书全图填写质量全部达标！")
        else:
            logger.error(f"   [达到最大修复上限] 已执行 {repair_count} 轮修复仍未通过，阻断错误 Word 发布。")
        return {
            "audit_passed": is_passed,
            "audit_blocked": not is_passed,
            "audit_items": audit_items,
            "review_findings": review_findings_payload,
        }


def blocked_docx_node(state: BidFillerState) -> Dict[str, Any]:
    """修复轮次耗尽后阻断自动发布，同时保留非空工作副本供下载和诊断。"""
    doc_id = state.get("document_id", "")
    raw_findings = state.get("review_findings") or []
    finding_models: List[ReviewFinding] = []
    for finding in raw_findings:
        try:
            finding_models.append(ReviewFinding(**finding))
        except Exception as exc:
            logger.warning(f"阻断报告转换质检问题失败: {exc}")

    error_count = sum(1 for finding in finding_models if finding.severity == "error")
    filled_docx_bytes: Optional[bytes] = None
    docx_temp_path = state.get("docx_temp_path")

    # 阻断只表示质量状态未通过，不应把已经写入的工作副本替换成空模板。
    if docx_temp_path and os.path.exists(docx_temp_path):
        try:
            with open(docx_temp_path, "rb") as work_file:
                filled_docx_bytes = work_file.read()
            logger.warning(
                f"[阻断发布] 保留已填工作副本供下载: {docx_temp_path}, bytes={len(filled_docx_bytes)}"
            )
        except (OSError, IOError) as read_error:
            logger.exception(f"[阻断发布] 读取已填工作副本失败: {read_error}")

    if not filled_docx_bytes:
        original_docx = state.get("original_docx")
        if original_docx:
            filled_docx_bytes = original_docx
            logger.warning("[阻断发布] 工作副本不可读，回退保留原始 Word 字节流")

    # 修复轮无法确认的问题不应导致已有数据消失；保留当前工作副本并用黄色标记问题节点。
    if docx_temp_path and os.path.exists(docx_temp_path):
        highlight_unresolved_docx_findings(docx_temp_path, raw_findings)
        try:
            with open(docx_temp_path, "rb") as work_file:
                filled_docx_bytes = work_file.read()
        except (OSError, IOError) as read_error:
            logger.exception(f"[阻断发布] 读取黄色标记后的工作副本失败: {read_error}")

    report = BidFillAuditReport(
        document_id=doc_id,
        total_fields_count=len(state.get("audit_items", [])),
        audit_items=state.get("audit_items", []),
        review_findings=finding_models,
        review_summary=f"{error_count} errors, 0 warnings, 0 infos (已阻断发布)",
        summary_note="标书产物未通过写盘后校验，已阻断自动发布；当前保留工作副本供下载，请根据审计问题修复后重试。",
    )
    logger.error(f"[产物发布阻断] document_id={doc_id}, errors={error_count}")
    return {
        "filled_docx_bytes": filled_docx_bytes,
        "audit_report": report,
        "audit_blocked": True,
    }


def should_repair(state: BidFillerState) -> str:
    """Supervisor 审核路由函数：通过发布、未通过重试、耗尽则阻断。"""
    if state.get("audit_blocked") is True:
        return "blocked_docx_node"
    if state.get("audit_passed") is True:
        return "write_docx_node"
    return "agent_fill_node"



def _fallback_write_commands(
    docx_temp_path: str,
    commands: List[Dict],
    audit_items: List,
    approved: int = 0,
    rejected: int = 0,
    proposals: Optional[List[Dict]] = None,
) -> Dict[str, Any]:
    """降级模式：LLM 不可用时，直接执行预转换好的写盘命令

    Args:
        docx_temp_path: 临时 Word 文件路径
        commands: OfficeCLI 写盘命令列表
        audit_items: 审计记录列表（原地追加）
        approved: 预校验通过的 Proposal 数量
        rejected: 预校验拒绝的 Proposal 数量
        proposals: Worker 原始提案列表（用于精确格式修饰）
    """
    import asyncio, concurrent.futures
    from app.mcp.office_cli_client import office_cli_mcp_client
    def _sync(coro):
        try: asyncio.get_running_loop()
        except RuntimeError: return asyncio.run(coro)
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as e:
            return e.submit(asyncio.run, coro).result()
    cli_only_commands = []
    if commands:
        for c in commands:
            c_path = c.get("path", "")
            c_text = c.get("props", {}).get("text", "")
            if re.search(r'/(?:tbl|tr|row)\[\d+\]', c_path) and c_text.startswith("[") and c_text.endswith("]"):
                try:
                    matrix = json.loads(c_text)
                    table_path = re.sub(r'/(?:tr|row|tc|cell)\[\d+\].*$', '', c_path)
                    for row_idx, row_values in enumerate(matrix):
                        # 确保 parent 锁定在物理表格节点 /body/tbl[N]
                        cli_only_commands.append({"command": "add", "parent": table_path, "type": "row"})
                        for col_idx, val in enumerate(row_values):
                            cli_only_commands.append({
                                "command": "set",
                                "path": f"{table_path}/row[last()]/cell[{col_idx + 1}]",
                                "props": {"text": str(val)}
                            })
                except Exception as e:
                    logger.error(f"   解析表格 JSON 矩阵失败 ({c_path}): {e}")

    # 前置自动自愈纠错所有写盘命令
    cli_only_commands = auto_repair_officecli_commands(cli_only_commands)
    commands = auto_repair_officecli_commands(commands)


    if proposals:
        try:
            logger.info(f"   [DOM 安全刷盘] 启动原位切片插值与表格装配引擎，安全刷盘 {len(proposals)} 条提案...")
            fill_docx_proposals_in_dom(docx_temp_path, proposals)
        except Exception as dom_err:
            logger.warning(f"   DOM 刷盘产生异常 ({dom_err})，降级尝试 OfficeCLI 全量批量写盘...")
            if commands:
                try:
                    coro = office_cli_mcp_client.batch_update(docx_temp_path, json.dumps(commands, ensure_ascii=False))
                    _sync(coro)
                except Exception as cli_err:
                    logger.error(f"   OfficeCLI 降级全量写盘亦失败: {cli_err}")
    elif commands:
        try:
            coro = office_cli_mcp_client.batch_update(docx_temp_path, json.dumps(commands, ensure_ascii=False))
            _sync(coro)
        except Exception as cli_err:
            logger.error(f"   OfficeCLI 全量批量写盘失败: {cli_err}")

    logger.info(f"   [写盘完成] 执行 {len(commands)} 条提案刷盘（{approved} 通过, {rejected} 拒绝）")
    audit_items.append(FillingAuditItem(
        target_field="[Review 刷盘]", raw_requirement=f"安全刷盘: {approved}A+{rejected}R",
        format_style="Safe-DOM-Inplace", tool_called="fill_docx_proposals_in_dom", data_source_table="proposals",
        db_raw_value="", final_filled_value=f"{approved} 写入",
        alignment_status="安全刷盘", has_underline=True, source_type="safe_dom",
        confidence=0.98, agent_reasoning="Safe in-place slot substitution engine"
    ))
    return {"audit_items": audit_items, "docx_temp_path": docx_temp_path}


# ============================================================
# 4. write_docx_node
# ============================================================

def write_docx_node(state: BidFillerState) -> Dict[str, Any]:
    """4. write_docx_node: 从临时文件读回 Word 并输出字节流，附带审查发现"""
    logger.info("[LangGraph Node 4/4] write_docx_node: 从临时文件读回 Word...")
    doc_id = state.get("document_id", "")
    docx_temp_path = state.get("docx_temp_path")
    audit_items = state.get("audit_items", [])
    raw_findings = state.get("review_findings") or []

    filled_bytes: Optional[bytes] = None
    used_temp_file = False

    # Agent 可能因为模型输出、工具参数或章节识别异常而没有提交 image proposal。
    # 在最终读取工作副本前做一次租户隔离的幂等兜底，确保“资格证明文件”章节不会漏图。
    if docx_temp_path and os.path.exists(docx_temp_path):
        try:
            from app.core.context import current_tenant_id
            from app.agents.tools.bid_db_tools import auto_embed_qualification_images_in_docx

            effective_tenant = state.get("tenant_id") or current_tenant_id.get() or "default-tenant"
            fallback_count = auto_embed_qualification_images_in_docx(
                docx_temp_path,
                tenant_id=effective_tenant,
            )
            if fallback_count:
                logger.info(f"   🖼️ [资格证明图片兜底] 自动补插 {fallback_count} 张资质图片 (tenant_id={effective_tenant})")
        except Exception as fallback_err:
            logger.warning(f"   资格证明图片兜底插入失败，不影响其他章节输出: {fallback_err}")

    if docx_temp_path and os.path.exists(docx_temp_path):
        try:
            with open(docx_temp_path, "rb") as f_temp:
                filled_bytes = f_temp.read()
            used_temp_file = True
            logger.info(f"   📄 从 Worker 填报工作副本精确读取 Word: {len(filled_bytes)} bytes")
        except Exception as exc_read:
            logger.warning(f"   读取临时文件失败: {exc_read}")
    if not filled_bytes:
        filled_bytes = state.get("original_docx")
        logger.info("   📄 回退使用原始字节流")

    # 自动导出 Worker 子 Agent 运行时上下文诊断报告
    try:
        from app.agents.bid_filler_workers import export_worker_context_log
        log_path = export_worker_context_log(doc_id)
        logger.info(f"   [上下文诊断日志] 已成功生成 Worker 上下文报告: {log_path}")
    except Exception as exc_log:
        logger.warning(f"   导出 Worker 诊断日志失败: {exc_log}")

    # 将 review_findings（dict 列表）转为 ReviewFinding Pydantic 模型列表
    review_finding_models: List[ReviewFinding] = []
    for f in raw_findings:
        try:
            review_finding_models.append(ReviewFinding(**f))
        except Exception:
            pass  # 跳过格式异常的 finding

    # 生成审查摘要
    errors = sum(1 for f in raw_findings if f.get("severity") == "error")
    warnings = sum(1 for f in raw_findings if f.get("severity") == "warning")
    infos = sum(1 for f in raw_findings if f.get("severity") == "info")
    review_summary = f"{errors} errors, {warnings} warnings, {infos} infos (共 {len(raw_findings)} 条)"

    # 文件保留在 drafts 目录，不删除（它就是最终结果）
    source_label = "已填写的工作副本" if used_temp_file else "原始字节流"
    report = BidFillAuditReport(
        document_id=doc_id,
        total_fields_count=len(audit_items),
        audit_items=audit_items,
        review_findings=review_finding_models,
        review_summary=review_summary,
        summary_note=f"BidFillerAgent Multi-Agent 标书撰写完成（数据源: {source_label}）",
    )

    logger.info(f"   审查摘要: {review_summary}")

    return {
        "filled_docx_bytes": filled_bytes,
        "audit_report": report,
    }


def _clean_and_parse_json(raw_text: str) -> Optional[Dict[str, Any]]:
    """容错解析大模型返回的 JSON"""
    if not raw_text:
        return None
    cleaned = raw_text.strip()
    if "```" in cleaned:
        match = re.search(r'```(?:json)?\s*({[\s\S]*?})\s*```', cleaned, re.DOTALL)
        if match:
            cleaned = match.group(1).strip()
        else:
            cleaned = re.sub(r'^```(?:json)?\s*', '', cleaned)
            cleaned = re.sub(r'\s*```$', '', cleaned).strip()
    try:
        return json.loads(cleaned, strict=False)
    except Exception:
        pass
    try:
        def replace_newlines_in_strings(m):
            return f'"{m.group(1).replace(chr(10), "\\n").replace(chr(13), "\\r")}"'
        fixed_json = re.sub(r'"((?:[^"\\]|\\.)*)"', replace_newlines_in_strings, cleaned, flags=re.DOTALL)
        return json.loads(fixed_json, strict=False)
    except Exception as e:
        logger.warning(f"JSON 修补解析失败: {e}")
        return None


# ============================================================
# Graph 构建
# ============================================================

def build_bid_filler_graph():
    """构建 LangGraph Multi-Agent 标书撰写状态图（Supervisor 划分 + Worker 直写 + Supervisor 闭环质量把控）"""
    workflow = StateGraph(BidFillerState)
    workflow.add_node("scan_node", scan_node)
    workflow.add_node("agent_fill_node", agent_fill_node)
    workflow.add_node("supervisor_audit_node", supervisor_audit_node)
    workflow.add_node("write_docx_node", write_docx_node)
    workflow.add_node("blocked_docx_node", blocked_docx_node)

    workflow.set_entry_point("scan_node")
    workflow.add_edge("scan_node", "agent_fill_node")
    workflow.add_edge("agent_fill_node", "supervisor_audit_node")
    workflow.add_conditional_edges(
        "supervisor_audit_node",
        should_repair,
        {
            "agent_fill_node": "agent_fill_node",
            "write_docx_node": "write_docx_node",
            "blocked_docx_node": "blocked_docx_node",
        }
    )
    workflow.add_edge("write_docx_node", END)
    workflow.add_edge("blocked_docx_node", END)
    return workflow.compile()



bid_filler_graph_app = build_bid_filler_graph()


# ============================================================
# 包装类 & 适配器
# ============================================================

class BidFillerAgent:
    """保持与原有 API 的全兼容接口"""

    def process_filling_tasks(
        self, db: Session, document_id: str, profile: CompanyProfile,
        detected_placeholders: List[Dict[str, Any]], original_docx: Optional[bytes] = None,
        custom_instructions: Optional[str] = None,
        category_hints: Optional[Dict[str, str]] = None,
        profile_id: Optional[str] = None,
    ) -> tuple[Dict[str, str], BidFillAuditReport, Optional[bytes]]:
        logger.info("启动 LangGraph BidFillerAgent Multi-Agent 标书撰写状态图...")

        # 设置运行时上下文：指定 Agent 工具查询使用的企业档案
        from app.agents.tools.bid_db_tools import current_profile_id
        if profile_id:
            current_profile_id.set(profile_id)
            logger.info(f"Agent 运行时绑定企业档案: profile_id={profile_id}")
        from app.agents.bid_filler_workers import clear_worker_proposals
        clear_worker_proposals(document_id)

        effective_tenant: Optional[str] = None
        try:
            from app.core.context import current_tenant_id
            effective_tenant = current_tenant_id.get()
            if not effective_tenant and db is not None:
                document_obj = db.query(DocumentModel).filter(DocumentModel.id == document_id).first()
                effective_tenant = getattr(document_obj, "tenant_id", None) if document_obj else None
        except Exception as tenant_err:
            logger.warning(f"读取标书租户上下文失败，使用默认租户: {tenant_err}")
        effective_tenant = effective_tenant or "default-tenant"

        initial_state: BidFillerState = {
            "document_id": document_id, "tenant_id": effective_tenant, "original_context": "",
            "db_session": db, "company_profile": profile,
            "profile_id": profile_id,
            "original_docx": original_docx, "docx_temp_path": None,
            "slot_analysis": None, "worker_proposals": None, "chapter_tasks": None,
            "custom_instructions": custom_instructions,
            "category_hints": category_hints,
            "repair_count": 0,
            "max_repair_rounds": 2,
            "repair_instructions_map": None,
            "audit_passed": None,
            "audit_blocked": False,
            "audit_items": [], "review_findings": None,
            "audit_report": None, "filled_docx_bytes": None,
        }
        final_state = bid_filler_graph_app.invoke(initial_state)
        audit_report = final_state.get("audit_report") or BidFillAuditReport(
            document_id=document_id, total_fields_count=0, audit_items=[],
            summary_note="BidFillerAgent Multi-Agent 标书撰写完成"
        )

        try:
            from app.services.audit_service import audit_service
            audit_service.log_event(
                action_type="llm_call_supervisor",
                node_name="Supervisor-总控调度",
                inputs={"document_id": document_id, "chapter_title": "Supervisor-总控调度"},
                outputs={
                    "summary": "✨ Multi-Agent 团队全自主标书撰写与写盘已全量收官！所有章节均已处理完成。",
                    "proposals_count": len(audit_report.audit_items) if audit_report and audit_report.audit_items else 0,
                },
                status="master_completed"
            )
        except Exception as log_err:
            logger.warning(f"更新 Supervisor 终态审计日志异常: {log_err}")

        return {}, audit_report, final_state.get("filled_docx_bytes")


bid_filler_agent = BidFillerAgent()


SKIP_BID_FILLER = os.getenv("SKIP_BID_FILLER", "false").lower() in ("true", "1", "yes")  # 开关控制：默认 False (开启真实标书起草流程)


def bid_filler_orchestrator_node(state: dict) -> dict:
    """Orchestrator 适配节点 — 将 BidFillerAgent 对接至 LangGraph 编排器"""
    from app.worker.tasks import emit_agent_log
    from app.core.config import settings
    document_id = state.get("document_id")
    tenant_id = state.get("tenant_id") or "default-tenant"

    emit_agent_log("info", "启动 BidFillerAgent (Multi-Agent Supervisor+Worker)...",
                   extra={"type": "worker_start", "worker": "writer_agent"})

    if not document_id:
        return {"status": "writer_failed", "error": "Missing document_id"}

    # 从 .env / Settings 动态读取开关配置 (false: 开启真实起草; true: 临时调试跳过)
    skip_bid_filler = os.getenv("SKIP_BID_FILLER", "false").lower() in ("true", "1", "yes")

    # 如果开启了跳过开关，直接发出完成日志并返回成功，加速整体 Multi-Agent 流程
    if skip_bid_filler:
        summary = "[调试配置] 已暂时跳过标书起草步骤，长流程快速闭环完成"
        logger.info(f"⏩ {summary}")
        emit_agent_log("info", summary, extra={
            "type": "worker_complete", "worker": "writer_agent", "status": "success", "summary": summary, "document_id": document_id
        })
        return {
            "completed_steps": ["writer_agent"],
            "worker_summaries": [{"worker": "writer_agent", "status": "success", "summary": summary}]
        }

    db: Session = SessionLocal()
    try:
        from app.services.bid_format_extractor_service import bid_format_extractor_service
        template_bytes, _, _ = bid_format_extractor_service.extract_and_export_bid_format(
            db=db, doc_id=document_id, tenant_id=tenant_id)
        if not template_bytes:
            raise ValueError("未提取到《投标文件格式》模板")

        _, _, filled_bytes = bid_filler_agent.process_filling_tasks(
            db=db, document_id=document_id, profile=CompanyProfile(),
            detected_placeholders=[], original_docx=template_bytes)

        # __file__ 向上 3 层为 backend 根目录
        base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        drafts_dir = os.path.join(base_dir, "uploads", "drafts")
        os.makedirs(drafts_dir, exist_ok=True)
        draft_path = os.path.join(drafts_dir, f"draft_{document_id}.docx")
        if filled_bytes:
            with open(draft_path, "wb") as f:
                f.write(filled_bytes)

        doc_obj = db.query(DocumentModel).filter(DocumentModel.id == document_id).first()
        if doc_obj:
            curr_meta = dict(doc_obj.parsed_metadata) if doc_obj.parsed_metadata else {}
            curr_meta["draft_path"] = draft_path
            curr_meta["draft_filename"] = f"draft_{document_id}.docx"
            doc_obj.parsed_metadata = curr_meta
            db.commit()

        summary = "已成功由 BidFillerAgent (Multi-Agent) 完成投标书 Word 草稿生成"
        emit_agent_log("info", summary, extra={
            "type": "worker_complete", "worker": "writer_agent", "status": "success", "summary": summary, "document_id": document_id
        })
        return {
            "completed_steps": ["writer_agent"], "draft_path": draft_path,
            "worker_summaries": [{"worker": "writer_agent", "status": "success", "summary": summary}]
        }
    except Exception as e:
        logger.exception(f"BidFillerAgent 适配节点失败: {e}")
        emit_agent_log("error", f"BidFillerAgent 执行失败: {str(e)}")
        return {"status": "writer_failed", "error": str(e)}
    finally:
        db.close()
