"""
投标文件格式提取与切片服务 (Bid Format Extractor Service)

负责定位并提取招标文件中的“投标文件格式/组成”章节：
1. 原生 DOCX 模式：对原始 Word (.docx/.doc) 文件进行底层 DOM 结构切片，保留 100% 原始格式与表格，并统一修改文字为黑色。
2. LLM 结构化重建模式：针对 PDF 格式，通过 LLM 识别定位并结合 DocxExporterService 重建规范 Word。
"""

import os
import re
import io
import copy
from typing import Tuple, Optional, List
from loguru import logger
from sqlalchemy.orm import Session
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor, Pt, Inches

from app.db.crud.document import document_crud
from app.services.extractor_service import ExtractorService
from app.services.docx_exporter_service import docx_exporter_service
from app.services.llm_service import LLMService
from app.schemas.bid_generator import (
    BidFormatStructure,
    BidFormatSection,
    ContentTypeEnum
)


class BidFormatExtractorService:
    """
    投标文件格式提取与导出核心业务服务
    """

    def __init__(self):
        self.extractor_service = ExtractorService()
        self.llm_service = LLMService()

        # 匹配“投标文件格式”大章标题正则表达式探照灯 (兼容 Markdown # / ** / ## 标记与内部空格)
        self.chapter_start_patterns = [
            re.compile(r'^[#\s\*]*第\s*[一二三四五六七八九十\d]+\s*[章篇部分卷节][\s\:\、\.\*]*(投标文件格式|应答文件格式|响应文件格式|投标文件组成|应答文件组成|格式及附件|投标文件格式要求|投标格式|响应格式)'),
            re.compile(r'^[#\s\*]*(投标文件格式|应答文件格式|响应文件格式|投标文件格式及附件|投标格式及要求)[\s\*]*$'),
            re.compile(r'^[#\s\*]*附\s*[件录][\s\:\、\.\*]*(投标文件格式|应答文件格式|响应文件格式|投标文件组成|投标格式)'),
        ]

        # 匹配下一个大章（用于判定“投标文件格式”章节的终止界限，兼容 Markdown # / ** 标记）
        self.chapter_next_patterns = [
            re.compile(r'^[#\s\*]*第\s*[一二三四五六七八九十\d]+\s*[章篇部分卷]'),
        ]

        # 章节正文特征仅用于区分目录与正文，不依赖固定章节编号或固定“三册”结构。
        self.format_body_markers = (
            "投标函",
            "应答函",
            "报价函",
            "授权委托",
            "法定代表人",
            "偏离表",
            "报价明细",
            "报价表",
            "承诺书",
            "资格审查",
            "商务响应",
            "技术响应",
            "格式附件",
        )

    def extract_and_export_bid_format(
        self, 
        db: Session, 
        doc_id: str,
        user_id: Optional[str] = None,
        tenant_id: Optional[str] = None
    ) -> Tuple[bytes, str, str]:
        """
        全流程处理方法：根据 doc_id 获取文件类型并执行切片提取与 Word 导出。

        :param db: 数据库 Session
        :param doc_id: 文档 ID
        :param user_id: 用户 ID (可选)
        :param tenant_id: 租户 ID (可选)
        :return: (docx_bytes, filename, extraction_mode)
        """
        # 1. 检索文档记录
        if user_id and tenant_id:
            doc_obj = document_crud.get_document_by_id(db, doc_id, user_id=user_id, tenant_id=tenant_id)
        else:
            doc_obj = document_crud.get_document_by_id_system(db, doc_id)

        if not doc_obj or not doc_obj.file_path or not os.path.exists(doc_obj.file_path):
            logger.error(f"提取投标文件格式失败：找不到文档记录或原文件不存在 (doc_id={doc_id})")
            raise FileNotFoundError("找不到原始招标文件记录或存储路径")

        file_path = doc_obj.file_path
        effective_tenant_id = tenant_id or getattr(doc_obj, "tenant_id", None)
        file_ext = os.path.splitext(file_path)[1].lower()
        base_name = os.path.splitext(os.path.basename(doc_obj.filename))[0]
        export_filename = f"{base_name}_投标文件格式模板.docx"

        # 2. 判断文件类型，优先使用原生 DOCX 切片模式
        if file_ext in ['.docx', '.doc']:
            try:
                target_docx_path = file_path
                if file_ext == '.doc':
                    logger.info(f"原生文件为 .doc，尝试使用 LibreOffice 转换为 .docx: {file_path}")
                    target_docx_path = self.extractor_service.convert_doc_to_docx(file_path)

                docx_bytes = self._slice_docx_natively(target_docx_path)
                if docx_bytes:
                    logger.info(f"原生 Word 切片成功！文件大小: {len(docx_bytes)} 字节")
                    return docx_bytes, export_filename, "native_docx"
            except Exception as e:
                logger.warning(f"原生 Word 切片未命中或执行异常，回退至 LLM 重建模式: {str(e)}")

        # 3. 回退模式 / PDF 模式：利用 ExtractorService 与 LLM 重建标准 Word
        logger.info(f"使用 LLM 结构化提取模式处理文件: {file_path}")
        docx_bytes, mode = self._extract_with_llm_and_rebuild(
            db,
            doc_obj,
            tenant_id=effective_tenant_id,
        )
        return docx_bytes, export_filename, mode

    def _is_toc_line(self, text: str, element=None) -> bool:
        """
        严密判定某个段落是否为目录页/导引线/目录项（TOC Line）
        """
        if not text:
            return False
        clean_txt = text.strip()

        # 1. 检查 XML 节点中是否包含 TOC / Hyperlink 目录特征
        if element is not None:
            try:
                xml_str = element.xml if hasattr(element, 'xml') else ""
                if 'w:hyperlink' in xml_str and '_Toc' in xml_str:
                    return True
                if 'w:pStyle' in xml_str and ('TOC' in xml_str or 'toc' in xml_str or '目录' in xml_str):
                    return True
                if 'w:fldSimple' in xml_str and 'TOC' in xml_str:
                    return True
                if 'w:instrText' in xml_str and 'TOC' in xml_str:
                    return True
            except Exception:
                pass

        # 2. 匹配目录导引线及页码 (如 ".......... -3-"、".......... 55"、".......... -55-"、"…… 40"、"...... 55页")
        if re.search(r'[\.….┈\-_]{2,}\s*[-–—\s]*\d+[-–—\s\.\)\]页]*$', clean_txt):
            return True

        # 3. 匹配包含多连点/制表符且末尾包含页码数字（兼容各种页码修饰符如 -55- 或 55页）
        if re.search(r'[\.….┈\-_]{2,}', clean_txt) and re.search(r'\d+[-–—\s\.\)\]页]*$', clean_txt):
            return True

        # 4. 匹配制表符或多连点后跟着页码数字（如 "第六章 投标文件格式  -55-"）
        if re.search(r'[\s\t]+[-–—\s]*\d+[-–—\s\.\)\]页]*$', clean_txt) and len(clean_txt) < 100:
            if re.search(r'[\.….┈\-_]', clean_txt) or '\t' in text:
                return True

        # 5. 纯目录卷标/目录标题行 (如 "第一卷"、"第二卷"、"第三卷"、"目  录")
        if re.search(r'^\s*(?:第[一二三四五六七八九十\d]+卷|目\s*录|Table\s*of\s*Contents)\s*$', clean_txt):
            return True

        return False

    def _is_real_next_main_chapter(self, text: str, element=None) -> bool:
        """
        判断某个段落是否为真正的下一个招标大章（如 第七章 评标办法 / 第七章 合同条款），
        避免误将“第六章 投标文件格式”内部的格式子项（如“格式七 授权书”、“附件七 承诺函”）判定为终点。
        """
        if not text:
            return False
        clean_txt = text.strip()

        # 如果属于目录行，直接排除
        if self._is_toc_line(clean_txt, element):
            return False

        # 匹配大章主标题模式，如 第七章、第八章 (兼容 Markdown # / ** / ## 标记与空格)
        main_chapter_pattern = re.compile(r'^[#\s\*]*第\s*[一二三四五六七八九十\d]+\s*[章篇部分卷]')
        if not main_chapter_pattern.search(clean_txt):
            return False

        # 排除包含格式附件关键词的内部子标题（如 格式、附件、表、样张、承诺、声明、证明、清单、函、明细、协议）
        internal_format_keywords = ['格式', '附件', '表', '样张', '承诺', '声明', '证明', '清单', '函', '明细', '协议', '响应', '授权']
        if any(kw in clean_txt for kw in internal_format_keywords):
            return False

        # 排除以子序号开头的条目，如 一、二、三、(一)、(1)
        if re.search(r'^\s*[\(（]?[一二三四五六七八九十\d]+[\)）\.\、]', clean_txt):
            return False

        return True

    def _slice_docx_natively(self, docx_path: str) -> Optional[bytes]:
        """
        核心方法：原生 Word DOM 节点裁剪算法。
        严密排除目录（TOC）与格式附件子标题干扰，精准提取“投标文件格式”正文全量内容与表格附件。
        """
        if not os.path.exists(docx_path):
            return None

        doc = Document(docx_path)
        body = doc._body._element

        start_index = -1
        end_index = -1

        # 遍历文本当中的所有子元素（包含 Paragraph 和 Table）
        children = list(body)
        
        # 建立段落文本与索引映射
        element_meta = []
        for idx, child in enumerate(children):
            tag_name = child.tag.split('}')[-1]
            text = ""
            if tag_name == 'p':
                text = "".join(child.itertext()).strip()
            element_meta.append({'index': idx, 'tag': tag_name, 'text': text, 'element': child})

        # 寻找起始大章（必须排除目录 TOC 行）
        candidate_starts = []
        for item in element_meta:
            txt = item['text']
            elem = item['element']
            if any(pat.search(txt) for pat in self.chapter_start_patterns):
                if not self._is_toc_line(txt, elem):
                    candidate_starts.append(item['index'])
                    logger.info(f"发现正文候选起始位置: line {item['index']} -> '{txt[:40]}'")

        if candidate_starts:
            # 锁定正文起始位置
            start_index = candidate_starts[0]
            logger.info(f"锁定投标文件格式正文起始位置: line {start_index}")
        else:
            logger.warning("未能在原生 Word 中匹配到非目录的'投标文件格式'正文起始位置")
            return None

        # 寻找结束界限（从 start_index + 1 开始，下一个真正的独立大章标题，或者文件末尾）
        for item in element_meta[start_index + 1:]:
            txt = item['text']
            elem = item['element']
            if self._is_real_next_main_chapter(txt, elem):
                end_index = item['index']
                logger.info(f"定位到下一个独立大章终止位置: line {end_index} -> '{txt[:40]}'")
                break

        if end_index == -1:
            end_index = len(children)
            logger.info("投标文件格式章节无后续主大章，全量延伸提取至文件末尾")

        # 提取切片范围元素
        target_elements = children[start_index:end_index]
        if not target_elements:
            return None

        # 【核心防护】：自动剪除切片头部残留的目录行、卷标行或导引线节点 (如 "第一卷"、"第一章 招标公告...-3-" 等)
        while target_elements:
            first_txt = "".join(target_elements[0].itertext()).strip()
            if self._is_toc_line(first_txt, target_elements[0]):
                logger.info(f"   ✂️ 自动剪除切片头部残留目录节点: '{first_txt[:50]}'")
                target_elements.pop(0)
            else:
                break

        if not target_elements:
            return None

        logger.info(f"成功裁剪投标文件格式正文切片！包含 {len(target_elements)} 个 DOM 元素节点")

        # 清空 body 中的非切片节点
        for child in list(body):
            if child not in target_elements and child.tag.endswith(('p', 'tbl')):
                body.remove(child)

        # 全量修改所有文字 Run 为黑色字体 (RGB 0,0,0)
        black_color = RGBColor(0, 0, 0)
        for p in doc.paragraphs:
            for run in p.runs:
                run.font.color.rgb = black_color

        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.color.rgb = black_color

        # 输出为字节流
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()

    def _extract_with_llm_and_rebuild(
        self,
        db: Session,
        doc_obj,
        tenant_id: Optional[str] = None,
    ) -> Tuple[bytes, str]:
        """
        LLM 提取模式：结合 ExtractorService 与 LLM 提取文本，并用 DocxExporterService 渲染 Word。

        :param tenant_id: 调用方显式传入的租户 ID；未传入时回退使用文档所属租户。
        :return: (docx_bytes, actual_mode) 其中 actual_mode 为 "llm_rebuilt" 或 "fallback_template"
        """
        # 显式保留租户上下文，避免线程池调用时 ContextVar 丢失而回退到全局模型配置。
        effective_tenant_id = tenant_id or getattr(doc_obj, "tenant_id", None)

        # 读取文本
        md_file_path = (
            doc_obj.parsed_metadata.get("md_file_path", "")
            if doc_obj and doc_obj.parsed_metadata
            else ""
        )
        doc_text = ""
        if md_file_path and os.path.exists(md_file_path):
            with open(md_file_path, "r", encoding="utf-8") as f:
                doc_text = f.read()
            logger.info(f"成功读取 Markdown 缓存文件: {md_file_path} (文本总长度: {len(doc_text)} 字符)")
        else:
            chunks = document_crud.get_document_chunks(db, doc_obj.id)
            doc_text = "\n\n".join([c.content for c in chunks]) if chunks else ""
            logger.info(f"从数据库切片提取文本完成 (共 {len(chunks) if chunks else 0} 个切片, 文本总长度: {len(doc_text)} 字符)")

        if not doc_text.strip():
            logger.warning("⚠️ [投标文件格式提取] 文档未提取到任何有效文本，使用备用基础模板数据构建")
            structure = self._build_fallback_structure(doc_obj.filename)
            return docx_exporter_service.export_bid_format_to_docx_bytes(structure), "fallback_template"

        # 正则快速定位文本范围
        target_text = self._slice_text_by_keywords(doc_text)
        if not target_text.strip():
            logger.warning("⚠️ [投标文件格式提取] 未截取到目标章节，跳过 LLM 调用并使用托底模板")
            structure = self._build_fallback_structure(doc_obj.filename)
            return docx_exporter_service.export_bid_format_to_docx_bytes(structure), "fallback_template"

        # 构建 Prompt 引导 LLM 输出结构化数据
        prompt = f"""你是一名资深招投标专家。请分析以下招标文件中的“投标文件格式/响应格式”部分文本，严格依据原文提取出完整的格式附件目录与样张模版。

【最高指令】:
1. 100% 忠实于【待分析文本】原文提取，严禁凭常识臆造或捏造原文不存在的附件名称、字段或内容。
2. 提取文本中出现的全部格式附件标题（如各类格式、附件、声明、承诺、样张等，严格以原文实际标题为准）。
3. 原文中的表格（无论以 Markdown 表格还是 HTML <table> 形式出现）必须完整保留其行列表格结构（转换为标准 Markdown 表格输出），原文中的填空下划线 `______` 必须完整保留。
4. 必须将原文中每个格式附件的完整正文、填空要素和表格内容原原本本提取并放入 `body_markdown`，严禁输出“原文未提供样张”等概括性文字。
5. 系统已将【待分析文本】严格截取在原文识别出的目标格式章节标题至下一独立大章之间，章节编号和标题以原文为准，禁止引入评审办法、收费标准、资格要求或合同附件等章节内容。

【待分析文本】:
{target_text[:40000]}

【提取要求与结构定义】:
`content_type` 字段可选值：'form_table' (表格样张/填报明细)、'text_template' (公文/承诺书/证明模板)、'checklist' (清单/目录)、'other' (其他格式附件)。
请返回合法 JSON 格式对象（只输出纯 JSON，不要包含任何前导或后置解释说明），严格符合以下数据结构定义：
{{
  "document_title": "{doc_obj.filename} - 投标文件格式模板",
  "source_chapter_name": "投标文件格式",
  "sections": [
     {{
        "section_title": "原文中的格式附件标题",
        "content_type": "text_template",
        "body_markdown": "原文中的模板正文内容或 Markdown 表格内容（完整保留填空下划线 ______）",
        "placeholders": ["从该格式中提炼出的待填空字段名"]
     }}
  ]
}}
"""
        try:
            if self.llm_service.is_configured_for_tenant(effective_tenant_id):
                logger.info(f"🚀 [投标文件格式提取] 正在调用 LLM 结构化提取招标文件格式 (待分析切片长度: {len(target_text[:40000])} 字符)...")
                parsed_json = self.llm_service.generate_structured_json(
                    prompt,
                    temperature=0.1,
                    tenant_id=effective_tenant_id,
                )
                # 若大模型直接返回了 sections 数组，自动包装为字典对象
                if isinstance(parsed_json, list):
                    parsed_json = {
                        "document_title": f"{doc_obj.filename} - 投标文件格式模板",
                        "source_chapter_name": "投标文件格式",
                        "sections": parsed_json
                    }
                structure = BidFormatStructure(**parsed_json)
                if not structure.sections:
                    logger.warning("⚠️ [投标文件格式提取] LLM 提取出的 sections 为空，自动降级至托底基础结构")
                    structure = self._build_fallback_structure(doc_obj.filename)
                    return docx_exporter_service.export_bid_format_to_docx_bytes(structure), "fallback_template"
                else:
                    section_names = [s.section_title for s in structure.sections]
                    logger.info(f"✅ [投标文件格式提取] LLM 结构化提取成功！共提取出 {len(structure.sections)} 个格式附件: {section_names}")
                    return docx_exporter_service.export_bid_format_to_docx_bytes(structure), "llm_rebuilt"
            else:
                logger.warning("⚠️ [投标文件格式提取] LLM 服务未配置，使用托底基础结构构建")
                structure = self._build_fallback_structure(doc_obj.filename)
                return docx_exporter_service.export_bid_format_to_docx_bytes(structure), "fallback_template"
        except Exception as e:
            logger.exception(f"❌ [投标文件格式提取] LLM 提取或解析过程发生异常: {str(e)}，正在触发托底基础结构降级构建")
            structure = self._build_fallback_structure(doc_obj.filename)
            return docx_exporter_service.export_bid_format_to_docx_bytes(structure), "fallback_template"

    def _slice_text_by_keywords(self, full_text: str) -> str:
        """
        在纯文本中截取“投标文件格式/应答文件格式”章节。

        先从所有同名标题中选择最像正文的候选项，再截取至下一个独立大章。
        未定位到目标章节时返回空字符串，禁止把整份招标文件交给 LLM 猜测。
        """
        if not full_text or not full_text.strip():
            logger.warning("⚠️ [投标文件格式提取] 输入文本为空，无法定位目标章节")
            return ""

        lines = full_text.splitlines()
        candidate_indices = [
            index
            for index, line in enumerate(lines)
            if any(pattern.search(line.strip()) for pattern in self.chapter_start_patterns)
        ]
        if not candidate_indices:
            logger.warning("⚠️ [投标文件格式提取] 未定位到“投标文件格式/应答文件格式”标题，跳过 LLM 调用")
            return ""

        # 优先使用目录中出现的章节身份（如“第九章”），再到正文查找同一章节，章节编号由原文动态决定。
        toc_chapter_keys = self._find_toc_target_chapter_keys(lines)
        if toc_chapter_keys:
            toc_matched_candidates = [
                index
                for index in candidate_indices
                if not self._is_toc_line(lines[index])
                and self._chapter_identity_key(lines[index]) in toc_chapter_keys
            ]
            if toc_matched_candidates:
                candidate_indices = toc_matched_candidates
                logger.info(f"🔍 [投标文件格式提取] 根据目录动态锁定目标章节: {sorted(toc_chapter_keys)}")

        start_idx = max(candidate_indices, key=lambda index: self._score_text_chapter_candidate(lines, index))
        if self._is_toc_line(lines[start_idx]):
            logger.warning(f"⚠️ [投标文件格式提取] 目标标题仅命中目录行: '{lines[start_idx].strip()}'，跳过 LLM 调用")
            return ""

        end_idx = len(lines)
        for index in range(start_idx + 1, len(lines)):
            line_str = lines[index].strip()
            if self._is_real_next_main_chapter(line_str):
                end_idx = index
                logger.info(f"🔍 [投标文件格式提取] 定位到下一个独立大章终止行 (第 {index + 1} 行): '{line_str[:50]}'")
                break

        slice_lines = lines[start_idx:end_idx]
        while slice_lines and self._is_toc_line(slice_lines[0]):
            slice_lines.pop(0)
        logger.info(f"🔍 [投标文件格式提取] 成功定位正文起始行 (第 {start_idx + 1} 行), 切片行数: {len(slice_lines)}")
        return "\n".join(slice_lines)

    def _score_text_chapter_candidate(self, lines: List[str], index: int) -> tuple[int, int]:
        """为章节标题候选项评分，优先选择包含格式正文特征的正文而非目录。"""
        line = lines[index].strip()
        context_before = "\n".join(lines[max(0, index - 8):index + 1])
        context_after = "\n".join(lines[index + 1:index + 36])
        score = 0
        if self._is_toc_line(line):
            score -= 100
        if re.search(r'目录|contents', context_before, re.IGNORECASE):
            score -= 30
        marker_hits = sum(marker in context_after for marker in self.format_body_markers)
        score += min(marker_hits, 4) * 8
        # 同分时取靠后的候选，避免目录中的同名标题遮蔽正文标题。
        return score, index

    def _chapter_identity_key(self, text: str) -> str:
        """提取章节编号作为动态匹配键，不绑定具体的章号。"""
        match = re.match(r'^[#\s\*]*(第\s*[一二三四五六七八九十\d]+\s*[章篇部分卷])', text.strip())
        if not match:
            return ""
        return re.sub(r'\s+', '', match.group(1))

    def _find_toc_target_chapter_keys(self, lines: List[str]) -> set[str]:
        """从目录行中提取目标格式章节身份，供正文定位动态复用。"""
        chapter_keys = {
            self._chapter_identity_key(line)
            for line in lines
            if self._is_toc_line(line)
            and any(pattern.search(line.strip()) for pattern in self.chapter_start_patterns)
        }
        return {key for key in chapter_keys if key}

    def _build_fallback_structure(self, filename: str) -> BidFormatStructure:
        """
        当未配置 LLM 或提取异常时的托底基础模板结构
        """
        base_title = os.path.splitext(filename)[0]
        return BidFormatStructure(
            document_title=f"{base_title} - 投标文件格式",
            source_chapter_name="应答文件格式",
            sections=[
                BidFormatSection(
                    section_title="附件一：投标函",
                    content_type=ContentTypeEnum.TEXT_TEMPLATE,
                    body_markdown="致：_____________________（招标人名称）\n\n1. 我方已仔细研究了_____________________（项目名称及招标编号）招标文件的全部内容，遵照招标文件要求，我方愿以人民币（大写）____________________（￥_________元）的投标总价，按合同约定实施和完成各项工作。\n\n2. 我方承诺本投标文件有效期为开标之日起_______天。\n\n投标人名称（盖章）：_____________________\n法定代表人或授权委托人（签字/盖章）：_____________________\n日期：______年___月___日",
                    placeholders=["招标人名称", "项目名称", "投标总价", "有效期天数"]
                ),
                BidFormatSection(
                    section_title="附件二：法定代表人授权委托书",
                    content_type=ContentTypeEnum.TEXT_TEMPLATE,
                    body_markdown="本授权声明：正式授权_________________（代理人姓名）为我方合法代理人，以我方名义签署、澄清、说明_____________________项目（招标编号：_____________）的投标文件，并处理一切与该项目投标有关的事宜。\n\n委托期限：自本授权书签署之日起至投标有效期届满止。\n\n法定代表人（签字/盖章）：_____________________\n身份证号码：_____________________\n授权委托人（签字/盖章）：_____________________\n身份证号码：_____________________\n投标人名称（盖大公章）：_____________________",
                    placeholders=["代理人姓名", "项目名称", "身份证号码"]
                ),
                BidFormatSection(
                    section_title="附件三：开标一览表（报价汇总表）",
                    content_type=ContentTypeEnum.FORM_TABLE,
                    body_markdown="| 项目名称 | 投标总价（元） | 工期/交货期 | 质量标准 | 备注 |\n| :--- | :--- | :--- | :--- | :--- |\n| _____________________ | ￥________________ | ______日历天 | 合格 | 详见分项报价表 |",
                    placeholders=["投标总价", "工期"]
                )
            ]
        )


# 单例初始化
bid_format_extractor_service = BidFormatExtractorService()
