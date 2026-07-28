"""
DocxTestFillerService - 结合 docx 技能规范的 Word 修改与填报测试服务

功能说明：
1. 提供基于 docx 技能的最佳实践：
   - 占位符原位替换与下划线格式 (w:u) 智能继承；
   - 表格双重宽度设置 (Table Width 与 Cell Width，单位 DXA，防止不同排版软件错位)；
   - 表格单元格内边距 (w:tcMar) 与清爽浅色背景 (w:shd) 配置；
   - 统一黑体/纯黑色字体 (RGBColor(0,0,0)) 渲染。
2. 包含模版生成与改写填报功能，用于单元测试与开发验证。
"""

import io
import re
from typing import Dict, Any, Optional
from loguru import logger
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls


class DocxTestFillerService:
    """ 结合 docx 技能规范的 Word 填报与格式修改测试服务类 """

    # 1 英寸 = 1440 DXA，标准 A4 打印区域宽度约为 9360 DXA (6.5 英寸)
    PAGE_CONTENT_WIDTH_DXA = 9360
    BLACK_COLOR = RGBColor(0, 0, 0)

    def __init__(self):
        logger.info("初始化 DocxTestFillerService 服务")

    def _check_underline(self, run) -> bool:
        """ 检查 Run 元素是否存在下划线（通过 python-docx 属性或原始 XML 属性判断） """
        if not run:
            return False

        try:
            if bool(run.underline) or bool(run.font.underline):
                return True
        except Exception:
            pass

        xml_str = run._element.xml if hasattr(run, '_element') else ''
        return 'w:u' in xml_str or 'underline' in xml_str.lower()

    def set_cell_margins(
        self,
        cell,
        top_dxa: int = 120,
        bottom_dxa: int = 120,
        left_dxa: int = 180,
        right_dxa: int = 180
    ) -> None:
        """
        设置单元格内边距 (w:tcMar)，遵循 docx 技能中对单元格间距控制的建议
        """
        if not cell:
            return

        try:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = OxmlElement('w:tcMar')
            for name, val in [('top', top_dxa), ('bottom', bottom_dxa), ('left', left_dxa), ('right', right_dxa)]:
                node = OxmlElement(f'w:{name}')
                node.set(qn('w:w'), str(val))
                node.set(qn('w:type'), 'dxa')
                tc_mar.append(node)
            tc_pr.append(tc_mar)
        except Exception as e:
            logger.error(f"设置单元格内边距失败: {str(e)}")

    def set_cell_shading(self, cell, fill_hex: str = "F2F4F7") -> None:
        """
        设置单元格背景颜色 (w:shd)，遵循 docx 技能规范（避免使用纯黑或黑白翻转）
        """
        if not cell:
            return

        try:
            shd_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
            cell._tc.get_or_add_tcPr().append(shd_elm)
        except Exception as e:
            logger.error(f"设置单元格背景颜色失败: {str(e)}")

    def set_table_borders(self, table, border_color: str = "CCCCCC") -> None:
        """
        为表格添加标准的浅灰色细边框 (w:tblBorders)
        """
        if not table:
            return

        try:
            tbl_pr = table._tbl.tblPr
            borders_elm = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
                f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
                f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
                f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
                f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
                f'</w:tblBorders>'
            )
            tbl_pr.append(borders_elm)
        except Exception as e:
            logger.error(f"设置表格边框失败: {str(e)}")

    def create_sample_docx(self) -> bytes:
        """
        生成一个测试用的标准招投标 Word 模版（包含带有下划线的占位段落与表格）
        """
        doc = Document()

        # 1. 标题
        h1 = doc.add_paragraph()
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = h1.add_run("投标书响应格式（测试模版）")
        r1.font.size = Pt(18)
        r1.font.bold = True
        r1.font.color.rgb = self.BLACK_COLOR

        # 2. 带有下划线占位符的正文段落
        p1 = doc.add_paragraph()
        r_label1 = p1.add_run("项目名称：")
        r_label1.font.color.rgb = self.BLACK_COLOR

        r_val1 = p1.add_run("____________________")
        r_val1.font.underline = True
        r_val1.font.color.rgb = self.BLACK_COLOR

        p2 = doc.add_paragraph()
        r_label2 = p2.add_run("投标人名称：")
        r_label2.font.color.rgb = self.BLACK_COLOR

        r_val2 = p2.add_run("____________________")
        r_val2.font.underline = True
        r_val2.font.color.rgb = self.BLACK_COLOR

        # 3. 数据表格（遵循 docx 技能的 DXA 双宽度要求）
        table = doc.add_table(rows=3, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.set_table_borders(table, "D3D3D3")

        # 比例：20%, 50%, 30% -> DXA: 1872, 4680, 2808 (总和 9360)
        col_widths_dxa = [1872, 4680, 2808]

        headers = ["序号", "标的名称", "响应报价(元)"]
        hdr_cells = table.rows[0].cells
        for idx, text in enumerate(headers):
            hdr_cells[idx].width = col_widths_dxa[idx]
            self.set_cell_shading(hdr_cells[idx], "EAECEF")
            self.set_cell_margins(hdr_cells[idx], 140, 140, 180, 180)
            p = hdr_cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.font.bold = True
            r.font.color.rgb = self.BLACK_COLOR

        # 数据行 1
        r1_cells = table.rows[1].cells
        for idx, val in enumerate(["1", "[未填标的A]", "______"]):
            r1_cells[idx].width = col_widths_dxa[idx]
            self.set_cell_margins(r1_cells[idx], 100, 100, 150, 150)
            p = r1_cells[idx].paragraphs[0]
            r = p.add_run(val)
            r.font.color.rgb = self.BLACK_COLOR
            if "___" in val:
                r.font.underline = True

        # 数据行 2
        r2_cells = table.rows[2].cells
        for idx, val in enumerate(["2", "[未填标的B]", "______"]):
            r2_cells[idx].width = col_widths_dxa[idx]
            self.set_cell_margins(r2_cells[idx], 100, 100, 150, 150)
            p = r2_cells[idx].paragraphs[0]
            r = p.add_run(val)
            r.font.color.rgb = self.BLACK_COLOR
            if "___" in val:
                r.font.underline = True

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    def fill_and_modify_docx(
        self,
        docx_bytes: bytes,
        fill_data: Dict[str, str],
        table_updates: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        修改与填报 Word 文档：
        1. 替换段落中的占位符，完美保留原始下划线 (w:u)；
        2. 修改并更新表格数据，套用 DXA 双宽度与间距规范；
        3. 全量文本确保为纯黑色 (RGBColor(0,0,0))。
        """
        if not docx_bytes:
            logger.error("传入的 docx_bytes 为空")
            raise ValueError("传入的 docx_bytes 不能为空")

        logger.info(f"开始执行 Word 修改与填报测试，数据条目数: {len(fill_data)}")
        doc = Document(io.BytesIO(docx_bytes))

        # 1. 段落填报与下划线样式保留
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if not text:
                continue

            for key, val in fill_data.items():
                if key in text:
                    # 检查原本的 runs 是否带有下划线
                    has_underline = False
                    for r in paragraph.runs:
                        if self._check_underline(r) or '___' in r.text:
                            has_underline = True
                            break

                    # 清空现有 runs，重新重构文本
                    paragraph.text = ""
                    r_prefix = paragraph.add_run(f"{key}：")
                    r_prefix.font.color.rgb = self.BLACK_COLOR

                    r_val = paragraph.add_run(str(val))
                    r_val.font.color.rgb = self.BLACK_COLOR
                    if has_underline:
                        r_val.font.underline = True
                    logger.debug(f"替换正文字段: [{key}] -> [{val}] (下划线: {has_underline})")

        # 2. 表格数据改写与双宽度规范
        if table_updates and len(doc.tables) > 0:
            target_table = doc.tables[0]
            col_widths_dxa = [1872, 4680, 2808]

            for row_idx, row in enumerate(target_table.rows):
                if row_idx == 0:  # 表头不修改
                    continue

                row_key = f"row_{row_idx}"
                if row_key in table_updates:
                    row_data = table_updates[row_key]  # [序号, 标的名称, 报价]
                    for col_idx, cell in enumerate(row.cells):
                        if col_idx < len(row_data):
                            new_val = str(row_data[col_idx])
                            # 校验原单元格下划线
                            has_underline = any(self._check_underline(r) or '___' in r.text for p in cell.paragraphs for r in p.runs)

                            cell.text = ""
                            cell.width = col_widths_dxa[col_idx]
                            self.set_cell_margins(cell, 100, 100, 150, 150)

                            p = cell.paragraphs[0]
                            r = p.add_run(new_val)
                            r.font.color.rgb = self.BLACK_COLOR
                            if has_underline:
                                r.font.underline = True

                            logger.debug(f"修改表格第 {row_idx} 行第 {col_idx} 列为: [{new_val}]")

        # 3. 导出生成的 Word 字节流
        out_stream = io.BytesIO()
        doc.save(out_stream)
        filled_bytes = out_stream.getvalue()
        logger.info(f"Word 修改完成，输出字节大小: {len(filled_bytes)} bytes")
        return filled_bytes


docx_test_filler_service = DocxTestFillerService()
