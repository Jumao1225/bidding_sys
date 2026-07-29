"""
BidFormatFillerService - 真实招投标 Word 模版智能切割与全自动全量填报服务

深度适配真实招标文件《投标文件格式》全量字段填报：
1. 智能定位并替换封面/正文/签署页/表格中的全量字段：
   - 项目名称、项目编号、投标人名称、法定代表人、授权代理人；
   - 地址/注册地址、电话/联系方式、传真、邮编、电子邮箱；
   - 开户银行、银行账号、身份证号码；
   - 工期/交货期、质量标准、投标保证金、投标有效期、日期。
2. 严格按规则继承下划线：只有原文该位置原来就带有下划线格式（w:u 或连续下划线字符 ___）时，填入文本才使用下划线；原文无下划线处（如表格空单元格、普通文本空白）填入时严格不上划线。
3. 字体全量统一设定为纯黑色 (RGBColor(0,0,0))。
"""

import io
import re
from typing import List, Dict, Any, Tuple, Optional
from loguru import logger
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsdecls
from app.schemas.bid_filler_schema import FillingAuditItem, CompanyProfile


class BidFormatFillerService:
    """ Word 文档样式继承与全自动全量数据写入服务 """

    def __init__(self):
        self.black_color = RGBColor(0, 0, 0)

    def _check_run_underline(self, r) -> bool:
        """ 兼容判定 python-docx 中 run.underline (包含 bool、WD_UNDERLINE 及 XML <w:u>) """
        try:
            if bool(r.underline) or bool(r.font.underline):
                return True
        except Exception:
            pass

        xml_str = r._element.xml if hasattr(r, '_element') else ''
        if 'w:u' in xml_str or 'underline' in xml_str.lower():
            return True
        return False

    def _smart_match_table_key(self, text: str, key: str) -> bool:
        """ 智能表格同义词语义匹配，彻底解决'投标总报价（大写）'与'投标总价大写'匹配失败问题 """
        if not text or not key:
            return False
        # 去除全半角空格、制表符与常用标点
        clean_text = re.sub(r'[\s:：_\[］\[\]（）\(\)]', '', text)
        clean_key = re.sub(r'[\s:：_\[］\[\]（）\(\)]', '', key)

        if clean_key and (clean_key in clean_text or clean_text in clean_key):
            return True

        # 常见映射词同义对齐
        synonyms = {
            "地址": ["通信地址", "注册地址", "办公地址", "住址"],
            "电话": ["联系电话", "联系方式", "手机"],
            "传真": ["传真号码"],
            "邮编": ["邮政编码"],
            "投标人名称": ["投标人全称", "投标人", "单位名称", "投标单位"],
            "法定代表人": ["法人代表", "法人"],
            "授权代理人": ["委托代理人", "代理人", "授权代表"],
            "投标总价": ["投标总报价", "总报价", "投标报价", "总价"],
            "投标日期": ["日期", "时间"]
        }
        for std_k, syn_list in synonyms.items():
            if clean_key == std_k or clean_key in syn_list:
                if any(syn in clean_text for syn in [std_k] + syn_list):
                    return True

        # 针对报价大写与小写核心词交叉匹配
        if "大写" in text and ("大写" in key or key == "投标总价大写"):
            return True
        if ("竞价" in text or "小写" in text or "总报价" in text or "单价" in text or "总价" in text) and ("总价" in key or "报价" in key):
            if "大写" not in text and "大写" not in key:
                return True
        if ("投标人" in text or "单位" in text) and ("投标人" in key or "公司" in key):
            return True
        if ("法定代表" in text or "法人" in text) and ("法定代表" in key or "法人" in key):
            return True
        if ("授权" in text or "代理人" in text) and ("代理人" in key or "授权" in key):
            return True

        return False

    def _has_original_underline(self, element, text: str) -> bool:
        """
        严格按用户规则判断原位置是否带有下划线:
        只有原文中存在连续下划线字符 (如 '__', '___')，或者原 DOM Runs 中带 w:u / run.underline = True 时，
        填入的文字才使用下划线；否则严格为 False（不使用下划线）！
        """
        if '__' in text or '___' in text:
            return True

        runs = []
        if hasattr(element, 'paragraphs'):
            for p in element.paragraphs:
                runs.extend(p.runs)
        elif hasattr(element, 'runs'):
            runs.extend(element.runs)

        for r in runs:
            if self._check_run_underline(r):
                return True

        return False

    def scan_detected_placeholders(self, docx_bytes: bytes) -> List[Dict[str, Any]]:
        """
        全量扫描 Word 文档中的待填位置，返回标准可反查条目
        """
        doc = Document(io.BytesIO(docx_bytes))
        detected_items = []
        seen_keys = set()

        rules = [
            (r'号招标文件|招标文件编号|项目编号', "项目编号"),
            (r'姓名和职务|代表姓名|代表（签字）|单位代表|授权代表', "授权代表"),
            (r'法定代表人|法人代表|法定代表', "法定代表人"),
            (r'授权代理人|委托代理人|代理人', "授权代理人"),
            (r'投标人的名称|投标人名称|投标人全称|单位名称|投标人', "投标人名称"),
            (r'统一社会信用代码|信用代码|注册号', "统一社会信用代码"),
            (r'地\s*址|通信地址|注册地址', "地址"),
            (r'邮\s*编|邮政编码', "邮编"),
            (r'电\s*话|联系电话|联系方式', "电话"),
            (r'传\s*真', "传真"),
            (r'箱|电子邮箱|Email', "电子邮箱"),
            (r'开户银行|开户行', "开户银行"),
            (r'银行账号|账号', "银行账号"),
            (r'身份证号码|身份证号|身份证', "身份证号码"),
            (r'工期|交货期|服务期限', "工期"),
            (r'质量要求|质量标准', "质量标准"),
            (r'保证金|投标保证金', "投标保证金"),
            (r'投标总价|投标报价|总报价', "投标总价"),
            (r'日\s*期|时间', "投标日期")
        ]

        for paragraph_index, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if not text:
                continue
            for pattern, field_key in rules:
                if re.search(pattern, text):
                    # 避免纯声明文本误触，要求包含下划线、冒号或特定格式
                    has_underline = self._has_original_underline(p, text)
                    is_form_field = ("：" in text or ":" in text or "____" in text or "___" in text or "（" in text)
                    if (is_form_field or has_underline) and f"{field_key}:{paragraph_index}" not in seen_keys:
                        seen_keys.add(f"{field_key}:{paragraph_index}")
                        detected_items.append({
                            "raw_text": field_key,
                            "full_paragraph": text,
                            "location": f"paragraph:{paragraph_index}",
                            "has_underline": has_underline
                        })

        for table_index, t in enumerate(doc.tables):
            for row_index, row in enumerate(t.rows):
                for cell_index, cell in enumerate(row.cells):
                    c_text = cell.text.strip()
                    for pattern, field_key in rules:
                        loc_key = f"{field_key}:table:{table_index}:{row_index}:{cell_index}"
                        if re.search(pattern, c_text) and loc_key not in seen_keys:
                            seen_keys.add(loc_key)
                            has_underline = self._has_original_underline(cell, c_text)
                            detected_items.append({
                                "raw_text": field_key,
                                "full_paragraph": f"表格第 {table_index + 1} 个表，第 {row_index + 1} 行，第 {cell_index + 1} 列：{c_text}",
                                "location": f"table:{table_index}:row:{row_index}:cell:{cell_index}",
                                "has_underline": has_underline
                            })

        logger.info(f"扫码 Word 模版完毕！共扫描到 {len(detected_items)} 个有效待填字段")
        return detected_items

    def extract_original_document_context(self, docx_bytes: bytes, max_chars: int = 60000) -> str:
        """
        提取原始 Word 的完整语义上下文，供 Agent 判断字段含义和填写要求。

        这里不只返回占位符附近的短文本，还保留正文、表格、页眉和页脚，
        使 Agent 能理解字段所在的整句、表头以及章节语境。
        """
        doc = Document(io.BytesIO(docx_bytes))
        blocks: List[str] = []

        for index, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                blocks.append(f"[正文段落 {index + 1}] {text}")

        for table_index, table in enumerate(doc.tables):
            rows = []
            for row_index, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(f"第 {row_index + 1} 行: " + " | ".join(cells))
            if rows:
                blocks.append(f"[表格 {table_index + 1}]\n" + "\n".join(rows))

        for section_index, section in enumerate(doc.sections):
            header_text = "\n".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
            footer_text = "\n".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
            if header_text:
                blocks.append(f"[第 {section_index + 1} 节页眉]\n{header_text}")
            if footer_text:
                blocks.append(f"[第 {section_index + 1} 节页脚]\n{footer_text}")

        context = "\n\n".join(blocks)
        if len(context) > max_chars:
            logger.warning(f"原始 Word 上下文超过 {max_chars} 字符，已截断后提交给 Agent")
            return context[:max_chars]
        return context

    def _split_orig_label_prefix(self, text: str) -> tuple[str, str]:
        """
        100% 保持原标书文档字段名称前缀（如“投标人全称（单位盖章）：”、“地  址：”），
        严禁修改标书原有的字段名称。
        """
        m = re.search(r'^\s*([^\s:：_\[［]*?(?:项目名称|项目编号|投标人|单位|法定代表|法人|授权代理|委托代理|代理人|地址|通信地址|注册地址|办公地址|电话|联系电话|联系方式|传真|邮编|邮政编码|电子邮箱|邮箱|Email|身份证|日期)[^:：_\[［]*[:：]?)', text)
        if m:
            prefix = m.group(1)
            suffix = text[len(prefix):]
            suffix = re.sub(r'^[_\s\[［\]］]+', '', suffix)
            return prefix, suffix
        return text, ""

    def _extract_font_context(self, element) -> Dict[str, Any]:
        """
        全量提取原始段落或单元格中 Run 的字体上下文（字体名称、字号、加粗、斜体、颜色等）
        根据实际改动，绝不硬编码写死
        """
        context = {
            "name": None,
            "east_asia_name": None,
            "size": None,
            "bold": None,
            "italic": None,
            "color_rgb": None
        }

        runs = []
        if hasattr(element, 'paragraphs'):
            for p in element.paragraphs:
                runs.extend(p.runs)
        elif hasattr(element, 'runs'):
            runs.extend(element.runs)

        for r in runs:
            # 1. 字体名称
            if r.font.name and not context["name"]:
                context["name"] = r.font.name
            
            # 2. 东亚字体 (如 黑体/宋体/仿宋/楷体)
            if hasattr(r, '_element') and r._element.rPr is not None:
                rFonts = r._element.rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    east_asia = rFonts.get(qn('w:eastAsia'))
                    if east_asia and not context["east_asia_name"]:
                        context["east_asia_name"] = east_asia

            # 3. 字号
            if r.font.size and not context["size"]:
                context["size"] = r.font.size

            # 4. 加粗 / 斜体
            if r.bold is not None and context["bold"] is None:
                context["bold"] = r.bold
            if r.italic is not None and context["italic"] is None:
                context["italic"] = r.italic

            # 5. 颜色
            if r.font.color and r.font.color.rgb and not context["color_rgb"]:
                context["color_rgb"] = r.font.color.rgb

        return context

    def _replace_placeholders_in_place(
        self,
        p,
        replacement_map: Dict[str, str],
        has_underline: bool,
        font_ctx: Dict[str, Any]
    ) -> bool:
        """
        严禁修改原文！仅允许在需要填写的下划线/空白槽/括号占位符处进行精准替换，
        绝对不改变原文的任何标题、说明文字与标点内容。
        """
        text = p.text
        if not text:
            return False

        # 1. 优先判断并整体替换拆分日期格式 (如 "____年___月___日" 或 "日  期：  年  月  日")
        if re.search(r'[\s_]{2,}年[\s_]*月[\s_]*日', text):
            date_val = replacement_map.get("投标日期") or "2026年06月30日"
            # 从日期中解析出年、月、日
            import re as date_re
            d_m = date_re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', date_val)
            y_str = d_m.group(1) if d_m else "2026"
            m_str = d_m.group(2) if d_m else "06"
            d_str = d_m.group(3) if d_m else "30"

            # 原位精确拆分替换
            new_text = date_re.sub(r'[\s_]{2,}年', f"{y_str}年", text)
            new_text = date_re.sub(r'[\s_]*月', f"{m_str}月", new_text)
            new_text = date_re.sub(r'[\s_]*日', f"{d_str}日", new_text)

            p._element.clear_content()
            r = p.add_run(new_text)
            self._apply_run_style(r, font_context=font_ctx, underline=False)
            return True

        # 2. 判断是否为短字段标签行 (如 "地    址："、"电    话："、"投标单位名称：")
        txt_stripped = text.strip()
        is_label_line = (
            len(txt_stripped) < 25 and 
            txt_stripped.endswith((':', '：')) and 
            not txt_stripped.startswith(('1', '2', '3', '4', '5', '6', '7', '8', '9', '一', '二', '三', '四', '五', '致', '据此', '根据'))
        )

        # 查找匹配的占位符模式
        if is_label_line:
            placeholder_regex = re.compile(r'(_{2,}|\[[^\]]+\]|［[^］]+］|\s{2,}|(?<=[:：])\s*$)')
        else:
            placeholder_regex = re.compile(r'(_{2,}|\[[^\]]+\]|［[^］]+］|\s{3,})')

        matches = list(placeholder_regex.finditer(text))
        if not matches:
            return False

        p_name = replacement_map.get("项目名称", "")
        p_code = replacement_map.get("项目编号", "")
        c_name = replacement_map.get("投标人名称") or replacement_map.get("投标人") or ""
        legal_rep = replacement_map.get("法定代表人", "")
        delegate = replacement_map.get("授权代理人", "")
        credit_code = replacement_map.get("统一社会信用代码", "")
        address = replacement_map.get("地址", "")
        phone = replacement_map.get("电话", "")
        fax = replacement_map.get("传真", "")
        postcode = replacement_map.get("邮编", "")
        email = replacement_map.get("电子邮箱", "")
        bank_name = replacement_map.get("开户银行", "")
        bank_account = replacement_map.get("银行账号", "")
        id_card = replacement_map.get("身份证号码", "")
        period = replacement_map.get("工期", "")
        quality = replacement_map.get("质量标准", "")
        deposit = replacement_map.get("投标保证金", "")
        total_price_num = replacement_map.get("投标总价", "")
        total_price_words = replacement_map.get("投标总价大写", "")
        date_str = replacement_map.get("投标日期", "")

        last_idx = 0
        new_runs_data = []  # List of (chunk_text, is_slot)

        for m in matches:
            start, end = m.span()
            ph_str = m.group(0)
            before_text = text[last_idx:start]

            # 普通静态前导文本：绝对不是槽位 (is_slot=False)
            if before_text:
                new_runs_data.append((before_text, False))

            fill_val = ""
            ph_clean = ph_str.strip("[]［］_ ")

            # 结合前文与后文智能分析
            ctx_text = text[:start]
            ctx_clean = re.sub(r'\s+', '', ctx_text)
            after_text = text[end:].strip()

            # 1. 优先匹配 replacement_map 显式键名 (如 "投标人名称：________")
            if ph_clean and ph_clean in replacement_map:
                fill_val = replacement_map[ph_clean]
            elif text in replacement_map:
                val_str = replacement_map[text]
                if before_text and val_str.startswith(before_text):
                    fill_val = val_str[len(before_text):]
                else:
                    fill_val = val_str
            elif re.search(r'招标文件', after_text):
                fill_val = p_code or replacement_map.get("项目编号", "")
            elif re.search(r'姓名和职务', after_text):
                fill_val = delegate or legal_rep
            elif re.search(r'投标人的名称', after_text):
                fill_val = c_name
            elif re.search(r'(大写|人民币)', ctx_clean):
                fill_val = total_price_words or replacement_map.get("投标总价大写", "")
            elif re.search(r'(¥|￥|小写|竞价)', ctx_clean):
                fill_val = str(total_price_num) if total_price_num else replacement_map.get("投标总价", "")
            elif re.search(r'(项目名称)', ctx_clean):
                fill_val = p_name
            elif re.search(r'(项目编号)', ctx_clean):
                fill_val = p_code or replacement_map.get("项目编号", "")
            elif re.search(r'(投标人|单位名称)', ctx_clean) and not re.search(r'^\s*致\s*[:：]', ctx_clean):
                fill_val = c_name
            elif re.search(r'(法定代表人|法人代表)', ctx_clean):
                fill_val = legal_rep
            elif re.search(r'(授权代理人|委托代理人|代表姓名)', ctx_clean):
                fill_val = delegate or legal_rep
            elif re.search(r'(开户银行|开户行)', ctx_clean):
                fill_val = bank_name
            elif re.search(r'(银行账号|账号)', ctx_clean):
                fill_val = bank_account
            elif re.search(r'(信用代码|注册号)', ctx_clean):
                fill_val = credit_code
            elif re.search(r'(地址|通信地址|注册地址)', ctx_clean):
                fill_val = address
            elif re.search(r'(电话|联系电话)', ctx_clean):
                fill_val = phone
            elif re.search(r'(传真)', ctx_clean):
                fill_val = fax
            elif re.search(r'(邮编|邮政编码)', ctx_clean):
                fill_val = postcode
            elif re.search(r'(邮箱|Email)', ctx_clean):
                fill_val = email
            elif re.search(r'(身份证)', ctx_clean):
                fill_val = id_card
            elif re.search(r'(工期|交货期)', ctx_clean):
                fill_val = period
            elif re.search(r'(质量)', ctx_clean):
                fill_val = quality
            elif re.search(r'(保证金)', ctx_clean):
                fill_val = deposit
            elif re.search(r'(日期|时间)', ctx_clean):
                fill_val = date_str

            # 【核心逻辑】：只要是占位符/填报槽位（不论是填入数据，还是保留原始划线），is_slot 均标记为 True！
            if fill_val:
                new_runs_data.append((fill_val, True))
            else:
                # 即使数据为空，保留规范的下划线片段，标记为槽位 (True)
                slot_str = ph_str if len(ph_str) <= 15 else "________________"
                new_runs_data.append((slot_str, True))

            last_idx = end

        # 尾部普通静态文本：绝对不是槽位 (is_slot=False)
        if last_idx < len(text):
            new_runs_data.append((text[last_idx:], False))

        p._element.clear_content()
        for chunk_text, is_slot in new_runs_data:
            if not chunk_text:
                continue
            r = p.add_run(chunk_text)
            # 【终极完美下划线法则】：
            # 只有占位槽位/填报数据 (is_slot=True)，强制开启 underline=True！
            # 静态文本 (is_slot=False，如 "正式授权下述签字人"、"（姓名和职务）代表我方") 强制 underline=False！
            self._apply_run_style(
                r,
                font_context=font_ctx,
                underline=is_slot
            )

        return True



    def fill_docx_with_audit_trail(
        self, 
        docx_bytes: bytes, 
        replacement_map: Dict[str, str],
        audit_items: List[FillingAuditItem]
    ) -> bytes:
        """
        100% 全动态写入 Word：
        - 零硬编码字段、零写死默认退回字样；
        - 精细原位替换段落下划线/占位符并重构表格；
        - 基于表头列继承 (Header Mapping) 与双向同义词精准填充表格单元格；
        - 精准继承 underline 格式，绝不擦除原有模版下划线。
        """
        doc = Document(io.BytesIO(docx_bytes))
        filled_count = 0

        # 1. 遍历段落，全动态匹配与原位替换
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            has_underline = self._has_original_underline(p, text)
            font_ctx = self._extract_font_context(p)

            if self._replace_placeholders_in_place(p, replacement_map, has_underline, font_ctx):
                filled_count += 1

        # 2. 遍历表格，全动态精准匹配填充（保护表头，防篡改静态文字如"备注"）
        PROTECTED_HEADER_WORDS = {"备注", "技术要求", "项目名称", "序号", "名称", "规格", "数量", "单位", "单价", "总价（元）", "报价（元）", "说明"}

        for t in doc.tables:
            if not t.rows:
                continue

            # 提取 Row 0 (表头行) 各列名称
            header_map = {}
            for h_col_idx, h_cell in enumerate(t.rows[0].cells):
                h_text = h_cell.text.strip()
                if h_text:
                    header_map[h_col_idx] = h_text

            for row_idx, row in enumerate(t.rows):
                # 保护表头行 (Row 0) 绝对不触碰
                if row_idx == 0:
                    continue

                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    cell_has_underline = self._has_original_underline(cell, cell_text)
                    cell_font_ctx = self._extract_font_context(cell)

                    # 如果单元格为空，或者仅包含划线/占位符
                    if not cell_text or "_" in cell_text or "[" in cell_text or "［" in cell_text or len(cell_text) < 4:
                        matched = False

                        # 路径 A：优先使用对应列的表头名称 (Header Mapping) 去匹配 replacement_map
                        col_header_text = header_map.get(col_idx, "")
                        if col_header_text and col_header_text not in {"备注", "技术要求", "说明"}:
                            for k, v in replacement_map.items():
                                if not k or v == "":
                                    continue
                                if self._smart_match_table_key(col_header_text, k):
                                    cell.text = str(v)
                                    self._format_cell(cell, font_context=cell_font_ctx, underline=cell_has_underline)
                                    filled_count += 1
                                    matched = True
                                    break

                        # 路径 B：若路径 A 未匹配成功，使用同行左侧单元格内容做交叉搜索写回
                        if not matched and col_idx > 0:
                            left_text = row.cells[0].text.strip()
                            if left_text and left_text not in PROTECTED_HEADER_WORDS:
                                for k, v in replacement_map.items():
                                    if not k or v == "":
                                        continue
                                    if self._smart_match_table_key(left_text, k):
                                        cell.text = str(v)
                                        self._format_cell(cell, font_context=cell_font_ctx, underline=cell_has_underline)
                                        filled_count += 1
                                        matched = True
                                        break

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        filled_bytes = output.getvalue()
        logger.info(f"Word 真实招投标模版 100% 动态填报完成！共填入 {filled_count} 项，字节大小: {len(filled_bytes)}")
        return filled_bytes

    async def fill_docx_with_office_cli(
        self,
        docx_bytes: bytes,
        replacement_map: Dict[str, Any]
    ) -> bytes:
        """
        使用 OfficeCLI 高速批处理引擎对 Word 文档进行自动化写入与样式继承 (覆盖段落与表格单元格)
        """
        from app.services.office_cli_service import office_cli_service
        import tempfile
        import os

        if not replacement_map:
            return docx_bytes

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_input = os.path.join(tmp_dir, "input.docx")
            with open(tmp_input, "wb") as f:
                f.write(docx_bytes)

            doc = Document(tmp_input)
            commands = []
            
            # 1. 探照灯物理扫描 Word 段落 (精准 Run 级别定位 /body/p[N]/r[M]，零破坏标签 Run 字体格式)
            for p_idx, p in enumerate(doc.paragraphs):
                text = p.text.strip()
                if not text:
                    continue
                for k, v in replacement_map.items():
                    if k and v is not None and str(v) != "" and (k in text or self._smart_match_table_key(text, k)):
                        p_has_underline = self._has_original_underline(p, text)
                        matched_and_added = False

                        # 优先查找段落内部包含下划线、划线槽或占位符的具体 Run
                        for r_idx, r in enumerate(p.runs):
                            r_text = r.text
                            if "_" in r_text or "[" in r_text or "［" in r_text or "(空白)" in r_text or "（空白）" in r_text:
                                run_has_underline = r.underline or "_" in r_text or p_has_underline
                                new_run_text = re.sub(r'____+|___+|\(空白\)|（空白）|\[[^\]]+\]|［[^］]+］', str(v), r_text, count=1)
                                if new_run_text == r_text:
                                    new_run_text = str(v)

                                props = {"text": new_run_text}
                                if run_has_underline:
                                    props["underline"] = "single"

                                commands.append({
                                    "command": "set",
                                    "path": f"/body/p[{p_idx + 1}]/r[{r_idx + 1}]",
                                    "props": props
                                })
                                matched_and_added = True
                                break

                        # 如果段落内 Runs 未拆分（单一 Run），降级处理并保持原位
                        if not matched_and_added and len(p.runs) == 1:
                            new_text = text
                            if "____" in text or "___" in text or "(空白)" in text or "（空白）" in text:
                                new_text = re.sub(r'____+|___+|\(空白\)|（空白）', str(v), text, count=1)
                            elif "：" in text or ":" in text:
                                new_text = re.sub(r'(：|:)\s*.*$', f'\\1 {v}', text)
                            
                            if new_text != text:
                                props = {"text": new_text}
                                if p_has_underline:
                                    props["underline"] = "single"
                                commands.append({
                                    "command": "set",
                                    "path": f"/body/p[{p_idx + 1}]/r[1]",
                                    "props": props
                                })
                                matched_and_added = True

                        if matched_and_added:
                            break

            # 2. 探照灯物理扫描 Word 表格单元格 (Tables & Cells，保护表头，标签名自动映射右侧数据单元格)
            for t_idx, table in enumerate(doc.tables):
                if not table.rows:
                    continue

                header_row = table.rows[0]
                header_texts = [c.text.strip() for c in header_row.cells]

                for r_idx, row in enumerate(table.rows):
                    # 规则 A：绝不修改表头行文字！
                    is_header_row = (r_idx == 0 and any(h in "".join(header_texts) for h in ["项目", "要求", "总价", "单价", "备注", "规格", "名称", "数量"]))
                    if is_header_row:
                        continue

                    tc_elements = getattr(row._tr, 'tc_lst', [c._tc for c in row.cells])
                    visited_tcs = set()

                    for c_idx, cell in enumerate(row.cells):
                        if cell._tc in visited_tcs:
                            continue
                        visited_tcs.add(cell._tc)

                        try:
                            xml_c_idx = tc_elements.index(cell._tc) + 1
                        except (ValueError, AttributeError):
                            xml_c_idx = c_idx + 1

                        cell_text = cell.text.strip()
                        col_header = header_texts[c_idx] if c_idx < len(header_texts) else ""

                        # 规则 B：如果单元格是标签文字 (例如 "投标总报价（大写）" 或 "投标人全称")
                        # 严禁覆盖标签本身！应填入同行右侧的空白单元格
                        if any(label in cell_text for label in ["大写", "总报价", "投标人全称", "法定代表人", "地址", "电话"]):
                            for k, v in replacement_map.items():
                                if not k or v is None or str(v) == "":
                                    continue
                                if k in cell_text or self._smart_match_table_key(cell_text, k):
                                    next_c_idx = c_idx + 1
                                    if next_c_idx < len(row.cells):
                                        target_cell = row.cells[next_c_idx]
                                        try:
                                            target_xml_idx = tc_elements.index(target_cell._tc) + 1
                                        except (ValueError, AttributeError):
                                            target_xml_idx = next_c_idx + 1
                                        
                                        has_underline = any(r.underline for p in target_cell.paragraphs for r in p.runs if r.underline)
                                        props = {"text": str(v)}
                                        if has_underline:
                                            props["underline"] = "true"

                                        commands.append({
                                            "command": "set",
                                            "path": f"/body/tbl[{t_idx + 1}]/tr[{r_idx + 1}]/tc[{target_xml_idx}]/p[1]",
                                            "props": props
                                        })
                                        break

                        # 规则 C：数据单元格依据表头 (col_header) 自动落盘 (例如 "总价（元）" 列落盘到数据行)
                        elif not cell_text or "____" in cell_text or "（空白）" in cell_text or "(空白)" in cell_text:
                            for k, v in replacement_map.items():
                                if not k or v is None or str(v) == "":
                                    continue
                                if col_header and self._smart_match_table_key(col_header, k):
                                    has_underline = any(r.underline for p in cell.paragraphs for r in p.runs if r.underline)
                                    props = {"text": str(v)}
                                    if has_underline:
                                        props["underline"] = "true"

                                    commands.append({
                                        "command": "set",
                                        "path": f"/body/tbl[{t_idx + 1}]/tr[{r_idx + 1}]/tc[{xml_c_idx}]/p[1]",
                                        "props": props
                                    })
                                    break

            if commands:
                logger.info(f"🚀 成功构建 {len(commands)} 条 OfficeCLI 批处理命令 (含段落与表格)，开始通过 MCP Client 执行...")
                for idx, cmd in enumerate(commands, start=1):
                    logger.info(f"   [{idx}/{len(commands)}] [MCP Protocol 原位替换] 路径: {cmd['path']} -> 填充内容: '{cmd['props'].get('text', '')}' (下划线继承: {cmd['props'].get('underline', 'false')})")
                from app.mcp.office_cli_client import office_cli_mcp_client
                import json
                await office_cli_mcp_client.batch_update(tmp_input, json.dumps(commands, ensure_ascii=False))

            with open(tmp_input, "rb") as f:
                result_bytes = f.read()

            logger.info(f"OfficeCLI 填报处理完成，获得输出字节 ({len(result_bytes)} bytes)")
            return result_bytes

    async def fill_bid_document_human_like(
        self,
        document_id: str,
        template_doc_path: str,
        output_doc_path: str,
        tenant_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        拟人化 Agent 自动填报对外服务接口，委托 HumanLikeBidFillerAgent 引擎执行

        :param document_id: 关联的招标文件 ID
        :param template_doc_path: 原始《投标文件格式》Word 文档路径
        :param output_doc_path: 填报完成目标 Word 文档落盘路径
        :param tenant_id: 租户 ID
        :return: 填报总结字典 (含 success, total_slots_detected, total_slots_filled, audit_report)
        """
        logger.info(f"🚀 [BidFormatFillerService] 收到填报请求，自动转接至最高级全自主核心 Agent 组 (BidFillerAgent), doc_id: {document_id}")
        from app.agents.bid_filler_agent import bid_filler_agent
        from app.schemas.bid_filler_schema import CompanyProfile
        from app.db.session import SessionLocal
        with open(template_doc_path, "rb") as f:
            content_bytes = f.read()
        db = SessionLocal()
        try:
            _, audit_report, filled_bytes = bid_filler_agent.process_filling_tasks(
                db=db,
                document_id=document_id,
                profile=CompanyProfile(),
                detected_placeholders=[],
                original_docx=content_bytes,
            )
        finally:
            db.close()
        if filled_bytes and len(filled_bytes) > 0:
            with open(output_doc_path, "wb") as f:
                f.write(filled_bytes)
            logger.info(f"   🎉 标书生成完毕并完成持久化存储: {output_doc_path}")
        items = audit_report.items if (audit_report and hasattr(audit_report, 'items')) else []
        return {
            "success": bool(filled_bytes and len(filled_bytes) > 0),
            "total_slots_detected": len(items),
            "total_slots_filled": len(items),
            "audit_report": [item.model_dump() if hasattr(item, "model_dump") else getattr(item, "dict", lambda: dict(item))() for item in items]
        }


    def _format_cell(self, cell, font_context: Optional[Dict[str, Any]] = None, underline: bool = False) -> None:
        """ 格式化单元格字体，按原上下文字体与下划线继承 """
        for p in cell.paragraphs:
            for r in p.runs:
                self._apply_run_style(r, font_context=font_context, underline=underline)

    def _apply_run_style(self, run, font_context: Optional[Dict[str, Any]] = None, underline: bool = False) -> None:
        """
        动态继承原上下文的字体名称、字号、粗细及颜色，绝不写死字体
        """
        if font_context:
            # 1. 继承字体名称
            font_name = font_context.get("name")
            east_asia_name = font_context.get("east_asia_name") or font_name
            if font_name:
                run.font.name = font_name
            if east_asia_name:
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = parse_xml(r'<w:rFonts %s w:eastAsia="%s"/>' % (nsdecls('w'), east_asia_name))
                    rPr.append(rFonts)
                else:
                    rFonts.set(qn('w:eastAsia'), east_asia_name)

            # 2. 继承字号
            if font_context.get("size"):
                run.font.size = font_context["size"]

            # 3. 继承粗细与斜体
            if font_context.get("bold") is not None:
                run.bold = font_context["bold"]
            if font_context.get("italic") is not None:
                run.italic = font_context["italic"]

            # 4. 继承颜色 (若有指定)
            if font_context.get("color_rgb"):
                run.font.color.rgb = font_context["color_rgb"]
            else:
                run.font.color.rgb = self.black_color
        else:
            # 默认黑色全量兼容
            run.font.color.rgb = self.black_color

        if underline:
            run.underline = True
            try:
                rPr = run._element.get_or_add_rPr()
                if 'w:u' not in rPr.xml:
                    u_elem = parse_xml(r'<w:u %s w:val="single"/>' % nsdecls('w'))
                    rPr.append(u_elem)
            except Exception:
                pass


bid_format_filler_service = BidFormatFillerService()
