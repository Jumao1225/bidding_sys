"""
Word (.docx) 导出渲染服务。

负责将结构化的投标文件格式数据转化为符合标准公文/商务样式的 .docx 文件。
遵照项目与用户最新规范：
1. 所有标题、正文、表格段落及表头文字全部统一采用纯黑色字体 (RGB: 0, 0, 0)。
2. 支持创建表格、排版标题、段落下划线与占位符高亮提示。
"""

import io
import logging
from typing import List, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

from app.schemas.bid_generator import BidFormatStructure, BidFormatSection

logger = logging.getLogger(__name__)


class DocxExporterService:
    """
    Word 文档渲染与导出服务类。
    支持纯黑字体 (RGB 0,0,0) 的标准标书模板生成。
    """

    def __init__(self):
        self.black_color = RGBColor(0, 0, 0)

    def _set_cell_margins(self, cell, top: int = 100, bottom: int = 100, left: int = 150, right: int = 150) -> None:
        """
        设置表格单元格内边距 (Padding)
        """
        try:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for m_name, m_val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
                node = OxmlElement(f'w:{m_name}')
                node.set(qn('w:w'), str(m_val))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)
        except Exception as e:
            logger.warning(f"设置单元格边距失败: {str(e)}")

    def _set_cell_background(self, cell, hex_color: str = "F5F5F5") -> None:
        """
        设置单元格背景颜色 (默认为极浅灰或透明)
        """
        try:
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
        except Exception as e:
            logger.warning(f"设置单元格背景色失败: {str(e)}")

    def _set_table_borders(self, table, color_hex: str = "CCCCCC") -> None:
        """
        为 Word 表格添加细线边框
        """
        try:
            tblPr = table._tbl.tblPr
            borders_elm = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
                f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
                f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
                f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
                f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
                f'</w:tblBorders>'
            )
            tblPr.append(borders_elm)
        except Exception as e:
            logger.warning(f"设置表格边框失败: {str(e)}")

    def export_bid_format_to_docx_bytes(self, structure: BidFormatStructure) -> bytes:
        """
        将结构化的投标文件格式数据生成为二进制 DOCX 文件流。

        :param structure: 投标文件格式全量结构数据
        :return: .docx 文件的 bytes 字节流
        """
        try:
            doc = Document()

            # ========== 1. 设置标准 A4 页边距 ==========
            section = doc.sections[0]
            section.page_width = Inches(8.27)    # A4 宽度 210mm
            section.page_height = Inches(11.69)  # A4 高度 297mm
            section.top_margin = Inches(1.0)     # 上边距 2.54cm
            section.bottom_margin = Inches(1.0)  # 下边距 2.54cm
            section.left_margin = Inches(1.25)   # 左边距 3.18cm
            section.right_margin = Inches(1.25)  # 右边距 3.18cm

            # ========== 2. 全局样式：纯黑字体 (RGB 0,0,0) ==========
            normal_style = doc.styles['Normal']
            normal_style.font.name = '宋体'
            normal_style.font.size = Pt(12)  # 小四
            normal_style.font.color.rgb = self.black_color
            normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # ========== 3. 页眉与页脚配置 ==========
            header = section.header
            hp = header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hrun = hp.add_run(f"{structure.document_title} - 投标文件格式框架")
            hrun.font.name = '宋体'
            hrun.font.size = Pt(9)
            hrun.font.color.rgb = self.black_color
            hrun._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            footer = section.footer
            fp = footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            frun = fp.add_run("--- 投标文件格式模板 (请按要求填写并盖章) ---")
            frun.font.name = '宋体'
            frun.font.size = Pt(9)
            frun.font.color.rgb = self.black_color
            frun._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # ========== 4. 文档大标题 ==========
            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_p.paragraph_format.space_before = Pt(12)
            title_p.paragraph_format.space_after = Pt(18)

            trun = title_p.add_run(structure.document_title)
            trun.font.name = '黑体'
            trun.font.size = Pt(18)  # 小二
            trun.font.bold = True
            trun.font.color.rgb = self.black_color
            trun._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

            # ========== 5. 逐个渲染提取出的格式子项 ==========
            for idx, sec in enumerate(structure.sections, 1):
                # 5.1 子项标题 (黑体三号/小三，纯黑加粗)
                sec_p = doc.add_paragraph()
                sec_p.paragraph_format.space_before = Pt(14)
                sec_p.paragraph_format.space_after = Pt(6)
                sec_p.paragraph_format.keep_with_next = True

                srun = sec_p.add_run(f"{sec.section_title}")
                srun.font.name = '黑体'
                srun.font.size = Pt(15)  # 三号
                srun.font.bold = True
                srun.font.color.rgb = self.black_color
                srun._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

                # 5.2 渲染正文或表格内容
                body_text = sec.body_markdown.strip() if sec.body_markdown else ""
                
                if body_text:
                    lines = body_text.split('\n')
                    in_table = False
                    table_rows_data = []

                    for line in lines:
                        clean_line = line.strip()
                        # 解析 Markdown 表格 (包含 | 的行)
                        if clean_line.startswith('|') and clean_line.endswith('|'):
                            # 过滤分隔行如 |---|---|
                            if set(clean_line.replace('|', '').replace('-', '').replace(':', '').strip()) == set():
                                continue
                            cells = [c.strip() for c in clean_line.split('|')[1:-1]]
                            table_rows_data.append(cells)
                            in_table = True
                        else:
                            # 若刚才在收集表格数据，先渲染表格
                            if in_table and table_rows_data:
                                self._render_table_from_data(doc, table_rows_data)
                                table_rows_data = []
                                in_table = False

                            if clean_line:
                                p = doc.add_paragraph()
                                p.paragraph_format.line_spacing = 1.5
                                p.paragraph_format.space_after = Pt(4)
                                self._add_formatted_text_to_paragraph(p, clean_line)

                    # 渲染残留未渲染的表格
                    if in_table and table_rows_data:
                        self._render_table_from_data(doc, table_rows_data)

                # 分隔空行
                doc.add_paragraph().paragraph_format.space_after = Pt(6)

            # ========== 6. 保存至 Byte 字节流并返回 ==========
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            docx_bytes = output.getvalue()
            logger.info(f"成功生成纯黑字体的 Word 标书格式文档，字节数: {len(docx_bytes)}")
            return docx_bytes

        except Exception as e:
            logger.exception(f"渲染生成 Word 文档发生严重异常: {str(e)}")
            raise RuntimeError(f"Word 文档渲染失败: {str(e)}")

    def _add_formatted_text_to_paragraph(self, paragraph, text: str) -> None:
        """
        解析段落中的下划线 (<u>...</u> 或 ______) 与斜体 (<i>...</i> 或 *...*) 标记并生成带样式的 Run
        """
        import re
        # 正则匹配 <u>...</u>, <i>...</i>, *...*, 以及连续下划线 ____
        pattern = re.compile(r'(<u>.*?</u>|<i>.*?</i>|\*[^\*]+\*|_{3,})')
        tokens = pattern.split(text)

        for token in tokens:
            if not token:
                continue

            run = paragraph.add_run()
            run.font.name = '宋体'
            run.font.size = Pt(12)
            run.font.color.rgb = self.black_color
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            if token.startswith('<u>') and token.endswith('</u>'):
                run.text = token[3:-4]
                run.underline = True
            elif token.startswith('<i>') and token.endswith('</i>'):
                run.text = token[3:-4]
                run.italic = True
            elif token.startswith('*') and token.endswith('*') and len(token) > 2:
                run.text = token[1:-1]
                run.italic = True
            elif token.startswith('_') and set(token) == {'_'}:
                run.text = token
                run.underline = True
            else:
                run.text = token

    def _render_table_from_data(self, doc: Document, rows_data: List[List[str]]) -> None:
        """
        根据表格矩阵数据渲染 Word 表格（纯黑字体、浅灰边框）
        """
        if not rows_data or not rows_data[0]:
            return

        col_count = len(rows_data[0])
        row_count = len(rows_data)

        table = doc.add_table(rows=row_count, cols=col_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table, color_hex="D3D3D3")

        for r_idx, row_cells_data in enumerate(rows_data):
            row = table.rows[r_idx]
            is_header = (r_idx == 0)

            for c_idx, cell_text in enumerate(row_cells_data):
                if c_idx < len(row.cells):
                    cell = row.cells[c_idx]
                    self._set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                    
                    if is_header:
                        self._set_cell_background(cell, hex_color="F2F2F2")

                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)

                    run = p.add_run(cell_text)
                    run.font.name = '黑体' if is_header else '宋体'
                    run.font.size = Pt(10.5)  # 五号
                    run.font.bold = is_header
                    run.font.color.rgb = self.black_color
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体' if is_header else '宋体')


# 单例初始化
docx_exporter_service = DocxExporterService()
