"""
BOM 成本测算与对标清单 Word 与 Excel 导出服务 (bom_export_service.py)

根据指定表头格式生成高保真 Word (.docx) 与 Excel (.xlsx) 文档：
表头：【序号 | 标的物名称 | 品牌、规格、型号 | 生产厂家 | 单位 | 数量 | 单价(元) | 总价(元) | 备注】
表尾：包含规范的小写与人民币大写总价合计。
"""

import io
from datetime import datetime
from typing import Any, Dict, List, Optional
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from loguru import logger
import openpyxl
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from app.utils.rmb_formatter import number_to_chinese_rmb


def _set_cell_background(cell, hex_color: str) -> None:
    """设置单元格背景颜色 (hex_color 如 'F1F5F9')"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_margins(cell, top: int = 100, bottom: int = 100, left: int = 120, right: int = 120) -> None:
    """设置单元格内边距 (单位: twips, 1pt = 20 twips)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in (('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)):
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None) -> None:
    """设置单元格边框"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    borders = {'w:top': top, 'w:bottom': bottom, 'w:left': left, 'w:right': right}
    for border_name, border_val in borders.items():
        if border_val:
            node = OxmlElement(border_name)
            node.set(qn('w:val'), border_val.get('val', 'single'))
            node.set(qn('w:sz'), str(border_val.get('sz', 4)))
            node.set(qn('w:space'), '0')
            node.set(qn('w:color'), border_val.get('color', 'CBD5E1'))
            tcBorders.append(node)
    tcPr.append(tcBorders)


def generate_bom_docx(
    document_title: str,
    items: List[Dict[str, Any]],
    total_cost: Optional[float] = None,
    budget_limit: Optional[str] = None,
    status_text: Optional[str] = None,
    analysis_summary: Optional[str] = None
) -> io.BytesIO:
    """
    生成 BOM 成本测算与对标清单的 Word (.docx) 文档。
    表头字段严格对齐：
    【序号 | 标的物名称 | 品牌、规格、型号 | 生产厂家 | 单位 | 数量 | 单价(元) | 总价(元) | 备注】

    :param document_title: 招标文件名称或项目名称
    :param items: BOM 成本明细列表
    :param total_cost: 实时预估总成本
    :param budget_limit: 最高投标限价或预算
    :param status_text: 预算控制状态
    :param analysis_summary: 专家评估指导意见
    :return: 包含 Word 二进制数据的 BytesIO 对象
    """
    logger.info(f"开始生成 BOM 成本测算 Word 文档: {document_title}, 共 {len(items)} 项")
    
    doc = Document()
    
    # 设置页边距为标准公文窄边距 (0.5 英寸左右，使 9 列排版充分舒展)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    # 计算实时总成本
    if total_cost is None:
        computed_total = 0.0
        for it in items:
            q = float(it.get('qty') or it.get('quantity') or 1)
            p = float(it.get('ref_price') or it.get('price') or 0)
            computed_total += q * p
        total_cost = computed_total

    total_cost_val = float(total_cost or 0.0)
    total_cost_upper = number_to_chinese_rmb(total_cost_val)

    # 1. 标题区
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("拟投入设备及 BOM 成本测算清单")
    title_run.font.name = "微软雅黑"
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # 2. 项目基本信息栏
    info_p = doc.add_paragraph()
    info_p.paragraph_format.space_after = Pt(10)
    info_p.paragraph_format.line_spacing = 1.2
    
    # 标书名称
    run_doc_lbl = info_p.add_run("关联招标文件：")
    run_doc_lbl.font.bold = True
    run_doc_lbl.font.size = Pt(9.5)
    run_doc_lbl.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    run_doc_val = info_p.add_run(f"{document_title}\n")
    run_doc_val.font.size = Pt(9.5)
    run_doc_val.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    # 测算时间与状态
    run_time_lbl = info_p.add_run("测算日期：")
    run_time_lbl.font.bold = True
    run_time_lbl.font.size = Pt(9.5)
    run_time_lbl.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    run_time_val = info_p.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M')}    ")
    run_time_val.font.size = Pt(9.5)
    
    if budget_limit:
        run_bgt_lbl = info_p.add_run("最高投标限价/预算：")
        run_bgt_lbl.font.bold = True
        run_bgt_lbl.font.size = Pt(9.5)
        run_bgt_lbl.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        run_bgt_val = info_p.add_run(f"{budget_limit}    ")
        run_bgt_val.font.size = Pt(9.5)

    if status_text:
        run_st_lbl = info_p.add_run("测算状态：")
        run_st_lbl.font.bold = True
        run_st_lbl.font.size = Pt(9.5)
        run_st_lbl.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        run_st_val = info_p.add_run(f"{status_text}\n")
        run_st_val.font.size = Pt(9.5)
        run_st_val.font.bold = True
        run_st_val.font.color.rgb = RGBColor(0x05, 0x96, 0x69) if "可控" in status_text else RGBColor(0xDC, 0x26, 0x26)

    # 3. BOM 明细表格（严格按图 9 列格式）
    col_widths = [0.45, 1.15, 1.30, 1.15, 0.45, 0.50, 0.75, 0.85, 1.30]
    headers = [
        "序号", "标的物名称", "品牌、规格、型号", "生产厂家", 
        "单位", "数量", "单价(元)", "总价(元)", "备注"
    ]

    total_rows = len(items) + 2  # 表头 1 行 + 数据 N 行 + 表尾合计 1 行
    table = doc.add_table(rows=total_rows, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    border_style = {'val': 'single', 'sz': 4, 'color': 'CBD5E1'}
    header_border = {'val': 'single', 'sz': 6, 'color': '94A3B8'}

    # 3.1 渲染表头
    hdr_row = table.rows[0]
    for c_idx, title in enumerate(headers):
        cell = hdr_row.cells[c_idx]
        cell.width = Inches(col_widths[c_idx])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_background(cell, "F1F5F9")
        _set_cell_margins(cell, top=120, bottom=120, left=80, right=80)
        _set_cell_borders(cell, top=header_border, bottom=header_border, left=border_style, right=border_style)
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        run.font.name = "微软雅黑"
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # 3.2 渲染数据行
    for r_idx, item in enumerate(items, start=1):
        row = table.rows[r_idx]
        name = str(item.get("name") or item.get("item_name") or "")
        
        brand = str(item.get("matched_brand") or item.get("brand") or "").strip()
        model = str(item.get("matched_model") or item.get("model") or "").strip()
        brand_spec_model_parts = []
        if brand and model:
            brand_spec_model_parts.append(f"{brand} {model}")
        elif brand:
            brand_spec_model_parts.append(brand)
        elif model:
            brand_spec_model_parts.append(model)
        else:
            brand_spec_model_parts.append("--")
        brand_spec_model_text = "\n".join(brand_spec_model_parts)

        manufacturer = str(item.get("matched_manufacturer") or item.get("manufacturer") or "--").strip()
        unit = str(item.get("unit") or "项")
        raw_qty = item.get("qty") if item.get("qty") is not None else item.get("quantity")
        qty_val = float(raw_qty) if raw_qty is not None else 1.0

        raw_price = item.get("ref_price") if item.get("ref_price") is not None else item.get("price")
        price_val = float(raw_price or 0.0)
        subtotal_val = float(item.get("subtotal") or (qty_val * price_val))

        # 备注列：严格使用前端 BOM 清单的备注 (remark) 字段
        remark_text = str(item.get("remark") or "").strip()

        is_parent = bool(item.get("isParent") or item.get("children"))
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        if is_parent:
            bg_color = "EFF6FF"

        row_data = [
            (str(r_idx), WD_ALIGN_PARAGRAPH.CENTER),
            (name, WD_ALIGN_PARAGRAPH.LEFT),
            (brand_spec_model_text, WD_ALIGN_PARAGRAPH.LEFT),
            (manufacturer, WD_ALIGN_PARAGRAPH.LEFT),
            (unit, WD_ALIGN_PARAGRAPH.CENTER),
            (f"{qty_val:g}", WD_ALIGN_PARAGRAPH.CENTER),
            (f"¥{price_val:,.2f}" if price_val > 0 else "--", WD_ALIGN_PARAGRAPH.RIGHT),
            (f"¥{subtotal_val:,.2f}" if subtotal_val > 0 else "¥0.00", WD_ALIGN_PARAGRAPH.RIGHT),
            (remark_text, WD_ALIGN_PARAGRAPH.LEFT),
        ]

        for c_idx, (text, align) in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = Inches(col_widths[c_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_background(cell, bg_color)
            _set_cell_margins(cell, top=70, bottom=70, left=70, right=70)
            _set_cell_borders(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)

            p = cell.paragraphs[0]
            p.alignment = align
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(text)
            run.font.name = "微软雅黑"
            run.font.size = Pt(8)
            if is_parent and c_idx in (1, 6, 7):
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
            else:
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # 3.3 渲染表尾统计合计行
    footer_row = table.rows[-1]
    cell_label = footer_row.cells[0]
    for c_idx in range(1, 6):
        cell_label.merge(footer_row.cells[c_idx])

    cell_unit_price = footer_row.cells[6]
    cell_amount = footer_row.cells[7]
    cell_remark = footer_row.cells[8]

    # 合计说明单元格 (前 6 列合并)
    _set_cell_background(cell_label, "F8FAFC")
    _set_cell_margins(cell_label, top=120, bottom=120, left=100, right=100)
    _set_cell_borders(cell_label, top=header_border, bottom=header_border, left=border_style, right=border_style)
    p_label = cell_label.paragraphs[0]
    p_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    run_tot_title = p_label.add_run("【合计】预估总成本：")
    run_tot_title.font.name = "微软雅黑"
    run_tot_title.font.size = Pt(9)
    run_tot_title.font.bold = True
    run_tot_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    run_tot_upper = p_label.add_run(f"人民币（大写）{total_cost_upper}")
    run_tot_upper.font.name = "微软雅黑"
    run_tot_upper.font.size = Pt(9)
    run_tot_upper.font.bold = True
    run_tot_upper.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    # 单价单元格
    _set_cell_background(cell_unit_price, "F8FAFC")
    _set_cell_margins(cell_unit_price, top=120, bottom=120, left=60, right=60)
    _set_cell_borders(cell_unit_price, top=header_border, bottom=header_border, left=border_style, right=border_style)
    p_unit = cell_unit_price.paragraphs[0]
    p_unit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_unit = p_unit.add_run("--")
    run_unit.font.size = Pt(8)
    run_unit.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    # 总价金额单元格 (小写)
    _set_cell_background(cell_amount, "EFF6FF")
    _set_cell_margins(cell_amount, top=120, bottom=120, left=80, right=80)
    _set_cell_borders(cell_amount, top=header_border, bottom=header_border, left=border_style, right=border_style)
    p_amount = cell_amount.paragraphs[0]
    p_amount.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run_tot_lower = p_amount.add_run(f"¥{total_cost_val:,.2f}")
    run_tot_lower.font.name = "微软雅黑"
    run_tot_lower.font.size = Pt(10)
    run_tot_lower.font.bold = True
    run_tot_lower.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    # 备注单元格
    _set_cell_background(cell_remark, "F8FAFC")
    _set_cell_margins(cell_remark, top=120, bottom=120, left=60, right=60)
    _set_cell_borders(cell_remark, top=header_border, bottom=header_border, left=border_style, right=border_style)
    p_rem = cell_remark.paragraphs[0]
    p_rem.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_rem = p_rem.add_run(status_text or "")
    run_rem.font.name = "微软雅黑"
    run_rem.font.size = Pt(7.5)
    run_rem.font.color.rgb = RGBColor(0x05, 0x96, 0x69) if "可控" in (status_text or "") else RGBColor(0x64, 0x74, 0x8B)

    # 4. 底部专家评估指导意见与附注
    if analysis_summary:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        expert_p = doc.add_paragraph()
        expert_p.paragraph_format.space_before = Pt(8)
        expert_p.paragraph_format.space_after = Pt(4)
        
        run_exp_title = expert_p.add_run("💡 专家评估指导意见：\n")
        run_exp_title.font.name = "微软雅黑"
        run_exp_title.font.size = Pt(9.5)
        run_exp_title.font.bold = True
        run_exp_title.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
        
        run_exp_body = expert_p.add_run(analysis_summary)
        run_exp_body.font.name = "微软雅黑"
        run_exp_body.font.size = Pt(8.5)
        run_exp_body.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    logger.info(f"成功生成 BOM Word 文档，大小: {doc_io.getbuffer().nbytes} 字节")
    return doc_io


def generate_bom_xlsx(
    document_title: str,
    items: List[Dict[str, Any]],
    total_cost: Optional[float] = None,
    budget_limit: Optional[str] = None,
    status_text: Optional[str] = None,
    analysis_summary: Optional[str] = None
) -> io.BytesIO:
    """
    生成标准 Excel (.xlsx) 工作簿文档。
    严格对齐 9 列标准格式：
    【序号 | 标的物名称 | 品牌、规格、型号 | 生产厂家 | 单位 | 数量 | 单价(元) | 总价(元) | 备注】
    包含单元格跨列合并、标准网格边框、数字货币格式与表尾大小写总价汇总。

    :param document_title: 招标文件名称
    :param items: BOM 成本明细列表
    :param total_cost: 预估总成本
    :param budget_limit: 最高限价/预算
    :param status_text: 预算状态
    :param analysis_summary: 专家指导意见
    :return: 包含 Excel 二进制数据的 BytesIO 对象
    """
    logger.info(f"开始生成 BOM 成本测算 Excel 文档: {document_title}, 共 {len(items)} 项")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "BOM成本测算清单"

    # 样式定义
    font_name = "微软雅黑"
    
    # 边框
    thin_border_side = Side(style='thin', color='CBD5E1')
    medium_border_side = Side(style='medium', color='94A3B8')
    double_border_side = Side(style='double', color='94A3B8')
    cell_border = Border(left=thin_border_side, right=thin_border_side, top=thin_border_side, bottom=thin_border_side)
    
    # 填充色
    header_fill = PatternFill(start_color='F1F5F9', end_color='F1F5F9', fill_type='solid')
    zebra_fill = PatternFill(start_color='F8FAFC', end_color='F8FAFC', fill_type='solid')
    footer_label_fill = PatternFill(start_color='F8FAFC', end_color='F8FAFC', fill_type='solid')
    footer_amount_fill = PatternFill(start_color='EFF6FF', end_color='EFF6FF', fill_type='solid')

    # 计算总成本
    if total_cost is None:
        computed_total = 0.0
        for it in items:
            q = float(it.get('qty') or it.get('quantity') or 1)
            p = float(it.get('ref_price') or it.get('price') or 0)
            computed_total += q * p
        total_cost = computed_total

    total_cost_val = float(total_cost or 0.0)
    total_cost_upper = number_to_chinese_rmb(total_cost_val)

    # 1. 标题行 (Row 1)
    ws.merge_cells('A1:I1')
    title_cell = ws['A1']
    title_cell.value = "拟投入设备及 BOM 成本测算清单"
    title_cell.font = Font(name=font_name, size=16, bold=True, color='1E293B')
    title_cell.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 36

    # 2. 项目信息行 (Row 2, 3)
    ws.merge_cells('A2:I2')
    info_cell_1 = ws['A2']
    info_cell_1.value = f"关联招标文件：{document_title}"
    info_cell_1.font = Font(name=font_name, size=9.5, color='334155', bold=True)
    info_cell_1.alignment = Alignment(horizontal='left', vertical='center')
    ws.row_dimensions[2].height = 20

    ws.merge_cells('A3:I3')
    info_cell_2 = ws['A3']
    now_str = datetime.now().strftime('%Y-%m-%d %H:%M')
    info_parts = [f"测算日期：{now_str}"]
    if budget_limit:
        info_parts.append(f"最高投标限价/预算：{budget_limit}")
    if status_text:
        info_parts.append(f"测算状态：{status_text}")
    info_cell_2.value = "   |   ".join(info_parts)
    info_cell_2.font = Font(name=font_name, size=9, color='64748B')
    info_cell_2.alignment = Alignment(horizontal='left', vertical='center')
    ws.row_dimensions[3].height = 20

    # 3. 表头行 (Row 5)
    headers = ["序号", "标的物名称", "品牌、规格、型号", "生产厂家", "单位", "数量", "单价(元)", "总价(元)", "备注"]
    ws.row_dimensions[5].height = 28

    for col_idx, header_text in enumerate(headers, start=1):
        c = ws.cell(row=5, column=col_idx, value=header_text)
        c.font = Font(name=font_name, size=9.5, bold=True, color='334155')
        c.fill = header_fill
        c.border = Border(left=thin_border_side, right=thin_border_side, top=medium_border_side, bottom=medium_border_side)
        c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    # 4. 数据行 (Row 6 ~ Row 6+N-1)
    start_row = 6
    for idx, item in enumerate(items, start=1):
        current_row = start_row + idx - 1
        ws.row_dimensions[current_row].height = 24

        name = str(item.get("name") or item.get("item_name") or "")
        
        brand = str(item.get("matched_brand") or item.get("brand") or "").strip()
        model = str(item.get("matched_model") or item.get("model") or "").strip()
        brand_parts = []
        if brand and model:
            brand_parts.append(f"{brand} {model}")
        elif brand:
            brand_parts.append(brand)
        elif model:
            brand_parts.append(model)
        else:
            brand_parts.append("--")
        brand_spec_model_text = " / ".join(brand_parts)

        manufacturer = str(item.get("matched_manufacturer") or item.get("manufacturer") or "--").strip()
        unit = str(item.get("unit") or "项")
        raw_qty = item.get("qty") if item.get("qty") is not None else item.get("quantity")
        qty_val = float(raw_qty) if raw_qty is not None else 1.0

        raw_price = item.get("ref_price") if item.get("ref_price") is not None else item.get("price")
        price_val = float(raw_price or 0.0)
        subtotal_val = float(item.get("subtotal") or (qty_val * price_val))

        # 备注列：严格使用前端 BOM 清单的备注 (remark) 字段
        remark_text = str(item.get("remark") or "").strip()

        # 写入 9 列数据
        c1 = ws.cell(row=current_row, column=1, value=idx)
        c2 = ws.cell(row=current_row, column=2, value=name)
        c3 = ws.cell(row=current_row, column=3, value=brand_spec_model_text)
        c4 = ws.cell(row=current_row, column=4, value=manufacturer)
        c5 = ws.cell(row=current_row, column=5, value=unit)
        c6 = ws.cell(row=current_row, column=6, value=qty_val)
        c7 = ws.cell(row=current_row, column=7, value=price_val)
        c8 = ws.cell(row=current_row, column=8, value=subtotal_val)
        c9 = ws.cell(row=current_row, column=9, value=remark_text)

        # 格式与对齐
        c1.alignment = Alignment(horizontal='center', vertical='center')
        c2.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
        c3.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
        c4.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
        c5.alignment = Alignment(horizontal='center', vertical='center')
        c6.alignment = Alignment(horizontal='center', vertical='center')
        c7.alignment = Alignment(horizontal='right', vertical='center')
        c8.alignment = Alignment(horizontal='right', vertical='center')
        c9.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)

        c6.number_format = '#,##0.##'
        c7.number_format = '¥#,##0.00'
        c8.number_format = '¥#,##0.00'

        is_zebra = (idx % 2 == 0)
        for c in (c1, c2, c3, c4, c5, c6, c7, c8, c9):
            c.font = Font(name=font_name, size=9, color='1E293B')
            c.border = cell_border
            if is_zebra:
                c.fill = zebra_fill

    # 5. 表尾统计行 (Row 6+N)
    footer_row_idx = start_row + len(items)
    ws.row_dimensions[footer_row_idx].height = 30

    # 合并 A ~ F 列 (Col 1 ~ 6)
    ws.merge_cells(start_row=footer_row_idx, start_column=1, end_row=footer_row_idx, end_column=6)
    foot_label_cell = ws.cell(row=footer_row_idx, column=1)
    foot_label_cell.value = f"【合计】预估总成本：人民币（大写）{total_cost_upper}"
    foot_label_cell.font = Font(name=font_name, size=9.5, bold=True, color='1D4ED8')
    foot_label_cell.alignment = Alignment(horizontal='left', vertical='center')
    foot_label_cell.fill = footer_label_fill

    # 单价列 G (Col 7)
    foot_unit_cell = ws.cell(row=footer_row_idx, column=7, value="--")
    foot_unit_cell.font = Font(name=font_name, size=9, color='94A3B8')
    foot_unit_cell.alignment = Alignment(horizontal='center', vertical='center')
    foot_unit_cell.fill = footer_label_fill

    # 总价列 H (Col 8)
    foot_amount_cell = ws.cell(row=footer_row_idx, column=8, value=total_cost_val)
    foot_amount_cell.font = Font(name=font_name, size=11, bold=True, color='1D4ED8')
    foot_amount_cell.alignment = Alignment(horizontal='right', vertical='center')
    foot_amount_cell.number_format = '¥#,##0.00'
    foot_amount_cell.fill = footer_amount_fill

    # 备注列 I (Col 9)
    foot_rem_cell = ws.cell(row=footer_row_idx, column=9, value=status_text or "")
    foot_rem_cell.font = Font(name=font_name, size=8.5, color='059669' if "可控" in (status_text or "") else '64748B')
    foot_rem_cell.alignment = Alignment(horizontal='left', vertical='center')
    foot_rem_cell.fill = footer_label_fill

    # 为表尾所有合并及独立单元格添加边框
    footer_border = Border(left=thin_border_side, right=thin_border_side, top=medium_border_side, bottom=double_border_side)
    for c_i in range(1, 10):
        cell_node = ws.cell(row=footer_row_idx, column=c_i)
        cell_node.border = footer_border

    # 6. 专家评估建议 (Row 6+N+2)
    if analysis_summary:
        exp_row_idx = footer_row_idx + 2
        ws.merge_cells(start_row=exp_row_idx, start_column=1, end_row=exp_row_idx, end_column=9)
        exp_cell = ws.cell(row=exp_row_idx, column=1)
        exp_cell.value = f"💡 专家评估指导意见：{analysis_summary}"
        exp_cell.font = Font(name=font_name, size=9, italic=True, color='475569')
        exp_cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
        ws.row_dimensions[exp_row_idx].height = 26

    # 7. 自动设置列宽
    col_width_defaults = {
        1: 8,   # 序号
        2: 24,  # 标的物名称
        3: 26,  # 品牌、规格、型号
        4: 24,  # 生产厂家
        5: 8,   # 单位
        6: 10,  # 数量
        7: 15,  # 单价(元)
        8: 18,  # 总价(元)
        9: 35,  # 备注
    }
    for col_i, default_w in col_width_defaults.items():
        col_letter = get_column_letter(col_i)
        ws.column_dimensions[col_letter].width = default_w

    doc_io = io.BytesIO()
    wb.save(doc_io)
    doc_io.seek(0)

    logger.info(f"成功生成 BOM Excel 文档，大小: {doc_io.getbuffer().nbytes} 字节")
    return doc_io
