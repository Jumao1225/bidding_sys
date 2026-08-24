"""
DocxSkillService - 封装 Codex Documents Skill (26.723.12215) 高级工具链服务

功能说明：
1. 接受全文修订痕迹 (accept_tracked_changes)：移除文档中的删除与替换红线标记，生成干净正文。
2. 自动插入与更新目录 (insert_table_of_contents)：在占位符或头部自动添加 Word TOC 动态目录。
3. 隐私与元数据清洗 (scrub_privacy_metadata)：抹去文档包含的作者信息、历史版本与隐藏编辑记录。
4. 提取全文批注 (extract_comments)：解析 DOCX 中插入的审阅批注并返回结构化列表。
"""

import io
import json
import os
import tempfile
import zipfile
import re
from typing import List, Dict, Any, Optional
from loguru import logger
from lxml import etree
from docx import Document
from docx.oxml import OxmlElement

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"

NS = {
    "w": W_NS,
    "rel": REL_NS,
    "ct": CT_NS,
    "cp": CP_NS,
    "dc": DC_NS
}


class DocxSkillService:
    """ 高级 Word 文档技能工具服务类 """

    def __init__(self):
        pass

    def accept_tracked_changes(self, docx_bytes: bytes) -> bytes:
        """
        接受全文修订痕迹 (w:ins / w:del / w:moveTo / w:moveFrom)，返回干净无痕迹的 Word 二进制数据
        """
        if not docx_bytes:
            logger.warning("接收到的 docx_bytes 为空，直接返回")
            return docx_bytes

        try:
            input_stream = io.BytesIO(docx_bytes)
            output_stream = io.BytesIO()

            with zipfile.ZipFile(input_stream, "r") as zin:
                with zipfile.ZipFile(output_stream, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                    for item in zin.infolist():
                        data = zin.read(item.filename)
                        if item.filename == "word/document.xml" or item.filename.startswith("word/header") or item.filename.startswith("word/footer"):
                            try:
                                root = etree.fromstring(data)
                                self._process_accept_revisions(root)
                                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                            except Exception as e:
                                logger.warning(f"处理修订痕迹 XML {item.filename} 失败: {e}")
                        zout.writestr(item, data)

            logger.info("成功完成 Word 文档修订痕迹全量接受处理")
            return output_stream.getvalue()

        except Exception as err:
            logger.exception("接受 Word 修订痕迹过程抛出异常")
            raise err

    def _process_accept_revisions(self, root: etree._Element) -> None:
        """ 遍历 XML 节点解包插入项并移除删除项 """
        # 处理插入节点 w:ins 及移动目标 w:moveTo -> 保留内部子节点
        for tag in (f"{{{W_NS}}}ins", f"{{{W_NS}}}moveTo"):
            for el in list(root.iter(tag)):
                parent = el.getparent()
                if parent is not None:
                    idx = parent.index(el)
                    for child in list(el):
                        parent.insert(idx, child)
                        idx += 1
                    parent.remove(el)

        # 处理删除节点 w:del 及移动来源 w:moveFrom -> 直接彻底删除
        for tag in (f"{{{W_NS}}}del", f"{{{W_NS}}}moveFrom"):
            for el in list(root.iter(tag)):
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)

    def insert_table_of_contents(self, docx_bytes: bytes, placeholder: str = "[[TOC]]") -> bytes:
        """
        在 Word 文档开头或包含 [[TOC]] 的占位符处自动插入 Word 动态目录域代码，
        并配置 updateFields=true，确保在 Word 打开时自动更新目录。
        """
        if not docx_bytes:
            return docx_bytes

        try:
            doc = Document(io.BytesIO(docx_bytes))
            toc_inserted = False

            # 检索是否包含占位符段落
            for p in doc.paragraphs:
                if placeholder in p.text:
                    p.text = "" # 清空占位符
                    self._add_toc_field_to_paragraph(p)
                    toc_inserted = True
                    logger.info(f"在占位符 {placeholder} 处成功插入目录")
                    break

            # 若未指定占位符或未找到，在文档顶部插一新段落并放入 TOC
            if not toc_inserted:
                first_p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
                new_p = first_p.insert_paragraph_before()
                # 添加标题
                title_p = first_p.insert_paragraph_before("目 录")
                title_p.style = "Heading 1"
                self._add_toc_field_to_paragraph(new_p)
                logger.info("已在文档开头成功追加动态目录域")

            out_stream = io.BytesIO()
            doc.save(out_stream)
            modified_bytes = out_stream.getvalue()

            # 修改 word/settings.xml，写入 updateFields 标志
            final_bytes = self._enable_update_fields_on_open(modified_bytes)
            return final_bytes

        except Exception as err:
            logger.exception("插入文档目录抛出异常")
            raise err

    def _add_toc_field_to_paragraph(self, p) -> None:
        """ 构建 Word TOC 域复合 XML (fldChar begin -> fldSimple/instrText -> fldChar end) """
        run = p.add_run()
        r_el = run._r
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(etree.QName(W_NS, 'fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(etree.QName('http://www.w3.org/XML/1998/namespace', 'space'), 'preserve')
        instrText.text = r'TOC \o "1-3" \h \z \u'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(etree.QName(W_NS, 'fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(etree.QName(W_NS, 'fldCharType'), 'end')

        r_el.append(fldChar1)
        r_el.append(instrText)
        r_el.append(fldChar2)
        r_el.append(fldChar3)

    def _enable_update_fields_on_open(self, docx_bytes: bytes) -> bytes:
        """ 给 word/settings.xml 追加 <w:updateFields w:val="true"/> """
        input_stream = io.BytesIO(docx_bytes)
        output_stream = io.BytesIO()

        with zipfile.ZipFile(input_stream, "r") as zin:
            with zipfile.ZipFile(output_stream, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                has_settings = False
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == "word/settings.xml":
                        has_settings = True
                        try:
                            root = etree.fromstring(data)
                            uf = root.find(f"{{{W_NS}}}updateFields")
                            if uf is None:
                                uf = etree.Element(f"{{{W_NS}}}updateFields", nsmap={"w": W_NS})
                                uf.set(f"{{{W_NS}}}val", "true")
                                root.insert(0, uf)
                            data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                        except Exception as e:
                            logger.warning(f"设置 updateFields 失败: {e}")
                    zout.writestr(item, data)

                if not has_settings:
                    # 手动创建一个基础 settings.xml
                    settings_xml = f'<w:settings xmlns:w="{W_NS}"><w:updateFields w:val="true"/></w:settings>'
                    zout.writestr("word/settings.xml", settings_xml.encode("utf-8"))

        return output_stream.getvalue()

    def scrub_privacy_metadata(self, docx_bytes: bytes) -> bytes:
        """
        擦除 Word 文档内置隐秘元数据（作者 creator、修改人 lastModifiedBy、rsid 痕迹 ID 等）
        """
        if not docx_bytes:
            return docx_bytes

        try:
            input_stream = io.BytesIO(docx_bytes)
            output_stream = io.BytesIO()

            with zipfile.ZipFile(input_stream, "r") as zin:
                with zipfile.ZipFile(output_stream, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                    for item in zin.infolist():
                        # 跳过自定义属性部分
                        if item.filename == "docProps/custom.xml":
                            continue

                        data = zin.read(item.filename)

                        # 清理核心属性核心字段
                        if item.filename == "docProps/core.xml":
                            try:
                                root = etree.fromstring(data)
                                for tag in ("creator", "lastModifiedBy"):
                                    for el in list(root.iter()):
                                        if el.tag.endswith(tag):
                                            el.text = ""
                                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                            except Exception as e:
                                logger.warning(f"清理 core.xml 属性失败: {e}")

                        # 清理正文及页眉页脚等的 rsid 离散标记
                        elif item.filename == "word/document.xml" or item.filename.startswith("word/header") or item.filename.startswith("word/footer"):
                            try:
                                root = etree.fromstring(data)
                                for el in root.iter():
                                    attrs_to_del = [k for k in el.attrib if "rsid" in k.lower()]
                                    for k in attrs_to_del:
                                        del el.attrib[k]
                                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                            except Exception as e:
                                logger.warning(f"清洗 rsid 失败: {e}")

                        zout.writestr(item, data)

            logger.info("成功完成 Word 文档隐私与元数据脱敏清洗")
            return output_stream.getvalue()

        except Exception as err:
            logger.exception(" Word 隐私清洗失败")
            raise err

    def extract_comments(self, docx_bytes: bytes) -> List[Dict[str, Any]]:
        """
        抓取 Word 文档 (word/comments.xml) 中的所有审阅批注信息
        返回结构化字典列表: [{"id": "1", "author": "[批注人]", "date": "YYYY-MM-DD", "text": "[批注内容]"}]
        """
        comments: List[Dict[str, Any]] = []
        if not docx_bytes:
            return comments

        try:
            input_stream = io.BytesIO(docx_bytes)
            with zipfile.ZipFile(input_stream, "r") as zin:
                if "word/comments.xml" not in zin.namelist():
                    logger.info("文档不包含 word/comments.xml，返回空批注列表")
                    return comments

                comments_xml = zin.read("word/comments.xml")
                root = etree.fromstring(comments_xml)

                for comment_el in root.findall(f"{{{W_NS}}}comment"):
                    c_id = comment_el.get(f"{{{W_NS}}}id", "")
                    author = comment_el.get(f"{{{W_NS}}}author", "未知作者")
                    date_str = comment_el.get(f"{{{W_NS}}}date", "")
                    
                    # 提取批注文字
                    text_parts = [t.text for t in comment_el.findall(f".//{{{W_NS}}}t") if t.text]
                    full_text = "".join(text_parts)

                    comments.append({
                        "id": c_id,
                        "author": author,
                        "date": date_str,
                        "text": full_text
                    })

            logger.info(f"成功提取到 {len(comments)} 条审阅批注")
            return comments

        except Exception as err:
            logger.exception("提取 Word 批注失败")
            return comments

    def strip_comments(self, docx_bytes: bytes) -> bytes:
        """
        彻底删除 Word 文档中的所有批注 (移除 word/comments.xml / word/commentsExtended.xml 节点与正文批注引用标记)
        """
        if not docx_bytes:
            return docx_bytes

        try:
            input_stream = io.BytesIO(docx_bytes)
            output_stream = io.BytesIO()

            with zipfile.ZipFile(input_stream, "r") as zin:
                with zipfile.ZipFile(output_stream, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                    for item in zin.infolist():
                        # 跳过批注相关 XML 文件
                        if item.filename in ("word/comments.xml", "word/commentsExtended.xml"):
                            continue

                        data = zin.read(item.filename)

                        # 清理正文及页眉页脚中的 commentRangeStart / commentRangeEnd / commentReference 标记
                        if item.filename == "word/document.xml" or item.filename.startswith("word/header") or item.filename.startswith("word/footer"):
                            try:
                                root = etree.fromstring(data)
                                for tag in ("commentRangeStart", "commentRangeEnd", "commentReference"):
                                    for el in list(root.iter(f"{{{W_NS}}}{tag}")):
                                        parent = el.getparent()
                                        if parent is not None:
                                            parent.remove(el)
                                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                            except Exception as e:
                                logger.warning(f"清除 XML 批注标记 {item.filename} 失败: {e}")

                        zout.writestr(item, data)

            logger.info("成功完成 Word 文档批注清空与剔除")
            return output_stream.getvalue()

        except Exception as err:
            logger.exception("清空 Word 批注失败")
            raise err


docx_skill_service = DocxSkillService()

